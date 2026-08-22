"""Проект НДВ: структура как у реального проекта (Методика № 581), gaps, API.

Эталоны структуры — «Книга 2. НДВ паркинг» (ООО «Альянс Консалтинг», 2025)
и «Проект НДВ ООО ПРОТЕЛЮКС» (2024–2030); требования экспертизы — заключение
ФБУЗ «Центр гигиены и эпидемиологии в г. Москве» № 77.01.06.Т.003352.10.25.
"""
import base64
import io
from decimal import Decimal

import openpyxl
import pytest
from docx import Document

from ecodoc.core import workspace
from ecodoc.core.models import Medium, NVOSObject, Pollutant, ReportContext
from ecodoc.development import ndv_project, volume_builder as vb
from ecodoc.gui import server

# дословные заголовки разделов реального проекта (содержание «Книги 2»)
REAL_SECTIONS = [
    "СОСТАВ ПРОЕКТА", "СВЕДЕНИЯ О РАЗРАБОТЧИКЕ И СПИСОК ИСПОЛНИТЕЛЕЙ", "АННОТАЦИЯ",
    "СОДЕРЖАНИЕ", "ВВЕДЕНИЕ",
    "1. ОБЩИЕ СВЕДЕНИЯ О ПРЕДПРИЯТИИ",
    "2. ХАРАКТЕРИСТИКА ПРЕДПРИЯТИЯ КАК ИСТОЧНИКА ЗАГРЯЗНЕНИЯ АТМОСФЕРЫ",
    "2.1 Краткая характеристика технологии производства и технологического оборудования",
    "2.2 Краткая характеристика существующих установок очистки газа",
    "2.3 Перспектива развития предприятия",
    "2.4 Перечень загрязняющих веществ, выбрасываемых в атмосферу",
    "2.5 Характеристика аварийных и залповых выбросов",
    "2.6 Параметры выбросов загрязняющих веществ в атмосферный воздух",
    "2.7 Обоснование полноты и достоверности исходных данных",
    "3. ПРОВЕДЕНИЕ РАСЧЕТОВ И ОПРЕДЕЛЕНИЕ НОРМАТИВОВ ДОПУСТИМЫХ ВЫБРОСОВ",
    "3.1 Критерии качества атмосферного воздуха",
    "3.2 Метеорологические характеристики и коэффициенты",
    "3.3 Результаты расчетов уровней загрязнения атмосферы",
    "3.3.2 Результаты расчетов максимальных приземных концентраций",
    "3.4 Мероприятия по снижению негативного воздействия выбросов",
    "3.5 Предложения по нормативам допустимых выбросов",
    "4. МЕРОПРИЯТИЯ ПО РЕГУЛИРОВАНИЮ ВЫБРОСОВ ПРИ НЕБЛАГОПРИЯТНЫХ МЕТЕОРОЛОГИЧЕСКИХ УСЛОВИЯХ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ЛИТЕРАТУРНЫХ ИСТОЧНИКОВ", "ПРИЛОЖЕНИЯ",
]
# таблицы реального проекта (номер + начало названия дословно)
REAL_TABLES = [
    "Таблица 2.4.1 – Перечень загрязняющих веществ, выбрасываемых в атмосферу",
    "Таблица 2.4.2 – Перечень загрязняющих веществ, подлежащих и не подлежащих государственному учету",
    "Таблица 2.4.3 – Перечень загрязняющих веществ, выбрасываемых в атмосферу и подлежащих государственному учету и нормированию",
    "Таблица 2.4.4 – Перечень источников выбросов и загрязняющих веществ, не подлежащих государственному учету",
    "Таблица 2.6 – Параметры выбросов загрязняющих веществ для расчета загрязнения атмосферы",
    "Таблица 3.2.1 – Метеорологические характеристики, коэффициенты, определяющие условия рассеивания",
    "Таблица 3.3 – Описание расчетной площадки",
    "Таблица 3.4 – Координаты и расположение расчетных точек",
    "Таблица 3.5.1 – Перечень источников, дающих наибольшие вклады в уровень загрязнения атмосферы (без учета фонового загрязнения). Расчет по максимально разовым концентрациям",
    "Таблица 3.5.2 – Перечень источников, дающих наибольшие вклады в уровень загрязнения атмосферы (без учета фонового загрязнения). Расчет по среднегодовым концентрациям",
    "Таблица 3.5.3 – Перечень источников, дающих наибольшие вклады в уровень загрязнения атмосферы (с учетом фонового загрязнения)",
    "Таблица 3.7 – Выбросы загрязняющих веществ на СП и срок достижения НДВ",
    "Таблица 3.8 – Нормативы выбросов загрязняющих веществ в атмосферный воздух по конкретным стационарным источникам выбросов и загрязняющим веществам",
    "Таблица 3.9 – Нормативы выбросов загрязняющих веществ в атмосферный воздух по объекту ОНВ",
    "Таблица 4.1 – Результаты расчета концентраций загрязняющих веществ, для обоснования перечня загрязняющих веществ, для которых производится уменьшение выбросов в период НМУ",
]
# приложения реального проекта (смысловые ключи)
REAL_APPENDICES = ["Ситуационный план", "перспективе развития", "аварийных и залповых",
                   "постановке на", "Исходные данные", "климатических", "фоновых",
                   "Расчет рассеивания", "ЕГРН", "Расчет выбросов", "СЗЗ", "НМУ"]


def _full_ctx() -> ReportContext:
    c = ReportContext()
    o = c.organization
    o.name = "Общество с ограниченной ответственностью «ПРОТЕЛЮКС»"
    o.short_name = "ООО «ПРОТЕЛЮКС»"
    o.inn, o.kpp, o.ogrn, o.okved, o.okpo = "4707038091", "470701001", "1154707000854", "10.89.4", "25853809"
    o.address = "188490, Ленинградская область, г. Ивангород, ул. Лесная, зд. 13А/1"
    o.director_name = "Брезгин М.Ю."
    c.period.year = 2025
    c.objects = [NVOSObject(code="41-0147-001234-П", name="Завод по производству биопротеина",
                            category="II", address="ЛО, г. Ивангород, ул. Лесная")]
    c.extra["emission_sources"] = [
        {"number": "0005", "name": "Котел 1", "kind": "организованный", "type": "Точечный",
         "workshop": "Котельная", "height": "37", "diameter": "0.35", "x1": "1266325.78",
         "y1": "371867.17", "speed": "3.51", "volume": "0.338", "temperature": "151",
         "hours_year": "8760",
         "method": "Методика определения выбросов ЗВ при сжигании топлива в котлах < 30 т/ч, 1999",
         "pollutants": [{"code": "0301", "name": "Азота диоксид", "conc": "1114.65",
                         "g_s": "0.2425783", "t_year": "0.728728"},
                        {"code": "0337", "name": "Углерода оксид", "g_s": "0.6188221",
                         "t_year": "2.227768"}]},
        {"number": "6002", "name": "Стоянка автотранспорта", "kind": "неорганизованный",
         "x1": "1266300", "y1": "371800", "x2": "1266320", "y2": "371800", "area_width": "10",
         "height": "2", "hours_year": "2000",
         "pollutants": [{"code": "0301", "name": "Азота диоксид", "g_s": "0.001", "t_year": "0.002"},
                        {"code": "0328", "name": "Углерод (Пигмент черный)", "g_s": "0.0001",
                         "t_year": "0.0005"}]},
    ]
    for code, name, mass in (("0301", "Азота диоксид", "0.730728"),
                             ("0337", "Углерода оксид", "2.227768"),
                             ("0328", "Углерод (Пигмент черный)", "0.0005")):
        p = Pollutant(code=code, name=name, mass_norm=Decimal(mass))
        p.medium = Medium.AIR
        c.pollutants.append(p)
    c.extra["ndv"] = {
        "developer": {"name": "ООО «Альянс Консалтинг»", "address": "СПб, пр. Большевиков, 7",
                      "director": "Краев Д.В."},
        "executors": [{"position": "инженер-эколог", "name": "Маляева Я.В."}],
        "years": [2025, 2026, 2027, 2028, 2029, 2030, 2031],
        "meteo": {"A": "160", "t_hot": "23,8", "t_cold": "-12,2", "u_star": "6,0", "relief": "1",
                  "wind_rose": {"С": 7, "СВ": 16, "В": 16, "ЮВ": 7, "Ю": 14, "ЮЗ": 18, "З": 17, "СЗ": 5}},
        "background": [{"code": "0301", "name": "Азота диоксид", "value": "0,055"}],
        "calc_points": [{"code": 1, "x": "489,70", "y": "570,10", "kind": "на границе жилой зоны",
                         "comment": "Многоквартирный дом, ул. Лесная, 5"}],
        "location": "Участок 47:21:0302003:258; с севера — производственная зона ПР-1.",
        "szz": {"class": "III классу", "size": "300", "status": "СЗЗ установлена (СЭЗ прилагается)."},
        "technology": "Завод по производству биопротеина: ферментация, сушка, котельная.",
        "gou": "PLAZKAT OXIZ 40000, эффективность 98,5 %.",
        "prospects": "Изменений до 2031 г. не планируется.",
        "emergency": "Аварийные и залповые выбросы отсутствуют.",
        "upraza": {"name": "УПРЗА «Эколог»", "version": "4.70", "rect": "3065×3065 м", "step": "100"},
        "summation_groups": [{"code": "6204", "members": "0301 Азота диоксид + 0330 Сера диоксид", "k": "1,6"}],
        "control_schedule": [{"source": "0005", "source_name": "Котел 1", "code": "0301",
                              "substance": "Азота диоксид", "norm_gs": "0,2425783",
                              "period": "1 раз в год", "method": "инструментальный",
                              "who": "аккредитованная лаборатория"}],
        "responsible": "инженер-эколог Маляева Я.В.",
        "results_analysis": "Максимальная концентрация азота диоксида на границе жилой зоны — 0,74 ПДК "
                            "с учетом фона; по остальным веществам < 0,1 ПДК.",
    }
    return c


def _text(path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def _xlsx_b64(sheets: dict[str, tuple[list, list]]) -> str:
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for title, (header, rows) in sheets.items():
        ws = wb.create_sheet(title)
        ws.append(header)
        for r in rows:
            ws.append(r)
    buf = io.BytesIO()
    wb.save(buf)
    return base64.b64encode(buf.getvalue()).decode()


# ── структура ────────────────────────────────────────────────────────────────

def test_sections_and_tables_match_real_project(tmp_path):
    path = ndv_project.build(_full_ctx(), ndv_project.NdvInputs(), tmp_path / "ндв.docx")
    text = _text(path)
    for s in REAL_SECTIONS:
        assert s in text, f"нет раздела: {s}"
    for t in REAL_TABLES:
        assert t in text, f"нет таблицы: {t}"
    for a in REAL_APPENDICES:
        assert a in text, f"нет приложения: {a}"
    # требования ФБУЗ/Методики сверх паркинга: контроль, застройка, группы суммации, фон
    assert "5. КОНТРОЛЬ ЗА СОБЛЮДЕНИЕМ НОРМАТИВОВ" in text
    assert "6. РАСЧЕТЫ ЗАГРЯЗНЕНИЯ АТМОСФЕРЫ С УЧЕТОМ ЗАСТРОЙКИ" in text
    assert "Таблица 2.4.5" in text and "6204" in text
    assert "Фоновые концентрации" in text and "0,055" in text


def test_table_headers_verbatim():
    """Шапки таблиц — дословно по формам Методики № 581 и «ПДВ-Эколог»."""
    h26 = ndv_project.HEADERS["2.6"]
    for col in ("Высота источника, м", "Диаметр устья, м", "Скорость выхода ГВС, м/с",
                "Температура ГВС, °C", "Наименование ГОУ", "Мощность выброса, г/с",
                "Валовый выброс, т/год"):
        assert col in h26
    assert ndv_project.HEADERS["3.4"] == ["Код", "X, м", "Y, м", "Высота, м", "Тип точки", "Комментарий"]
    assert "Процент вклада" in ndv_project.HEADERS["3.5.1"]
    assert ndv_project.HEADERS["3.2.1"] == ["Наименование характеристик", "Величина"]
    assert "Класс опасности" in ndv_project.HEADERS["2.4.1"]


def test_pollutant_table_has_pdk_class_and_totals(tmp_path):
    path = ndv_project.build(_full_ctx(), ndv_project.NdvInputs(), tmp_path / "ндв.docx")
    text = _text(path)
    # ПДК м/р азота диоксида 0,2, класс 3; сумма по 0301 = 0,2425783 + 0,001
    assert "0,2" in text and "0,2435783" in text
    assert "Всего веществ:" in text and "в том числе твердых: 1" in text


def test_norm_tables_by_years(tmp_path):
    ctx = _full_ctx()
    path = ndv_project.build(ctx, ndv_project.NdvInputs(), tmp_path / "ндв.docx")
    text = _text(path)
    for y in (2025, 2031):
        assert f"{y} г., г/с" in text and f"{y} г., т/год" in text
    assert "НДВ/ВРВ" in text and "Всего по объекту:" in text
    assert ndv_project.norm_years(ctx) == list(range(2025, 2032))
    ctx.extra["ndv"].pop("years")
    assert ndv_project.norm_years(ctx) == list(range(2025, 2032))   # 7 лет с отчётного


def test_ecolog_exports_land_in_right_tables(tmp_path):
    src = ndv_project.NdvInputs(
        dispersion_header=["Код", "Вещество", "Cмакс, доли ПДК"],
        dispersion_table=[["0301", "Азота диоксид", 0.74]],
        points_header=["Код", "X", "Y", "Высота", "Тип", "Комментарий"],
        points_table=[[1, 100.1, 28.4, 2, "на границе жилой зоны", "д. 5"]],
        contrib_header=["Код", "Вещество", "Cмакс", "Площ.", "Цех", "Источн.", "Цех", "%", "X", "Y"],
        contrib_table=[["0301", "Азота диоксид", 0.74, 1, 1, "0005", "Котельная", 100, 1, 2]],
        appendices=["ситплан.pdf"])
    text = _text(ndv_project.build(_full_ctx(), src, tmp_path / "ндв.docx"))
    assert "0.74" in text and "д. 5" in text and "Котельная" in text
    assert "[место для вставки: ситплан.pdf]" in text


def test_typical_texts_are_marked(tmp_path):
    text = _text(ndv_project.build(ReportContext(), ndv_project.NdvInputs(), tmp_path / "п.docx"))
    assert "(типовая формулировка)" in text
    assert "[требуется:" in text


# ── gaps ─────────────────────────────────────────────────────────────────────

def test_gaps_empty_context_lists_everything():
    g = ndv_project.gaps(ReportContext(), ndv_project.NdvInputs())
    joined = " | ".join(g)
    for key in ("организации", "объект ОНВ", "источников выбросов", "перечня загрязняющих",
                "период нормирования", "метеохарактеристик", "фоновых", "расчётных точек",
                "результатов расчёта рассеивания", "вкладов", "разработчике", "СЗЗ",
                "газоочистных", "плана-графика контроля"):
        assert key in joined, key


def test_gaps_full_context_only_upraza_left():
    g = ndv_project.gaps(_full_ctx(), ndv_project.NdvInputs())
    assert all(("рассеивания" in x) or ("вкладов" in x) for x in g), g
    src = ndv_project.NdvInputs(dispersion_table=[[1]], contrib_table=[[1]])
    assert ndv_project.gaps(_full_ctx(), src) == []


def test_gaps_source_without_parameters():
    ctx = _full_ctx()
    ctx.extra["emission_sources"][0].pop("diameter")
    g = ndv_project.gaps(ctx, ndv_project.NdvInputs())
    assert any("№0005" in x and "diameter" in x for x in g)


# ── volume_builder: делегирование и разбор книги «Эколога» ───────────────────

def test_volume_builder_delegates_and_keeps_other_types(tmp_path):
    ctx = _full_ctx()
    src = vb.VolumeSources(points_table=[[1, 1, 2, 2, "жилая", ""]])
    assert "Таблица 3.8" in _text(vb.build("ndv", ctx, src, tmp_path / "a.docx"))
    for vt, title in (("nds", "НОРМАТИВОВ ДОПУСТИМЫХ СБРОСОВ"), ("szz", "САНИТАРНО-ЗАЩИТНОЙ ЗОНЫ")):
        assert title in _text(vb.build(vt, ctx, src, tmp_path / f"{vt}.docx"))
    assert vb.gaps("ndv", ctx, src)
    assert vb.gaps("nds", ctx, src) == ["нет таблицы источников", "нет результатов расчёта/обоснования"]


def test_ingest_ecolog_workbook_classifies_sheets(tmp_path):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for title, header in (("Максимальные концентрации", ["Код", "Вещество", "Cмакс"]),
                          ("Расчетные точки", ["Код", "X", "Y", "Высота", "Тип", "Комментарий"]),
                          ("Вклады источников", ["Код", "Вещество", "Источник", "%"]),
                          ("Лист1", ["что-то", "ещё"])):
        ws = wb.create_sheet(title)
        ws.append(header)
        ws.append([1, 2, 3])
    p = tmp_path / "res.xlsx"
    wb.save(p)
    src = vb.VolumeSources()
    notes = vb.ingest_ecolog_workbook(p, src)
    assert src.dispersion_table and src.points_table and src.contrib_table
    assert any("не распознан" in n for n in notes)


# ── API ──────────────────────────────────────────────────────────────────────

@pytest.fixture()
def site(tmp_path, monkeypatch):
    monkeypatch.setenv("ECODOC_RESULTS", str(tmp_path / "res"))
    workspace.add_org("ОРГ")
    workspace.add_site("ОРГ", "Пл")
    ctx = workspace.load_context("ОРГ", "Пл")
    full = _full_ctx()
    ctx.organization = full.organization
    ctx.objects = full.objects
    ctx.period.year = 2025
    ctx.pollutants = full.pollutants
    ctx.extra.update(full.extra)
    workspace.save_context("ОРГ", "Пл", ctx)
    return ("ОРГ", "Пл")


def test_api_volume_ndv_full_project_with_multisheet_results(site):
    org, st = site
    out = server.api_volume({}, {
        "org": org, "site": st, "vtype": "ndv",
        "dispersion_file": {"name": "результаты.xlsx", "b64": _xlsx_b64({
            "Максимальные концентрации": (["Код", "Вещество", "Cмакс, доли ПДК"],
                                          [["0301", "Азота диоксид", 0.74]]),
            "Расчетные точки": (["Код", "X", "Y", "Высота", "Тип", "Комментарий"],
                                [[1, 10, 20, 2, "жилая зона", "дом 5"]]),
            "Вклады источников": (["Код", "Вещество", "Cмакс", "Источник", "%"],
                                  [["0301", "Азота диоксид", 0.74, "0005", 100]])})}})
    assert out["path"].endswith("проект_НДВ_2025.docx")
    text = _text(out["path"])
    assert "ПРОЕКТ НОРМАТИВОВ ДОПУСТИМЫХ ВЫБРОСОВ" in text
    assert "0.74" in text and "дом 5" in text
    assert out["gaps"] == []
    assert not any(n.startswith("⚠") for n in out["notes"])


def test_api_volume_ndv_without_upraza_reports_gaps(site):
    org, st = site
    out = server.api_volume({}, {"org": org, "site": st, "vtype": "ndv"})
    assert any("рассеивания нет" in n for n in out["notes"])
    assert any("рассеивания" in g for g in out["gaps"])
    assert "Котел 1" in _text(out["path"])              # таблица 2.6 из инвентаризации
