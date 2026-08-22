"""Тесты полноценного Плана мероприятий НМУ (приказы МПР № 651 и № 662).

Проверяем: форма строго по рекомендуемому образцу № 662 (ровно 9 пунктов,
дословные названия, строка номеров граф «1…6»), перечень п. 7 только по
контролируемым веществам (пп. 3–4 № 651), графа 6 НЕ считается по процентам
№ 651 (они — про вклады в приземные концентрации), запреты — дословно п. 9
№ 651, числа с запятой и 7 знаками, пометки «[требуется: …]» = gaps().
"""
from decimal import Decimal

from ecodoc.core.models import (Medium, NVOSObject, Organization, Pollutant,
                                ReportContext)
from ecodoc.development import nmu


def _ctx(code="41-0247-001234-П", address="Ленинградская обл., г. Кингисепп, "
         "ул. Заводская, 1", region_code="", **nmu_extra) -> ReportContext:
    """Полный контекст; по умолчанию объект в Ленобласти (47) — федеральный
    образец № 662. Почему не СПб-код: регион 78 переключает форму."""
    ctx = ReportContext(
        organization=Organization(name="ООО «Завод»", inn="7801234564",
                                  ogrn="1027800000000",
                                  director_name="Иванов И.И."))
    ctx.objects.append(NVOSObject(code=code, name="Промплощадка № 1",
                                  category="II", address=address,
                                  region_code=region_code))
    ctx.extra["emission_sources"] = [
        {"number": "0001", "name": "Котельная", "kind": "организованный",
         "workshop": "Котельная", "site": "Котёл № 1",
         "pollutants": [
             {"code": "0301", "name": "Азота диоксид", "g_s": "0.5",
              "t_year": "1.2"},
             {"code": "0337", "name": "Углерода оксид", "g_s": "1.0",
              "t_year": "3.0"}]},
        {"number": "0002", "name": "Сварочный пост", "kind": "организованный",
         "pollutants": [
             {"code": "0123", "name": "Железа оксид", "g_s": "0.02",
              "t_year": "0.05"}]},
    ]
    ctx.pollutants.append(Pollutant(name="Азота диоксид", code="0301",
                                    medium=Medium.AIR,
                                    mass_norm=Decimal("1.2")))
    if nmu_extra:
        ctx.extra["nmu"] = nmu_extra
    return ctx


# контролируемое вещество по расчёту рассеивания — только NO2 (п. 3 № 651)
_CTRL = {"controlled_substances": ["0301"],
         "control_points": [{"name": "КТ-1 (граница жилой зоны)",
                             "code": "0301",
                             "sources": [{"number": "0001",
                                          "contribution_pct": 62}]}]}


def _text(path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def _tables(path) -> list[list[list[str]]]:
    from docx import Document
    return [[[c.text for c in r.cells] for r in t.rows]
            for t in Document(str(path)).tables]


def test_sections_verbatim_662(tmp_path):
    """Форма по образцу № 662: грифы, ровно 9 пунктов с дословными названиями,
    п. 6 при специализированном прогнозе — со степенями НМУ, реквизиты
    действующих НПА; пп. 10–11 со сквозной нумерацией отсутствуют."""
    p = nmu.generate(_ctx(forecast_kind="специализированный", **_CTRL),
                     tmp_path / "n.docx")
    text = _text(p)
    assert "УТВЕРЖДЕНО" in text and "СОГЛАСОВАНО" in text
    for fragment in (
            nmu.P1, nmu.P2, nmu.P3, nmu.P4, nmu.P5, nmu.P6, nmu.P7,
            nmu.P8, nmu.P9,
            "(при наличии) индивидуального предпринимателя, осуществляющего "
            "хозяйственную и (или) иную деятельность",
            "4. Категория объекта, оказывающего негативное воздействие "
            "на окружающую среду: II.",
            "5. Код объекта, оказывающего негативное воздействие "
            "на окружающую среду: 41-0247-001234-П.",
            "(общий или специализированный): специализированный "
            "(степени НМУ: 1, 2, 3)."):
        assert fragment in text, fragment
    # за формой не продолжается нумерация образца
    assert "10. " not in text and "11. " not in text
    assert "Сводный расчёт" not in text
    assert "Пояснительная записка" in text
    # действующая база, а не отменённый № 811
    assert "от 26.11.2025 № 651" in text
    assert "от 28.11.2025 № 662" in text
    assert "№ 96-ФЗ" in text
    assert "№ 811" not in text
    # три степени НМУ при специализированном прогнозе, без «целевого %»
    for mode in (1, 2, 3):
        assert nmu.MODES[mode] in text
    assert "целевое снижение" not in text


def test_table_header_and_numbering_row(tmp_path):
    """Таблица п. 7: заголовки граф дословно из образца и строка «1…6»."""
    p = nmu.generate(_ctx(forecast_kind="общий", **_CTRL), tmp_path / "n.docx")
    t7 = [t for t in _tables(p) if t[0] == nmu.HEADER]
    assert t7, "таблица п. 7 с заголовками образца не найдена"
    assert t7[0][1] == ["1", "2", "3", "4", "5", "6"]
    assert nmu.HEADER[1].startswith("Номер источника (источников)")
    assert nmu.HEADER[4] == "Величины выбросов до мероприятия г/с"
    assert nmu.HEADER[5] == "Величины выбросов после мероприятия г/с"


def test_only_controlled_substances(tmp_path):
    """Пп. 3–4 № 651: в перечне только контролируемые вещества и источники
    с ними — CO и сварочный пост (Fe) в таблицу не попадают."""
    ctx = _ctx(forecast_kind="общий", **_CTRL)
    rows, typical = nmu.measure_rows(ctx, 0)
    assert typical and len(rows) == 1
    assert rows[0][1] == "0001" and rows[0][3] == "0301 Азота диоксид"
    text = _text(nmu.generate(ctx, tmp_path / "n.docx"))
    assert "Углерода оксид" not in text
    assert "Железа оксид" not in text and "0002" not in text
    # контрольная точка и вклад источника попадают в п. 8
    assert "КТ-1 (граница жилой зоны)" in text and "№0001 — 62 %" in text


def test_no_controlled_list_means_placeholder(tmp_path):
    """Без перечня контролируемых веществ таблица по всем источникам НЕ
    строится — только пометка «[требуется: …]» и та же строка в gaps()."""
    ctx = _ctx(forecast_kind="общий")
    assert nmu.measure_rows(ctx, 0) == ([], False)
    text = _text(nmu.generate(ctx, tmp_path / "n.docx"))
    assert f"[{nmu.CONTROLLED_REQUIRED}]" in text
    assert nmu.CONTROLLED_REQUIRED in nmu.gaps(ctx)
    assert "Азота диоксид" not in text
    assert "пп. 3–5 приказа № 651" in text


def test_after_not_computed_by_651_percent(tmp_path):
    """Графа 6 не вычисляется по нормативным процентам № 651: для типовых
    мероприятий — пометка технологу; 0,4250000 (0,5×0,85) нигде нет."""
    ctx = _ctx(forecast_kind="специализированный", **_CTRL)
    for mode in (1, 2, 3):
        rows, _ = nmu.measure_rows(ctx, mode)
        assert rows and all(r[5] == nmu.AFTER_REQUIRED for r in rows)
    text = _text(nmu.generate(ctx, tmp_path / "n.docx"))
    assert "0,5000000" in text                  # до, г/с
    assert "0,4250000" not in text and "0,3000000" not in text
    assert nmu.AFTER_REQUIRED in text
    assert any("графа 6" in g for g in nmu.gaps(ctx))
    # проценты № 651 в пояснительной записке — про вклады в концентрации
    assert "снижение вкладов в приземные концентрации" in text
    assert "величина выброса × процент" not in text


def test_user_measures_after_from_technologist(tmp_path):
    """Графа 6 берётся только из данных технолога по мероприятию:
    reduction_pct — снижение этим мероприятием, after — явная величина;
    мероприятие без того и другого — пометка."""
    ctx = _ctx(forecast_kind="специализированный", **_CTRL, measures=[
        {"mode": 1, "text": "Перевод котлов на природный газ",
         "reduction_pct": 25, "source": "0001"},
        {"mode": 2, "text": "Снижение нагрузки котлов до 50 %",
         "after": "0.21", "source": "0001"},
        {"mode": 3, "text": "Остановка резервного котла", "source": "0001"}])
    r1, _ = nmu.measure_rows(ctx, 1)
    r2, _ = nmu.measure_rows(ctx, 2)
    r3, _ = nmu.measure_rows(ctx, 3)
    assert r1[0][5] == "0,3750000"              # 0,5 × 0,75 — процент технолога
    assert r2[0][5] == "0,2100000"              # явная величина после
    assert r3[0][5] == nmu.AFTER_REQUIRED
    text = _text(nmu.generate(ctx, tmp_path / "n.docx"))
    assert "Перевод котлов на природный газ" in text
    assert "0,3750000" in text and "0,2100000" in text


def test_general_forecast_single_table(tmp_path):
    """При общем прогнозе — одна таблица, требование № 651 п. 10 —
    не менее 20 % снижения вкладов; степеней нет."""
    text = _text(nmu.generate(_ctx(forecast_kind="общий", **_CTRL),
                              tmp_path / "n.docx"))
    assert "не менее чем на 20 %" in text and "п. 10 требований" in text
    assert nmu.MODES[3] not in text


def test_regulated_targets():
    """Снижение вкладов по приказу № 651: прочие 20/15/20/40, для
    регулируемых видов деятельности (ТЭК/ЖКХ) — 15/5/10/20."""
    ctx = _ctx()
    assert [nmu.target_pct(ctx, m) for m in (0, 1, 2, 3)] == [20, 15, 20, 40]
    ctx_reg = _ctx(regulated=True)
    assert [nmu.target_pct(ctx_reg, m) for m in (0, 1, 2, 3)] == [15, 5, 10, 20]
    assert not hasattr(nmu, "efficiency_rows")   # сводного расчёта по % нет


def test_prohibitions_verbatim_651(tmp_path):
    """Запреты при НМУ — дословно п. 9 № 651: регламенты, ГПУ, залповые
    с оговоркой, ПНР и испытания; «продувки и чистки» из № 811 нет."""
    text = _text(nmu.generate(_ctx(**_CTRL), tmp_path / "n.docx"))
    assert ("соблюдаются технологические регламенты работ всех производств, "
            "оборудования и установок, а также запрещаются остановки "
            "газопылеулавливающих сооружений для выполнения профилактических "
            "работ, залповые выбросы вредных веществ в атмосферный воздух "
            "(кроме случаев, когда уже проводятся технологические операции "
            "по подготовке к проведению залповых выбросов), запрещается "
            "проведение пусконаладочных работ и испытаний оборудования") in text
    assert "продувк" not in text.lower()
    assert "с изменением технологического режима" not in text


def test_number_format_and_dots():
    """Числа — запятая и 7 знаков, как в принятых планах; ноль — «0»;
    точка в конце не дублируется."""
    assert nmu._fmt(Decimal("0.064")) == "0,0640000"
    assert nmu._fmt(Decimal("0")) == "0"
    assert nmu._fmt(None) == "—"
    assert nmu._fmt(Decimal("1.23456789")) == "1,2345679"
    assert nmu._dot("Иванов И.И.") == "Иванов И.И."
    assert nmu._dot("Иванов") == "Иванов."


def test_responsible_by_units(tmp_path):
    """Ответственные по структурным подразделениям (п. 8 № 651) — таблицей;
    строка по-старому тоже принимается; без всего — пометка в gaps()."""
    ctx = _ctx(**_CTRL, responsible=[
        {"unit": "Котельная", "position": "Начальник", "name": "Петров П.П."},
        {"unit": "Сварочный участок", "position": "Мастер",
         "name": "Сидоров С.С."}])
    text = _text(nmu.generate(ctx, tmp_path / "n.docx"))
    assert "Структурное подразделение" in text
    assert "Петров П.П." in text and "Сидоров С.С." in text
    assert nmu.responsible_list(_ctx(responsible="Главный инженер Козлов К.К.")) \
        == [{"unit": "", "position": "", "name": "Главный инженер Козлов К.К."}]
    empty = ReportContext()
    assert any("подразделениям" in g for g in nmu.gaps(empty))
    assert "И.И.." not in _text(nmu.generate(_ctx(**_CTRL), tmp_path / "d.docx"))


def test_sample_hint_in_gaps():
    """Пока в Формы/Разработка/НМУ нет образца — подсказка в gaps()."""
    gs = nmu.gaps(_ctx(**_CTRL))
    if not nmu._user_sample_exists():
        assert nmu.SAMPLE_HINT in gs
    else:
        assert nmu.SAMPLE_HINT not in gs


def test_empty_ctx_placeholders(tmp_path):
    """Пустая база: документ собирается, все дыры помечены «[требуется…»."""
    ctx = ReportContext()
    p = nmu.generate(ctx, tmp_path / "n.docx")
    text = _text(p)
    assert p.exists() and p.stat().st_size > 1000
    assert "[требуется" in text
    problems = nmu.gaps(ctx)
    assert any("источники выбросов" in g for g in problems)
    assert any("вид прогноза" in g for g in problems)
    assert any("рассеивания" in g for g in problems)
    assert nmu.CONTROLLED_REQUIRED in problems


def test_gaps_match_document_marks(tmp_path):
    """Каждая строка gaps() дословно печатается в Плане (раздел «Чего не
    хватает») — эколог видит список и в GUI, и в самом документе."""
    for ctx in (ReportContext(), _ctx(), _ctx(forecast_kind="общий", **_CTRL),
                _ctx(forecast_kind="специализированный",
                     controlled_substances=["0301"])):
        text = _text(nmu.generate(ctx, tmp_path / "g.docx"))
        for g in nmu.gaps(ctx):
            assert g in text, g


def test_no_measures_required_branch(tmp_path):
    """П. 5 требований № 662: превышений ПДК нет — план фиксирует отсутствие
    необходимости мероприятий, таблицы степеней не выводятся."""
    ctx = _ctx(no_measures_required=True)
    text = _text(nmu.generate(ctx, tmp_path / "n.docx"))
    assert "не требуется" in text
    assert "п. 5 требований" in text
    assert nmu.MODES[1] not in text


# ── региональные варианты (сверены с принятыми планами Москвы и СПб) ────────

def test_region_detection():
    """Регион: из region_code, иначе из кода ОНВОС («45-0177» → 77,
    «40-0178» → 78); ЛО (47) и пустой контекст — федеральный образец."""
    from ecodoc.development import nmu_regions as R
    assert R.region_code(_ctx()) == "47" and R.profile(_ctx()) is None
    assert R.profile(_ctx(code="45-0177-012344-П"))["key"] == "msk"
    assert R.profile(_ctx(code="40-0178-001234-П"))["key"] == "spb"
    assert R.profile(_ctx(code="", region_code="77"))["key"] == "msk"
    assert R.profile(ReportContext()) is None


def test_moscow_variant_like_approved_plan(tmp_path):
    """Москва (77): пп. 1–9 № 662 + таблица на 8 граф дословно как в плане,
    согласованном ДПиООС 24.04.2026 («Степень опасности НМУ»,
    «Структурное подразделение (цех)»); гриф СОГЛАСОВАНО — Департамент;
    приложения по 231-ПП: пояснительная записка п. 17, журнал прил. 1
    (6 граф дословно + примечания), план-график контроля прил. 3, состав
    заявления п. 16; альбомный лист."""
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from ecodoc.development import nmu_regions as R
    ctx = _ctx(code="45-0177-012344-П", address="105005, г. Москва, "
               "пер. Волховский, вл. 11, стр. 1, 2",
               forecast_kind="общий", below_0_1_pdk=True,
               processes="Учебные мастерские: пайка, слесарные работы.",
               **_CTRL)
    p = nmu.generate(ctx, tmp_path / "msk.docx")
    text = _text(p)
    tables = _tables(p)
    assert Document(str(p)).sections[0].orientation == WD_ORIENT.LANDSCAPE
    for frag in (nmu.P1, nmu.P6, nmu.P8, nmu.P9):
        assert frag in text
    t7 = [t for t in tables if t[0] == R.MSK_HEADER]
    assert t7 and t7[0][1] == [str(i) for i in range(1, 9)]
    row = t7[0][2]
    assert row[1] == "Общий вид прогноза НМУ"
    assert row[2] == "Котельная. Котёл № 1" and row[3] == "0001"
    assert row[5] == "0301 Азота диоксид" and row[6] == "0,5000000"
    assert ("Департамент природопользования и охраны окружающей среды "
            "города Москвы") in text
    assert "№ 231-ПП" in text and "п. 17 Порядка" in text
    # пояснительная записка «< 0,1 ПДК» — 3 раздела по п. 17
    for sec in R.MSK_NOTE_BELOW_01:
        assert sec[1:40] in text
    assert R.MSK_NOTE_FULL[1][1:40] not in text
    assert "Учебные мастерские" in text
    journal = [t for t in tables if t[0] == R.MSK_JOURNAL_HEADER]
    assert journal and journal[0][1] == ["1", "2", "3", "4", "5", "6"]
    assert R.MSK_JOURNAL_NOTES[0] in text
    assert [t for t in tables if t[0] == R.MSK_SCHEDULE_HEADER]
    for item in R.MSK_APPLICATION:
        assert item in text
    # федеральные таблицы по степеням не дублируются
    assert not [t for t in tables if t[0] == nmu.HEADER]
    # gaps: план-график — в списке, подразделения заданы; всё — в документе
    gs = nmu.gaps(ctx)
    assert any("план-график" in g for g in gs)
    assert not any("структурное подразделение (цех) для" in g for g in gs)
    for g in gs:
        assert g in text, g


def test_moscow_full_note_when_not_below_01(tmp_path):
    """Москва без признака «< 0,1 ПДК»: полный состав записки (4 раздела),
    пометка про признак — в gaps() и в документе."""
    from ecodoc.development import nmu_regions as R
    ctx = _ctx(code="45-0177-012344-П", forecast_kind="общий", **_CTRL)
    text = _text(nmu.generate(ctx, tmp_path / "m.docx"))
    for sec in R.MSK_NOTE_FULL:
        assert sec[1:40] in text
    assert any("below_0_1_pdk" in g for g in nmu.gaps(ctx))
    assert "below_0_1_pdk" in text


def test_spb_variant_like_accepted_list(tmp_path):
    """СПб (78): таблица на 9 граф дословно как в принятом перечне
    (с «Достигаемым экологическим эффектом, %»), одна таблица на все
    степени («1 степени опасности» …), блок «Должностное лицо,
    ответственное за проведение мероприятий: ФИО, подпись», гриф
    СОГЛАСОВАНО — Комитет по природопользованию; эффект — целое число
    из величины «после» (0,5 → 0,375 = 25 %), без «после» — пометка."""
    from ecodoc.development import nmu_regions as R
    ctx = _ctx(code="40-0178-001234-П", address="СПб, ул. Заводская, 1",
               forecast_kind="специализированный", **_CTRL,
               responsible=[{"unit": "Котельная", "position": "Начальник",
                             "name": "Петров П.П."}],
               measures=[{"mode": 1, "text": "Перевод котлов на газ",
                          "reduction_pct": 25, "source": "0001"},
                         {"mode": 2, "text": "Остановка котла № 2",
                          "after": "0", "source": "0001"}])
    p = nmu.generate(ctx, tmp_path / "spb.docx")
    text = _text(p)
    tables = _tables(p)
    t7 = [t for t in tables if t[0] == R.SPB_HEADER]
    assert t7 and t7[0][1] == [str(i) for i in range(1, 10)]
    body = t7[0][2:]
    assert [r[1] for r in body] == ["1 степени опасности",
                                    "2 степени опасности",
                                    "3 степени опасности"]
    assert body[0][8] == "25" and body[1][8] == "100" and body[1][7] == "0"
    assert body[2][8] == R.EFFECT_REQUIRED
    assert body[0][2] == "Котельная. Котёл № 1"
    assert R.SPB_RESPONSIBLE_LABEL in text and "Петров П.П." in text
    assert "ФИО, подпись" in text
    assert ("Комитет по природопользованию, охране окружающей среды "
            "и обеспечению экологической безопасности") in text
    assert "№ 86-р" in text
    gs = nmu.gaps(ctx)
    assert any("эффект" in g for g in gs)
    for g in gs:
        assert g in text, g


def test_regional_unit_gap(tmp_path):
    """Нет цеха у источника с контролируемым веществом — пометка
    «[требуется: структурное подразделение (цех)]» в таблице и в gaps()."""
    from ecodoc.development import nmu_regions as R
    ctx = _ctx(code="40-0178-001234-П", forecast_kind="общий", **_CTRL)
    ctx.extra["emission_sources"][0].pop("workshop")
    ctx.extra["emission_sources"][0].pop("site")
    rows = nmu.regional_rows(ctx, R.profile(ctx))
    assert rows[0][2] == R.UNIT_REQUIRED
    assert any("структурное подразделение (цех) для источников 0001" in g
               for g in nmu.gaps(ctx))
    assert R.UNIT_REQUIRED in _text(nmu.generate(ctx, tmp_path / "u.docx"))
