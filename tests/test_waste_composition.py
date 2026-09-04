"""Справка о компонентном составе отхода (по ООС/ПНООЛР) — проект для
протокола; приоритет источников состава в _merge_passports; origin и
агрегатное состояние из справочника в паспорте."""
from docx import Document

from ecodoc.core.models import NVOSObject, ReportContext, WasteFlow
from ecodoc.development import waste_composition as wcp
from ecodoc.development import waste_passport as wp
from ecodoc.intake import candidates

LAMP, TBO = "47110101521", "73310001724"


def _ctx():
    ctx = ReportContext()
    ctx.organization.name = "ООО «Ромашка»"
    ctx.organization.inn = "7800000000"
    ctx.organization.address = "СПб, ул. Пример, 1"
    ctx.objects = [NVOSObject(code="40-0178-001234-П", name="Площадка")]
    ctx.period.year = 2025
    ctx.wastes = [WasteFlow(fkko_code=LAMP, name="Лампы ртутные", hazard_class=1),
                  WasteFlow(fkko_code=TBO, name="Мусор от офисных помещений",
                            hazard_class=4)]
    ctx.extra["waste_passports"] = [{
        "fkko": TBO, "name": "Мусор от офисных помещений", "hazard_class": 4,
        "origin": "уборка офисных помещений",
        "aggregate_state": "Смесь твердых материалов (включая волокна) и изделий",
        "components": [{"name": "бумага", "percent": "60"},
                       {"name": "пластик", "percent": "25"},
                       {"name": "пищевые отходы", "percent": "10"}],
        "_src": "ООС.pdf (лист 12)", "_kind": "oos"}]
    return ctx


def _text(path) -> str:
    d = Document(path)
    return "\n".join(p.text for p in d.paragraphs) + "\n" + "\n".join(
        c.text for t in d.tables for r in t.rows for c in r.cells)


def test_generate_creates_docx_with_components_and_control(tmp_path):
    ctx = _ctx()
    st = candidates.Store(tmp_path / "site")
    st.add(candidates.Candidate(key=f"wastes[fkko={TBO}].generated", value="2",
                                file="ООС.pdf", page=12))
    st.save()
    made = wcp.generate(ctx, tmp_path / "out", site_dir=tmp_path / "site")
    assert [p.name for p in made] == [f"Состав_{TBO}.docx"]
    text = _text(made[0])
    assert "ООО «Ромашка»" in text and "ИНН 7800000000" in text
    assert "40-0178-001234-П" in text
    assert "7 33 100 01 72 4" in text and "IV" in text
    assert "уборка офисных помещений" in text
    assert "Смесь твердых материалов" in text
    assert "бумага" in text and "пластик" in text and "60" in text
    assert "95" in text                                    # итого 60+25+10
    assert "сумма состава ≠ 100 %" in text                 # контроль
    assert "ООС.pdf (лист 12)" in text                     # источник
    assert "Ответственный за обращение с отходами" in text


def test_generate_sum_ok_no_warning(tmp_path):
    ctx = _ctx()
    ctx.extra["waste_passports"][0]["components"].append(
        {"name": "стекло", "percent": "5"})
    (path,) = wcp.generate(ctx, tmp_path)
    text = _text(path)
    assert "сходится" in text and "≠ 100" not in text


def test_generate_skips_without_components_and_v_class_needs_components(tmp_path):
    ctx = _ctx()
    ctx.wastes.append(WasteFlow(fkko_code="82220101215", name="Лом бетона",
                                hazard_class=5))
    made = wcp.generate(ctx, tmp_path)
    assert len(made) == 1
    ctx.extra["waste_passports"].append({
        "fkko": "82220101215", "name": "Лом бетона", "hazard_class": 5,
        "components": [{"name": "бетон", "percent": "100"}], "_src": "ООС.pdf"})
    made = wcp.generate(ctx, tmp_path)
    assert len(made) == 2


def test_generate_empty_returns_nothing(tmp_path):
    ctx = _ctx()
    ctx.extra["waste_passports"] = []
    assert wcp.generate(ctx, tmp_path) == []
    assert not (tmp_path / "Состав_73310001724.docx").exists()


def test_gaps_lists_classes_1_4_without_components():
    g = wcp.gaps(_ctx())
    assert len(g) == 1
    assert "4 71 101 01 52 1" in g[0] and "нет состава" in g[0]
    assert "ООС/ПНООЛР" in g[0]


def test_devdoc_branch_registered_and_reports_error_when_empty(tmp_path):
    from ecodoc.core import serialize, workspace
    from ecodoc.gui import server
    workspace.add_org("Орг")
    workspace.add_site("Орг", "Пл")
    ctx = _ctx()
    serialize.to_json(ctx, workspace.site_dir("Орг", "Пл") / "context.json")
    out = server.api_devdoc({}, {"org": "Орг", "site": "Пл",
                                 "kind": "waste-composition"})
    assert out["files"] == [f"Состав_{TBO}.docx"]
    assert out["path"].endswith("состав_отходов") and out["gaps"]
    ctx.extra["waste_passports"] = []
    serialize.to_json(ctx, workspace.site_dir("Орг", "Пл") / "context.json")
    out = server.api_devdoc({}, {"org": "Орг", "site": "Пл",
                                 "kind": "waste-composition"})
    assert "ООС/ПНООЛР" in out["error"]


# ── приоритет источников состава в _merge_passports ────────────────────
def test_merge_passports_priority_and_origin():
    from ecodoc.ai.analyzer import ExtractionReport, _merge_passports
    ctx = ReportContext()
    rep = ExtractionReport()
    oos = {"waste_passports": [{
        "fkko": TBO, "name": "Мусор от офисных помещений", "hazard_class": 4,
        "origin": "уборка помещений", "aggregate_state": "смесь",
        "components": [{"name": "бумага", "percent": "50"}]}]}
    _merge_passports(ctx, oos, "ООС.pdf (лист 12)", rep)
    item = ctx.extra["waste_passports"][0]
    assert item["_kind"] == "oos" and item["origin"] == "уборка помещений"
    assert item["aggregate_state"] == "смесь"

    # паспорт важнее ООС — состав заменяется, происхождение не затирается
    passport = {"waste_passports": [{
        "fkko": TBO, "origin": "другое",
        "components": [{"name": "бумага", "percent": "55"},
                       {"name": "пластик", "percent": "45"}]}]}
    _merge_passports(ctx, passport, "паспорт ТБО.pdf (лист 1)", rep)
    item = ctx.extra["waste_passports"][0]
    assert item["_kind"] == "passport" and len(item["components"]) == 2
    assert item["origin"] == "уборка помещений"
    assert any("состав заменён" in a.value for a in rep.accepted)

    # ООС после паспорта — состав из паспорта остаётся
    _merge_passports(ctx, oos, "ПНООЛР.pdf", rep)
    assert len(ctx.extra["waste_passports"][0]["components"]) == 2
    assert ctx.extra["waste_passports"][0]["_kind"] == "passport"


def test_merge_passports_old_record_without_kind_keeps_passport():
    """Старая запись без _kind: класс — по _src; паспорт не затирается ООС."""
    from ecodoc.ai.analyzer import ExtractionReport, _merge_passports
    ctx = ReportContext()
    ctx.extra["waste_passports"] = [{
        "fkko": TBO, "name": "x", "hazard_class": 4,
        "components": [{"name": "бумага", "percent": "50"}], "_src": "П.о.о ТБО.pdf"}]
    rep = ExtractionReport()
    _merge_passports(ctx, {"waste_passports": [{
        "fkko": TBO, "components": [{"name": "картон", "percent": "50"}]}]},
        "ООС.pdf", rep)
    item = ctx.extra["waste_passports"][0]
    assert item["components"][0]["name"] == "бумага" and item["_kind"] == "passport"


# ── паспорт берёт origin/агрегатное состояние из справочника ───────────
def test_passport_uses_origin_and_aggregate_state_from_store(tmp_path):
    ctx = _ctx()
    (path,) = [p for p in wp.generate(ctx, tmp_path) if TBO in p.name]
    text = _text(path)
    assert "уборка офисных помещений" in text
    assert "Смесь твердых материалов (включая волокна) и изделий" in text
    assert "бумага" in text
