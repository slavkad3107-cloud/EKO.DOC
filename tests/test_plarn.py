"""Тесты ПЛАРН (ecodoc/development/plarn.py) — план по ПП РФ № 2451
в редакции ПП от 02.06.2026 № 684 (действует с 01.09.2026; сверено
22.08.2026 по normativ.kontur.ru/document/1/506401 и base.garant.ru/400170332).

Проверяем: состав разделов по п. 5 Правил, расчёт максимального расчётного
объёма «100 % наибольшей ёмкости» и пороги 3 т / 0,5 т, наполнение таблиц из
данных, пометки «[требуется: …]» при пустых данных и их совпадение с gaps(),
реквизиты НПА в тексте, подбор отходов ликвидации по каталогу ФККО.
"""
import json
from decimal import Decimal

import pytest

from ecodoc.core.models import (NVOSObject, Organization, ReportContext,
                                WasteAct)
from ecodoc.development import plarn


def _ctx(full=True):
    ctx = ReportContext(organization=Organization(
        name="ООО «Терминал»", inn="7801234564", ogrn="1027800000000",
        address="СПб, Нефтяная дорога, 1", director_name="Иванов И.И."))
    ctx.objects = [NVOSObject(code="78-0178-004321-П", name="Склад ГСМ",
                              category="II", address="СПб, Нефтяная дорога, 1")]
    if full:
        ctx.extra["plarn"] = {
            "tanks": [
                {"name": "РВС-100 № 1", "volume_m3": 100,
                 "product": "дизельное топливо", "density": "0.84"},
                {"name": "РГС-25 № 2", "volume_m3": 25,
                 "product": "бензин АИ-92", "density": "0.75"},
            ],
            "asf": {"name": "ПАСФ «Спасатель»", "contract": "№ 12-ЛРН",
                    "date": "01.02.2026", "certificate": "серия 00 № 1234"},
            "finance": {"kind": "договор страхования",
                        "details": "полис № 555", "amount": "1000000"},
            "operations": ["приём", "хранение", "выдача дизельного топлива"],
            "notify": [{"to": "ЕДДС района", "phone": "112"}],
            "sorbent": "торфяной сорбент «Сорбойл»",
            "waste_transporter": {"name": "ООО «ЭкоТранс»",
                                  "license": "№ 78-00123 от 01.01.2025"},
        }
    return ctx


def _docx_text(path):
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def test_sections_follow_pp2451():
    """Состав плана — по п. 5 Правил № 2451 в ред. ПП № 684: подпункты
    «а»–«л», «о» на месте, «м»/«н» (утратили силу) отсутствуют, приложения
    по п. 6 — отдельным разделом."""
    letters = [l for _k, l, _t in plarn.SECTIONS if l]
    assert letters == ["а", "б", "в", "г", "д", "е", "ж", "з", "и", "к",
                       "л", "о"]
    joined = " ".join(t for _k, _l, t in plarn.SECTIONS).lower()
    for need in ["полное наименование и сокращённое наименование",
                 "идентификационный номер налогоплательщика",
                 "источниках разливов", "максимальные расчётные объёмы",
                 "зоны распространения", "первоочередных действий",
                 "локализации", "и (или) привлечённых", "оповещения",
                 "вида сорбента", "размещение и транспортировка",
                 "финансовое обеспечение", "приложения к плану",
                 "календарные планы"]:
        assert need in joined, f"в составе разделов нет «{need}»"
    assert "объём работ" not in joined and "стоимости единицы" not in joined
    assert "2451" in plarn.NPA and "46" in plarn.NPA and "684" in plarn.NPA


def test_max_spill_is_100_percent_of_largest_tank():
    """П. 7 Правил: 100 % объёма одной наибольшей ёмкости, масса по плотности."""
    ctx = _ctx()
    spill = plarn.max_spill(ctx)
    assert spill["tank"] == "РВС-100 № 1"
    assert spill["volume_m3"] == Decimal("100")
    assert spill["mass_t"] == Decimal("100") * Decimal("0.84")
    required, reason = plarn.plan_required(ctx)
    assert required is True
    assert "2451" in reason and "84" in reason


def test_below_threshold_plan_not_required():
    """Меньше 3 т на суше — обязанность разработки плана не наступает."""
    ctx = _ctx()
    ctx.extra["plarn"]["tanks"] = [
        {"name": "бочка", "volume_m3": 2, "product": "масло",
         "density": "0.9"}]                       # 1,8 т < 3 т
    required, reason = plarn.plan_required(ctx)
    assert required is False
    assert "3" in reason
    # на водном объекте порог 0,5 т — та же ёмкость уже обязывает
    ctx.extra["plarn"]["on_water"] = True
    required_w, _ = plarn.plan_required(ctx)
    assert required_w is True


def test_generate_fills_tables_from_data(tmp_path):
    """Документ наполняется из данных: ёмкости, АСФ, финансы, НПА, нормативы."""
    ctx = _ctx()
    p = plarn.generate(ctx, tmp_path / "plarn.docx")
    assert p.exists() and p.stat().st_size > 1000
    text = _docx_text(p)
    for need in ["РВС-100 № 1", "дизельное топливо", "ПАСФ «Спасатель»",
                 "№ 12-ЛРН", "договор страхования", "2451", "7-ФЗ",
                 "100 процентов объёма одной наибольшей ёмкости",
                 "пп. «л» п. 7", "Сорбойл", "ЭкоТранс", "78-00123",
                 "координат (широта, долгота)", "30 минут"]:
        assert need in text, f"в документе нет «{need}»"
    # нормативы локализации — 4 часа вода / 6 часов суша (пп. «ж» п. 5)
    assert "в течение 4 часов" in text and "6 часов" in text
    # пп. «а» п. 5 (ред. № 684): полное/сокращённое наименование, ИНН, адрес
    assert "полное наименование" in text and "ИНН 7801234564" in text
    # адресаты оповещения по п. 25 — в тексте даже при заданных телефонах
    assert "ЦУКС" in plarn.ALERT_ADDRESSEES[0]


def test_empty_data_gives_marks_and_gaps(tmp_path):
    """Без данных документ собирается, но честно размечен «[требуется…]»."""
    ctx = _ctx(full=False)
    p = plarn.generate(ctx, tmp_path / "plarn.docx")
    text = _docx_text(p)
    assert "[требуется" in text
    problems = plarn.gaps(ctx)
    assert any("tanks" in g for g in problems)          # нет ёмкостей
    assert any("АСФ" in g for g in problems)            # нет АСФ
    assert any("финансовое обеспечение" in g.lower() for g in problems)
    required, _ = plarn.plan_required(ctx)
    assert required is None                             # сравнить не с чем


def test_gaps_match_text_marks(tmp_path):
    """Каждая пометка gaps() вида «требуется: …» есть в тексте документа."""
    ctx = _ctx(full=False)
    text = _docx_text(plarn.generate(ctx, tmp_path / "plarn.docx"))
    for g in plarn.gaps(ctx):
        if g.startswith("требуется: "):
            assert f"[{g}]" in text, f"пометки нет в тексте: {g}"


def test_approval_and_drills_in_validate_and_text(tmp_path):
    """Согласование в ТЕРРИТОРИАЛЬНЫЙ орган РПН (16 раб. дней), учения до
    утверждения и не реже 1 раза в 5 лет (п. 8 в ред. № 684), уведомление
    об утверждении за 14 календарных дней (п. 24)."""
    ctx = _ctx()
    problems = plarn.gaps(ctx)
    assert any("учения" in g and "Росприроднадзора" in g for g in problems)
    issues = plarn.validate(ctx)
    assert any("учения" in i.message for i in issues)
    text = _docx_text(plarn.generate(ctx, tmp_path / "plarn.docx"))
    assert "комплексные учения" in text
    assert "территориальный орган Росприроднадзора" in text
    assert "одного раза в 5 лет" in text and "3 года" not in text
    assert "16 рабочих дней" in text and "14 календарных дней" in text
    assert plarn.DRILLS_EVERY_YEARS == 5


def test_sorbent_and_transporter_are_gaps_when_missing(tmp_path):
    """Пп. «л» п. 5 и п. 6 Правил (ред. № 684): без вида сорбента и лицензии
    перевозчика отходов — пометка в тексте и строка в gaps()."""
    ctx = _ctx()
    del ctx.extra["plarn"]["sorbent"]
    del ctx.extra["plarn"]["waste_transporter"]
    problems = plarn.gaps(ctx)
    assert any("сорбента" in g for g in problems)
    assert any("транспортировке отходов" in g for g in problems)
    text = _docx_text(plarn.generate(ctx, tmp_path / "plarn.docx"))
    assert f"[требуется: {plarn.REQ_SORBENT}]" in text
    assert f"[требуется: {plarn.REQ_TRANSPORTER}]" in text
    assert "Приложения к плану" in text


def test_liquidation_wastes_from_base_and_fkko(tmp_path, monkeypatch):
    """Отходы ликвидации: нефтесодержащие из базы + типовые из каталога ФККО."""
    cat = tmp_path / "fkko_test.json"
    cat.write_text(json.dumps({
        "updated": "2099-01-01", "source": "test", "codes": {
            "93110001393": {"name": "грунт, загрязненный нефтью или "
                                    "нефтепродуктами (содержание нефти или "
                                    "нефтепродуктов 15 % и более)",
                            "class": 3},
            "40613001313": {"name": "отходы минеральных масел моторных",
                            "class": 3},
        }}, ensure_ascii=False), encoding="utf-8")
    monkeypatch.setenv("ECODOC_FKKO", str(cat))

    ctx = _ctx()
    ctx.waste_acts = [WasteAct(name="песок, загрязненный нефтью (менее 15 %)",
                               fkko_code="9 19 201 02 39 4", hazard_class=4)]
    rows = plarn.liquidation_wastes(ctx)
    sources = {r["source"] for r in rows}
    assert any(r["fkko"] == "91920102394" for r in rows)      # из базы
    assert any(r["fkko"] == "93110001393" for r in rows)      # из каталога
    assert "отходы объекта (из базы)" in sources
    text = _docx_text(plarn.generate(ctx, tmp_path / "plarn.docx"))
    assert "93110001393" in text and "91920102394" in text
