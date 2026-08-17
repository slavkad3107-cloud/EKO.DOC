"""Программа ПЭК (приказ № 109): состав разделов, таблицы из данных, пробелы.

Проверяем главные обещания модуля:
  * разделы — дословно по п. 2 Требований (приложение 1 к приказу № 109);
  * таблицы наполняются из базы (источники, вещества, отходы, лаборатория);
  * периодичности сточных вод — по категории объекта (п. 9.2.2);
  * пустые данные дают пометки «[требуется: …]», и gaps() совпадает с ними;
  * реквизиты НПА (№ 109 и редакция 2025 г.) присутствуют в тексте.
"""
from decimal import Decimal

from ecodoc.core.models import (Medium, NVOSObject, Organization, Pollutant,
                                ReportContext, WasteFlow)
from ecodoc.development import pek_program


def _ctx(category="III"):
    ctx = ReportContext(organization=Organization(
        name="ООО Тест", inn="7801234564", ogrn="1027800000000",
        address="СПб, Тестовая ул., 1", director_name="Иванов И.И."))
    ctx.objects = [NVOSObject(code="78-0178-001234-П", name="Площадка",
                              category=category, address="СПб")]
    ctx.period.year = 2026
    return ctx


def _rich_ctx():
    ctx = _ctx()
    ctx.pollutants = [
        Pollutant(name="Азота диоксид", code="0301", medium=Medium.AIR,
                  mass_norm=Decimal("1.5")),
        Pollutant(name="Взвешенные вещества", code="", medium=Medium.WATER,
                  mass_norm=Decimal("0.2")),
    ]
    ctx.wastes = [WasteFlow(fkko_code="4 06 110 01 31 3",
                            name="Отработанное масло", hazard_class=3,
                            generated=Decimal("0.8"))]
    ctx.extra["emission_sources"] = [
        {"number": "0001", "name": "Труба котельной", "kind": "организованный",
         "pollutants": [{"code": "0301", "name": "Азота диоксид"}]},
        {"number": "6001", "name": "Стоянка техники", "kind": "неорганизованный",
         "pollutants": [{"code": "2704", "name": "Бензин"}]},
    ]
    ctx.extra["water"] = {"discharge": [{"receiver": "р. Охта",
                                         "volume": 12.5}]}
    ctx.extra["pek"] = {
        "program_date": "01.02.2026", "markers": ["0301"],
        "air_inventory_date": "отчёт от 10.01.2025",
        "water_permit": "решение № 78-В от 01.03.2025",
        "responsible": "инженер-эколог Петрова А.А.",
        "labs": [{"name": "ООО «Экоанализ»", "address": "СПб",
                  "certificate": "RA.RU.21ЭК01"}],
        "surface_water": [{"water_body": "р. Охта", "location": "500 м ниже",
                           "substance": "БПК5", "period": "ежеквартально"}],
    }
    return ctx


def _docx_text(path):
    """Весь текст docx: абзацы + ячейки таблиц."""
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def test_sections_match_npa_and_requisites(tmp_path):
    """Все девять разделов п. 2 Требований и реквизиты приказа — в тексте."""
    text = _docx_text(pek_program.generate(_rich_ctx(), tmp_path / "p.docx"))
    for _, title in pek_program.SECTIONS:
        assert title in text, f"нет раздела: {title}"
    assert "18.02.2022 № 109" in text
    assert "12.05.2025 № 262" in text          # действующая редакция 2025 г.
    assert "10.01.2002 № 7-ФЗ" in text


def test_air_plan_built_from_sources(tmp_path):
    """План-график выбросов: источники, метод по типу источника, маркеры."""
    ctx = _rich_ctx()
    text = _docx_text(pek_program.generate(ctx, tmp_path / "p.docx"))
    assert "0001 Труба котельной" in text
    assert "Азота диоксид" in text
    assert "инструментальный" in text
    assert "расчётный" in text                  # неорганизованный источник
    assert "1 раз в год" in text                # типовая периодичность
    rows = pek_program.plan_air(ctx)
    marked = [r for r in rows if r["marker"] == "да"]
    assert marked and marked[0]["code"] == "0301"


def test_water_frequency_by_category(tmp_path):
    """Периодичность сточных вод — из п. 9.2.2 по категории объекта."""
    ctx3 = _rich_ctx()                          # III категория
    text3 = _docx_text(pek_program.generate(ctx3, tmp_path / "p3.docx"))
    assert "не менее 1 раза в квартал" in text3
    assert "Токсичность (биотестирование)" in text3
    assert "не реже 2 раз в год" in text3       # очистные, п. 9.2.4

    ctx1 = _rich_ctx()
    ctx1.objects[0].category = "I"
    rows = pek_program.plan_water(ctx1)
    usual = [r for r in rows if "Токсичность" not in r["name"]]
    assert all(r["frequency"] == "не менее 1 раза в месяц" for r in usual)


def test_waste_section_from_data(tmp_path):
    """Отходы попадают в раздел 4 и в подраздел 9.3 (учёт по № 1028)."""
    text = _docx_text(pek_program.generate(_rich_ctx(), tmp_path / "p.docx"))
    assert "40611001313" in text                # код ФККО нормализован (norm_fkko)
    assert "Отработанное масло" in text
    assert "№ 1028" in text
    assert "№ 1030" in text                     # мониторинг ОРО упомянут


def test_empty_ctx_gets_marks_and_gaps(tmp_path):
    """Пустая база: документ собирается, пробелы видны и в тексте, и в gaps()."""
    ctx = ReportContext()
    problems = pek_program.gaps(ctx)
    assert problems                             # пробелы найдены
    text = _docx_text(pek_program.generate(ctx, tmp_path / "p.docx"))
    assert "[требуется:" in text
    assert any("наименование организации" in g for g in problems)
    assert any("нечем наполнять" in g for g in problems)


def test_gaps_equal_text_marks(tmp_path):
    """Инвариант: каждая строка gaps() напечатана в документе как «[…]»."""
    for ctx in (ReportContext(), _ctx(), _rich_ctx()):
        text = _docx_text(pek_program.generate(ctx, tmp_path / "g.docx"))
        for g in pek_program.gaps(ctx):
            assert f"[{g}]" in text, f"пометка не напечатана: {g}"


def test_legacy_points_still_supported(tmp_path):
    """Старый ручной ввод точек (extra.pek.points) не теряется."""
    ctx = _ctx()
    ctx.extra["pek"] = {"points": [
        {"medium": "воздух", "point": "ист.0001", "indicators": "NO2",
         "frequency": "1 раз в квартал"}]}
    ctx.pollutants = [Pollutant(name="Азота диоксид", code="0301",
                                medium=Medium.AIR)]
    text = _docx_text(pek_program.generate(ctx, tmp_path / "p.docx"))
    assert "ист.0001" in text and "NO2" in text and "1 раз в квартал" in text


def test_iv_category_warning():
    """Для IV категории программа не нужна — говорим об этом, а не молчим."""
    ctx = _ctx(category="IV")
    ctx.wastes = [WasteFlow(fkko_code="4 06 110 01 31 3", name="Масло",
                            hazard_class=3, generated=Decimal("0.1"))]
    assert any("IV категории" in g for g in pek_program.gaps(ctx))
