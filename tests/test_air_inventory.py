"""Отчёт об инвентаризации выбросов по приказу Минприроды № 871.

Проверяется структура по приложению 4 (разделы 1–7), дословные шапки таблиц
приложений 1–3, сводка по веществам и gaps() на полном/пустом контексте.
"""
from decimal import Decimal

import openpyxl
import pytest
from docx import Document

from ecodoc.core.models import Medium, NVOSObject, Pollutant, ReportContext
from ecodoc.development import air_inventory as inv


@pytest.fixture()
def full_ctx():
    """Полный контекст — как у реального отчёта (котельная + площадка)."""
    c = ReportContext()
    o = c.organization
    o.name = "Общество с ограниченной ответственностью «ПРОТЕЛЮКС»"
    o.short_name = "ООО «ПРОТЕЛЮКС»"
    o.inn, o.ogrn, o.okved, o.okpo, o.oktmo = "4707038091", "1154707000854", "10.89.4", "25853809", "41621102001"
    o.address = "188490, Ленинградская область, г. Ивангород, ул. Лесная, зд. 13А/1"
    o.director_name, o.phone, o.email = "Брезгин М.Ю.", "8921-449-47-05", "eco@protelux.ru"
    c.period.year = 2025
    c.objects = [NVOSObject(code="41-0147-001234-П", name="Завод по производству биопротеина",
                            category="II", address="ЛО, г. Ивангород, ул. Лесная", oktmo="41621102001")]
    c.extra["emission_sources"] = [
        {"number": "0005", "name": "Котел 1", "kind": "организованный", "type": "Точечный",
         "height": "37", "diameter": "0.35", "x1": "1266325.78", "y1": "371867.17",
         "speed": "3.51", "volume": "0.338", "temperature": "151",
         "hours_day": "24", "hours_year": "8760",
         "method": "Методика определения выбросов ЗВ в атмосферу при сжигании топлива в котлах "
                   "производительностью менее 30 тонн пара в час, М., 1999",
         "pollutants": [{"code": "0301", "name": "Азота диоксид", "conc": "1114.65",
                         "g_s": "0.2425783", "t_year": "0.728728"},
                        {"code": "0337", "name": "Углерода оксид", "conc": "2843.49",
                         "g_s": "0.6188221", "t_year": "2.227768"},
                        {"code": "0703", "name": "Бенз/а/пирен", "g_s": "0.0000001", "t_year": "0.000001"}]},
        {"number": "6002", "name": "Стоянка автотранспорта", "kind": "неорганизованный",
         "x1": "1266300", "y1": "371800", "x2": "1266320", "y2": "371800", "area_width": "10",
         "height": "2", "hours_day": "8", "hours_year": "2000",
         "method": "Методика проведения инвентаризации выбросов ЗВ в атмосферу для "
                   "автотранспортных предприятий (расчетным методом), М., 1998",
         "pollutants": [{"code": "0301", "name": "Азота диоксид", "g_s": "0.001", "t_year": "0.002"},
                        {"code": "0328", "name": "Углерод (Пигмент черный)", "g_s": "0.0001", "t_year": "0.0005"}]},
    ]
    for code, name, mass in (("0301", "Азота диоксид", "0.730728"),
                             ("0337", "Углерода оксид", "2.227768"),
                             ("0703", "Бенз/а/пирен", "0.000001"),
                             ("0328", "Углерод (Пигмент черный)", "0.0005")):
        p = Pollutant(code=code, name=name, mass_norm=Decimal(mass))
        p.medium = Medium.AIR
        c.pollutants.append(p)
    c.extra["air_inventory"] = {
        "developer": {"name": "ООО «Альянс Консалтинг»", "address": "СПб, пр. Большевиков, 7",
                      "director": "Краев Д.В.", "requisites": "ИНН 7811751716"},
        "executors": [{"position": "инженер-эколог", "name": "Маляева Я.В.", "phone": "8921-449-47-05"}],
        "date": "апрель 2025 г.", "responsible": "инженер-эколог Маляева Я.В.",
        "production": "Завод по производству биопротеина; котельная 3 котла на природном газе.",
        "surroundings": "Промзона; ближайшая жилая застройка — 300 м к югу.",
        "szz": "Класс III по СанПиН 2.2.1/2.1.1.1200-03, СЗЗ 300 м.",
        "map": "карта-схема.pdf", "has_gou": False,
        "hazard_classes": {"0301": "3", "0337": "4", "0703": "1", "0328": "3"},
    }
    return c


def _sheet(path, name):
    return [[c.value for c in row] for row in openpyxl.load_workbook(path)[name].iter_rows()]


def _docx_text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


# ── шапки таблиц — дословно по приказу № 871 ─────────────────────────────

def test_table_headers_match_order_871():
    assert len(inv.T32_HEADER) == 26          # табл. 3.2 Порядка — 26 граф
    assert inv.T32_HEADER[:5] == ["№ ИЗАВ", "Тип ИЗАВ", "Наименование ИЗАВ",
                                  "Число ИЗАВ, объединенных под одним номером",
                                  "Высота источника, м"]
    assert "Скорость выхода ГВС, м/с" in inv.T32_HEADER
    assert "Температура ГВС, град. С" in inv.T32_HEADER
    assert "Выбрасываемые в атмосферу вещества: мощность выброса, г/с" in inv.T32_HEADER
    assert inv.T32_HEADER[-2] == "Итого за год выброс вещества источником, т/год"
    assert len(inv.T31_HEADER) == 17 and inv.T31_HEADER[0] == "№ цеха"
    assert len(inv.T36_HEADER) == 11
    assert inv.T36_HEADER[-2:] == ["Коэффициент обеспеченности, %: нормативный",
                                   "Коэффициент обеспеченности, %: фактический"]
    assert len(inv.T37_HEADER) == 10
    assert inv.T37_HEADER[2] == "Количество загрязняющих веществ, отходящих от источников выделения"
    assert inv.T37_HEADER[-1] == "Всего выброшено в атмосферный воздух"
    assert len(inv.T11_HEADER) == 6 and len(inv.T12_HEADER) == 7 and len(inv.T13_HEADER) == 5
    assert len(inv.T21_HEADER) == 15 and inv.T21_HEADER[1] == "Дата"
    assert [n for n, _ in inv.SECTIONS] == ["1", "2", "3", "4", "5", "6", "7"]
    assert inv.SECTIONS[2][1] == "Карта-схема территории объекта ОНВ"


# ── сводка по веществам ──────────────────────────────────────────────────

def test_substance_totals_sum_sources(full_ctx):
    rows = {r["code"]: r for r in inv.substance_totals(full_ctx)}
    assert rows["0301"]["t_year"] == pytest.approx(0.730728)   # котёл + стоянка
    assert rows["0301"]["t_org"] == pytest.approx(0.728728)    # только организованный
    assert rows["0301"]["g_s"] == pytest.approx(0.2435783)
    assert rows["0328"]["solid"] and not rows["0301"]["solid"]
    # твёрдые идут первыми (как в табл. 3.7 «Эколога»)
    codes = [r["code"] for r in inv.substance_totals(full_ctx)]
    assert codes.index("0328") < codes.index("0301")


def test_t37_totals(full_ctx):
    rows, totals = inv.rows_t37(full_ctx)
    assert totals[0][0] == "Всего:"
    assert float(totals[0][-1]) == pytest.approx(0.730728 + 2.227768 + 0.000001 + 0.0005, abs=1e-6)
    assert float(totals[1][-1]) == pytest.approx(0.000501, abs=1e-6)   # твёрдых: сажа + БаП


# ── генерация: xlsx + docx ───────────────────────────────────────────────

def test_generate_full_report(full_ctx, tmp_path):
    out = inv.generate(full_ctx, tmp_path / "инв.xlsx")
    wb = openpyxl.load_workbook(out)
    for name in ("Титул", "Источники", "Вещества", "Перечень ЗВ (табл.1)", "Табл.3.1 ИВ",
                 "Табл.3.2 ИЗАВ", "Табл.3.6 ГОУ", "Табл.3.7 Суммарные", "Табл.2.1 Измерения",
                 "Табл.1.1-1.3 Нестац", "Чего не хватает"):
        assert name in wb.sheetnames, name
    t32 = _sheet(out, "Табл.3.2 ИЗАВ")
    assert t32[1][:3] == ["№ ИЗАВ", "Тип ИЗАВ", "Наименование ИЗАВ"]
    assert t32[2][0] == "1" and t32[2][25] == "26"          # строка номеров граф
    row = t32[3]
    assert row[0] == "0005" and row[4] == "37" and row[5] == "0.35" and row[17] == "151"
    assert row[19] == "0301" and row[22] == "0.2425783" and row[24] == "0.728728"
    assert t32[4][0] in ("", None) and t32[4][19] == "0337"          # второе вещество того же ИЗАВ

    docx_path = out.with_suffix(".docx")
    assert docx_path.exists()
    text = _docx_text(docx_path)
    for num, title in inv.SECTIONS:
        assert f"{num}. {title}" in text
    for sub in ("1.1. Реквизиты объекта ОНВ", "1.2.3. Сведения о количестве, характеристиках и эффективности ГОУ",
                "1.5. Размеры и границы санитарно-защитной зоны объекта ОНВ",
                "4.2. Источники выбросов загрязняющих веществ", "4.4. Суммарные выбросы по объекту ОНВ"):
        assert sub in text
    assert "СВЕДЕНИЯ О РАЗРАБОТЧИКЕ И СПИСОК ИСПОЛНИТЕЛЕЙ" in text
    assert "ООО «Альянс Консалтинг»" in text and "Маляева Я.В." in text
    assert "2 стационарных источников" in text and "1 организованных и 1 неорганизованных" in text
    assert "Таблица 1. Полный перечень загрязняющих веществ" in text
    assert "Таблица 7.1 – Режимы работы ИЗАВ" in text
    assert "Нестационарность выбросов отсутствует" in text
    assert "Газоочистные установки на объекте ОНВ отсутствуют" in text
    assert "[типовой текст]" in text                          # типовые тексты помечены
    assert "Приложение 9. Протоколы замеров" in text
    assert "1154707000854" in text                            # ОГРН в реквизитах


def test_gaps_full_context_is_clean(full_ctx):
    assert inv.gaps(full_ctx) == []


def test_gaps_point_to_order_871(full_ctx):
    full_ctx.extra["air_inventory"].pop("szz")
    full_ctx.extra["air_inventory"].pop("map")
    full_ctx.organization.ogrn = ""
    full_ctx.extra["emission_sources"][0].pop("speed")
    text = " | ".join(inv.gaps(full_ctx))
    assert "СЗЗ" in text and "п. 35 «д»" in text
    assert "карты-схемы" in text and "п. 37" in text
    assert "ОГРН" in text
    assert "№0005" in text and "скорость ГВС" in text


def test_gaps_mismatch_sources_vs_pollutants(full_ctx):
    full_ctx.pollutants[0].mass_norm = Decimal("5")           # 0301: 5 т вместо 0.73
    text = " | ".join(inv.gaps(full_ctx))
    assert "0301" in text and "≠ валовой массе" in text


def test_gaps_empty_context():
    text = " | ".join(inv.gaps(ReportContext()))
    assert "не найдены источники выбросов" in text
    assert "не заданы вещества" in text
    assert "разработчик" in text and "ОГРН" in text


def test_generate_empty_context_has_placeholders(tmp_path):
    out = inv.generate(ReportContext(), tmp_path / "пусто.xlsx")
    text = _docx_text(out.with_suffix(".docx"))
    assert "[требуется:" in text
    assert "4. " + inv.SECTIONS[3][1] in text


def test_gou_goes_to_t36_and_t37(full_ctx):
    full_ctx.extra["air_inventory"]["has_gou"] = True
    full_ctx.extra["air_inventory"]["gou"] = [{
        "izav": "0005", "emitter": "Котел 1", "name": "Циклон ЦН-15",
        "eff_design": "85", "eff_actual": "80", "coef_norm": "100", "coef_actual": "100",
        "pollutants": [{"code": "0328", "name": "Углерод (Пигмент черный)"}]}]
    t36 = inv.rows_t36(full_ctx)
    assert t36[0][4] == "Циклон ЦН-15" and t36[0][7] == "80"
    rows, _ = inv.rows_t37(full_ctx)
    sazha = next(r for r in rows if r[0] == "0328")
    assert float(sazha[5]) == pytest.approx(0.0005 / 0.2, abs=1e-6)   # поступает на очистку
    assert float(sazha[-1]) == pytest.approx(0.0005, abs=1e-6)        # выброшено
