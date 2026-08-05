"""Тесты расчёта класса опасности и справочников."""
from pathlib import Path

from ecodoc.development.hazard_class import Component, calculate
from ecodoc.core.refdata import substances, common_wastes


def test_hazard_class_boundaries():
    # K = 50000/100 + 950000/1e6 = 500.95 → 10³ ≥ K > 10² → III класс
    r = calculate([Component("Нефтепродукты", 50000, 100),
                   Component("Песок", 950000, 1_000_000)])
    assert r.hazard_class == 3
    assert 500 < r.k_total < 502


def test_hazard_class_official_scale():
    """Приложение № 1 к Критериям (пр. № 158, ранее № 536): верхняя граница
    диапазона включается, нижняя — нет. K=10⁴ — это II класс, K=10 — V."""
    def cls(k):
        return calculate([Component("х", k, 1)]).hazard_class
    assert cls(10) == 5 and cls(10.1) == 4          # K ≤ 10 → V
    assert cls(100) == 4 and cls(101) == 3          # 10² ≥ K > 10 → IV
    assert cls(1_000) == 3 and cls(1_001) == 2      # 10³ ≥ K > 10² → III
    assert cls(10_000) == 2 and cls(10_001) == 1    # 10⁴ ≥ K > 10³ → II
    assert cls(500_000) == 1                        # 10⁶ ≥ K > 10⁴ → I


def test_hazard_class_high():
    r = calculate([Component("Ртуть", 1000, 1)])   # K = 1000 → III (граница)
    assert r.hazard_class == 3
    r2 = calculate([Component("Оч.опасное", 100000, 1)])  # K = 1e5 → I
    assert r2.hazard_class == 1


def test_hazard_zero_wi_skipped():
    r = calculate([Component("Инерт", 1_000_000, 0)])
    assert any("Wi" in w for w in r.warnings)
    assert r.hazard_class == 5


def test_reference_loaded():
    subs = substances()
    assert len(subs) >= 20
    no2 = next(s for s in subs if s["code"] == "0301")
    assert no2["pdk_mr"] == 0.2
    assert len(common_wastes()) >= 5


def test_hazard_calc_document(tmp_path):
    """Расчёт оформляется документом: таблица компонентов, K, вывод."""
    from docx import Document

    from ecodoc.development.hazard_class import generate
    path = generate([Component("Нефтепродукты", 50000, 100),
                     Component("Песок", 950000, 1_000_000)],
                    tmp_path / "расчёт.docx",
                    waste_name="Грунт, загрязнённый нефтепродуктами",
                    fkko="93110001394", org_name="ООО «Тест»",
                    basis="протокол КХА № 12-25, Wi — прил. к пр. № 536")
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    assert "приказ Минприроды России от 31.03.2025 № 158" in text
    assert "10⁶ ≥ K > 10⁴" in text                  # официальные границы
    assert "93110001394" in text and "протокол КХА № 12-25" in text
    assert "3 классу" in text                       # K ≈ 501 → III класс
    assert "Нефтепродукты" in cells and "50000" in cells


def test_hazard_doc_v_class_biotest_note(tmp_path):
    from docx import Document

    from ecodoc.development.hazard_class import generate
    path = generate([Component("Песок", 1_000_000, 1_000_000)],
                    tmp_path / "v.docx", waste_name="Песок чистый")
    text = "\n".join(p.text for p in Document(str(path)).paragraphs)
    assert "5 классу" in text and "биотестированием" in text


def test_api_hazard_class_saves_document(tmp_path, monkeypatch):
    monkeypatch.setenv("ECODOC_RESULTS", str(tmp_path / "res"))
    from ecodoc.gui import server
    out = server.api_hazard_class({}, {
        "components": [{"name": "Нефтепродукты", "ci": 50000, "wi": 100}],
        "save": 1, "waste_name": "Грунт: тест/1", "fkko": "93110001394"})
    assert out["hazard_class"] == 3                 # K = 500 → III
    p = Path(out["path"])
    assert p.exists() and p.suffix == ".docx"
    assert "/" not in p.name.replace("расчёт_класса_", "", 1)  # имя очищено
