"""Тесты генераторов документов разработки (НМУ, программа ПЭК)."""
from ecodoc.core.models import NVOSObject, Organization, ReportContext
from ecodoc.development import nmu, pek_program


def _ctx():
    ctx = ReportContext(organization=Organization(
        name="ООО Тест", inn="7801234564", director_name="Иванов И.И."))
    ctx.objects = [NVOSObject(code="78-0178-001234-П", name="Площадка",
                              category="III", address="СПб")]
    return ctx


def _docx_text(path):
    """Весь текст docx: абзацы + ячейки таблиц."""
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def test_nmu_generates(tmp_path):
    ctx = _ctx()
    ctx.extra["nmu"] = {"measures": [
        {"mode": 1, "text": "Усилить контроль ГОУ", "reduction_pct": 15}]}
    p = nmu.generate(ctx, tmp_path / "nmu.docx")
    assert p.exists() and p.stat().st_size > 1000
    text = _docx_text(p)
    assert "НМУ" in text and "Усилить контроль ГОУ" in text


def test_pek_program_generates(tmp_path):
    ctx = _ctx()
    ctx.extra["pek"] = {"points": [
        {"medium": "воздух", "point": "ист.0001", "indicators": "NO2",
         "frequency": "1 раз в квартал"}]}
    p = pek_program.generate(ctx, tmp_path / "pek.docx")
    assert p.exists()
    from docx import Document
    text = "\n".join(par.text for par in Document(str(p)).paragraphs)
    assert "экологического контроля" in text


def test_devdoc_flag_in_registry():
    from ecodoc.core import registry
    registry.load_all()
    assert getattr(registry.get("nmu"), "devdoc", False)
    assert getattr(registry.get("pek-program"), "devdoc", False)
