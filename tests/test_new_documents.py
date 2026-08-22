"""Новые документы: инвентаризации, ПНООЛР, 4-ОС, ТУ."""
from decimal import Decimal

import openpyxl
import pytest

from ecodoc.core.models import (Medium, NVOSObject, Pollutant, ReportContext,
                                WasteAct, WasteFlow)


@pytest.fixture()
def ctx():
    c = ReportContext()
    c.organization.name = "ИП Миних Елена Анатольевна"
    c.organization.inn = "780600114472"
    c.organization.ogrn = "307784705100221"
    c.period.year = 2025
    c.objects = [NVOSObject(code="41-0247-005048-П", address="Промзона Янино")]
    c.wastes = [
        WasteFlow(fkko_code="47110101521", name="Лампы ртутные", hazard_class=1,
                  generated=Decimal("0.052"), transferred=Decimal("0.052")),
        WasteFlow(fkko_code="73310001724", name="Мусор офисный", hazard_class=4,
                  generated=Decimal("1.9"), transferred=Decimal("1.9"),
                  placed_norm=Decimal("1.9")),
    ]
    c.waste_acts = [WasteAct(fkko_code="47110101521", name="Лампы ртутные",
                             mass=Decimal("0.052"), operation="обезвреживание",
                             receiver="ООО «Меркурий»", date="15.02.2025")]
    c.extra["waste_passports"] = [{
        "fkko": "47110101521", "name": "Лампы ртутные", "hazard_class": 1,
        "components": [{"name": "стекло", "percent": "92"}]}]
    c.extra["emission_sources"] = [{
        "number": "0001", "name": "Котельная", "kind": "организованный",
        "_src": "ООС.pdf",
        "pollutants": [{"code": "0301", "name": "Азота диоксид",
                        "g_s": "0.05", "t_year": "0.412"}]}]
    air = Pollutant(code="0301", name="Азота диоксид", mass_norm=Decimal("0.412"))
    air.medium = Medium.AIR
    c.pollutants = [air]
    return c


def _cells(path, sheet):
    ws = openpyxl.load_workbook(path)[sheet]
    return [[c.value for c in row] for row in ws.iter_rows()]


# ── инвентаризация отходов ───────────────────────────────────────────────

def test_waste_inventory_collects_from_all_sources(ctx):
    from ecodoc.development.waste_inventory import collect
    rows = {r["fkko"]: r for r in collect(ctx)}
    lamp = rows["47110101521"]
    assert lamp["hazard"] == 1 and lamp["generated"] == pytest.approx(0.052)
    assert lamp["operations"] == ["обезвреживание"]      # из акта
    assert lamp["receivers"] == ["ООО «Меркурий»"]
    assert "стекло" in lamp["composition"] and lamp["passport"]
    assert not rows["73310001724"]["passport"]           # паспорта нет


def test_waste_inventory_reports_gaps(ctx):
    from ecodoc.development.waste_inventory import gaps
    text = " | ".join(gaps(ctx))
    assert "нет паспорта отхода" in text                 # для мусора IV класса
    assert "Мусор офисный" in text


def test_waste_inventory_document(ctx, tmp_path):
    from ecodoc.development.waste_inventory import generate
    out = generate(ctx, tmp_path / "инв.xlsx")
    wb = openpyxl.load_workbook(out)
    assert wb.sheetnames == ["Титул", "Перечень отходов", "Чего не хватает"]
    flat = [str(v) for row in _cells(out, "Перечень отходов") for v in row if v]
    assert "47110101521" in flat and "Лампы ртутные" in flat
    assert "ООО «Меркурий»" in flat


# ── инвентаризация выбросов ──────────────────────────────────────────────

def test_air_inventory_sources_and_substances(ctx, tmp_path):
    from ecodoc.development.air_inventory import generate, gaps, sources
    assert sources(ctx)[0]["number"] == "0001"
    out = generate(ctx, tmp_path / "инв_воздух.xlsx")
    flat = [str(v) for row in _cells(out, "Источники") for v in row if v]
    assert "Котельная" in flat and "0.412 т/год" in " ".join(flat)
    subs = [str(v) for row in _cells(out, "Вещества") for v in row if v]
    assert "0301" in subs and "Азота диоксид" in subs
    # источники и вещества есть — замечаний про них нет; остальные пробелы
    # (параметры ИЗАВ, разработчик, СЗЗ…) — предмет tests/test_air_inventory.py
    text = " | ".join(gaps(ctx))
    assert "не найдены источники" not in text and "не заданы вещества" not in text
    assert out.with_suffix(".docx").exists()             # текстовый отчёт рядом


def test_air_inventory_complains_without_sources(tmp_path):
    from ecodoc.development.air_inventory import gaps
    empty = ReportContext()
    text = " | ".join(gaps(empty))
    assert "не найдены источники выбросов" in text
    assert "не заданы вещества" in text


# ── ПНООЛР ───────────────────────────────────────────────────────────────

def test_pnoolr_norms_and_limits(ctx, tmp_path):
    from ecodoc.development.pnoolr import generate, rows
    data = {r["fkko"]: r for r in rows(ctx)}
    assert data["73310001724"]["norm"] == pytest.approx(1.9)
    assert data["73310001724"]["limit"] == pytest.approx(1.9)   # размещено
    assert data["47110101521"]["limit"] is None                 # не размещался
    out = generate(ctx, tmp_path / "пноолр.xlsx")
    flat = [str(v) for row in _cells(out, "Нормативы и лимиты") for v in row if v]
    assert "ИТОГО" in flat and "Лампы ртутные" in flat
    gaps_text = " ".join(str(v) for row in _cells(out, "Чего не хватает")
                         for v in row if v)
    assert "пишется экологом" in gaps_text                # честно про разделы


# ── 4-ОС ────────────────────────────────────────────────────────────────

def test_oos4_registered_and_validates(ctx):
    from ecodoc.core import registry
    registry.load_all()
    rep = registry.get("4-oos")(ctx)
    msgs = " | ".join(i.message for i in rep.validate())
    assert "не заданы текущие затраты" in msgs            # extra.oos4 пуст


def test_oos4_print_matches_blank(ctx, tmp_path):
    """Форма № 4-ОС: бланк альбома форм 2026 г. (приказ № 346, изменения
    № 566 и № 734), платы за НВОС нет, целые тысячи рублей (N(18,0))."""
    from ecodoc.core import registry
    registry.load_all()
    ctx.organization.okpo = "0123456789"
    ctx.organization.director_position = "Директор"
    ctx.organization.director_name = "Миних Е.А."
    ctx.extra["oos4"] = {"costs": {"air": "120.5", "waste": "340", "water": 0,
                                   "noise": "7.7"},
                         "services": {"water": "55", "03": "ignored",
                                      "other": "12.4"}}
    rep = registry.get("4-oos")(ctx)
    assert rep.title.startswith("4-ОС")            # официальный индекс формы
    out = rep.render_print(tmp_path / "4oos.xlsx")
    rows = _cells(out, "Форма")
    flat = [str(v) for row in rows for v in row if v is not None]
    # строки бланка, а не выдуманные 101-106
    assert "01" in flat and "02" in flat and "10" in flat
    assert not any(v in ("101", "106") for v in flat)
    # целые тысячи рублей: 120.5→121, 7.7→8; строка 01 = сумма округлённых
    assert "121" in flat and "340" in flat and "8" in flat
    assert "120.5" not in flat and "7.7" not in flat
    assert "469" in flat                           # 121 + 340 + 0 + 8
    # раздела о плате за НВОС в действующей форме нет
    assert not any("Плата за негативное" in v for v in flat)
    # бланк: заголовок таблицы и блоки граф
    joined = " ".join(flat)
    assert "Выполнение работ по охране окружающей среды, тысяча рублей" in joined
    assert "Для собственных нужд (кроме предоставления специализированных "            "природоохранных услуг)" in joined
    assert "Специализированные природоохранные услуги" in joined
    assert "№ строки" in flat                      # не «N строки»
    # два блока по 8 граф: нумерация А, Б, 3–10, 11–18
    num_row = next(r for r in rows if r and str(r[0]) == "А")
    nums = [str(v) for v in num_row if v is not None]
    assert nums == ["А", "Б"] + [str(i) for i in range(3, 19)]
    # строка 01 и отдельная строка «в том числе:» — как в бланке и XML-шаблоне
    r01 = next(r for r in rows if r and r[0] == "Охрана окружающей среды – всего")
    assert r01[1] == "01" and r01[2] == 469
    assert r01[10] == 55 + 12                      # гр. 11 = сумма строк услуг
    assert any(r and r[0] == "в том числе:" for r in rows)
    # строки 02–10: наименования, значения в гр. 3 и гр. 11
    by_code = {}
    for r in rows:
        if len(r) > 10 and str(r[1]) in {f"{i:02d}" for i in range(2, 11)}:
            by_code[str(r[1])] = r
    assert len(by_code) == 9
    assert not any(str(r[0]).startswith("в том числе") for r in by_code.values())
    assert by_code["02"][0] == ("охрана атмосферного воздуха и предотвращение "
                                "изменения климата")
    assert by_code["03"][0] == "обращение со сточными водами"
    assert by_code["04"][0] == "обращение с отходами"
    assert by_code["05"][0] == ("защита и экологическая реабилитация земель, "
                                "поверхностных и подземных водных объектов")
    assert by_code["06"][0] == "снижение шумового и вибрационного воздействия"
    assert by_code["06"][2] == 8                   # шум именно в строке 06
    assert by_code["07"][0] == ("сохранение биоразнообразия и охрана природных "
                                "территорий")
    assert by_code["08"][0] == ("радиационная безопасность окружающей среды "
                                "(за исключением мер по предотвращению аварий "
                                "и катастроф)")
    assert by_code["09"][0] == ("научно-исследовательская деятельность и "
                                "разработки в области охраны окружающей среды")
    assert by_code["10"][0] == ("другие направления деятельности в области "
                                "охраны окружающей среды")
    # услуги: по ключу «water» → стр. 03 гр. 11, «other» → стр. 10 гр. 11;
    # ключ совпадающего кода «03» не перебивает ключ «water»
    assert by_code["03"][10] == 55 and by_code["10"][10] == 12
    assert by_code["03"][2] is None                # гр. 3 пустая при нуле
    # графы 4–10 и 12–18 программа не заполняет
    assert all(by_code["02"][i] is None for i in list(range(3, 10)) + list(range(11, 18)))
    # подписной блок бланка и контроль
    assert "Должностное лицо, ответственное за предоставление первичных "            "статистических данных" in joined
    assert "(должность)" in flat and "(Ф.И.О.)" in flat and "(подпись)" in flat
    assert "Директор" in flat and "Миних Е.А." in flat
    assert "стр. 01 = стр. 02" in joined and "≥ 100" in joined
    # титул: форма, ОКУД, приказ об утверждении + изменения 566/734, сроки
    titul = " ".join(str(v) for row in _cells(out, "Титул") for v in row if v)
    assert "4-ОС" in titul and "0609030" in titul
    assert "от 22.07.2025 № 346" in titul
    assert "от 10.10.2025 № 566" in titul and "от 19.12.2025 № 734" in titul
    assert "с 1-го рабочего дня по 25 января после отчетного периода" in titul
    assert "СВЕДЕНИЯ О ТЕКУЩИХ ЗАТРАТАХ НА ОХРАНУ ОКРУЖАЮЩЕЙ СРЕДЫ" in titul
    assert "0123456789" in titul and "идентификационный номер" in titul


def test_oos4_threshold_and_okved_rules(ctx):
    """Приказ № 566 (порог 100 тыс. руб., не суммируется) и № 734 (ОКВЭД
    по строкам 04/07)."""
    from ecodoc.core import registry
    registry.load_all()
    ctx.organization.okpo = "0123456789"
    # 60 + 60 = 120 > 100, но виды затрат не суммируются → предупреждение
    ctx.extra["oos4"] = {"costs": {"air": "60"}, "paid_services": "60"}
    msgs = " | ".join(i.message for i in registry.get("4-oos")(ctx).validate())
    assert "№ 566" in msgs and "более 100 тысяч рублей" in msgs
    assert "не заданы текущие затраты" not in msgs
    # оплата услуг > 100 сама по себе даёт право/обязанность сдавать форму
    ctx.extra["oos4"] = {"costs": {"air": "60"}, "paid_services": "101"}
    msgs = " | ".join(i.message for i in registry.get("4-oos")(ctx).validate())
    assert "№ 566" not in msgs
    # № 734: оператор по обработке отходов (38.32.11) — строка 04
    ctx.organization.okved = "38.32.11"
    ctx.extra["oos4"] = {"costs": {"waste": "500"}}
    msgs = " | ".join(i.message for i in registry.get("4-oos")(ctx).validate())
    assert "№ 734" in msgs and "строке 04" in msgs and "38.32.1" in msgs
    # тот же ОКВЭД, но затраты не по отходам — молчит
    ctx.extra["oos4"] = {"costs": {"air": "500"}}
    msgs = " | ".join(i.message for i in registry.get("4-oos")(ctx).validate())
    assert "№ 734" not in msgs
    # рыбоводство 03.22.1 — строка 07
    ctx.organization.okved = "03.22.1"
    ctx.extra["oos4"] = {"costs": {"air": "500"}, "services": {"bio": "200"}}
    msgs = " | ".join(i.message for i in registry.get("4-oos")(ctx).validate())
    assert "строке 07" in msgs and "03.22" in msgs
    # итог услуг, не совпадающий с суммой строк, — предупреждение
    ctx.organization.okved = ""
    ctx.extra["oos4"] = {"costs": {"air": "500"},
                         "services": {"water": "100", "total": "999"}}
    msgs = " | ".join(i.message for i in registry.get("4-oos")(ctx).validate())
    assert "контроль 1" in msgs and "999" in msgs


# ── ТУ ───────────────────────────────────────────────────────────────────

def test_tu_letter(ctx, tmp_path):
    from docx import Document

    from ecodoc.development.tu_waste import generate
    out = generate(ctx, tmp_path / "ту.docx", receiver="ООО «Полигон»",
                   purpose="размещения на полигоне")
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "ООО «Полигон»" in text and "размещения на полигоне" in text
    assert "780600114472" in text and "41-0247-005048-П" in text
    table = doc.tables[0]
    cells = [c.text for row in table.rows for c in row.cells]
    assert "47110101521" in cells and "Лампы ртутные" in cells


# ── реестр и API ─────────────────────────────────────────────────────────

def test_all_new_documents_registered():
    from ecodoc.core import registry
    registry.load_all()
    reports = registry.all_reports()
    for code in ("waste-inventory", "air-inventory", "pnoolr", "tu-waste",
                 "waste-passport", "4-oos"):
        assert code in reports, code
    for code in ("waste-inventory", "air-inventory", "pnoolr", "tu-waste"):
        assert getattr(reports[code], "devdoc", False), code


def test_api_devdoc_generates_new_documents(ctx, tmp_path, monkeypatch):
    from ecodoc.core import workspace
    from ecodoc.gui import server
    monkeypatch.setenv("ECODOC_RESULTS", str(tmp_path / "res"))
    workspace.add_org("ТЕСТ")
    workspace.add_site("ТЕСТ", "Пл")
    workspace.save_context("ТЕСТ", "Пл", ctx)
    for kind in ("waste-inventory", "air-inventory", "pnoolr", "tu-waste"):
        out = server.api_devdoc({}, {"org": "ТЕСТ", "site": "Пл", "kind": kind})
        assert "path" in out, (kind, out)
        assert out["path"]
