"""Паспорт отходов I–IV класса (приказ № 1026 до 01.09.2026, № 286 — с неё) — форма и наполнение."""
from ecodoc.core.models import ReportContext, WasteFlow
from ecodoc.development import waste_passport as wp


def _ctx():
    ctx = ReportContext()
    ctx.organization.name = "ИП Миних Елена Анатольевна"
    ctx.organization.inn = "780600114472"
    ctx.organization.director_name = "Миних Е.А."
    ctx.period.year = 2025
    ctx.wastes = [
        WasteFlow(fkko_code="47110101521", name="Лампы ртутные", hazard_class=1),
        WasteFlow(fkko_code="73310001724", name="Мусор от офисных помещений",
                  hazard_class=4),
        WasteFlow(fkko_code="73310002725", name="Отход V класса", hazard_class=5),
    ]
    return ctx


def test_passport_only_for_classes_1_4(tmp_path):
    """V класс паспорта не требует — на него документ не создаётся."""
    paths = wp.generate(_ctx(), tmp_path)
    assert len(paths) == 2
    names = {p.name for p in paths}
    assert "Паспорт_47110101521.docx" in names
    assert not any("73310002725" in n for n in names)


def test_passport_without_components_does_not_crash(tmp_path):
    """Регресс: при ПУСТОМ составе подставляется строка-плейсхолдер, а размер
    таблицы считался без неё → IndexError на генерации."""
    ctx = _ctx()
    ctx.wastes = [ctx.wastes[1]]              # отход без известного состава
    (path,) = wp.generate(ctx, tmp_path)
    from docx import Document
    rows = [r for t in Document(path).tables for r in t.rows]
    assert any("КХА" in c.text for r in rows for c in r.cells)


def test_passport_takes_components_from_ai_store(tmp_path):
    """Состав, извлечённый ИИ из загруженных паспортов (extra.waste_passports),
    попадает в таблицу «Сведения об отходах»."""
    ctx = _ctx()
    ctx.extra["waste_passports"] = [{
        "fkko": "47110101521", "name": "Лампы ртутные", "hazard_class": 1,
        "components": [{"name": "стекло", "percent": "92"},
                       {"name": "ртуть", "percent": "0.02"}]}]
    paths = wp.generate(ctx, tmp_path)
    lamp = next(p for p in paths if "47110101521" in p.name)
    from docx import Document
    text = "\n".join(c.text for t in Document(lamp).tables
                     for r in t.rows for c in r.cells)
    assert "стекло" in text and "92" in text
    assert "ртуть" in text


# ─────────────── две редакции формы: № 1026 (до 01.09.2026) и № 286 ───────────────
import datetime as dt

from docx import Document


def _texts(path):
    """Все ячейки таблиц + абзацы документа — для дословной сверки строк."""
    doc = Document(path)
    cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
    paras = [p.text for p in doc.paragraphs]
    return cells, paras


def _ul_ctx():
    ctx = _ctx()
    ctx.organization.name = 'Общество с ограниченной ответственностью «ТЕХНОСТРОЙ»'
    ctx.organization.short_name = 'ООО «ТЕХНОСТРОЙ»'
    ctx.organization.inn = "7800000000"
    ctx.organization.address = "СПб, ул. Тестовая, 1"
    ctx.wastes = [WasteFlow(fkko_code="7 23 101 01 39 4",
                            name="Осадок мойки", hazard_class=4)]
    return ctx


def test_form_switch_by_date():
    """п. 3 приказа № 286 — вступает в силу с 01.09.2026; до этого — № 1026."""
    assert wp.form_for(dt.date(2026, 8, 31)) == wp.FORM_1026
    assert wp.form_for(dt.date(2026, 9, 1)) == wp.FORM_286
    assert wp.form_for(dt.date(2027, 1, 1)) == wp.FORM_286


def test_form_1026_rows_verbatim(tmp_path):
    """До 01.09.2026 — старая форма, дословно как принятые паспорта ТЕХНОСТРОЙ."""
    (path,) = wp.generate(_ul_ctx(), tmp_path, approved_date=dt.date(2025, 12, 19))
    cells, paras = _texts(path)
    assert "Наименование вида отходов по ФККО" in cells
    assert ("Происхождение отходов (указывается наименование технологического "
            "процесса, в результате которого образовался отход, или процесса, "
            "в результате которого товар (продукция) утратил свои "
            "потребительские свойства, с указанием наименования исходного "
            "товара)") in cells
    assert ("Способ определения химического и (или) компонентного состава "
            "вида отходов (указывается согласно документации и (или) с "
            "использованием количественного химического анализа)") in cells
    assert ("Фамилия, имя, отчество (при наличии) индивидуального "
            "предпринимателя или полное наименование юридического лица") in cells
    assert "Сокращенное наименование юридического лица" in cells
    assert "Индивидуальный номер налогоплательщика (ИНН)" in cells
    assert "Место нахождения" in cells
    assert "Адрес (адреса) фактического осуществления деятельности" in cells
    assert "включенных в Федеральный классификационный каталог отходов" in paras
    assert "М.П." in paras and "место печати" not in paras
    assert "Генеральный директор" in paras
    # «Сведения о лице» у ЮЛ — 8 строк + заголовок
    person = Document(path).tables[-1]
    assert len(person.rows) == 9


def test_form_286_rows_verbatim(tmp_path):
    """С 01.09.2026 — приложение № 2 к приказу № 286, дословно."""
    ctx = _ul_ctx()
    ctx.extra["waste_details"] = {"7 23 101 01 39 4": {
        "site_address": ["СПб, пл. 1", "СПб, пл. 2"]}}   # п. 6 Порядка — несколько адресов
    (path,) = wp.generate(ctx, tmp_path, approved_date="01.09.2026")
    cells, paras = _texts(path)
    assert ("Наименование вида отходов по федеральному классификационному "
            "каталогу отходов (далее - ФККО)") in cells
    assert "Код вида отходов по ФККО" in cells
    assert ("Происхождение и условия образования отходов (указывается "
            "наименование технологического процесса, в результате которого "
            "образовался отход, или процесса, в результате которого товар "
            "(продукция) утратил (утратила) свои потребительские свойства, "
            "с указанием наименования исходного товара)") in cells
    assert ("Способ определения химического и (или) компонентного состава "
            "вида отходов (указывается согласно технологическим регламентам, "
            "техническим условиям, стандартам, руководствам по эксплуатации, "
            "проектной документации и (или) с использованием количественного "
            "химического анализа)") in cells
    assert ("Фамилия, имя, отчество (при наличии) индивидуального "
            "предпринимателя или полное и (или) сокращенное наименования "
            "юридического лица") in cells
    assert "Сокращенное наименование юридического лица" not in cells
    assert "Идентификационный номер налогоплательщика (ИНН)" in cells
    assert "Индивидуальный номер налогоплательщика (ИНН)" not in cells
    assert ("Адрес регистрации индивидуального предпринимателя по месту "
            "жительства или адрес юридического лица в пределах места "
            "нахождения юридического лица") in cells
    assert "Почтовый адрес" in cells
    assert "Адрес (адреса) места (мест) образования отходов" in cells
    assert "СПб, пл. 1; СПб, пл. 2" in cells
    assert 'Общество с ограниченной ответственностью «ТЕХНОСТРОЙ» (ООО «ТЕХНОСТРОЙ»)' in cells
    assert "включенных в федеральный классификационный каталог отходов" in paras
    assert "место печати" in paras and "М.П." not in paras
    assert "Руководитель юридического лица (индивидуальный предприниматель)" in paras
    # «Сведения о лице» — 7 строк новой формы + заголовок
    person = Document(path).tables[-1]
    assert len(person.rows) == 8
    assert [r.cells[0].text for r in person.rows[1:]] == [
        ("Фамилия, имя, отчество (при наличии) индивидуального предпринимателя "
         "или полное и (или) сокращенное наименования юридического лица"),
        "Идентификационный номер налогоплательщика (ИНН)",
        "Код по Общероссийскому классификатору предприятий и организаций (ОКПО)",
        "Код по Общероссийскому классификатору видов экономической деятельности (ОКВЭД)",
        ("Адрес регистрации индивидуального предпринимателя по месту жительства "
         "или адрес юридического лица в пределах места нахождения юридического лица"),
        "Почтовый адрес",
        "Адрес (адреса) места (мест) образования отходов",
    ]


def test_approved_date_from_details_and_default_today(tmp_path, monkeypatch):
    """Дата берётся из waste_details['approved_date'], а без неё — сегодня."""
    ctx = _ul_ctx()
    ctx.extra["waste_details"] = {"7 23 101 01 39 4": {"approved_date": "2026-09-15"}}
    (path,) = wp.generate(ctx, tmp_path / "a")
    assert "место печати" in _texts(path)[1]

    class _D(dt.date):
        @classmethod
        def today(cls):
            return cls(2026, 8, 21)
    monkeypatch.setattr(wp._dt, "date", _D)
    (path,) = wp.generate(_ul_ctx(), tmp_path / "b")
    assert "М.П." in _texts(path)[1]


def test_aggregate_state_by_fkko_code(tmp_path):
    """9–10 знаки кода ФККО → агрегатное состояние (как в эталонах ТЕХНОСТРОЙ)."""
    assert wp.aggregate_state("7 23 101 01 39 4") == "Прочие дисперсные системы"
    assert wp.aggregate_state("8 22 401 01 21 4") == "Кусковая форма"
    assert wp.aggregate_state("73310001724") == (
        "Смесь твердых материалов (включая волокна) и изделий")
    assert wp.aggregate_state("47110101521") == "Изделия из нескольких материалов"
    assert wp.aggregate_state("4711010152") == ""          # неполный код
    assert wp.aggregate_state("7 23 101 01 88 4") == ""    # нет такой пары

    ctx = _ul_ctx()
    (path,) = wp.generate(ctx, tmp_path / "auto", approved_date=dt.date(2026, 1, 1))
    assert "Прочие дисперсные системы" in _texts(path)[0]
    # ручной ввод приоритетнее
    ctx.extra["waste_details"] = {"7 23 101 01 39 4": {wp._AGG: "шлам (ручной)"}}
    (path,) = wp.generate(ctx, tmp_path / "manual", approved_date=dt.date(2026, 1, 1))
    cells = _texts(path)[0]
    assert "шлам (ручной)" in cells and "Прочие дисперсные системы" not in cells


# ─────────── сверка со всеми паспортами ПОО (ТЕХНОСТРОЙ: ТБО, ЛКМ, паркинг) ───────────

def _docx_text(path):
    from docx import Document
    d = Document(path)
    parts = [p.text for p in d.paragraphs]
    parts += [c.text for t in d.tables for r in t.rows for c in r.cells]
    return "\n".join(parts)


def test_passport_composition_from_kha_protocol_with_requisites(tmp_path):
    """Принятый паспорт ТБО ТЕХНОСТРОЙ (ПОО/П.о.о ТБО.pdf): состав — 11
    компонентов из протокола ИЦ ООО «ТАСИС» № 20002.25-1-Отх от 27.02.2025
    (аттестат РОСС RU.0001.21АУ50, МИ М-27-2023), протокол подшит к паспорту,
    дата в грифе «28» февраля 2025 г., способ — «количественный морфологический
    анализ отхода». У нас: состав берётся из extra.lab_results (КХА), реквизиты
    протокола/лаборатории печатаются листом «Основание…», дата — в грифе."""
    ctx = _ctx()
    ctx.wastes = [WasteFlow(fkko_code="7 33 100 01 72 4",
                            name="Мусор от офисных и бытовых помещений организаций "
                                 "несортированный (исключая крупногабаритный)",
                            hazard_class=4)]
    ctx.extra["lab_results"] = [{
        "kind": "КХА", "protocol_no": "20002.25-1-Отх", "date": "27.02.2025",
        "lab": "ИЦ ООО «ТАСИС»", "lab_attestation": "№ РОСС RU.0001.21АУ50",
        "method": "М-27-2023",
        "object": "мусор от офисных и бытовых помещений организаций несортированный",
        "substances": [{"name": "Бумага, картон", "value": "31,5", "unit": "%"},
                       {"name": "Древесина", "value": "2,1", "unit": "%"},
                       {"name": "Полиэтилентерефталат", "value": "11,3", "unit": "%"}]}]
    ctx.extra["waste_details"] = {"7 33 100 01 72 4": {
        "origin": "жизнедеятельность работников",
        "method": "количественный морфологический анализ отхода",
        "approved_date": "28.02.2025"}}
    (path,) = wp.generate(ctx, tmp_path)
    text = _docx_text(path)
    assert "«28» февраля 2025 г." in text                      # дата в грифе
    assert "количественный морфологический анализ отхода" in text
    assert "Смесь твердых материалов (включая волокна) и изделий" in text  # …72 4
    # состав из протокола, в порядке убывания
    assert text.index("Бумага, картон") < text.index("Полиэтилентерефталат") < text.index("Древесина")
    # лист «Основание…» с реквизитами протокола и лаборатории
    assert "Основание для определения" in text
    assert "№ 20002.25-1-Отх от 27.02.2025" in text
    assert "ИЦ ООО «ТАСИС»" in text and "РОСС RU.0001.21АУ50" in text
    assert "М-27-2023" in text


def test_passport_protocol_from_waste_details_and_blank_date(tmp_path):
    """Ручные реквизиты протокола (waste_details[код]['protocol']) печатаются;
    без даты утверждения гриф остаётся бланком «____»; без протокола третьего
    листа нет."""
    ctx = _ctx()
    ctx.wastes = [ctx.wastes[1]]
    ctx.extra["waste_details"] = {"73310001724": {
        "components": [{"name": "Бумага", "percent": "60"}],
        "protocol": {"number": "7/25", "date": "01.02.2025", "lab": "ООО Лаб",
                     "lab_attestation": "RA.RU.21XX"}}}
    (path,) = wp.generate(ctx, tmp_path)
    text = _docx_text(path)
    assert "«____» ________________ 20____ г." in text
    assert "№ 7/25 от 01.02.2025" in text and "RA.RU.21XX" in text
    ctx.extra["waste_details"]["73310001724"].pop("protocol")
    (path,) = wp.generate(ctx, tmp_path / "b")
    assert "Основание для определения" not in _docx_text(path)


def test_v_class_confirmation_by_biotest(tmp_path):
    """V класс: паспорт не делается, но по акту/протоколу биотестирования
    (ПОО: «акт о.п_биотест 5 кл» ТАСИС № 20102.25-1-4 от 27.01.2025, цель —
    подтверждение V класса; протокол БИО ЦЭД) печатается справка о
    подтверждении V класса со ссылкой на акт, лабораторию и аттестат."""
    ctx = _ctx()
    ctx.extra["lab_results"] = [{
        "kind": "биотест", "protocol_no": "20102.25-1-4", "date": "27.01.2025",
        "lab": "ИЦ ООО «ТАСИС»", "lab_attestation": "№ РОСС RU.0001.21АУ50",
        "fkko": "73310002725"}]
    paths = wp.generate_v_class(ctx, tmp_path)
    assert [p.name for p in paths] == ["V_класс_73310002725.docx"]
    text = _docx_text(paths[0])
    assert "V КЛАССУ" in text and "№ 20102.25-1-4 от 27.01.2025" in text
    assert "ТАСИС" in text and "№ 158" in text and "биотестирование" in text
    # паспортов на V класс по-прежнему нет
    assert not any("73310002725" in p.name for p in wp.generate(ctx, tmp_path / "p"))
    # без акта биотестирования справка не выдумывается
    ctx.extra["lab_results"] = []
    assert wp.generate_v_class(ctx, tmp_path / "n") == []
