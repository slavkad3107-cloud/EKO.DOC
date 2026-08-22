"""Раздел ООС (ПП РФ № 87, п. 25): структура тома как у настоящих разделов.

Эталон структуры — АК-01-25-ООС1 (ООО «Интегра», 2026) и 03-26-1-ООС
«Выборжец» (прошёл экспертизу): заголовки разделов и таблиц проверяются
ДОСЛОВНО (oos_structure), расчёты отходов — по типовым формулам томов.
"""
from decimal import Decimal

import pytest
from docx import Document

from ecodoc.core.models import (Medium, NVOSObject, Pollutant, ReportContext,
                                WasteFlow)
from ecodoc.development import oos, oos_structure, oos_waste_calc as wc


@pytest.fixture()
def ctx():
    c = ReportContext()
    c.organization.name = "ООО «Тест»"
    c.organization.inn = "7806001144"
    c.organization.address = "СПб, ул. Тестовая, 1"
    c.period.year = 2026
    c.objects = [NVOSObject(code="41-0247-005048-П", category="III",
                            address="Промзона Янино", name="Склад")]
    c.wastes = [WasteFlow(fkko_code="73310001724", name="Мусор офисный",
                          hazard_class=4, generated=Decimal("1.9"),
                          transferred=Decimal("1.9"))]
    air = Pollutant(code="0301", name="Азота диоксид", mass_norm=Decimal("0.412"))
    air.medium = Medium.AIR
    c.pollutants = [air]
    c.extra["emission_sources"] = [{"number": "0001", "name": "Котельная",
                                    "pollutants": [{"code": "301", "g_s": "0.05",
                                                    "t_year": "0.412"}]}]
    return c


@pytest.fixture()
def full_ctx(ctx):
    """Полный контекст: всё, что приходит из ПОС/ИЭИ/УПРЗА/ИОС через extra.oos."""
    ctx.extra["oos"] = {
        "project": {"code": "АК-01-25", "title": "Гостиничный комплекс", "stage": "П",
                    "designer": "ООО «Интегра»", "cadastral": "23:30:0301001:905",
                    "description": "Здание гостиницы на 134 номера.\nПодвал технический.",
                    "months": 24, "workers": 39, "itr": 8, "shifts": 2, "days": 720,
                    "tep": [{"name": "Площадь застройки", "unit": "м2", "value": 3417.13}],
                    "land_tep": [{"name": "Площадь земельного участка", "m2": 8178}],
                    "machinery": [{"area": "Земляные работы", "name": "Экскаватор",
                                   "brand": "CAT 320DL", "spec": "ковш 0,5 м3", "qty": 1}]},
        "iei": {"report": "АК-01-25-ИЭИ", "soil_category": "чистая",
                "radiation": "в пределах нормы", "oopt": "отсутствуют",
                "flora_fauna": "Редкие виды не встречены.",
                "background": [{"name": "Диоксид азота", "value": 0.054, "unit": "мг/м3"}],
                "noise": [{"place": "Т1", "leq": 42.6, "lmax": 55, "minutes": 30, "pdu": 55}]},
        "climate": {"A": 200, "t_warm": 25.2, "t_cold": -0.3, "u5": 8.5,
                    "wind": {"С": 12, "В": 21}},
        "construction": {
            "sources": [{"number": "6001", "kind": "передвижной", "name": "Земляные работы",
                         "height": 5, "pollutants": [{"code": "301", "name": "Азота диоксид",
                                                      "g_s": 0.085}]},
                        {"number": "0001", "kind": "стационарный", "name": "Работа ДГУ",
                         "height": 2, "organized": True}],
            "pollutants": [{"code": "301", "name": "Азота диоксид", "criterion": "ПДК м/р",
                            "value": 0.2, "hazard": 3, "g_s": 0.085, "t_year": 0.115694},
                           {"code": "337", "name": "Углерод оксид", "criterion": "ПДК м/р",
                            "value": 5, "hazard": 4, "g_s": 0.152, "t_year": 0.169}],
            "points": [{"code": 1, "x": -284.43, "y": -211.51, "h": 2,
                        "type": "на границе жилой зоны"}],
            "skip": [{"code": "703", "name": "Бенз/а/пирен", "sum": 0}],
            "concentrations": [{"code": "301", "name": "Азота диоксид",
                                "values": [0.41, 0.33]}],
            "noise_sources": [{"n": "001", "name": "Экскаватор", "x": 0, "y": 0, "h": 1.5,
                               "spectrum": "…", "leq": 80, "lmax": 85}],
            "noise_results": [{"n": 1, "name": "Жилая застройка", "leq": 48.4, "lmax": 53.6}],
            "noise_bg": [{"n": 1, "name": "Жилая застройка", "bg": 42.6, "leq": 48.4,
                          "total": 49.4}],
            "materials": [{"name": "Бетон", "kind": "бетон", "qty": 4027, "unit": "м3"},
                          {"name": "Песок", "kind": "песок", "qty": 5040, "unit": "м3"},
                          {"name": "Стальные конструкции", "kind": "металл", "qty": 9.6,
                           "unit": "т"},
                          {"name": "Неизвестный материал", "qty": 10}],
            "electrodes_t": 0.469, "wheel_wash": {"cars_per_day": 3, "water_m3": 0.3},
            "soil_excess_m3": 6208},
        "operation": {
            "points": [{"x": 1233407.13, "y": 513768.59, "h": 2, "type": "на границе жилой зоны",
                        "comment": "Расчётная точка 001"}],
            "concentrations": [{"code": "301", "name": "Азота диоксид", "values": [0.24, 0.4]}],
            "noise_sources": [{"n": "001", "name": "В2", "spectrum": "…", "leq": 79.1, "r": 1}],
            "noise_bg": [{"n": 1, "name": "Жилая застройка", "bg": 43, "const": 43, "var": 28,
                          "total": 46}],
            "water_balance": [{"consumer": "Хозяйственно-бытовые нужды", "in_day": 130.79,
                               "in_year": 47738, "out_day": 130.79, "out_year": 47738,
                               "to": "КОС → вывоз"}],
            "wastes_norm": [{"name": "Мусор от офисных и бытовых помещений организаций "
                             "несортированный (исключая крупногабаритный)",
                             "fkko": "73310001724", "hazard": 4, "count": 58,
                             "count_unit": "чел.", "norm_m3": 2.8, "density": 0.2}],
            "lamps": [{"name": "Светильники ARS/R", "count": 445, "life_h": 30000,
                       "hours": 4380, "mass_kg": 2}],
            "waste_handling": {"73310001724": "передача региональному оператору"}},
        "pek": {"category": "III", "periodicity": "1 раз в год"},
    }
    return ctx


def _text(path):
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def _captions(path):
    return [p.text for p in Document(str(path)).paragraphs if p.text.startswith("Таблица ")]


# ── структура ──────────────────────────────────────────────────────────
def test_structure_matches_real_volume():
    """Разделы 1–11 эталона АК-01-25 — дословно, включая длинное название ПЭК."""
    titles = oos_structure.section_titles()
    for must in ("Введение", "Характеристика природных условий района",
                 "Проектные решения", "Охрана земельных ресурсов",
                 "Охрана атмосферного воздуха",
                 "Оценка воздействия на атмосферный воздух в период строительства",
                 "Предложения по нормативам ПДВ",
                 "Разработка мероприятий по охране и рациональному использованию "
                 "водных ресурсов",
                 "Охрана растительного и животного мира",
                 "Акустическое воздействие объекта на прилегающую территорию",
                 "Разработка мероприятий по сбору, использованию, обезвреживанию, "
                 "транспортировке и размещению опасных отходов",
                 "Перечень и характеристики строительных отходов",
                 "Перечень и характеристики отходов, образующихся в период эксплуатации",
                 "Перечень затрат на реализацию природоохранных мероприятий и "
                 "компенсационных выплат",
                 "План-график ПЭКиМ", "Список литературы", "Приложения"):
        assert must in titles, must
    assert any(t.startswith("Программа производственного экологического контроля "
                            "(мониторинга)") for t in titles)


def test_all_tables_of_real_volume_present(full_ctx, tmp_path):
    """Каждая таблица эталона (шапка дословно) есть в сгенерированном томе."""
    path = oos.generate(full_ctx, tmp_path / "оос.docx")
    caps = _captions(path)
    for num, title, _head in oos_structure.TABLES.values():
        assert any(c.startswith(f"Таблица {num} – ") for c in caps), (num, title)
    text = _text(path)
    for head in ("№ и вид источника выброса*", "Суммарный выброс вещества г/с",
                 "Удельный норматив образования отхода, %",
                 "Ставка платы за 1 тонну загрязняющих веществ руб./тонну",
                 "Норматив выброса, г/с", "Сведения об организации, производящей контроль"):
        assert head in text, head
    assert len(caps) >= len(oos_structure.TABLES)


def test_headings_numbered_like_real_volume(full_ctx, tmp_path):
    text = _text(oos.generate(full_ctx, tmp_path / "оос.docx"))
    assert "5.1. Оценка воздействия на атмосферный воздух в период строительства" in text
    assert "9.3. Перечень и характеристики отходов, образующихся в период эксплуатации" in text
    assert "ПРОЕКТНАЯ ДОКУМЕНТАЦИЯ" in text and "Том 8" in text
    assert "Аннотация" in text and "Оглавление" in text and "Содержание тома" in text
    for i, a in enumerate(oos_structure.APPENDICES_TEXT, start=1):
        assert f"Приложение {i}. {a}" in text
    assert "Карта-схема с обозначением источников выбросов ЗВ" in text


def test_data_flow_into_tables(full_ctx, tmp_path):
    text = _text(oos.generate(full_ctx, tmp_path / "оос.docx"))
    assert "ООО «Тест»" in text and "41-0247-005048-П" in text
    assert "Гостиничный комплекс" in text and "АК-01-25-ООС" in text
    assert "Азота диоксид" in text and "0.412" in text and "0.05" in text   # т/год и г/с
    assert "7 33 100 01 72 4" in text                                         # ФККО с пробелами
    assert "CAT 320DL" in text and "23:30:0301001:905" in text
    assert "130.79" in text                                                  # баланс воды
    assert "0.054" in text                                                   # фон
    assert "не превышают ПДК" in text                                        # вывод по рассеиванию
    assert "АК-01-25-ИЭИ" in text and "не воспроизводятся" in text           # ссылка на ИЭИ


def test_iei_not_reproduced(full_ctx, tmp_path):
    """Решение пользователя: раздел 2 ссылается на отчёт ИЭИ, а не пересказывает его."""
    text = _text(oos.generate(full_ctx, tmp_path / "оос.docx"))
    assert "см. отчёт ИЭИ" in text


# ── расчёт отходов ─────────────────────────────────────────────────────
def test_construction_waste_formulas_match_real_volume():
    """Числа из АК-01-25: бетон 4027 м3 × 0,3 % → 12,081 м3 / 28,994 т; песок 5040 м3 ×
    1,5 % → 75,6 м3 / 113,4 т; ТКО 39 раб. × 0,22 × 24/12 × 0,18 → 3,089 т; огарки
    0,469 т × 15 % → 0,070 т; мойка колёс 0,9 м3 × 720 дн → 0,292 / 6,966 т."""
    mats = wc.material_waste([{"name": "Бетон", "kind": "бетон", "qty": 4027},
                              {"name": "Песок", "kind": "песок", "qty": 5040}])
    assert (mats[0]["m3"], mats[0]["t"]) == (Decimal("12.081"), Decimal("28.994"))
    assert (mats[1]["m3"], mats[1]["t"]) == (Decimal("75.600"), Decimal("113.400"))
    tko = wc.tko_construction(39, 8, 24)
    assert tko[0]["t"] == Decimal("3.089") and tko[1]["t"] == Decimal("3.168")
    assert wc.electrodes(Decimal("0.469"))["t"] == Decimal("0.070")
    ww = wc.wheel_wash(Decimal(3), Decimal("0.3"), 720)
    assert ww[0]["t"] == Decimal("0.292") and ww[1]["t"] == Decimal("6.966")
    assert wc.soil_excess(Decimal(6208))["t"] == Decimal("9932.800")
    assert wc.cesspool(47, 1440)["m3"] == Decimal("0.372")


def test_operation_waste_by_norm_and_lamps():
    rows = wc.by_norm([{"name": "ТКО", "fkko": "73310001724", "count": 58,
                        "count_unit": "чел.", "norm_m3": 2.8, "density": 0.2}])
    assert rows[0]["m3"] == Decimal("162.400") and rows[0]["t"] == Decimal("32.480")
    lamps = wc.lamps([{"name": "ARS/R", "count": 445, "life_h": 30000, "hours": 4380,
                       "mass_kg": 2}])
    assert lamps[0]["replaced"] == Decimal("64.970") and lamps[0]["t"] == Decimal("0.130")


def test_construction_waste_summary_by_class(full_ctx):
    cw = oos.construction_wastes(full_ctx)
    codes = {r["fkko"] for r in cw["summary"]}
    assert {"73310001724", "72310101394", "82220101215", "81910001495",
            "46101001205", "91910001205", "81110001495", "73210001304"} <= codes
    unknown = [r for r in cw["materials"] if r["note"]]
    assert unknown and "Неизвестный материал" in unknown[0]["note"]


def test_operation_wastes_merge_context_and_norms(full_ctx):
    ow = oos.operation_wastes(full_ctx)
    by = {r["fkko"]: r for r in ow["summary"]}
    assert by["73310001724"]["t"] == Decimal("32.480")       # расчёт по нормативу — приоритет
    assert by["73310001724"]["handling"] == "передача региональному оператору"
    assert "48241100525" in by                                 # лампы


# ── пробелы и устойчивость ─────────────────────────────────────────────
def test_gaps_name_external_data(ctx):
    gaps = oos.gaps(ctx)
    assert any("рассеивания" in g for g in gaps)
    assert any("инженерно-экологических" in g for g in gaps)
    assert any("акустический" in g for g in gaps)
    assert any("ведомость потребности" in g for g in gaps)
    assert not any("наименование организации" in g for g in gaps)


def test_gaps_shrink_with_full_context(full_ctx):
    import copy
    bare = copy.deepcopy(full_ctx)
    bare.extra.pop("oos")                     # фикстуры делят один объект
    assert len(oos.gaps(full_ctx)) < len(oos.gaps(bare))
    assert not any("рассеивания" in g for g in oos.gaps(full_ctx))


def test_gaps_report_empty_object():
    gaps = oos.gaps(ReportContext())
    assert any("наименование организации" in g for g in gaps)
    assert any("не заведён объект" in g for g in gaps)
    assert any("нечем наполнять" in g for g in gaps)


def test_document_generated_even_without_data(tmp_path):
    path = oos.generate(ReportContext(), tmp_path / "пустой.docx")
    text = _text(path)
    assert "[требуется" in text and oos.TITLE in text
    assert len(_captions(path)) >= len(oos_structure.TABLES)   # все таблицы даже пустые


def test_water_discharges_listed_when_exist(ctx, tmp_path):
    w = Pollutant(code="1000", name="Взвешенные вещества", mass_norm=Decimal("0.3"))
    w.medium = Medium.WATER
    ctx.pollutants.append(w)
    text = _text(oos.generate(ctx, tmp_path / "оос.docx"))
    assert "Взвешенные вещества" in text and "0.3" in text


def test_typical_wording_is_marked(ctx, tmp_path):
    text = _text(oos.generate(ctx, tmp_path / "оос.docx"))
    assert "типовая формулировка" in text


def test_costs_tables_use_rates(full_ctx, tmp_path):
    text = _text(oos.generate(full_ctx, tmp_path / "оос.docx"))
    assert "Таблица 10.1.1 – " in text and "Таблица 10.2.1 – " in text
    assert "2409-р" in text                                    # ставки 2026–2030


def test_registered_as_dev_document():
    from ecodoc.core import registry
    registry.load_all()
    form = registry.all_reports()["oos"]
    assert form.domain == "development" and form.devdoc is True
