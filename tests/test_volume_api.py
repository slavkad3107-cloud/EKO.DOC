"""Сборка тома НДВ/НДС/СЗЗ: обвязка обмена с УПРЗА через API."""
import base64
import io

import openpyxl
import pytest
from docx import Document

from ecodoc.core import workspace
from ecodoc.gui import server


def _xlsx_b64(header, rows) -> str:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(header)
    for r in rows:
        ws.append(r)
    buf = io.BytesIO()
    wb.save(buf)
    return base64.b64encode(buf.getvalue()).decode()


@pytest.fixture()
def site(tmp_path, monkeypatch):
    monkeypatch.setenv("ECODOC_RESULTS", str(tmp_path / "res"))
    workspace.add_org("ОРГ")
    workspace.add_site("ОРГ", "Пл")
    ctx = workspace.load_context("ОРГ", "Пл")
    ctx.organization.name = "ООО «Тест»"
    ctx.organization.inn = "7806001144"
    ctx.period.year = 2025
    ctx.extra["emission_sources"] = [
        {"number": "0001", "name": "Котельная", "kind": "организованный",
         "pollutants": [{"code": "0301", "name": "Азота диоксид",
                         "g_s": "0.05", "t_year": "0.412"}]}]
    workspace.save_context("ОРГ", "Пл", ctx)
    return ("ОРГ", "Пл")


def _text(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def test_volume_with_ecolog_exports(site):
    org, st = site
    out = server.api_volume({}, {
        "org": org, "site": st, "vtype": "ndv",
        "sources_file": {"name": "источники.xlsx", "b64": _xlsx_b64(
            ["№", "Наименование", "H, м"], [["0001", "Котельная", 12]])},
        "dispersion_file": {"name": "рассеивание.xlsx", "b64": _xlsx_b64(
            ["Вещество", "Cмакс, доли ПДК"], [["Азота диоксид", 0.42]])}})
    assert out["path"].endswith(".docx")
    text = _text(out["path"])
    assert "НОРМАТИВОВ ДОПУСТИМЫХ ВЫБРОСОВ" in text
    assert "Котельная" in text and "0.42" in text
    assert not any(n.startswith("⚠") for n in out["notes"])


def test_volume_without_exports_uses_inventory_and_warns(site):
    org, st = site
    out = server.api_volume({}, {"org": org, "site": st, "vtype": "ndv"})
    text = _text(out["path"])
    assert "Котельная" in text                        # из инвентаризации
    assert any("из инвентаризации" in n for n in out["notes"])
    assert any("рассеивания нет" in n for n in out["notes"])


def test_volume_bad_file_reported_not_fatal(site):
    org, st = site
    out = server.api_volume({}, {
        "org": org, "site": st, "vtype": "nds",
        "dispersion_file": {"name": "битый.xlsx",
                            "b64": base64.b64encode("мусор".encode()).decode()}})
    assert out["path"].endswith(".docx")
    assert any("не прочитан" in n for n in out["notes"])
