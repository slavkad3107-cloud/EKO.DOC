"""ДВОС (ст. 31.2 ФЗ-7, приказ МПР от 19.03.2025 № 117): состав, данные,
пробелы. Форма — приложение № 1 к приказу № 117, сверена дословно 22.08.2026
по normativ.kontur.ru (documentId=493064)."""
from decimal import Decimal

import pytest
from docx import Document

from ecodoc.core.models import (Medium, NVOSObject, Pollutant, ReportContext,
                                WasteFlow)
from ecodoc.development import dvos


@pytest.fixture()
def ctx():
    c = ReportContext()
    c.organization.name = "ООО «Тест»"
    c.organization.inn = "7806001144"
    c.organization.ogrn = "1047855175785"
    c.organization.okved = "38.11 Сбор неопасных отходов"
    c.organization.address = "СПб, ул. Тестовая, 1"
    c.period.year = 2026
    c.objects = [NVOSObject(code="40-0178-001234-П", category="II",
                            address="Промзона Парнас")]
    air = Pollutant(code="0301", name="Азота диоксид",
                    mass_norm=Decimal("0.412"))
    air.medium = Medium.AIR
    water = Pollutant(code="1000", name="Взвешенные вещества",
                      mass_norm=Decimal("0.3"), mass_over=Decimal("0.05"))
    water.medium = Medium.WATER
    c.pollutants = [air, water]
    c.wastes = [WasteFlow(fkko_code="73310001724", name="Мусор офисный",
                          hazard_class=4, generated=Decimal("1.9"),
                          transferred=Decimal("1.9"))]
    c.extra["dvos"] = {
        "products": [{"name": "Изделия пластмассовые", "volume": 120,
                      "unit": "т"}],
        "measures": [{"name": "Установка циклона", "start": "2024",
                      "end": "2025", "cost": 500, "funding": "собственные",
                      "result": "снижение выбросов"}],
        "accidents": [],
        "authority": "Северо-Западное межрегиональное управление "
                     "Росприроднадзора",
        "executor": "эколог Петрова А.А., +7 812 000-00-00, eco@test.ru",
        "pollutant_details": {
            "0301": {"class": "3", "source": "ИЗАВ № 0001 котельная",
                     "gs": "0.0131"},
            "1000": {"class": "4", "source": "выпуск № 1",
                     "nds": "10.45", "water_body": "р. Охта"},
        },
    }
    c.extra["pek"] = {"approved_date": "10.01.2026",
                      "responsible": "ген. директор Иванов И.И.",
                      "authority": "СЗМУ Росприроднадзора",
                      "last_report_date": "20.03.2026"}
    return c


def _text(path):
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts += [c.text for c in row.cells]
    return "\n".join(parts)


def test_all_form_sections_present(ctx, tmp_path):
    """Все разделы формы приказа № 117 (дословно), шапка «В ___», поля
    титула, приложения и сноски — в документе."""
    text = _text(dvos.generate(ctx, tmp_path / "двос.docx"))
    for _key, title in dvos.SECTIONS:
        assert title.format(y_from=2019, y_to=2025) in text, title
    assert "произошедших за 2019 - 2025 годы" in text
    assert "19.03.2025 № 117" in text          # действующая форма
    assert "31.2" in text and "7-ФЗ" in text   # закон-основание
    assert "11.10.2018 № 509" in text          # явно помечен как заменённый
    # шапка и титул — поля формы
    assert "В Северо-Западное межрегиональное управление" in text
    assert "организационно-правовая форма юридического лица" in text
    assert "Код основного вида экономической деятельности" in text
    assert "Наименование основного вида экономической деятельности" in text
    assert "Декларация составлена на ___ листах" in text
    assert "Исполнитель, ответственный за представление Декларации" in text
    assert "Петрова А.А." in text
    assert "Руководитель юридического лица/индивидуальный предприниматель" \
        in text
    # полей, которых нет в форме, на титуле быть не должно
    assert "Период действия декларации" not in text
    assert "ИНН / ОГРН" not in text
    # приложения и сноски формы
    assert "Приложениями к Декларации являются" in text
    assert "расчет нормативов допустимых сбросов" in text
    assert "квоты выбросов" in text
    assert "2909-р" in text and "СанПиН 1.2.3685-21" in text


def test_table_columns_match_form(ctx, tmp_path):
    """Графы таблиц — дословно по форме № 117."""
    doc = Document(str(dvos.generate(ctx, tmp_path / "двос.docx")))
    heads = [[c.text for c in t.rows[0].cells] for t in doc.tables]
    assert dvos.COLS_I in heads and dvos.COLS_II in heads
    assert dvos.COLS_III_ACC in heads and dvos.COLS_III_INC in heads
    assert dvos.COLS_IV in heads and dvos.COLS_V in heads
    assert dvos.COLS_VI_1 in heads and dvos.COLS_VI_2 in heads
    assert len(dvos.COLS_I) == 5 and len(dvos.COLS_II) == 7
    assert len(dvos.COLS_III_ACC) == 7 and len(dvos.COLS_IV) == 8
    assert len(dvos.COLS_V) == 9 and len(dvos.COLS_VI_1) == 9
    assert "Код производимой продукции (товара)" in dvos.COLS_I
    assert "Источники финансирования" in dvos.COLS_II
    assert "грамм/секунду" in dvos.COLS_IV[4]
    assert "миллиграмм/кубический дециметр" in dvos.COLS_V[5]
    assert "ГРОРО" in dvos.COLS_VI_1[6] and "ГРОРО" in dvos.COLS_VI_1[8]
    # строка с номерами граф 1..N под шапкой — как в бланке
    t_iv = doc.tables[[i for i, h in enumerate(heads)
                       if h == dvos.COLS_IV][0]]
    assert [c.text for c in t_iv.rows[1].cells] == [str(i) for i in
                                                    range(1, 9)]
    # раздел IV: класс, источник, г/с, всего / в пределах / с превышением
    assert [c.text for c in t_iv.rows[2].cells] == [
        "1", "Азота диоксид", "3", "ИЗАВ № 0001 котельная", "0.0131",
        "0.412", "0.412", "0"]
    t_v = doc.tables[[i for i, h in enumerate(heads)
                      if h == dvos.COLS_V][0]]
    assert [c.text for c in t_v.rows[2].cells] == [
        "1", "р. Охта", "Взвешенные вещества", "4", "выпуск № 1", "10.45",
        "0.35", "0.3", "0.05"]


def test_tables_filled_from_context(ctx, tmp_path):
    text = _text(dvos.generate(ctx, tmp_path / "двос.docx"))
    assert "ООО «Тест»" in text and "40-0178-001234-П" in text
    assert "38.11" in text                                   # ОКВЭД на титуле
    assert "Азота диоксид" in text and "0.412" in text       # раздел IV
    assert "Взвешенные вещества" in text and "0.35" in text  # раздел V, сумма
    assert "Мусор офисный" in text and "73310001724" in text # раздел VI
    assert "Изделия пластмассовые" in text                   # раздел I
    assert "Установка циклона" in text                       # раздел II
    assert "Иванов И.И." in text and "10.01.2026" in text    # раздел VII
    assert "СЗМУ Росприроднадзора" in text and "20.03.2026" in text
    assert "собственные" in text                             # источники фин.


def test_period_is_seven_years(ctx, tmp_path):
    assert dvos.declared_period(ctx) == (2026, 2032)
    assert dvos.accident_years(ctx) == (2019, 2025)
    text = _text(dvos.generate(ctx, tmp_path / "двос.docx"))
    assert "2026–2032" in text
    assert "на следующие семь лет" in text
    assert "В случае изменения в течение семи лет с даты подачи" in text


def test_waste_section_uses_placement_columns(ctx, tmp_path):
    """Раздел VI: «передано НА РАЗМЕЩЕНИЕ» — только хранение/захоронение, а
    не всё переданное; без № ГРОРО при размещении — пробел."""
    w = ctx.wastes[0]
    w.transferred_burial = Decimal("1.2")      # из 1.9 т на захоронение 1.2
    rows = dvos.rows_waste(ctx)
    assert rows[0]["to_placement"] == Decimal("1.2")
    assert rows[0]["transferred"] == Decimal("1.9")
    assert any("ГРОРО" in g for g in dvos.gaps(ctx))
    ctx.extra["dvos"]["groro"] = {"73310001724": {"other": "78-00012-З-00592-250914"}}
    assert not any("ГРОРО" in g for g in dvos.gaps(ctx))
    text = _text(dvos.generate(ctx, tmp_path / "двос.docx"))
    assert "78-00012-З-00592-250914" in text


def test_missing_details_are_gaps(ctx, tmp_path):
    """Без класса/источника/г/с по веществу и без органа приёма — пробелы."""
    del ctx.extra["dvos"]["pollutant_details"]
    del ctx.extra["dvos"]["authority"]
    gaps = dvos.gaps(ctx)
    assert any("класс опасности" in g and "г/с" in g for g in gaps)
    assert any("водного объекта" in g for g in gaps)
    assert any("органа, уполномоченного на приём" in g for g in gaps)
    text = _text(dvos.generate(ctx, tmp_path / "двос.docx"))
    for g in gaps:
        if g.startswith("требуется:"):
            assert f"[{g}]" in text, g


def test_accidents_empty_list_means_none_happened(ctx, tmp_path):
    """Пустой список аварий — подтверждённое «не было», не пробел."""
    text = _text(dvos.generate(ctx, tmp_path / "двос.docx"))
    assert "не зафиксировано" in text
    assert not any("авариях и инцидентах" in g for g in dvos.gaps(ctx))


def test_accidents_absent_is_a_gap(ctx, tmp_path):
    del ctx.extra["dvos"]["accidents"]
    assert dvos.accidents(ctx) is None
    assert any("авариях и инцидентах" in g for g in dvos.gaps(ctx))
    text = _text(dvos.generate(ctx, tmp_path / "двос.docx"))
    assert "авариях и инцидентах" in text and "[требуется:" in text


def test_gaps_markers_match_document(ctx, tmp_path):
    """Каждый пробел «требуется: …» продублирован пометкой в тексте."""
    text = _text(dvos.generate(ctx, tmp_path / "двос.docx"))
    for g in dvos.gaps(ctx):
        if g.startswith("требуется:"):
            assert f"[{g}]" in text, g


def test_wrong_category_reported(ctx):
    ctx.objects[0].category = "III"
    assert any("II категории" in g for g in dvos.gaps(ctx))


def test_empty_context_still_generates_with_markers(tmp_path):
    empty = ReportContext()
    text = _text(dvos.generate(empty, tmp_path / "пусто.docx"))
    assert "[требуется" in text
    assert dvos.TITLE.upper() in text
    gaps = dvos.gaps(empty)
    assert any("наименование организации" in g for g in gaps)
    assert any("объект НВОС" in g for g in gaps)
    assert any("нечем наполнять" in g for g in gaps)
