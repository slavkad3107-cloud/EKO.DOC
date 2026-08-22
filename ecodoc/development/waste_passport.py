"""Паспорт отходов I–IV классов опасности — типовая форма (приложение № 2).

Две редакции формы, выбираются по дате утверждения паспорта:
  * до 01.09.2026 — приказ Минприроды России от 08.12.2020 № 1026
    (воспроизведена по принятым паспортам ООО «ТЕХНОСТРОЙ», 2025);
  * с 01.09.2026 — приказ Минприроды России от 15.05.2026 № 286
    (рег. Минюст 02.06.2026 № 86840; п. 2 — № 1026 утратил силу, п. 3 —
    действует с 01.09.2026 по 01.09.2032). Формулировки строк — дословно
    по приложению № 2 к № 286 (consultant.ru, cons_doc_LAW_535862).

Лист 1 — гриф «УТВЕРЖДАЮ» + таблица «Сведения об отходах», лист 2 —
таблица «Сведения о лице, которое образовало отходы».

Данные берутся из ReportContext:
  * ctx.wastes — перечень отходов (код ФККО, наименование, класс опасности);
  * ctx.extra['waste_details'][<код>] — состав и реквизиты отхода (ручной ввод);
  * ctx.extra['waste_passports'] — то же, но из ИИ-разбора приложенных паспортов
    и протоколов КХА (fallback, если waste_details не заполнен).

Паспорт оформляется только на отходы I–IV класса: для V класса он не нужен —
требуется подтверждение отнесения к V классу (протокол биотестирования).

Сверено 22.08.2026 со ВСЕМИ паспортами из «Формы/Разработка/ПОО» (ТЕХНОСТРОЙ:
ТБО, банки ЛКМ, осадки мойки, раствор, теплоизоляция): состав паспорта — это
результаты протокола КХА/морфологии ИЦ ООО «ТАСИС» (№ 20002.25-1-Отх от
27.02.2025, аттестат № РОСС RU.0001.21АУ50, МИ М-27-2023), сам протокол
подшит к паспорту, дата в грифе заполнена («28» февраля 2025 г.), способ
определения — «количественный морфологический анализ отхода». Поэтому:
  * дата утверждения печатается в грифе, если известна;
  * реквизиты протокола и лаборатории (extra['waste_details'][код]
    ['protocol'] = {number, date, lab, lab_attestation, method_doc} либо
    extra['lab_results'] вида {kind: "КХА", protocol_no, date, lab, object,
    substances}) печатаются третьим листом «Основание…» — у типовой формы
    строки для них нет, а в принятых паспортах протокол приложен;
  * состав без ручного ввода берётся из extra['lab_results'] (substances →
    компоненты), как из протокола;
  * для V класса generate_v_class() делает справку об отнесении к V классу
    по акту/протоколу биотестирования (в ПОО: «акт о.п_биотест 5 кл»,
    «БИО_26312_25т…» — 3 пробы ТАСИС, цель «подтверждение V класса»).
"""
from __future__ import annotations

import datetime as _dt
from pathlib import Path

from ecodoc.core.models import Organization, ReportContext, WasteFlow

_AGG = "агрегатное состояние и физическая форма"
_TITLE = "ПАСПОРТ ОТХОДОВ I - IV КЛАССОВ ОПАСНОСТИ,"

# Дата смены формы: п. 3 приказа № 286 — «вступает в силу с 1 сентября
# 2026 г.», п. 2 — № 1026 признан утратившим силу. Паспорт, утверждаемый
# с этой даты, оформляется уже по новой типовой форме.
SWITCH_DATE = _dt.date(2026, 9, 1)
FORM_1026 = "1026"
FORM_286 = "286"

# ── формулировки по редакциям: ключ → {форма: текст}. Тексты взяты дословно
# из приложений № 2 к приказам (№ 1026 сверен с эталонами ТЕХНОСТРОЙ,
# № 286 — с consultant.ru); менять «для красоты» нельзя ──
_TEXT = {
    # подзаголовок: в № 286 «федеральный» со строчной
    "subtitle": {
        FORM_1026: "включенных в Федеральный классификационный каталог отходов",
        FORM_286: "включенных в федеральный классификационный каталог отходов",
    },
    # под датой в грифе: «М.П.» (№ 1026) против «место печати» (№ 286)
    "seal": {FORM_1026: "М.П.", FORM_286: "место печати"},
    "name": {
        FORM_1026: "Наименование вида отходов по ФККО",
        FORM_286: "Наименование вида отходов по федеральному классификационному "
                  "каталогу отходов (далее - ФККО)",
    },
    "origin": {
        FORM_1026: ("Происхождение отходов (указывается наименование "
                    "технологического процесса, в результате которого "
                    "образовался отход, или процесса, в результате которого "
                    "товар (продукция) утратил свои потребительские свойства, "
                    "с указанием наименования исходного товара)"),
        FORM_286: ("Происхождение и условия образования отходов (указывается "
                   "наименование технологического процесса, в результате "
                   "которого образовался отход, или процесса, в результате "
                   "которого товар (продукция) утратил (утратила) свои "
                   "потребительские свойства, с указанием наименования "
                   "исходного товара)"),
    },
    "method": {
        FORM_1026: ("Способ определения химического и (или) компонентного "
                    "состава вида отходов (указывается согласно документации "
                    "и (или) с использованием количественного химического "
                    "анализа)"),
        FORM_286: ("Способ определения химического и (или) компонентного "
                   "состава вида отходов (указывается согласно "
                   "технологическим регламентам, техническим условиям, "
                   "стандартам, руководствам по эксплуатации, проектной "
                   "документации и (или) с использованием количественного "
                   "химического анализа)"),
    },
    # «Сведения о лице»: в № 286 полное и сокращённое наименования — одной
    # строкой, ИНН — «Идентификационный», адреса переименованы
    "person": {
        FORM_1026: ("Фамилия, имя, отчество (при наличии) индивидуального "
                    "предпринимателя или полное наименование юридического лица"),
        FORM_286: ("Фамилия, имя, отчество (при наличии) индивидуального "
                   "предпринимателя или полное и (или) сокращенное "
                   "наименования юридического лица"),
    },
    "inn": {
        FORM_1026: "Индивидуальный номер налогоплательщика (ИНН)",
        FORM_286: "Идентификационный номер налогоплательщика (ИНН)",
    },
    "location": {
        FORM_1026: "Место нахождения",
        FORM_286: ("Адрес регистрации индивидуального предпринимателя по месту "
                   "жительства или адрес юридического лица в пределах места "
                   "нахождения юридического лица"),
    },
    "site": {
        FORM_1026: "Адрес (адреса) фактического осуществления деятельности",
        FORM_286: "Адрес (адреса) места (мест) образования отходов",
    },
}

_L_COMP = ("Химический и (или) компонентный состав (указывается в порядке "
           "убывания содержания компонентов)")
_L_HAZARD = ("Класс опасности по степени негативного воздействия на "
             "окружающую среду")
# шапка грифа по типовой форме № 286 (в № 1026 подставляется должность
# подписанта — так оформлены принятые паспорта ТЕХНОСТРОЙ)
_L_HEAD_286 = "Руководитель юридического лица (индивидуальный предприниматель)"

_DEFAULT_METHOD = "количественный химический анализ отхода"

# Агрегатное состояние и физическая форма по 9–10-му знакам кода ФККО.
# Источник: структура кода ФККО (приказ Росприроднадзора от 22.05.2017 № 242,
# вводная часть; та же кодировка — Порядок ведения государственного кадастра
# отходов). Значения «как в каталоге» — так их печатают принятые паспорта
# (эталон: «…39 4» → «Прочие дисперсные системы», «…21 4» → «Кусковая форма»).
AGG_STATE_BY_CODE = {
    "00": "Не требует определения агрегатного состояния и физической формы",
    "10": "Жидкое",
    "20": "Твердое",
    "21": "Кусковая форма",
    "22": "Стружка",
    "23": "Волокно",
    "29": "Прочие формы твердых веществ",
    "30": "Дисперсные системы",
    "31": "Жидкое в жидком (эмульсия)",
    "32": "Твердое в жидком (суспензия)",
    "33": "Твердое в жидком (паста)",
    "39": "Прочие дисперсные системы",
    "40": "Твердые сыпучие материалы",
    "41": "Порошок",
    "42": "Пыль",
    "43": "Опилки",
    "49": "Прочие сыпучие материалы",
    "50": "Изделия из твердых материалов, за исключением волокон",
    "51": "Изделие из одного материала",
    "52": "Изделия из нескольких материалов",
    "53": "Изделия, содержащие жидкость",
    "54": "Изделия, содержащие газ",
    "60": "Изделия из волокон",
    "61": "Изделие из одного волокна",
    "62": "Изделия из нескольких волокон",
    "70": "Смеси твердых материалов и изделий",
    "71": "Смесь твердых материалов (включая волокна)",
    "72": "Смесь твердых материалов (включая волокна) и изделий",
}


def aggregate_state(fkko_code: str) -> str:
    """Агрегатное состояние по 9–10-му знакам 11-значного кода ФККО.

    Пустая строка — если код неполный или пара знаков не из таблицы:
    выдумывать значение нельзя, пусть в документе останется плейсхолдер.
    """
    from ecodoc.core.waste_agg import norm_fkko

    code = norm_fkko(fkko_code or "")
    if len(code) != 11:
        return ""
    return AGG_STATE_BY_CODE.get(code[8:10], "")


def form_for(approved_date: _dt.date | None = None) -> str:
    """Какая редакция формы действует на дату утверждения паспорта."""
    day = approved_date or _dt.date.today()
    return FORM_286 if day >= SWITCH_DATE else FORM_1026


def _parse_date(value) -> _dt.date | None:
    """Дата из параметра/поля: date, datetime или строка ДД.ММ.ГГГГ / ГГГГ-ММ-ДД."""
    if isinstance(value, _dt.datetime):
        return value.date()
    if isinstance(value, _dt.date):
        return value
    if isinstance(value, str) and value.strip():
        for fmt in ("%d.%m.%Y", "%Y-%m-%d"):
            try:
                return _dt.datetime.strptime(value.strip(), fmt).date()
            except ValueError:
                continue
    return None


def _details(ctx: ReportContext, w: WasteFlow) -> dict:
    """Сведения об отходе: ручной ввод главнее, ИИ-разбор — как запасной."""
    from ecodoc.core.waste_agg import norm_fkko

    extra = ctx.extra if isinstance(ctx.extra, dict) else {}
    d = dict((extra.get("waste_details") or {}).get(w.fkko_code) or {})
    code = norm_fkko(w.fkko_code or "")
    if not d.get("components"):
        for p in extra.get("waste_passports") or []:
            if code and norm_fkko(str(p.get("fkko") or "")) == code:
                d.setdefault("components", p.get("components") or [])
                break
    lab = lab_result_for(ctx, w, kinds=("КХА", "хим", "морф"))
    if lab:
        # как в принятых паспортах ТЕХНОСТРОЙ: состав = результаты протокола
        if not d.get("components") and lab.get("substances"):
            d["components"] = [{"name": x.get("name", ""),
                                "percent": x.get("value", "")}
                               for x in lab["substances"] if isinstance(x, dict)]
        d.setdefault("protocol", {
            "number": lab.get("protocol_no", ""), "date": lab.get("date", ""),
            "lab": lab.get("lab", ""),
            "lab_attestation": lab.get("lab_attestation", ""),
            "method_doc": lab.get("method", "")})
    return d


def lab_result_for(ctx: ReportContext, w: WasteFlow, kinds=()) -> dict | None:
    """Протокол из extra['lab_results'] для отхода: по коду ФККО или по
    наименованию в поле object (ИИ пишет туда «вид отхода» из протокола)."""
    from ecodoc.core.waste_agg import norm_fkko

    extra = ctx.extra if isinstance(ctx.extra, dict) else {}
    code = norm_fkko(w.fkko_code or "")
    name = (w.name or "").strip().lower()
    for lab in extra.get("lab_results") or []:
        if not isinstance(lab, dict):
            continue
        kind = str(lab.get("kind") or "").lower()
        if kinds and not any(k.lower() in kind for k in kinds):
            continue
        target = str(lab.get("object") or "")
        tcode = norm_fkko(lab.get("fkko") or "")
        if (code and (tcode == code or code in norm_fkko(target))) or (
                name and len(name) > 8 and name[:40] in target.lower()):
            return lab
    return None


def generate(ctx: ReportContext, out_dir: str | Path,
             approved_date=None) -> list[Path]:
    """Сгенерировать по одному .docx-паспорту на каждый отход I–IV класса.

    approved_date — дата утверждения паспорта (по ней выбирается редакция
    формы); приоритет: аргумент → extra['waste_details'][код]['approved_date']
    → ctx.extra['passport_approved_date'] → сегодня.
    """
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    extra = ctx.extra if isinstance(ctx.extra, dict) else {}

    for w in ctx.wastes:
        try:
            hazard = int(w.hazard_class)
        except (TypeError, ValueError):
            continue
        if not 1 <= hazard <= 4:
            continue  # паспорт нужен только для I–IV класса
        path = out_dir / f"Паспорт_{(w.fkko_code or 'без кода').replace(' ', '')}.docx"
        d = _details(ctx, w)
        day = (_parse_date(approved_date) or _parse_date(d.get("approved_date"))
               or _parse_date(extra.get("passport_approved_date")))
        _build(ctx.organization, w, hazard, d, form_for(day), day).save(path)
        paths.append(path)
    return paths


_MONTHS_GEN = ("января", "февраля", "марта", "апреля", "мая", "июня", "июля",
               "августа", "сентября", "октября", "ноября", "декабря")


def date_line(day: _dt.date | None) -> str:
    """«28» февраля 2025 г. — как в подписанных паспортах; без даты — бланк."""
    if not day:
        return "«____» ________________ 20____ г."
    return f"«{day.day:02d}» {_MONTHS_GEN[day.month - 1]} {day.year} г."


def generate_v_class(ctx: ReportContext, out_dir: str | Path) -> list[Path]:
    """Справка о подтверждении V класса по биотестированию — по одному .docx
    на отход V класса, у которого есть extra['waste_details'][код]['biotest']
    = {number, date, lab, lab_attestation, conclusion} или запись
    extra['lab_results'] с kind «биотест». Паспорт на V класс не нужен
    (п. 1 Порядка), но отнесение к V классу подтверждается биотестированием
    (ст. 14 ФЗ-89, приказ МПР № 158); в ПОО пользователя это «акт приёмки
    проб на токсичность методом биотестирования» ТАСИС с целью
    «подтверждение V класса опасности» и протокол БИО ЦЭД."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Pt

    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    extra = ctx.extra if isinstance(ctx.extra, dict) else {}
    paths: list[Path] = []
    for w in ctx.wastes:
        try:
            if int(w.hazard_class) != 5:
                continue
        except (TypeError, ValueError):
            continue
        d = dict((extra.get("waste_details") or {}).get(w.fkko_code) or {})
        bio = d.get("biotest")
        if not bio:
            lab = lab_result_for(ctx, w, kinds=("биотест",))
            if lab:
                bio = {"number": lab.get("protocol_no", ""), "date": lab.get("date", ""),
                       "lab": lab.get("lab", ""),
                       "lab_attestation": lab.get("lab_attestation", ""),
                       "conclusion": lab.get("conclusion", "")}
        if not bio:
            continue
        doc = Document()
        st = doc.styles["Normal"]
        st.font.name, st.font.size = "Times New Roman", Pt(12)
        h = doc.add_paragraph()
        h.alignment = AL.CENTER
        h.add_run("СВЕДЕНИЯ О ПОДТВЕРЖДЕНИИ ОТНЕСЕНИЯ ОТХОДА К V КЛАССУ "
                  "ОПАСНОСТИ").bold = True
        doc.add_paragraph(f"{ctx.organization.name}")
        rows = [
            ("Наименование вида отходов по ФККО", w.name or "—"),
            ("Код вида отходов по ФККО", w.fkko_code or "—"),
            ("Класс опасности", "V"),
            ("Способ подтверждения", "биотестирование водной вытяжки отхода "
             "(ст. 14 Федерального закона от 24.06.1998 № 89-ФЗ; Критерии, "
             "утв. приказом Минприроды России от 31.03.2025 № 158)"),
            ("Акт (протокол) биотестирования",
             " ".join(x for x in (f"№ {bio.get('number')}" if bio.get("number") else "",
                                  f"от {bio.get('date')}" if bio.get("date") else "")
                      if x) or "—"),
            ("Испытательная лаборатория", bio.get("lab") or "—"),
            ("Реквизиты аттестата аккредитации", bio.get("lab_attestation") or "—"),
            ("Вывод", bio.get("conclusion") or "отход не оказывает острого "
             "токсического действия на тест-объекты — V класс опасности подтверждён"),
        ]
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        for i, (k, v) in enumerate(rows):
            _set(t.cell(i, 0), k, AL.LEFT)
            _set(t.cell(i, 1), str(v), AL.LEFT)
        path = out_dir / f"V_класс_{(w.fkko_code or 'без кода').replace(' ', '')}.docx"
        doc.save(path)
        paths.append(path)
    return paths


def _build(org: Organization, w: WasteFlow, hazard: int, d: dict,
           form: str = FORM_1026, day: _dt.date | None = None):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.enum.text import WD_BREAK
    from docx.shared import Cm, Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # интервал после абзаца по умолчанию раздвигает шапку так, что таблица
    # «Сведения об отходах» перестаёт помещаться на лист — гасим его
    style.paragraph_format.space_after = Pt(0)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(2), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(1.2), Cm(1.2)

    _approval_block(doc, org, AL, Cm, Pt, form, day)

    for text, size in ((_TITLE, 13), (_TEXT["subtitle"][form], 12)):
        p = doc.add_paragraph()
        p.alignment = AL.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.add_run(text).font.size = Pt(size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    comps = _sorted_components(d.get("components") or [])
    method = d.get("method") or _DEFAULT_METHOD
    basis = d.get("basis") or ""
    if basis and basis not in method:
        method = f"{method} ({basis})"

    # агрегатное состояние: ручной ввод главнее, иначе — по 9–10 знакам кода
    agg = d.get(_AGG) or aggregate_state(w.fkko_code) or "‹указать по ФККО›"
    rows = [
        (_TEXT["name"][form], w.name or "—"),
        ("Код вида отходов по ФККО", w.fkko_code or "—"),
        (_TEXT["origin"][form],
         d.get("origin") or "‹указать технологический процесс›"),
        ("__COMP__", comps),
        (_TEXT["method"][form], method),
        ("Агрегатное состояние и физическая форма", agg),
        (_L_HAZARD, _roman(hazard)),
    ]
    _info_table(doc, "Сведения об отходах", rows, AL, Cm)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    _person_table(doc, org, d, AL, Cm, form)
    proto = d.get("protocol")
    if isinstance(proto, dict) and any(proto.values()):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        _protocol_table(doc, proto, comps, AL, Cm)
    return doc


def _protocol_table(doc, proto: dict, comps: list, AL, Cm):
    """Третий лист «Основание для определения состава»: реквизиты протокола
    КХА/морфологии и лаборатории. В типовой форме (№ 1026/№ 286) такой строки
    нет, но в принятых паспортах ТЕХНОСТРОЙ протокол ИЦ ООО «ТАСИС» подшит к
    паспорту — здесь его реквизиты, чтобы паспорт был проверяем без поиска
    бумажного протокола."""
    p = doc.add_paragraph()
    p.alignment = AL.CENTER
    p.add_run("Основание для определения химического и (или) компонентного "
              "состава отхода").bold = True
    rows = [("Протокол исследований (измерений)",
             " ".join(x for x in (f"№ {proto.get('number')}" if proto.get("number") else "",
                                  f"от {proto.get('date')}" if proto.get("date") else "")
                      if x) or "—"),
            ("Испытательная лаборатория (центр)", proto.get("lab") or "—"),
            ("Уникальный номер записи об аккредитации (аттестат)",
             proto.get("lab_attestation") or "—"),
            ("Методика измерений", proto.get("method_doc") or "—")]
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    _fix_widths(t, (Cm(9.3), Cm(8.2)))
    for i, (k, v) in enumerate(rows):
        _set(t.cell(i, 0), k, AL.JUSTIFY)
        _set(t.cell(i, 1), str(v), AL.LEFT)
    if comps:
        doc.add_paragraph()
        t2 = doc.add_table(rows=len(comps) + 1, cols=3)
        t2.style = "Table Grid"
        _fix_widths(t2, (Cm(1.5), Cm(11), Cm(5)))
        for j, h in enumerate(("№", "Наименование компонента (по протоколу)",
                               "Содержание, %")):
            _set(t2.cell(0, j), h, AL.CENTER)
        for i, c in enumerate(comps, 1):
            _set(t2.cell(i, 0), str(i), AL.CENTER)
            _set(t2.cell(i, 1), str(c.get("name", "")), AL.LEFT)
            _set(t2.cell(i, 2), _fmt_pct(c.get("percent", "")), AL.CENTER)


def _approval_block(doc, org: Organization, AL, Cm, Pt, form: str = FORM_1026,
                    day: _dt.date | None = None):
    """Гриф «УТВЕРЖДАЮ» в правом верхнем углу листа; дата — как в подписанных
    паспортах («28» февраля 2025 г.), без даты — бланк для рукописной."""
    head = _L_HEAD_286 if form == FORM_286 else org.official_title
    for text in ("УТВЕРЖДАЮ", head, org.short_name or org.name):
        p = doc.add_paragraph()
        p.alignment = AL.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.add_run(text).bold = True

    doc.add_paragraph()  # место для подписи и печати

    t = doc.add_table(rows=2, cols=2)
    t.alignment = 2  # WD_TABLE_ALIGNMENT.RIGHT
    _fix_widths(t, (Cm(5.5), Cm(5.5)))
    for row, (left, right, bold) in enumerate((
            ("", org.director_name or "", True),
            ("(подпись)", "(расшифровка)", False))):
        for col, text in enumerate((left, right)):
            cell = t.cell(row, col)
            p = cell.paragraphs[0]
            p.alignment = AL.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.add_run(text).bold = bold
            if row == 0:
                _bottom_border(cell)

    for text, bold in ((date_line(day), True),
                       (_TEXT["seal"][form], False), ("(при наличии)", False)):
        p = doc.add_paragraph()
        p.alignment = AL.RIGHT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        if not bold:
            run.font.size = Pt(10)
    doc.add_paragraph()


def _info_table(doc, caption: str, rows, AL, Cm):
    """Таблица «Сведения об отходах» (3 колонки: строка состава вложенная)."""
    # строк: обычные — по одной; состав — заголовок + компоненты, причём при
    # ПУСТОМ составе ниже подставляется строка-плейсхолдер (max(1, ...)),
    # иначе таблица окажется на строку короче и падает IndexError
    n = sum(1 + max(1, len(v)) if k == "__COMP__" else 1 for k, v in rows) + 1
    t = doc.add_table(rows=n, cols=3)
    t.style = "Table Grid"
    _fix_widths(t, (Cm(9.3), Cm(4.7), Cm(3.5)))

    head = t.rows[0]
    head.cells[0].merge(head.cells[2]).paragraphs[0].alignment = AL.CENTER
    head.cells[0].paragraphs[0].add_run(caption)

    i = 1
    for label, value in rows:
        if label == "__COMP__":
            first = i
            _set(t.cell(i, 0), _L_COMP, AL.JUSTIFY)
            _set(t.cell(i, 1), "Наименование компонента", AL.CENTER)
            _set(t.cell(i, 2), "Содержание, %", AL.CENTER)
            i += 1
            if not value:
                value = [{"name": "‹состав из протокола КХА›", "percent": ""}]
            for c in value:
                _set(t.cell(i, 1), str(c.get("name", "")), AL.LEFT)
                _set(t.cell(i, 2), _fmt_pct(c.get("percent", "")), AL.CENTER)
                i += 1
            t.cell(first, 0).merge(t.cell(i - 1, 0))
        else:
            _set(t.cell(i, 0), label, AL.JUSTIFY)
            cell = t.cell(i, 1).merge(t.cell(i, 2))
            _set(cell, str(value), AL.LEFT)
            i += 1
    return t


def _person_table(doc, org: Organization, d: dict, AL, Cm,
                  form: str = FORM_1026):
    # адрес(а) места образования: п. 6 Порядка № 286 допускает один паспорт
    # на несколько адресов (ТКО/ОИТ III–IV кл.) — список печатаем через «; »
    site = d.get("site_address") or "‹адрес площадки›"
    if isinstance(site, (list, tuple)):
        site = "; ".join(str(a) for a in site if a) or "‹адрес площадки›"

    if form == FORM_286:
        # одна строка «полное и (или) сокращенное наименования»
        name = org.name
        if org.short_name and not org.is_individual and org.short_name != org.name:
            name = f"{org.name} ({org.short_name})"
        rows = [(_TEXT["person"][form], name)]
    else:
        rows = [(_TEXT["person"][form], org.name),
                ("Сокращенное наименование юридического лица", org.short_name)]
        if org.is_individual:
            rows.pop(1)
    rows += [
        (_TEXT["inn"][form], org.inn),
        ("Код по Общероссийскому классификатору предприятий и организаций "
         "(ОКПО)", org.okpo),
        ("Код по Общероссийскому классификатору видов экономической "
         "деятельности (ОКВЭД)", org.okved),
        (_TEXT["location"][form], org.address),
        ("Почтовый адрес", d.get("postal_address") or org.address),
        (_TEXT["site"][form], site),
    ]

    t = doc.add_table(rows=len(rows) + 1, cols=2)
    t.style = "Table Grid"
    _fix_widths(t, (Cm(9.3), Cm(8.2)))
    head = t.rows[0]
    head.cells[0].merge(head.cells[1]).paragraphs[0].alignment = AL.CENTER
    head.cells[0].paragraphs[0].add_run("Сведения о лице, которое образовало отходы")

    for i, (label, value) in enumerate(rows, start=1):
        _set(t.cell(i, 0), label, AL.JUSTIFY)
        _set(t.cell(i, 1), str(value or "—"), AL.LEFT)
    return t


def _set(cell, text: str, align):
    p = cell.paragraphs[0]
    p.alignment = align
    p.add_run(text)


def _fix_widths(table, widths):
    """Жёсткая ширина колонок: без неё Word/LibreOffice «схлопывают» подписи.

    Ширину надо проставить до объединения ячеек и продублировать в разметке
    таблицы (tblLayout fixed), иначе автоподбор её игнорирует.
    """
    from docx.oxml.ns import qn

    table.autofit = False
    layout = table._tbl.tblPr.makeelement(qn("w:tblLayout"), {})
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
    for column, width in zip(table.columns, widths):
        column.width = width


def _bottom_border(cell):
    """Линия под ячейкой — для строки подписи в грифе «УТВЕРЖДАЮ»."""
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    bottom = borders.makeelement(qn("w:bottom"), {})
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    tc_pr.append(borders)


def _sorted_components(comps: list) -> list:
    """Форма требует состав в порядке убывания содержания компонентов."""
    def key(c):
        try:
            return -float(str(c.get("percent", "")).replace(",", ".").strip())
        except (TypeError, ValueError):
            return 0.0
    return sorted([c for c in comps if isinstance(c, dict)], key=key)


def _fmt_pct(v) -> str:
    s = str(v).strip().replace(".", ",")
    return s


def _roman(hazard_class) -> str:
    return {1: "I", 2: "II", 3: "III", 4: "IV", 5: "V"}.get(int(hazard_class),
                                                            str(hazard_class))
