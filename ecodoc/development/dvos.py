"""Декларация о воздействии на окружающую среду (ДВОС) — объекты II категории.

Правовая основа (проверено по состоянию на август 2026):

* ст. 31.2 Федерального закона от 10.01.2002 № 7-ФЗ «Об охране окружающей
  среды» — обязанность лиц, ведущих деятельность на объектах II категории
  (кроме имеющих КЭР), представлять декларацию; состав сведений — п. 3;
  одновременно с декларацией представляются РАСЧЁТЫ нормативов допустимых
  выбросов и сбросов (п. 4); подаётся один раз в СЕМЬ лет при условии
  неизменности технологических процессов основных производств, качественных
  и количественных характеристик выбросов, сбросов и стационарных
  источников (п. 6);
* форма и порядок заполнения — приказ Минприроды России от 19.03.2025 № 117
  (зарегистрирован Минюстом России 14.04.2025, рег. № 81831; вступил в силу
  01.09.2025, действует шесть лет — до 01.09.2031). Он ЗАМЕНИЛ приказ
  от 11.10.2018 № 509 (в ред. приказа от 23.06.2020 № 383), который с
  01.09.2025 утратил силу — реквизиты № 509 в новых декларациях не указывать.

Состав формы по приказу № 117 (приложение № 1; сверено дословно 22.08.2026
по normativ.kontur.ru, documentId=493064): шапка «В ___ (орган приёма)»,
титульный лист (код объекта НВОС, наименование/ОПФ/адрес, код И наименование
основного вида деятельности, «составлена на __ листах, приложений __»,
обязательство вносить изменения, исполнитель, руководитель, дата, печать) и
разделы:
  I   — виды и объём производимой продукции (товара);
  II  — информация о реализации природоохранных мероприятий;
  III — данные об авариях и инцидентах за предыдущие семь лет — ДВЕ
        таблицы (1 — аварии, 2 — инциденты) по 7 граф;
  IV  — масса выбросов: вещество, класс опасности (СанПиН 1.2.3685-21),
        данные об источнике, г/с, т/год всего / в пределах НДВ / с
        превышением НДВ (8 граф);
  V   — масса сбросов: водный объект, вещество, класс, источник, НДС
        мг/дм³, т/год всего / в пределах / с превышением (9 граф);
  VI  — отходы: таблица 1 за отчётный год и таблица 2 на следующие семь
        лет, по 9 граф (образовано, размещено на собственных ОРО + № ГРОРО,
        передано НА РАЗМЕЩЕНИЕ другим лицам + № ГРОРО);
  VII — программа ПЭК: кем и когда утверждена, орган представления отчёта
        ПЭК, дата последнего отчёта;
  «Приложениями к Декларации являются»: расчёт НДВ, расчёт НДС, квоты
  выбросов (при эксперименте по квотированию); сноски <1>–<5>.

Что берётся из базы: организация и объект — из карточек; выбросы —
ctx.pollutants (среда «воздух»); сбросы — ctx.pollutants (среда «вода») +
выпуски из ctx.extra['water']['discharge']; отходы — ctx.wastes /
ctx.waste_acts (через инвентаризацию отходов); ПЭК — ctx.extra['pek'].
Продукция, мероприятия и аварии машиной не выдумываются: они читаются из
ctx.extra['dvos'] ({"products": [...], "measures": [...], "accidents": [...]}),
а при их отсутствии в документе остаётся пометка «[требуется: …]», и та же
строка возвращается из gaps(ctx) — пользователь видит список до генерации.
"""
from __future__ import annotations

from datetime import date
from decimal import Decimal
from pathlib import Path

from ecodoc.core.models import Medium, ReportContext

TITLE = "Декларация о воздействии на окружающую среду"

NPA = ("приказом Минприроды России от 19.03.2025 № 117 (зарегистрирован "
       "Минюстом России 14.04.2025, рег. № 81831; действует с 01.09.2025, "
       "взамен приказа от 11.10.2018 № 509) в соответствии со статьёй 31.2 "
       "Федерального закона от 10.01.2002 № 7-ФЗ «Об охране окружающей среды»")

# Разделы формы приказа № 117 — заголовки ДОСЛОВНО по приложению № 1
# (сверено 22.08.2026 по normativ.kontur.ru, documentId=493064).
SECTIONS: list[tuple[str, str]] = [
    ("products", "Раздел I. Виды и объем производимой продукции (товара)"),
    ("measures", "Раздел II. Информация о реализации природоохранных "
                 "мероприятий"),
    ("accidents", "Раздел III. Данные об авариях и инцидентах, повлекших "
                  "негативное воздействие на окружающую среду, произошедших "
                  "за {y_from} - {y_to} годы"),
    ("air", "Раздел IV. Масса выбросов загрязняющих веществ в атмосферный "
            "воздух"),
    ("water", "Раздел V. Масса сбросов загрязняющих веществ в водные объекты"),
    ("waste", "Раздел VI. Сведения об образовании и размещении отходов "
              "производства и потребления (далее - отходы)"),
    ("pek", "Раздел VII. Информация о программе производственного "
            "экологического контроля"),
]

# Графы таблиц — дословно по форме № 117 (с переносами, как в бланке).
COLS_I = ["N п/п", "Наименование производимой продукции (товара)",
          "Код производимой продукции (товара)", "Единица измерения",
          "Объем (количество) производимой продукции (товара)"]
COLS_II = ["N п/п", "Наименование мероприятия", "Срок выполнения — начало",
           "Срок выполнения — конец", "Объем финансирования, тысяч рублей",
           "Источники финансирования", "Результат мероприятия"]
COLS_III_ACC = ["N п/п", "Дата возникновения аварии",
                "Дата ликвидации последствий аварии",
                "Краткая характеристика аварии, причины возникновения",
                "Краткая характеристика негативного воздействия на "
                "окружающую среду при аварии",
                "Размер причиненного вреда окружающей среде, тысяч рублей",
                "Основные мероприятия по локализации и ликвидации "
                "последствий аварии"]
COLS_III_INC = ["N п/п", "Дата возникновения инцидента",
                "Дата ликвидации инцидента",
                "Краткая характеристика инцидента, причины возникновения",
                "Краткая характеристика негативного воздействия на "
                "окружающую среду при инциденте",
                "Размер вреда, причиненного окружающей среде, тысяч рублей",
                "Основные мероприятия по локализации и ликвидации "
                "последствий инцидента"]
COLS_IV = ["N п/п", "Наименование загрязняющего вещества <1>",
           "Класс опасности <2>", "Данные об источнике выбросов",
           "Масса выбросов загрязняющих веществ в атмосферный воздух — "
           "грамм/секунду",
           "Масса выбросов — тонн/год — всего",
           "Масса выбросов — тонн/год — в том числе в пределах нормативов "
           "допустимых выбросов",
           "Масса выбросов — тонн/год — с превышением нормативов допустимых "
           "выбросов"]
COLS_V = ["N п/п", "Наименование водного объекта",
          "Наименование загрязняющего вещества <1>", "Класс опасности",
          "Данные об источнике сбросов",
          "Норматив допустимого сброса, миллиграмм/кубический дециметр",
          "Масса сбросов загрязняющих веществ, тонн/год — всего",
          "Масса сбросов — в том числе в пределах нормативов допустимых "
          "сбросов",
          "Масса сбросов — с превышением нормативов допустимых сбросов"]
COLS_VI_1 = ["N п/п", "Код отхода по федеральному классификационному "
             "каталогу отходов (далее - ФККО) <3>",
             "Наименование отхода по ФККО <4>",
             "Класс опасности отхода по ФККО", "Образовано, тонн/год",
             "Размещено на собственных объектах размещения отходов, "
             "тонн/год — количество",
             "Размещено на собственных объектах — номер объекта размещения "
             "отходов в государственном реестре объектов размещения отходов "
             "(далее - ГРОРО) <4>",
             "Передано на размещение другим индивидуальным предпринимателям "
             "или юридическим лицам, тонн/год — количество",
             "Передано на размещение другим лицам — номер объекта размещения "
             "отходов в ГРОРО <5> (за исключением твердых коммунальных "
             "отходов)"]
COLS_VI_2 = ["N п/п", "Код отхода по ФККО <4>", "Наименование отхода по ФККО "
             "<4>", "Класс опасности отхода по ФККО <4>",
             "Образование, тонн/год",
             "Размещение на собственных объектах размещения отходов, "
             "тонн/год — количество",
             "Размещение на собственных объектах — номер объекта размещения "
             "отходов в ГРОРО <5>",
             "Передача на размещение другим индивидуальным предпринимателям "
             "или юридическим лицам, тонн/год — количество",
             "Передача на размещение другим лицам — номер объекта размещения "
             "отходов в ГРОРО <6> (за исключением твердых коммунальных "
             "отходов)"]

# Сноски к форме — дословно (сокращены перечни изменяющих приказов).
FOOTNOTES = [
    "<1> Перечень загрязняющих веществ, в отношении которых применяются меры "
    "государственного регулирования в области охраны окружающей среды, "
    "утвержденный распоряжением Правительства Российской Федерации от "
    "20 октября 2023 г. N 2909-р.",
    "<2> Классы опасности загрязняющих веществ в составе выбросов согласно "
    "СанПиН 1.2.3685-21 (постановление Главного государственного санитарного "
    "врача Российской Федерации от 28 января 2021 г. N 2, с изменениями от "
    "30 декабря 2022 г. N 24; действуют до 1 марта 2027 г.).",
    "<3> Приказ Росприроднадзора от 22 мая 2017 г. N 242 «Об утверждении "
    "Федерального классификационного каталога отходов» с изменениями.",
    "<4> Глава III Порядка ведения государственного кадастра отходов, "
    "утвержденного приказом Минприроды России от 30 сентября 2011 г. N 792 "
    "(с изменениями, приказ от 19 апреля 2023 г. N 211).",
    "<5> Пункт 23.6 методических указаний по разработке проектов нормативов "
    "образования отходов и лимитов на их размещение, утвержденных приказом "
    "Минприроды России от 7 декабря 2020 г. N 1021 (с изменениями, приказ от "
    "30 октября 2024 г. N 634; действует до 1 января 2027 г.).",
]

# Приложения к декларации — абзацы формы дословно
ATTACHMENTS = [
    "расчет нормативов допустимых выбросов загрязняющих веществ в "
    "атмосферный воздух;",
    "расчет нормативов допустимых сбросов загрязняющих веществ в водные "
    "объекты;",
    "утвержденные квоты выбросов (в случае установления таких квот в период "
    "проведения эксперимента по квотированию выбросов загрязняющих веществ "
    "в атмосферный воздух с 1 января 2020 г. по 31 декабря 2026 г.) в "
    "соответствии с частью 7 статьи 5 Федерального закона от 26 июля 2019 г. "
    "N 195-ФЗ.",
]

# чего машина взять не может — единые формулировки для текста и gaps()
EXTERNAL = {
    "authority": "наименование органа, уполномоченного на приём декларации "
                 "(территориальный орган Росприроднадзора или орган "
                 "исполнительной власти субъекта РФ) — "
                 "extra['dvos']['authority']",
    "okved_name": "наименование основного вида экономической деятельности "
                  "по ЕГРЮЛ/ЕГРИП (п. 16-17 Порядка) — "
                  "extra['dvos']['okved_name']",
    "products": "виды и объём производимой продукции (товара) с кодами по "
                "ОКПД2 — extra['dvos']['products'] (данные бухгалтерского "
                "учёта, максимальные значения за период действия "
                "декларации, п. 19-20 Порядка)",
    "measures": "перечень реализованных природоохранных мероприятий по "
                "действующему плану мероприятий по охране окружающей среды "
                "(extra['dvos']['measures'], п. 21 Порядка)",
    "accidents": "данные об авариях и инцидентах за предыдущие 7 лет "
                 "(extra['dvos']['accidents']; если их не было — укажите "
                 "пустой список)",
    "air_details": "по каждому веществу раздела IV — класс опасности "
                   "(СанПиН 1.2.3685-21), данные об источнике выбросов и "
                   "максимальный разовый выброс, г/с — из расчёта НДВ "
                   "(extra['dvos']['pollutant_details'][код])",
    "water_details": "по каждому веществу раздела V — наименование водного "
                     "объекта, класс опасности, данные об источнике сбросов "
                     "и НДС, мг/дм³ — из расчёта НДС "
                     "(extra['dvos']['pollutant_details'][код])",
    "ndv": "расчёты нормативов допустимых выбросов (и сбросов — при их "
           "наличии), прилагаемые к декларации (п. 4 ст. 31.2 ФЗ-7; "
           "п. 12-13 Порядка)",
    "groro": "номера объектов размещения отходов в ГРОРО (для размещаемых "
             "отходов; extra['dvos']['groro'][код ФККО])",
    "pek": "программа производственного экологического контроля: ФИО "
           "утвердившего должностного лица и дата утверждения "
           "(extra['pek']['responsible'], ['approved_date'])",
    "pek_report": "орган, в который представляется отчёт ПЭК, и дата "
                  "представления последнего отчёта "
                  "(extra['pek']['authority'], ['last_report_date'])",
    "executor": "исполнитель, ответственный за представление декларации: "
                "должность, ФИО, телефон, e-mail "
                "(extra['dvos']['executor'])",
}

# ОПФ по префиксу наименования — только очевидные случаи, иначе [требуется]
_OPF = {
    "ООО": "Общество с ограниченной ответственностью",
    "АО": "Акционерное общество",
    "ПАО": "Публичное акционерное общество",
    "ЗАО": "Закрытое акционерное общество",
    "МУП": "Муниципальное унитарное предприятие",
    "ГУП": "Государственное унитарное предприятие",
    "ИП": "Индивидуальный предприниматель",
}


def _num(value) -> str:
    # п. 7 Порядка: при отсутствии показателя в графе ставится прочерк
    if value in (None, ""):
        return "-"
    try:
        d = Decimal(str(value).replace(",", "."))
    except Exception:
        return str(value)
    return f"{d.normalize():f}" if d else "0"


def _dec(value) -> Decimal:
    try:
        return Decimal(str(value or 0).replace(",", "."))
    except Exception:
        return Decimal("0")


def _dash(value) -> str:
    """Текстовая графа: пусто — прочерк (п. 7 Порядка)."""
    v = str(value or "").strip()
    return v if v else "-"


def opf(ctx: ReportContext) -> str:
    """Организационно-правовая форма — поле титульного листа."""
    org = ctx.organization
    if org.is_individual:
        return "Индивидуальный предприниматель"
    name = (org.name or org.short_name or "").strip()
    for prefix, full in _OPF.items():
        if name.startswith(prefix + " ") or name.startswith(prefix + "«"):
            return full
        if full.lower() in name.lower():
            return full
    return "[требуется: организационно-правовая форма]"


def _dvos_extra(ctx: ReportContext) -> dict:
    d = (ctx.extra or {}).get("dvos", {})
    return d if isinstance(d, dict) else {}


def _details(ctx: ReportContext, code: str) -> dict:
    """Справочные данные по веществу (класс, источник, г/с, НДС, водный
    объект) из extra['dvos']['pollutant_details'][код]."""
    det = _dvos_extra(ctx).get("pollutant_details") or {}
    d = det.get(code) if isinstance(det, dict) else None
    return d if isinstance(d, dict) else {}


def okved_split(ctx: ReportContext) -> tuple[str, str]:
    """Код и наименование основного вида деятельности — два поля титула.

    В базе ОКВЭД может лежать как «38.11 Сбор неопасных отходов» — делим
    по первому пробелу; наименование можно задать явно в
    extra['dvos']['okved_name'].
    """
    raw = (ctx.organization.okved or "").strip()
    code, _, name = raw.partition(" ")
    name = name.strip() or str(_dvos_extra(ctx).get("okved_name") or "")
    return code.strip(), name


def products(ctx: ReportContext) -> list[dict]:
    """Раздел I: продукция из extra['dvos']['products'] — не выдумывается."""
    out = []
    for p in _dvos_extra(ctx).get("products", []) or []:
        if isinstance(p, dict):
            out.append({"name": str(p.get("name") or ""),
                        "code": str(p.get("code") or ""),       # ОКПД2
                        "volume": p.get("volume"),
                        "unit": str(p.get("unit") or "")})
        elif isinstance(p, str) and p.strip():
            out.append({"name": p.strip(), "code": "", "volume": None,
                        "unit": ""})
    return out


def measures(ctx: ReportContext) -> list[dict]:
    """Раздел II: реализованные природоохранные мероприятия."""
    out = []
    for m in _dvos_extra(ctx).get("measures", []) or []:
        if isinstance(m, dict):
            out.append({"name": str(m.get("name") or ""),
                        "start": str(m.get("start") or ""),
                        "end": str(m.get("end") or ""),
                        "cost": m.get("cost"),
                        "funding": str(m.get("funding") or ""),
                        "result": str(m.get("result") or "")})
        elif isinstance(m, str) and m.strip():
            out.append({"name": m.strip(), "start": "", "end": "",
                        "cost": None, "funding": "", "result": ""})
    return out


def accidents(ctx: ReportContext) -> list[dict] | None:
    """Раздел III: аварии/инциденты за 7 лет (kind = «авария»/«инцидент»).

    None — сведения не заводились (пометка «[требуется]»); пустой список —
    пользователь подтвердил, что аварий и инцидентов не было.
    """
    d = _dvos_extra(ctx)
    if "accidents" not in d:
        return None
    out = []
    for a in d.get("accidents") or []:
        if isinstance(a, dict):
            kind = str(a.get("kind") or "авария").strip().lower()
            out.append({"kind": "инцидент" if "инцидент" in kind else "авария",
                        "date": str(a.get("date") or ""),
                        "end_date": str(a.get("end_date") or ""),
                        "description": str(a.get("description") or ""),
                        "impact": str(a.get("impact") or ""),
                        "damage": a.get("damage"),
                        "measures": str(a.get("measures") or "")})
        elif isinstance(a, str) and a.strip():
            out.append({"kind": "авария", "date": "", "end_date": "",
                        "description": a.strip(), "impact": "", "damage": None,
                        "measures": ""})
    return out


def rows_air(ctx: ReportContext) -> list[dict]:
    """Раздел IV: выбросы по веществам, т/год + справочные графы формы."""
    out = []
    for p in ctx.pollutants:
        if p.medium != Medium.AIR:
            continue
        norm = _dec(p.mass_norm)
        # графа 8 формы — «с превышением нормативов допустимых выбросов»:
        # у объектов II категории ВСВ нет, лимитную корзину считаем превышением
        over = _dec(p.mass_limit) + _dec(p.mass_over)
        det = _details(ctx, p.code or "")
        out.append({"code": p.code or "", "name": p.name or "",
                    "hazard": str(det.get("class") or ""),
                    "source_info": str(det.get("source") or ""),
                    "gs": det.get("gs"),
                    "norm": norm, "over": over, "total": norm + over})
    return out


def rows_water(ctx: ReportContext) -> list[dict]:
    """Раздел V: сбросы по веществам, т/год + водный объект, класс, НДС."""
    bodies = water_bodies(ctx)
    out = []
    for p in ctx.pollutants:
        if p.medium != Medium.WATER:
            continue
        norm = _dec(p.mass_norm)
        over = _dec(p.mass_limit) + _dec(p.mass_over)
        det = _details(ctx, p.code or "")
        body = str(det.get("water_body") or "") or \
            (bodies[0] if len(bodies) == 1 else "")
        out.append({"code": p.code or "", "name": p.name or "",
                    "water_body": body,
                    "hazard": str(det.get("class") or ""),
                    "source_info": str(det.get("source") or ""),
                    "nds": det.get("nds"),
                    "norm": norm, "over": over, "total": norm + over})
    return out


def water_bodies(ctx: ReportContext) -> list[str]:
    """Приёмники сточных вод — из выпусков 2-ТП (водхоз), если заведены."""
    water = (ctx.extra or {}).get("water", {}) or {}
    out = []
    for d in water.get("discharge", []) or []:
        if isinstance(d, dict) and d.get("receiver"):
            name = str(d["receiver"])
            if name not in out:
                out.append(name)
    return out


def rows_waste(ctx: ReportContext) -> list[dict]:
    """Раздел VI: образование и размещение отходов по видам.

    Графы формы: «Размещено на собственных ОРО» (placed_norm + placed_over)
    и «Передано НА РАЗМЕЩЕНИЕ другим лицам» (transferred_storage +
    transferred_burial) — не всё переданное, а только на хранение/захоронение.
    Номера ГРОРО — из extra['dvos']['groro'][код ФККО] ({"own": …,
    "other": …}); нет — прочерк и пометка.
    """
    from ecodoc.core.waste_agg import norm_fkko
    from ecodoc.development.waste_inventory import collect

    placed: dict[str, Decimal] = {}
    to_place: dict[str, Decimal] = {}
    for w in ctx.wastes:
        code = norm_fkko(w.fkko_code)
        placed[code] = placed.get(code, Decimal("0")) + \
            _dec(w.placed_norm) + _dec(w.placed_over)
        to_place[code] = to_place.get(code, Decimal("0")) + \
            _dec(w.transferred_storage) + _dec(w.transferred_burial)
    groro = _dvos_extra(ctx).get("groro") or {}
    if not isinstance(groro, dict):
        groro = {}

    out = []
    for r in collect(ctx):
        g = groro.get(r["fkko"]) or {}
        g = g if isinstance(g, dict) else {"own": str(g), "other": ""}
        out.append({"fkko": r["fkko"], "name": r["name"],
                    "hazard": r["hazard"] or 0,
                    "generated": _dec(r["generated"]),
                    "placed": placed.get(r["fkko"], Decimal("0")),
                    "groro_own": str(g.get("own") or ""),
                    "to_placement": to_place.get(r["fkko"], Decimal("0")),
                    "groro_other": str(g.get("other") or ""),
                    "transferred": _dec(r["transferred"])})
    return out


def pek_info(ctx: ReportContext) -> dict:
    """Раздел VII: поля формы — кем и когда утверждена программа ПЭК, орган
    представления отчёта ПЭК, дата последнего отчёта (из extra['pek'])."""
    pek = (ctx.extra or {}).get("pek", {}) or {}
    if not isinstance(pek, dict):
        pek = {}
    return {"approved_date": str(pek.get("approved_date") or ""),
            "responsible": str(pek.get("responsible") or ""),
            "authority": str(pek.get("authority") or ""),
            "last_report_date": str(pek.get("last_report_date") or "")}


def declared_period(ctx: ReportContext) -> tuple[int, int]:
    """Период действия декларации: 7 лет с года подачи (п. 6 ст. 31.2)."""
    year = ctx.period.year or date.today().year
    return year, year + 6


def accident_years(ctx: ReportContext) -> tuple[int, int]:
    """Семь лет, предшествующих подаче (п. 22 Порядка) — для заголовка III."""
    year = ctx.period.year or date.today().year
    return year - 7, year - 1


def gaps(ctx: ReportContext) -> list[str]:
    """Чего не хватает для декларации — тот же список, что и пометки в тексте."""
    out: list[str] = []
    org = ctx.organization
    ex = _dvos_extra(ctx)
    if not (org.name or org.short_name):
        out.append("не заполнено наименование организации")
    code, name = okved_split(ctx)
    if not code:
        out.append("не указан код основного вида деятельности (ОКВЭД) — "
                   "поле титульного листа")
    elif not name:
        out.append(f"требуется: {EXTERNAL['okved_name']}")
    if not ex.get("authority"):
        out.append(f"требуется: {EXTERNAL['authority']}")
    if not ctx.objects:
        out.append("не заведён объект НВОС: декларация подаётся по коду объекта")
    else:
        for o in ctx.objects:
            if not o.code:
                out.append(f"объект «{o.name or '?'}»: не указан код объекта НВОС")
            if o.category and o.category.strip().upper() not in ("II", "2"):
                out.append(f"объект {o.code or o.name}: категория "
                           f"«{o.category}» — декларация о воздействии "
                           f"представляется для объектов II категории "
                           f"(п. 1 ст. 31.2 ФЗ-7)")
    if not products(ctx):
        out.append(f"требуется: {EXTERNAL['products']}")
    if not measures(ctx):
        out.append(f"требуется: {EXTERNAL['measures']}")
    if accidents(ctx) is None:
        out.append(f"требуется: {EXTERNAL['accidents']}")
    air = rows_air(ctx)
    if air and any(not (r["hazard"] and r["source_info"]
                        and r["gs"] not in (None, "")) for r in air):
        out.append(f"требуется: {EXTERNAL['air_details']}")
    water = rows_water(ctx)
    if water and any(not (r["water_body"] and r["hazard"] and r["source_info"]
                          and r["nds"] not in (None, "")) for r in water):
        out.append(f"требуется: {EXTERNAL['water_details']}")
    wastes = rows_waste(ctx)
    if not (air or water or wastes):
        out.append("нет ни выбросов, ни сбросов, ни отходов — разделы IV–VI "
                   "нечем наполнять: загрузите исходные данные")
    if any((r["placed"] and not r["groro_own"])
           or (r["to_placement"] and not r["groro_other"]) for r in wastes):
        out.append(f"требуется: {EXTERNAL['groro']}")
    out.append(f"требуется: {EXTERNAL['ndv']}")
    pek = pek_info(ctx)
    if not (pek["approved_date"] and pek["responsible"]):
        out.append(f"требуется: {EXTERNAL['pek']}")
    if not (pek["authority"] and pek["last_report_date"]):
        out.append(f"требуется: {EXTERNAL['pek_report']}")
    if not ex.get("executor"):
        out.append(f"требуется: {EXTERNAL['executor']}")
    return out


def generate(ctx: ReportContext, out_path: str | Path) -> Path:
    """Собрать декларацию (.docx) по форме приложения № 1 к приказу № 117."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Cm, Pt

    org = ctx.organization
    ex = _dvos_extra(ctx)
    obj = ctx.objects[0] if ctx.objects else None
    y1, y2 = declared_period(ctx)
    a_from, a_to = accident_years(ctx)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(3), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(2), Cm(2)

    # ── титульный лист — поля формы № 117 дословно и в том же порядке ────
    head = doc.add_paragraph()
    head.alignment = AL.RIGHT
    head.add_run("В " + (str(ex.get("authority"))
                         if ex.get("authority")
                         else f"[требуется: {EXTERNAL['authority']}]")
                 + "\n(наименование федерального органа исполнительной "
                 "власти или органа исполнительной власти субъекта "
                 "Российской Федерации, уполномоченных на осуществление "
                 "приема декларации о воздействии на окружающую среду)")
    h = doc.add_paragraph()
    h.alignment = AL.CENTER
    run = h.add_run(TITLE.upper())
    run.bold = True
    run.font.size = Pt(14)
    sub = doc.add_paragraph()
    sub.alignment = AL.CENTER
    sub.add_run(f"Форма утверждена {NPA}.").italic = True

    okved_code, okved_name = okved_split(ctx)
    doc.add_paragraph()
    _table(doc, ["Поле титульного листа (форма № 117)", "Значение"], [
        ["код объекта, оказывающего негативное воздействие на окружающую "
         "среду",
         obj.code if obj and obj.code else "[требуется: код объекта НВОС]"],
        ["полное или сокращенное (при наличии) наименование юридического "
         "лица, фамилия, имя, отчество (при наличии) индивидуального "
         "предпринимателя",
         org.name or org.short_name or "[требуется: наименование организации]"],
        ["организационно-правовая форма юридического лица", opf(ctx)],
        ["адрес юридического лица в пределах места нахождения юридического "
         "лица или место жительства индивидуального предпринимателя",
         org.address or "[требуется: адрес организации]"],
        ["Код основного вида экономической деятельности",
         okved_code or "[требуется: код ОКВЭД]"],
        ["Наименование основного вида экономической деятельности",
         okved_name or f"[требуется: {EXTERNAL['okved_name']}]"],
        ["Декларация составлена на ___ листах, количество приложений ___",
         "[заполняется при распечатке]"],
    ])
    doc.add_paragraph(
        "В случае изменения в течение семи лет с даты подачи декларации о "
        "воздействии на окружающую среду (далее - Декларация) технологических "
        "процессов основных производств, качественных и количественных "
        "характеристик выбросов, сбросов загрязняющих веществ и стационарных "
        "источников в Декларацию будут внесены изменения в порядке, "
        "установленном законодательством Российской Федерации в области "
        "охраны окружающей среды.")
    execu = ex.get("executor")
    doc.add_paragraph(
        "Исполнитель, ответственный за представление Декларации, "
        + (str(execu) if execu else f"[требуется: {EXTERNAL['executor']}]")
        + "\n(должность, фамилия, имя, отчество (при наличии), телефон, факс "
        "(при наличии), адрес электронной почты (при наличии))")
    doc.add_paragraph(
        "Руководитель юридического лица/индивидуальный предприниматель "
        f"_____________ {org.director_name or ''}\n"
        f"«__» ______________ {y1} г.          место подписи и печати "
        "(при наличии)")
    doc.add_page_break()

    for key, title in SECTIONS:
        title = title.format(y_from=a_from, y_to=a_to)
        head = doc.add_paragraph()
        head.add_run(title).bold = True

        if key == "products":
            rows = products(ctx)
            if rows:
                _table(doc, COLS_I,
                       [[str(i), _dash(r["name"]), _dash(r["code"]),
                         _dash(r["unit"]), _num(r["volume"])]
                        for i, r in enumerate(rows, start=1)])
            else:
                _table(doc, COLS_I, [])
                doc.add_paragraph(f"[требуется: {EXTERNAL['products']}]")

        elif key == "measures":
            rows = measures(ctx)
            if rows:
                _table(doc, COLS_II,
                       [[str(i), _dash(r["name"]), _dash(r["start"]),
                         _dash(r["end"]), _num(r["cost"]), _dash(r["funding"]),
                         _dash(r["result"])]
                        for i, r in enumerate(rows, start=1)])
            else:
                _table(doc, COLS_II, [])
                doc.add_paragraph(f"[требуется: {EXTERNAL['measures']}]")

        elif key == "accidents":
            rows = accidents(ctx)
            # форма делит раздел III на две таблицы: 1 — аварии, 2 — инциденты
            for n, (kind, cols) in enumerate(
                    (("авари", COLS_III_ACC), ("инцидент", COLS_III_INC)),
                    start=1):
                label = ("авариях" if n == 1 else "инцидентах")
                doc.add_paragraph(
                    f"{n}. Данные об {label}, повлекших негативное "
                    f"воздействие на окружающую среду, произошедших за "
                    f"{a_from} - {a_to} годы")
                sel = [r for r in (rows or []) if r["kind"].startswith(kind)]
                _table(doc, cols,
                       [[str(i), _dash(r["date"]), _dash(r["end_date"]),
                         _dash(r["description"]), _dash(r["impact"]),
                         _num(r["damage"]), _dash(r["measures"])]
                        for i, r in enumerate(sel, start=1)])
            if rows is None:
                doc.add_paragraph(f"[требуется: {EXTERNAL['accidents']}]")
            elif not rows:
                doc.add_paragraph(
                    "Аварий и инцидентов, повлекших негативное воздействие "
                    "на окружающую среду, за предыдущие семь лет "
                    "не зафиксировано (в графах проставлены прочерки, "
                    "п. 7 Порядка).")

        elif key == "air":
            rows = rows_air(ctx)
            _table(doc, COLS_IV,
                   [[str(i), _dash(r["name"]), _dash(r["hazard"]),
                     _dash(r["source_info"]), _num(r["gs"]),
                     _num(r["total"]), _num(r["norm"]), _num(r["over"])]
                    for i, r in enumerate(rows, start=1)])
            if rows:
                if any(not (r["hazard"] and r["source_info"]
                            and r["gs"] not in (None, "")) for r in rows):
                    doc.add_paragraph(
                        f"[требуется: {EXTERNAL['air_details']}]")
            else:
                doc.add_paragraph(
                    "Раздел IV заполняется при негативном воздействии на "
                    "атмосферный воздух (п. 9 Порядка); данные о выбросах "
                    "не заведены.")
            doc.add_paragraph(f"[требуется: {EXTERNAL['ndv']}]")

        elif key == "water":
            rows = rows_water(ctx)
            _table(doc, COLS_V,
                   [[str(i), _dash(r["water_body"]), _dash(r["name"]),
                     _dash(r["hazard"]), _dash(r["source_info"]),
                     _num(r["nds"]), _num(r["total"]), _num(r["norm"]),
                     _num(r["over"])]
                    for i, r in enumerate(rows, start=1)])
            if rows:
                if any(not (r["water_body"] and r["hazard"]
                            and r["source_info"]
                            and r["nds"] not in (None, "")) for r in rows):
                    doc.add_paragraph(
                        f"[требуется: {EXTERNAL['water_details']}]")
            else:
                doc.add_paragraph(
                    "Раздел V заполняется при негативном воздействии на "
                    "водные объекты (п. 10 Порядка); данные о сбросах "
                    "не заведены.")

        elif key == "waste":
            rows = rows_waste(ctx)
            doc.add_paragraph(
                "1. Сведения об образовании и размещении отходов за "
                "отчетный год")
            _table(doc, COLS_VI_1,
                   [[str(i), _dash(r["fkko"]), _dash(r["name"]),
                     str(r["hazard"] or "-"), _num(r["generated"]),
                     _num(r["placed"]), _dash(r["groro_own"]),
                     _num(r["to_placement"]), _dash(r["groro_other"])]
                    for i, r in enumerate(rows, start=1)])
            doc.add_paragraph(
                "Данные — за календарный год, предшествующий году "
                "представления Декларации (п. 25 Порядка).")
            doc.add_paragraph(
                "2. Сведения об образовании и размещении отходов на "
                "следующие семь лет")
            doc.add_paragraph(
                f"Период действия Декларации — {y1}–{y2} гг.; показатели "
                "приняты равными отчётному году при неизменности "
                "технологических процессов (основание — проект нормативов "
                "образования отходов, п. 23.5-23.6 методических указаний "
                "№ 1021).")
            _table(doc, COLS_VI_2,
                   [[str(i), _dash(r["fkko"]), _dash(r["name"]),
                     str(r["hazard"] or "-"), _num(r["generated"]),
                     _num(r["placed"]), _dash(r["groro_own"]),
                     _num(r["to_placement"]), _dash(r["groro_other"])]
                    for i, r in enumerate(rows, start=1)])
            if any((r["placed"] and not r["groro_own"])
                   or (r["to_placement"] and not r["groro_other"])
                   for r in rows):
                doc.add_paragraph(f"[требуется: {EXTERNAL['groro']}]")
            if not rows:
                doc.add_paragraph(
                    "Отходы производства и потребления: данные не заведены "
                    "(загрузите справки-акты или заполните вкладку ОТХОДЫ).")

        elif key == "pek":
            info = pek_info(ctx)
            ok_pek = info["approved_date"] and info["responsible"]
            doc.add_paragraph(
                "Программа производственного экологического контроля "
                "утверждена "
                + (f"{info['responsible']} «{info['approved_date']}»"
                   if ok_pek else f"[требуется: {EXTERNAL['pek']}]")
                + "\n(фамилия, имя, отчество (при наличии) должностного "
                "лица; дата)")
            doc.add_paragraph(
                "Наименование территориального органа Федеральной службы "
                "по надзору в сфере природопользования или органа "
                "исполнительной власти субъекта Российской Федерации, в "
                "который представляется отчет об организации и о "
                "результатах осуществления производственного "
                "экологического контроля: "
                + (info["authority"] or "[требуется]"))
            doc.add_paragraph(
                "Дата представления последнего отчета об организации и "
                "результатах осуществления производственного "
                "экологического контроля: "
                + (info["last_report_date"] or "«__» ________ 20__ года"))
            if not (info["authority"] and info["last_report_date"]):
                doc.add_paragraph(f"[требуется: {EXTERNAL['pek_report']}]")
        doc.add_paragraph()

    doc.add_paragraph("Приложениями к Декларации являются:")
    for a in ATTACHMENTS:
        doc.add_paragraph(a)
    doc.add_paragraph()
    for f in FOOTNOTES:
        doc.add_paragraph(f).runs[0].font.size = Pt(9)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def _table(doc, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, text in enumerate(header):
        table.rows[0].cells[i].text = text
    # строка с номерами граф — как в бланке (1, 2, 3, …)
    cells = table.add_row().cells
    for i in range(len(header)):
        cells[i].text = str(i + 1)
    for r in rows:
        cells = table.add_row().cells
        for i, text in enumerate(r):
            cells[i].text = text
    if not rows:
        cells = table.add_row().cells
        for i in range(len(header)):
            cells[i].text = "-"
