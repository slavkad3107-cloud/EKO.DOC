"""Программа производственного экологического контроля (ПЭК).

Кто обязан: юридические лица и ИП, эксплуатирующие объекты I–III категорий,
разрабатывают и утверждают программу ПЭК (п. 2 ст. 67 Федерального закона
от 10.01.2002 № 7-ФЗ «Об охране окружающей среды»). Для IV категории
программа не требуется.

Состав программы задаёт НПА, а не привычка: приказ Минприроды России от
18.02.2022 № 109 «Об утверждении требований к содержанию программы
производственного экологического контроля, порядка и сроков представления
отчёта об организации и о результатах осуществления производственного
экологического контроля» (зарег. Минюстом 25.02.2022 № 67461, действует
с 01.09.2022 до 01.09.2028). Действующая редакция — с изменениями,
внесёнными приказами Минприроды России от 24.03.2023 № 150 (побочные
продукты производства), от 13.11.2024 № 659, от 07.05.2025 № 255 и от
12.05.2025 № 262 (искусственные грунты из органической части ТКО;
изменения действуют с 01.09.2025).

Разделы программы — дословно по п. 2 Требований (приложение 1 к приказу),
см. SECTIONS. Раздел 9 состоит из подразделов по средам (п. 9–9.5).

Откуда берутся данные (ничего не выдумывается):
  * организация и объекты НВОС      — ctx.organization / ctx.objects
                                      (ОКВЭД — org.okved, регион — o.region_code);
  * источники выбросов              — ctx.extra['emission_sources']
                                      (те же, что в инвентаризации № 871;
                                      у веществ источника — g_s / t_year);
  * вещества (воздух и вода)        — ctx.pollutants (medium=air/water);
  * отходы                          — ctx.wastes / ctx.waste_acts
                                      (через waste_inventory.collect);
  * забор воды и выпуски            — ctx.extra['water'] (intake/discharge;
                                      у выпуска могут быть pollutants
                                      [{code,name,mass}], treatment_code,
                                      treatment_capacity — как в 2-ТП водхоз);
  * побочные продукты, искусственные
    грунты                          — ctx.extra['ppp'] / ['artificial_soil'].

Ключи ctx.extra['pek'] (всё необязательно; чего нет — «[требуется: …]»):
  Раздел 1 (п. 3 Требований):
    program_date / approved_date — дата утверждения Программы
                                   («09.02.2025» — попадает и в гриф);
    legal_form              — организационно-правовая форма (ОПФ); если не
                              задана — выводится из названия организации
                              (ООО/АО/ПАО/ИП…), иначе [требуется];
    supervision_level       — «федеральный» / «региональный» — уровень
                              государственного экологического надзора;
    authority               — наименование уполномоченного органа, в который
                              направляется отчёт (если не задано — берётся из
                              справочника AUTHORITIES по region_code объекта и
                              уровню надзора);
    responsible             — должностное лицо, ответственное за отчёт.
  Раздел 2 (п. 4):
    air_inventory_date      — дата/реквизиты инвентаризации выбросов;
    air_inventory_next      — срок следующей инвентаризации;
    air_inventory_revision  — условия/сроки корректировки её данных;
    markers                 — перечень маркерных веществ (коды или названия).
  Раздел 3 (п. 5):
    water_permit            — реквизиты права пользования водным объектом;
    treatment_facilities    — очистные сооружения (текст);
    water_scheme            — схемы систем водопотребления и водоотведения;
    flow_meters             — средства измерения расхода сброса:
                              [{name, error, verification}];
    water_accounting_terms  — сроки ведения учёта сточных вод;
    nds                     — НДС по веществу: {код или название: значение}.
  Раздел 4 (п. 6):
    oro                     — [{name, groro, inventory_date, inventory_next}];
    oro_none=True           — ОРО на объекте отсутствуют (явное подтверждение).
  Раздел 7 (п. 7):
    unit, responsible, duties, staff_count, rights_duties (текст или список;
    если нет — печатается ТИПОВОЙ текст по ст. 67 ФЗ-7 с пометкой),
    order_no / order_date   — приказ о назначении ответственного.
  Раздел 8 (п. 8):
    lab (строка) или labs — [{name, address, certificate, scope}];
    scope — область аккредитации (или ссылка на реестр ФСА).
  Раздел 9 (п. 9):
    methods                 — методики (методы) измерений по веществу:
                              {код или название: «ПНД Ф …»};
    air_frequency, water_frequency, air_monitoring, surface_water,
    points (совместимость со старым вводом).

Периодичности, прямо установленные Требованиями (значения по умолчанию):
  * сточные воды (п. 9.2.2): объекты I и II категорий — не менее 1 раза
    в месяц, по показателю токсичности — не менее 1 раза в квартал;
    объекты III категории — не менее 1 раза в квартал;
  * проверки работы очистных сооружений (п. 9.2.4) — не реже 2 раз в год;
  * периодичность контроля ВЫБРОСОВ приказом № 109 не фиксируется —
    её определяет владелец объекта в плане-графике (п. 9.1.1); по
    умолчанию ставим «1 раз в год» и честно говорим об этом в gaps().

Чего машина не делает: не считает рассеивание (какие источники дают
< 0,1 ПДКмр на границе участка и могут не включаться в план-график,
п. 9.1.2), не выбирает створы наблюдений за водным объектом (п. 9.2.3),
не назначает маркерные вещества, не определяет уровень надзора. Там, где
нужны эти данные, в документе остаётся пометка «[требуется: …]», и та же
строка попадает в gaps().
"""
from __future__ import annotations

import re
from collections import OrderedDict
from datetime import date
from decimal import Decimal
from pathlib import Path

from ecodoc.core.models import Medium, ReportContext

TITLE = "Программа производственного экологического контроля"

NPA = ("приказ Минприроды России от 18.02.2022 № 109 (в редакции приказов "
       "от 24.03.2023 № 150, от 13.11.2024 № 659, от 07.05.2025 № 255, "
       "от 12.05.2025 № 262)")

# Разделы программы — дословно по п. 2 Требований (приложение 1 к № 109).
SECTIONS: list[tuple[str, str]] = [
    ("general", "Общие положения"),
    ("air_inv", "Сведения об инвентаризации выбросов загрязняющих веществ "
                "в атмосферный воздух и их источников"),
    ("water_inv", "Сведения об инвентаризации сбросов загрязняющих веществ "
                  "в окружающую среду и их источников"),
    ("waste_inv", "Сведения об инвентаризации отходов производства и "
                  "потребления и объектов их размещения"),
    ("ppp", "Сведения о побочных продуктах производства"),
    ("soil", "Сведения о произведённых из органической части твёрдых "
             "коммунальных отходов искусственных грунтах"),
    ("staff", "Сведения о подразделениях и (или) должностных лицах, "
              "отвечающих за осуществление производственного "
              "экологического контроля"),
    ("lab", "Сведения о собственных и (или) привлекаемых испытательных "
            "лабораториях (центрах), аккредитованных в соответствии с "
            "законодательством Российской Федерации об аккредитации в "
            "национальной системе аккредитации"),
    ("plan", "Сведения о периодичности и методах осуществления "
             "производственного экологического контроля, местах отбора "
             "проб и методиках (методах) измерений"),
]

DEFAULT_AIR_FREQ = "1 раз в год"          # НПА не фиксирует — типовая практика
FREQ_WATER_I_II = "не менее 1 раза в месяц"        # п. 9.2.2
FREQ_WATER_III = "не менее 1 раза в квартал"       # п. 9.2.2
FREQ_TOXICITY = "не менее 1 раза в квартал"        # п. 9.2.2
FREQ_TREATMENT = "не реже 2 раз в год"             # п. 9.2.4

# Уполномоченные органы, куда направляется отчёт ПЭК (п. 3 Требований), по
# коду субъекта РФ объекта. Почему справочник маленький: наименования органов
# субъектов меняются при реорганизациях, и выдумывать их для 85 регионов
# нельзя — заведены только те, что проверены (СПб, ЛО, Москва, МО).
# Федеральный надзор — территориальный орган Росприроднадзора, региональный —
# орган исполнительной власти субъекта.
AUTHORITIES: dict[str, dict[str, str]] = {
    "78": {"федеральный": "Северо-Западное межрегиональное управление "
                          "Росприроднадзора",
           "региональный": "Комитет по природопользованию, охране окружающей "
                           "среды и обеспечению экологической безопасности "
                           "Санкт-Петербурга"},
    "47": {"федеральный": "Северо-Западное межрегиональное управление "
                          "Росприроднадзора",
           "региональный": "Комитет государственного экологического надзора "
                           "Ленинградской области"},
    "77": {"федеральный": "Центральное межрегиональное управление "
                          "Росприроднадзора",
           "региональный": "Департамент природопользования и охраны "
                           "окружающей среды города Москвы"},
    "50": {"федеральный": "Центральное межрегиональное управление "
                          "Росприроднадзора",
           "региональный": "Министерство экологии и природопользования "
                           "Московской области"},
}

# ОПФ по аббревиатуре в начале наименования — только бесспорные случаи.
LEGAL_FORMS: dict[str, str] = {
    "ООО": "Общество с ограниченной ответственностью",
    "АО": "Акционерное общество",
    "ПАО": "Публичное акционерное общество",
    "ЗАО": "Закрытое акционерное общество",
    "ОАО": "Открытое акционерное общество",
    "НАО": "Непубличное акционерное общество",
    "ИП": "Индивидуальный предприниматель",
    "МУП": "Муниципальное унитарное предприятие",
    "ГУП": "Государственное унитарное предприятие",
    "ФГУП": "Федеральное государственное унитарное предприятие",
    "ГБУ": "Государственное бюджетное учреждение",
    "МБУ": "Муниципальное бюджетное учреждение",
    "СНТ": "Садоводческое некоммерческое товарищество",
    "ТСЖ": "Товарищество собственников жилья",
}

# Типовые права и обязанности ответственного за ПЭК (ст. 67 ФЗ-7; п. 7
# Требований). Печатаются только при отсутствии pek.rights_duties и ПОМЕЧАЮТСЯ
# как типовые, чтобы их не приняли за сведения из приказа организации.
DEFAULT_RIGHTS_DUTIES: list[str] = [
    "организовывать производственный экологический контроль на объекте и "
    "координировать природоохранную деятельность подразделений;",
    "организовывать инвентаризацию (корректировку инвентаризации) выбросов "
    "и сбросов загрязняющих веществ и их источников, инвентаризацию отходов;",
    "организовывать контроль за выбросами, сбросами загрязняющих веществ и "
    "обращением с отходами, в том числе приборный и лабораторный контроль "
    "с привлечением аккредитованных лабораторий;",
    "обеспечивать ведение учёта в области обращения с отходами и учёта "
    "сточных вод, подготовку и представление отчёта об организации и о "
    "результатах осуществления ПЭК (ежегодно до 25 марта), статистической "
    "отчётности и расчёта платы за негативное воздействие;",
    "контролировать выполнение предписаний уполномоченных органов и "
    "природоохранных мероприятий, приостанавливать работы, ведущиеся с "
    "нарушением природоохранных требований;",
    "запрашивать у подразделений сведения, необходимые для осуществления ПЭК, "
    "и вносить руководителю предложения по устранению нарушений.",
]

MONTHS_GEN = ("января", "февраля", "марта", "апреля", "мая", "июня", "июля",
              "августа", "сентября", "октября", "ноября", "декабря")


# ────────────────────────────────────────────────────────────────────────
# доступ к данным
# ────────────────────────────────────────────────────────────────────────
def _pek(ctx: ReportContext) -> dict:
    pek = (ctx.extra or {}).get("pek", {}) if isinstance(ctx.extra, dict) else {}
    return pek if isinstance(pek, dict) else {}


def _water_extra(ctx: ReportContext) -> dict:
    w = (ctx.extra or {}).get("water", {}) if isinstance(ctx.extra, dict) else {}
    return w if isinstance(w, dict) else {}


def worst_category(ctx: ReportContext) -> str:
    """Самая строгая категория среди объектов — по ней дефолтные периодичности."""
    from ecodoc.calendar.engine import category_of
    cats = {category_of(o) for o in ctx.objects}
    for c in ("I", "II", "III", "IV"):
        if c in cats:
            return c
    return ""


def air_pollutants(ctx: ReportContext) -> list:
    return [p for p in ctx.pollutants if p.medium == Medium.AIR]


def water_pollutants(ctx: ReportContext) -> list:
    return [p for p in ctx.pollutants if p.medium == Medium.WATER]


def _legacy_points(ctx: ReportContext, medium: str) -> list[dict]:
    """Старый ручной ввод точек контроля (extra.pek.points) — не теряем его.

    medium — подстрока для поля medium точки: «возд», «вод», «отход».
    """
    out = []
    for p in _pek(ctx).get("points", []) or []:
        if isinstance(p, dict) and medium in str(p.get("medium", "")).lower():
            out.append(p)
    return out


def _lookup(mapping, code: str, name: str) -> str:
    """Значение из словаря pek.methods / pek.nds по коду ИЛИ названию вещества.

    Коды сравниваем без ведущих нулей («0301» == «301») — в выгрузках
    «Эколога» и в приказах они пишутся по-разному.
    """
    if not isinstance(mapping, dict):
        return ""
    norm = {}
    for k, v in mapping.items():
        ks = str(k).strip().lower()
        norm[ks] = v
        if ks.isdigit():
            norm[ks.lstrip("0") or "0"] = v
    for key in (code, name):
        ks = str(key or "").strip().lower()
        if not ks:
            continue
        if ks in norm:
            return str(norm[ks])
        if ks.isdigit() and (ks.lstrip("0") or "0") in norm:
            return str(norm[ks.lstrip("0") or "0"])
    return ""


def legal_form(ctx: ReportContext) -> str:
    """ОПФ: из pek.legal_form, иначе по аббревиатуре в начале наименования."""
    pek, org = _pek(ctx), ctx.organization
    if pek.get("legal_form"):
        return str(pek["legal_form"])
    for text in (org.name, org.short_name):
        m = re.match(r"\s*([А-ЯЁ]{2,4})\b", text or "")
        if m and m.group(1) in LEGAL_FORMS:
            return LEGAL_FORMS[m.group(1)]
    if org.is_individual:
        return LEGAL_FORMS["ИП"]
    return ""


def supervision_level(ctx: ReportContext) -> str:
    """Уровень надзора — только из pek.supervision_level (машина его не выводит:
    он зависит от критериев ПП № 1096, а не от категории объекта)."""
    lvl = str(_pek(ctx).get("supervision_level") or "").strip().lower()
    if "фед" in lvl:
        return "федеральный"
    if "рег" in lvl:
        return "региональный"
    return ""


def authority(ctx: ReportContext) -> str:
    """Уполномоченный орган для отчёта ПЭК: pek.authority, иначе справочник."""
    pek = _pek(ctx)
    if pek.get("authority"):
        return str(pek["authority"])
    lvl = supervision_level(ctx)
    from ecodoc.core.nvos import subject_code
    # region_code в базе — префикс ОКТМО (40 = СПб); справочник органов — по субъекту (78)
    regions = {subject_code(o.region_code) or subject_code(o.code) for o in ctx.objects}
    regions.discard("")
    if lvl and len(regions) == 1:
        return AUTHORITIES.get(next(iter(regions)), {}).get(lvl, "")
    return ""


def _inventory_source_rows(ctx: ReportContext) -> list[dict]:
    """Источник × вещество из инвентаризации (для раздела 2 и плана 9.1)."""
    from ecodoc.development.air_inventory import sources
    out = []
    for s in sources(ctx):
        label = " ".join(x for x in (s["number"], s["name"]) if x).strip()
        for p in s["pollutants"]:
            out.append({"source": label, "kind": s["kind"],
                        "code": str(p.get("code") or ""),
                        "name": str(p.get("name") or ""),
                        "g_s": p.get("g_s"), "t_year": p.get("t_year")})
    return out


def _marker_fn(ctx: ReportContext):
    markers = {str(m).strip().lower() for m in _pek(ctx).get("markers", []) or []}

    def marker(code: str, name: str) -> str:
        c = code.lower()
        return "да" if (c in markers or (c.lstrip("0") or "0") in markers
                        or name.lower() in markers) else "—"
    return marker


# ────────────────────────────────────────────────────────────────────────
# план-графики (то, что становится таблицами)
# ────────────────────────────────────────────────────────────────────────
def plan_air(ctx: ReportContext) -> list[dict]:
    """План-график контроля стационарных источников выбросов (п. 9.1.1).

    Графы по НПА: подразделение/источник, загрязняющее вещество (в т.ч.
    маркерное), норматив/мощность выброса (г/с, т/год — по нему вещество и
    попадает в план-график), метод контроля (инструментальный/расчётный),
    периодичность, место отбора проб, методика (метод) измерений.
    Источник данных — инвентаризация (extra.emission_sources); если
    источников нет, а вещества есть — строки по веществам с пометкой, что
    нужен номер источника.
    """
    from ecodoc.development.air_inventory import sources
    pek = _pek(ctx)
    marker = _marker_fn(ctx)
    freq = str(pek.get("air_frequency") or "") or DEFAULT_AIR_FREQ
    methods = pek.get("methods")

    def methodology(method: str, code: str, name: str) -> str:
        m = _lookup(methods, code, name)
        if m:
            return m
        # для расчётного метода «методика» — это методика расчёта выбросов,
        # её тоже нужно указать (п. 9.1.3), но выдумывать нельзя
        return "[требуется: аттестованная методика измерений]" \
            if method == "инструментальный" \
            else "[требуется: методика расчёта выбросов]"

    def norm(g_s, t_year) -> str:
        parts = []
        if g_s not in (None, ""):
            parts.append(f"{_num(g_s)} г/с")
        if t_year not in (None, ""):
            parts.append(f"{_num(t_year)} т/год")
        return "; ".join(parts) or "[требуется: НДВ]"

    rows: list[dict] = []
    srcs = sources(ctx)
    for s in srcs:
        label = " ".join(x for x in (s["number"], s["name"]) if x).strip()
        # неорганизованный источник — классический случай расчётного метода
        # (п. 9.1.3: нет практической возможности инструментальных измерений)
        method = ("расчётный" if "неорг" in s["kind"].lower()
                  else "инструментальный")
        subs = s["pollutants"] or [{}]
        for p in subs:
            code, name = str(p.get("code") or ""), str(p.get("name") or "")
            rows.append({
                "source": label or "[требуется: номер источника выбросов]",
                "code": code or "—",
                "name": name or "[требуется: вещества источника]",
                "marker": marker(code, name),
                "norm": norm(p.get("g_s"), p.get("t_year")),
                "method": method,
                "frequency": freq,
                "place": "устье источника" if method == "инструментальный"
                         else "по данным инвентаризации",
                "methodology": methodology(method, code, name),
            })
    if not srcs:
        for p in air_pollutants(ctx):
            rows.append({
                "source": "[требуется: номер источника выбросов]",
                "code": p.code or "—", "name": p.name or "—",
                "marker": marker(p.code or "", p.name or ""),
                "norm": "[требуется: НДВ]",
                "method": "инструментальный", "frequency": freq,
                "place": "устье источника",
                "methodology": methodology("инструментальный", p.code or "",
                                           p.name or ""),
            })
    for p in _legacy_points(ctx, "возд"):
        rows.append({"source": str(p.get("point", "")), "code": "—",
                     "name": str(p.get("indicators", "")), "marker": "—",
                     "norm": str(p.get("norm", "") or "—"),
                     "method": str(p.get("method", "инструментальный")),
                     "frequency": str(p.get("frequency", freq)),
                     "place": str(p.get("point", "")),
                     "methodology": str(p.get("methodology", "") or
                                        "[требуется: методика измерений]")})
    return rows


def plan_water(ctx: ReportContext) -> list[dict]:
    """План-график контроля сточных вод (п. 9.2.1–9.2.2).

    Графы: выпуск, место отбора проб, вещество, НДС, периодичность,
    аттестованная методика измерений. Периодичность по умолчанию — из НПА
    по категории объекта; показатель токсичности добавляется отдельной
    строкой (п. 9.2.2).
    """
    pek = _pek(ctx)
    cat = worst_category(ctx)
    if pek.get("water_frequency"):
        freq = str(pek["water_frequency"])
    elif cat in ("I", "II"):
        freq = FREQ_WATER_I_II
    elif cat == "III":
        freq = FREQ_WATER_III
    else:
        freq = "[требуется: категория объекта — от неё зависит периодичность]"

    outs = [d for d in _water_extra(ctx).get("discharge", []) or []
            if isinstance(d, dict)]
    outlets = [str(d.get("receiver") or "") for d in outs]
    outlet = "; ".join(x for x in outlets if x) or \
        "[требуется: выпуск / место отбора проб сточных вод]"
    # место отбора проб — отдельная графа (п. 9.2.1); в выгрузке 2-ТП его
    # нет, поэтому берём sampling_point выпуска, иначе просим
    places = [str(d.get("sampling_point") or "") for d in outs]
    place = "; ".join(x for x in places if x) or \
        "[требуется: место отбора проб]"
    methods, nds = pek.get("methods"), pek.get("nds")

    def row(code, name, frequency):
        return {"outlet": outlet, "place": place, "code": code or "—",
                "name": name or "—",
                "nds": _lookup(nds, code, name) or "[требуется: НДС]",
                "frequency": frequency,
                "methodology": _lookup(methods, code, name)
                or "[требуется: аттестованная методика измерений]"}

    rows = [row(p.code, p.name, freq) for p in water_pollutants(ctx)]
    if rows:
        t = row("", "Токсичность (биотестирование)", FREQ_TOXICITY)
        t["nds"] = "—"
        rows.append(t)
    for p in _legacy_points(ctx, "вод"):
        rows.append({"outlet": str(p.get("point", "")),
                     "place": str(p.get("point", "")), "code": "—",
                     "name": str(p.get("indicators", "")), "nds": "—",
                     "frequency": str(p.get("frequency", freq)),
                     "methodology": str(p.get("methodology", "") or
                                        "[требуется: методика измерений]")})
    return rows


def rows_waste(ctx: ReportContext) -> list[dict]:
    """Перечень отходов — общий с инвентаризацией отходов (один источник правды)."""
    from ecodoc.development.waste_inventory import collect
    return collect(ctx)


# ────────────────────────────────────────────────────────────────────────
# пробелы в данных
# ────────────────────────────────────────────────────────────────────────
def _gap_map(ctx: ReportContext) -> "OrderedDict[str, list[str]]":
    """Пробелы по разделам. gaps() — это же списком; generate() печатает
    каждую строку в своём разделе как «[…]», поэтому текст и список всегда
    совпадают (никакого второго источника правды)."""
    m: "OrderedDict[str, list[str]]" = OrderedDict(
        (k, []) for k, _ in SECTIONS)
    org, pek = ctx.organization, _pek(ctx)
    have_air = bool(air_pollutants(ctx)) or bool(
        (ctx.extra or {}).get("emission_sources"))
    have_water = bool(water_pollutants(ctx)) or bool(
        _water_extra(ctx).get("discharge"))
    have_waste = bool(ctx.wastes or ctx.waste_acts)

    if not (org.name or org.short_name):
        m["general"].append("требуется: наименование организации (п. 3 Требований)")
    if not org.inn:
        m["general"].append("требуется: ИНН организации (п. 3 Требований)")
    if not legal_form(ctx):
        m["general"].append("требуется: организационно-правовая форма "
                            "(п. 3 Требований)")
    if not org.okved:
        m["general"].append("требуется: ОКВЭД (основной вид деятельности)")
    if not ctx.objects:
        m["general"].append("требуется: объект НВОС — код, категория и адрес "
                            "по свидетельству о постановке на учёт")
    else:
        for o in ctx.objects:
            if not worst_category(ctx):
                m["general"].append(
                    f"требуется: категория объекта {o.code or o.name} "
                    "(программа ПЭК — для I–III категорий)")
                break
    if worst_category(ctx) == "IV":
        m["general"].append("объект IV категории — программа ПЭК для него "
                            "не разрабатывается (ст. 67 ФЗ-7)")
    if not supervision_level(ctx):
        m["general"].append("требуется: уровень государственного экологического "
                            "надзора (федеральный/региональный) — от него "
                            "зависит орган, куда сдаётся отчёт ПЭК")
    if not authority(ctx):
        m["general"].append("требуется: наименование уполномоченного органа, в "
                            "который направляется отчёт ПЭК (п. 3 Требований; "
                            "для регионов вне справочника — задайте "
                            "pek.authority)")
    if not (pek.get("program_date") or pek.get("approved_date")):
        m["general"].append("требуется: дата утверждения Программы "
                            "(п. 3 Требований)")
    if not (have_air or have_water or have_waste):
        m["general"].append("нет ни выбросов, ни сбросов, ни отходов — "
                            "программу нечем наполнять: загрузите исходные "
                            "данные")

    if have_air:
        if not pek.get("air_inventory_date"):
            m["air_inv"].append("требуется: дата и реквизиты инвентаризации "
                                "выбросов и её последней корректировки "
                                "(п. 4 Требований)")
        if not (pek.get("air_inventory_next") or pek.get("air_inventory_revision")):
            m["air_inv"].append("требуется: сроки проведения инвентаризации "
                                "выбросов и корректировки её данных "
                                "(п. 4 Требований)")
        if not (ctx.extra or {}).get("emission_sources"):
            m["air_inv"].append("требуется: перечень стационарных источников "
                                "выбросов — загрузите инвентаризацию по "
                                "приказу № 871 или проект НДВ")
        if not pek.get("markers"):
            m["plan"].append("требуется: перечень маркерных веществ для "
                             "плана-графика контроля выбросов (п. 9.1.1)")
        if not pek.get("air_frequency"):
            m["plan"].append("периодичность контроля выбросов принята типовой "
                             f"({DEFAULT_AIR_FREQ}) — НПА её не фиксирует; "
                             "уточните в плане-графике при утверждении")
        # п. 9.1.1: в план-график включаются вещества, для которых
        # установлены НДВ — без норматива строка плана не обоснована
        no_ndv = sorted({r["name"] or r["code"] for r in plan_air(ctx)
                         if r["norm"].startswith("[требуется")})
        if no_ndv:
            m["plan"].append("требуется: нормативы допустимых выбросов (г/с, "
                             "т/год) для веществ плана-графика — без НДВ "
                             "вещество не включается в план (п. 9.1.1): "
                             + ", ".join(no_ndv))
        if not pek.get("methods"):
            m["plan"].append("требуется: аттестованные методики (методы) "
                             "измерений по веществам плана-графика "
                             "(п. 9 Требований)")
        m["plan"].append("требуется: результаты расчёта рассеивания — источники "
                         "с вкладом менее 0,1 ПДКмр на границе участка можно "
                         "не включать в план-график (п. 9.1.2)")

    if have_water:
        if not pek.get("water_permit"):
            m["water_inv"].append("требуется: реквизиты договора водопользования "
                                  "или решения о предоставлении водного объекта "
                                  "в пользование (п. 5 Требований)")
        outs = [d for d in _water_extra(ctx).get("discharge", []) or []
                if isinstance(d, dict)]
        if not any(d.get("pollutants") for d in outs):
            m["water_inv"].append("требуется: масса сброса по каждому веществу "
                                  "ПО КАЖДОМУ ВЫПУСКУ (п. 5 Требований)")
        if not (pek.get("treatment_facilities")
                or any(d.get("treatment_code") for d in outs)):
            m["water_inv"].append("требуется: сведения об очистных сооружениях "
                                  "(п. 5 Требований)")
        if not pek.get("water_scheme"):
            m["water_inv"].append("требуется: сведения о схемах систем "
                                  "водопотребления и водоотведения "
                                  "(п. 5 Требований)")
        if not pek.get("flow_meters"):
            m["water_inv"].append("требуется: средства измерения расхода сброса "
                                  "— наименование, погрешность, свидетельство "
                                  "о поверке (п. 5 Требований)")
        if not pek.get("water_accounting_terms"):
            m["water_inv"].append("требуется: сроки ведения учёта сточных вод "
                                  "(п. 5 Требований)")
        if not pek.get("nds"):
            m["plan"].append("требуется: нормативы допустимых сбросов по "
                             "показателям плана-графика (п. 9.2.1)")
        if not pek.get("methods") and not have_air:
            m["plan"].append("требуется: аттестованные методики (методы) "
                             "измерений по веществам плана-графика "
                             "(п. 9 Требований)")
        if not pek.get("surface_water"):
            m["plan"].append("требуется: программа наблюдений за водным объектом "
                             "и его водоохранной зоной — фоновый и контрольный "
                             "створы (п. 9.2.3)")

    if have_waste:
        if not pek.get("oro") and not pek.get("oro_none"):
            m["waste_inv"].append("требуется: сведения об эксплуатируемых "
                                  "объектах размещения отходов по ГРОРО либо "
                                  "подтверждение их отсутствия (п. 6 Требований)")
        for o in pek.get("oro") or []:
            if isinstance(o, dict) and not (o.get("inventory_date")
                                            or o.get("inventory_next")):
                m["waste_inv"].append(
                    f"требуется: сведения об инвентаризации ОРО "
                    f"«{o.get('name') or o.get('groro') or '—'}» и сроки её "
                    "проведения (п. 6 Требований)")

    if not (pek.get("responsible") or pek.get("unit") or org.director_name):
        m["staff"].append("требуется: подразделение и (или) должностное лицо, "
                          "отвечающее за ПЭК (п. 7 Требований)")
    else:
        if not pek.get("staff_count"):
            m["staff"].append("требуется: численность сотрудников подразделения "
                              "и (или) должностных лиц ПЭК (п. 7 Требований)")
        if not pek.get("rights_duties"):
            m["staff"].append("права и обязанности приведены ТИПОВЫЕ (ст. 67 "
                              "ФЗ-7) — замените текстом из должностной "
                              "инструкции/приказа (п. 7 Требований)")
        if not (pek.get("order_no") or pek.get("order_date")):
            m["staff"].append("требуется: реквизиты приказа о назначении "
                              "ответственного за ПЭК")
    labs = pek.get("labs") or ([{"name": pek["lab"]}] if pek.get("lab") else [])
    if not labs:
        m["lab"].append("требуется: наименование, адрес и реквизиты аттестата "
                        "аккредитации испытательной лаборатории "
                        "(п. 8 Требований)")
    elif not all(isinstance(l, dict) and l.get("scope") for l in labs):
        m["lab"].append("требуется: область аккредитации лаборатории "
                        "(п. 8 Требований; по реестру ФСА)")
    return m


def gaps(ctx: ReportContext) -> list[str]:
    """Чего не хватает для программы — тот же список, что и пометки в тексте."""
    return [g for lst in _gap_map(ctx).values() for g in lst]


# ────────────────────────────────────────────────────────────────────────
# документ
# ────────────────────────────────────────────────────────────────────────
def _num(value) -> str:
    if value in (None, "", 0):
        return "—"
    try:
        d = Decimal(str(value).replace(",", "."))
    except Exception:
        return str(value)
    return f"{d.normalize():f}" if d else "—"


def _dec(value) -> Decimal:
    try:
        return Decimal(str(value).replace(",", ".")) if value not in (None, "") \
            else Decimal(0)
    except Exception:
        return Decimal(0)


def approval_stamp_date(ctx: ReportContext) -> str:
    """Строка даты в грифе «УТВЕРЖДАЮ».

    Если задана pek.program_date — она целиком («09» февраля 2025 г.): п. 3
    Требований знает одну дату утверждения, и в грифе должна быть она же.
    Иначе — пустой шаблон с текущим годом (а не отчётным: программа
    утверждается сейчас, а не в отчётном году).
    """
    pek = _pek(ctx)
    raw = str(pek.get("program_date") or pek.get("approved_date") or "").strip()
    m = re.match(r"(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{4})", raw)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), m.group(3)
        if 1 <= mo <= 12:
            return f"«{d:02d}» {MONTHS_GEN[mo - 1]} {y} г."
    m = re.match(r"(\d{4})-(\d{1,2})-(\d{1,2})", raw)          # ISO
    if m:
        y, mo, d = m.group(1), int(m.group(2)), int(m.group(3))
        if 1 <= mo <= 12:
            return f"«{d:02d}» {MONTHS_GEN[mo - 1]} {y} г."
    if raw:
        return raw                      # дата в свободной форме — как задана
    return f"«___» __________ {date.today().year} г."


def _table(doc, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, text in enumerate(header):
        table.rows[0].cells[i].text = text
    for r in rows:
        cells = table.add_row().cells
        for i, text in enumerate(r):
            cells[i].text = str(text)
    if not rows:
        table.add_row().cells[0].text = "[требуется: данные не заведены]"


def _marks(doc, items: list[str]) -> None:
    """Пометки раздела: та же строка, что в gaps(), в квадратных скобках."""
    for g in items:
        p = doc.add_paragraph()
        p.add_run(f"[{g}]").italic = True


def _sub(doc, text: str) -> None:
    doc.add_paragraph().add_run(text).bold = True


def _sec_general(doc, ctx: ReportContext) -> None:
    """Раздел 1 — таблица реквизитов по п. 3 Требований (как в эталоне ТХС)."""
    org, pek = ctx.organization, _pek(ctx)
    req = "[требуется]"
    lvl = supervision_level(ctx)
    rows = [
        ["Полное наименование организации", org.name or req],
        ["Сокращённое наименование", org.short_name or "—"],
        ["Организационно-правовая форма", legal_form(ctx) or req],
        ["Юридический адрес", org.address or req],
        ["ОГРН", org.ogrn or req],
        ["ИНН", org.inn or req],
        ["ОКВЭД (основной вид деятельности)", org.okved or req],
        ["Руководитель организации",
         " ".join(x for x in (org.official_title, org.director_name) if x)
         or req],
    ]
    if ctx.objects:
        for o in ctx.objects:
            rows += [
                ["Наименование объекта", o.name or req],
                ["Место нахождения объекта", o.address or req],
                ["Категория объекта", o.category or req],
                ["Код объекта, присвоенный при постановке на учёт",
                 o.code or req],
            ]
    else:
        rows.append(["Объект НВОС", "[требуется: объект по свидетельству о "
                                    "постановке на учёт]"])
    rows += [
        ["Уровень надзора", lvl.capitalize() if lvl else req],
        ["Наименование уполномоченного органа, в который направляется отчёт "
         "об организации и о результатах осуществления ПЭК",
         authority(ctx) or req],
        ["Должностное лицо, ответственное за подготовку отчёта об "
         "организации и о результатах осуществления ПЭК",
         pek.get("responsible") or org.director_name or req],
        ["Дата утверждения Программы",
         pek.get("program_date") or pek.get("approved_date") or req],
    ]
    _table(doc, ["Показатель", "Сведения"], rows)
    doc.add_paragraph(
        "Отчёт об организации и о результатах осуществления ПЭК "
        "представляется ежегодно до 25 марта года, следующего за отчётным, "
        f"в {authority(ctx) or 'уполномоченный орган [требуется]'} "
        "(приложение 2 к приказу № 109).")


def _sec_air_inv(doc, ctx: ReportContext, have_air: bool) -> None:
    """Раздел 2 — п. 4 Требований: масса по каждому веществу по КАЖДОМУ
    источнику и по объекту, маркерные вещества, сроки инвентаризации."""
    pek = _pek(ctx)
    if not have_air:
        doc.add_paragraph(
            "Стационарные источники выбросов на объекте отсутствуют / "
            "данные о выбросах не заведены — раздел не заполняется.")
        return
    if pek.get("air_inventory_date"):
        doc.add_paragraph("Инвентаризация выбросов проведена: "
                          f"{pek['air_inventory_date']}.")
    marker = _marker_fn(ctx)
    src_rows = _inventory_source_rows(ctx)
    if src_rows:
        _sub(doc, "Перечень источников выбросов и загрязняющих веществ по "
                  "каждому источнику (по данным инвентаризации)")
        _table(doc, ["Источник выбросов", "Тип", "Код", "Загрязняющее "
                     "вещество", "Маркерное", "г/с", "т/год"],
               [[r["source"] or "[требуется: номер источника]",
                 r["kind"] or "—", r["code"] or "—", r["name"] or "—",
                 marker(r["code"], r["name"]), _num(r["g_s"]),
                 _num(r["t_year"])] for r in src_rows])
    # итог по объекту: вещества базы (ctx.pollutants) — главный источник
    # масс; если их нет, суммируем т/год по источникам
    subs = air_pollutants(ctx)
    _sub(doc, "Суммарные выбросы по объекту в целом")
    if subs:
        _table(doc, ["№", "Код", "Загрязняющее вещество", "Маркерное",
                     "Выброс, т/год"],
               [[str(n), p.code or "—", p.name or "—",
                 marker(p.code or "", p.name or ""),
                 _num(p.mass_norm + p.mass_limit + p.mass_over)]
                for n, p in enumerate(subs, start=1)])
    else:
        totals: "OrderedDict[tuple, Decimal]" = OrderedDict()
        for r in src_rows:
            key = (r["code"], r["name"])
            totals[key] = totals.get(key, Decimal(0)) + _dec(r["t_year"])
        _table(doc, ["№", "Код", "Загрязняющее вещество", "Маркерное",
                     "Выброс, т/год"],
               [[str(n), c or "—", nm or "—", marker(c, nm), _num(t)]
                for n, ((c, nm), t) in enumerate(totals.items(), start=1)])
    _sub(doc, "Сроки проведения инвентаризации выбросов и их стационарных "
              "источников, корректировки её данных")
    if pek.get("air_inventory_next") or pek.get("air_inventory_revision"):
        if pek.get("air_inventory_next"):
            doc.add_paragraph("Срок проведения следующей инвентаризации: "
                              f"{pek['air_inventory_next']}.")
        if pek.get("air_inventory_revision"):
            doc.add_paragraph("Корректировка данных инвентаризации: "
                              f"{pek['air_inventory_revision']}.")
    # пробел печатается ниже из gap-карты


def _sec_water_inv(doc, ctx: ReportContext, have_water: bool) -> None:
    """Раздел 3 — п. 5 Требований: масса по веществу по КАЖДОМУ выпуску,
    очистные, схемы водопотребления/водоотведения, приборы учёта, сроки."""
    pek, w = _pek(ctx), _water_extra(ctx)
    if not have_water:
        doc.add_paragraph(
            "Сбросы загрязняющих веществ в окружающую среду "
            "отсутствуют / данные не заведены — раздел не заполняется.")
        return
    if pek.get("water_permit"):
        doc.add_paragraph("Право пользования водным объектом: "
                          f"{pek['water_permit']}.")
    outs = [d for d in w.get("discharge", []) or [] if isinstance(d, dict)]
    if outs:
        _sub(doc, "Выпуски сточных вод")
        _table(doc, ["№", "Выпуск (приёмник)", "Объём водоотведения, "
                     "тыс. м³/год"],
               [[str(n), d.get("receiver") or "—", _num(d.get("volume"))]
                for n, d in enumerate(outs, start=1)])
    per_outlet = [(d, [p for p in (d.get("pollutants") or [])
                       if isinstance(p, dict)]) for d in outs]
    if any(pl for _, pl in per_outlet):
        _sub(doc, "Масса сброса по каждому загрязняющему веществу по каждому "
                  "выпуску")
        _table(doc, ["Выпуск", "Код", "Загрязняющее вещество", "Сброс, т/год"],
               [[d.get("receiver") or "—", str(p.get("code") or "—"),
                 str(p.get("name") or "—"), _num(p.get("mass"))]
                for d, pl in per_outlet for p in pl])
    subs = water_pollutants(ctx)
    _sub(doc, "Суммарный сброс по объекту в целом")
    _table(doc, ["№", "Код", "Загрязняющее вещество", "Сброс, т/год"],
           [[str(n), p.code or "—", p.name or "—",
             _num(p.mass_norm + p.mass_limit + p.mass_over)]
            for n, p in enumerate(subs, start=1)])

    _sub(doc, "Сведения об очистных сооружениях")
    if pek.get("treatment_facilities"):
        doc.add_paragraph(str(pek["treatment_facilities"]))
    else:
        tf = [d for d in outs if d.get("treatment_code")]
        if tf:
            _table(doc, ["Выпуск", "Код типа очистного сооружения (2-ТП "
                         "водхоз)", "Мощность, тыс. м³/год"],
                   [[d.get("receiver") or "—", str(d.get("treatment_code")),
                     _num(d.get("treatment_capacity"))] for d in tf])
    _sub(doc, "Сведения о схемах систем водопотребления и водоотведения")
    if pek.get("water_scheme"):
        doc.add_paragraph(str(pek["water_scheme"]))
    _sub(doc, "Средства измерения расхода сброса сточных вод")
    fm = [x for x in (pek.get("flow_meters") or []) if isinstance(x, dict)]
    if fm:
        _table(doc, ["Наименование средства измерения", "Погрешность",
                     "Свидетельство о поверке"],
               [[x.get("name") or "—", x.get("error") or "[требуется]",
                 x.get("verification") or "[требуется]"] for x in fm])
    _sub(doc, "Сроки ведения учёта сточных вод")
    if pek.get("water_accounting_terms"):
        doc.add_paragraph(str(pek["water_accounting_terms"]))
    # пробелы по пустым блокам печатаются ниже из gap-карты


def _sec_waste_inv(doc, ctx: ReportContext) -> None:
    """Раздел 4 — п. 6 Требований: отходы, ОРО по ГРОРО, их инвентаризация."""
    pek = _pek(ctx)
    rows = rows_waste(ctx)
    if not rows:
        doc.add_paragraph("Отходы производства и потребления не "
                          "заведены — раздел не заполняется.")
    else:
        doc.add_paragraph(
            f"Отходы, образующиеся в процессе деятельности (по ФККО) — "
            f"{len(rows)} вид(ов):")
        _table(doc, ["№", "Наименование отхода", "Код ФККО",
                     "Класс опасности", "Образование, т/год"],
               [[str(n), r["name"] or "—", r["fkko"] or "—",
                 str(r["hazard"] or "—"), _num(r["generated"])]
                for n, r in enumerate(rows, start=1)])
    _sub(doc, "Сведения об объектах размещения отходов, их инвентаризации и "
              "сроках её проведения")
    oro = [o for o in (pek.get("oro") or []) if isinstance(o, dict)]
    if oro:
        _table(doc, ["№", "Объект размещения отходов", "№ в ГРОРО",
                     "Дата инвентаризации", "Срок следующей инвентаризации"],
               [[str(n), str(o.get("name", "—")), str(o.get("groro", "—")),
                 str(o.get("inventory_date") or "[требуется]"),
                 str(o.get("inventory_next") or "[требуется]")]
                for n, o in enumerate(oro, start=1)])
    elif pek.get("oro_none"):
        doc.add_paragraph(
            "Объекты размещения отходов на объекте отсутствуют; отходы "
            "передаются на размещение, утилизацию и обезвреживание сторонним "
            "организациям, имеющим лицензии. Инвентаризация объектов "
            "размещения отходов не проводится.")
    # иначе — пробел из gap-карты


def _sec_staff(doc, ctx: ReportContext) -> None:
    """Раздел 7 — п. 7 Требований: подразделения/лица, полномочия,
    численность, права и обязанности, приказ о назначении."""
    org, pek = ctx.organization, _pek(ctx)
    unit = pek.get("unit") or ""
    resp = pek.get("responsible") or org.director_name or ""
    if not (unit or resp):
        return                                  # пробел из gap-карты
    _table(doc, ["Подразделение", "Должностное лицо", "Полномочия",
                 "Численность, чел."],
           [[unit or "—", resp or "—",
             pek.get("duties") or "организация и осуществление "
             "ПЭК, подготовка отчёта о результатах ПЭК",
             str(pek.get("staff_count") or "[требуется]")]])
    if pek.get("order_no") or pek.get("order_date"):
        doc.add_paragraph(
            "Ответственный за осуществление ПЭК назначен приказом "
            + (f"№ {pek['order_no']} " if pek.get("order_no") else "")
            + (f"от {pek['order_date']}" if pek.get("order_date") else "")
            + ".")
    _sub(doc, "Сведения о правах и обязанностях руководителей, сотрудников "
              "подразделений")
    rd = pek.get("rights_duties")
    if rd:
        items = rd if isinstance(rd, (list, tuple)) else [rd]
        for it in items:
            doc.add_paragraph(str(it))
    else:
        # типовой текст — помечаем явно, это не сведения организации
        doc.add_paragraph("(типовой перечень по ст. 67 Федерального закона "
                          "№ 7-ФЗ) Ответственный за осуществление ПЭК "
                          "обязан и вправе:")
        for it in DEFAULT_RIGHTS_DUTIES:
            doc.add_paragraph(f"— {it}")


def _sec_lab(doc, ctx: ReportContext) -> None:
    """Раздел 8 — п. 8 Требований: лаборатории с областью аккредитации."""
    pek = _pek(ctx)
    labs = pek.get("labs") or ([{"name": pek["lab"]}] if pek.get("lab") else [])
    _table(doc, ["№", "Лаборатория (центр)", "Адрес", "Аттестат аккредитации",
                 "Область аккредитации"],
           [[str(n), l.get("name") or "—", l.get("address") or "—",
             l.get("certificate") or "[требуется: аттестат]",
             l.get("scope") or "[требуется: область аккредитации]"]
            for n, l in enumerate(labs, start=1) if isinstance(l, dict)])


def generate(ctx: ReportContext, out_path: str | Path) -> Path:
    """Собрать программу ПЭК (.docx) по составу приказа № 109."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Cm, Pt

    org, pek = ctx.organization, _pek(ctx)
    gap = _gap_map(ctx)
    have_air = bool(air_pollutants(ctx)) or bool(
        (ctx.extra or {}).get("emission_sources"))
    have_water = bool(water_pollutants(ctx)) or bool(
        _water_extra(ctx).get("discharge"))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(3), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(2), Cm(2)

    # ── титул: гриф утверждения обязателен — дату требует п. 3; дата в
    # грифе та же, что в разделе 1 (pek.program_date) ─────────────────────
    ap = doc.add_paragraph()
    ap.alignment = AL.RIGHT
    ap.add_run("УТВЕРЖДАЮ\n"
               f"{org.official_title or 'Руководитель'}\n"
               f"_____________ {org.director_name or '[Ф.И.О.]'}\n"
               f"{approval_stamp_date(ctx)}")
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = AL.CENTER
    run = t.add_run(TITLE.upper())
    run.bold = True
    run.font.size = Pt(14)
    sub = doc.add_paragraph()
    sub.alignment = AL.CENTER
    sub.add_run(org.name or org.short_name or "[требуется: наименование "
                "организации (п. 3 Требований)]")
    basis = doc.add_paragraph()
    basis.alignment = AL.CENTER
    basis.add_run(
        "Разработана в соответствии со ст. 67 Федерального закона от "
        f"10.01.2002 № 7-ФЗ и требованиями к содержанию программы "
        f"производственного экологического контроля, утверждёнными "
        f"приказом Минприроды России от 18.02.2022 № 109 (в редакции "
        f"приказов от 24.03.2023 № 150, от 13.11.2024 № 659, "
        f"от 07.05.2025 № 255, от 12.05.2025 № 262)").italic = True
    doc.add_page_break()

    # ── содержание ───────────────────────────────────────────────────────
    doc.add_paragraph().add_run("СОДЕРЖАНИЕ").bold = True
    for i, (_, title) in enumerate(SECTIONS, start=1):
        doc.add_paragraph(f"{i}. {title}")
    doc.add_page_break()

    for i, (key, title) in enumerate(SECTIONS, start=1):
        head = doc.add_paragraph()
        head.add_run(f"{i}. {title}").bold = True

        if key == "general":
            _sec_general(doc, ctx)

        elif key == "air_inv":
            _sec_air_inv(doc, ctx, have_air)

        elif key == "water_inv":
            _sec_water_inv(doc, ctx, have_water)

        elif key == "waste_inv":
            _sec_waste_inv(doc, ctx)

        elif key == "ppp":
            ppp = [p for p in ((ctx.extra or {}).get("ppp") or [])
                   if isinstance(p, dict)]
            if ppp:
                _table(doc, ["№", "Побочный продукт производства",
                             "Образование, т/год"],
                       [[str(n), p.get("name") or "—", _num(p.get("formed"))]
                        for n, p in enumerate(ppp, start=1)])
            else:
                doc.add_paragraph("Побочные продукты производства не "
                                  "образуются (сведений в базе нет).")

        elif key == "soil":
            soil = [s for s in ((ctx.extra or {}).get("artificial_soil") or [])
                    if isinstance(s, dict)]
            if soil:
                _table(doc, ["№", "Искусственный грунт", "Производство, т/год"],
                       [[str(n), s.get("name") or "—", _num(s.get("formed"))]
                        for n, s in enumerate(soil, start=1)])
            else:
                doc.add_paragraph("Искусственные грунты из органической части "
                                  "ТКО не производятся (сведений в базе нет).")

        elif key == "staff":
            _sec_staff(doc, ctx)

        elif key == "lab":
            _sec_lab(doc, ctx)

        elif key == "plan":
            # 9.1 воздух
            p91 = doc.add_paragraph()
            p91.add_run(f"{i}.1. Производственный контроль в области охраны "
                        "атмосферного воздуха").bold = True
            if have_air:
                doc.add_paragraph(
                    "План-график контроля стационарных источников выбросов "
                    "(п. 9.1.1 Требований): загрязняющие вещества, в том "
                    "числе маркерные, в отношении которых установлены "
                    "нормативы допустимых выбросов. В план-график не "
                    "включаются источники, выброс от которых по результатам "
                    "рассеивания не превышает 0,1 ПДКмр на границе "
                    "земельного участка (п. 9.1.2); расчётные методы "
                    "применяются при отсутствии аттестованных методик или "
                    "практической возможности инструментальных измерений "
                    "(п. 9.1.3).")
                _table(doc, ["Источник выбросов", "Код", "Загрязняющее вещество",
                             "Маркерное", "Норматив (мощность) выброса",
                             "Метод контроля", "Периодичность",
                             "Место отбора проб", "Методика (метод) измерений"],
                       [[r["source"], r["code"], r["name"], r["marker"],
                         r["norm"], r["method"], r["frequency"], r["place"],
                         r["methodology"]]
                        for r in plan_air(ctx)])
                mon = [x for x in (pek.get("air_monitoring") or [])
                       if isinstance(x, dict)]
                if mon:
                    doc.add_paragraph("План-график наблюдений за загрязнением "
                                      "атмосферного воздуха:")
                    _table(doc, ["Точка", "Вещество", "Периодичность"],
                           [[x.get("point_no") or x.get("address") or "—",
                             x.get("substance") or "—", x.get("period") or "—"]
                            for x in mon])
                else:
                    doc.add_paragraph(
                        "Наблюдения за загрязнением атмосферного воздуха "
                        "проводятся, если объект включён в перечень, "
                        "предусмотренный п. 3 ст. 23 Федерального закона "
                        "от 04.05.1999 № 96-ФЗ (сведений о включении нет).")
            else:
                doc.add_paragraph("Выбросы отсутствуют — контроль в области "
                                  "охраны атмосферного воздуха не ведётся.")

            # 9.2 вода
            p92 = doc.add_paragraph()
            p92.add_run(f"{i}.2. Производственный контроль в области охраны "
                        "и использования водных объектов").bold = True
            if have_water:
                doc.add_paragraph(
                    "План-график контроля состава и свойств сточных вод "
                    "(п. 9.2.1 Требований): показатели по нормативам "
                    "допустимого сброса, периодичность, места отбора проб, "
                    "аттестованные методики (методы) измерений. "
                    "Периодичность по п. 9.2.2 Требований: объекты I и II "
                    f"категорий — {FREQ_WATER_I_II}, объекты III категории — "
                    f"{FREQ_WATER_III}; по показателю токсичности — "
                    f"{FREQ_TOXICITY}.")
                _table(doc, ["Выпуск", "Место отбора проб", "Код",
                             "Показатель", "НДС", "Периодичность",
                             "Методика (метод) измерений"],
                       [[r["outlet"], r["place"], r["code"], r["name"],
                         r["nds"], r["frequency"], r["methodology"]]
                        for r in plan_water(ctx)])
                doc.add_paragraph(
                    "Проверки работы очистных сооружений проводятся с "
                    f"периодичностью {FREQ_TREATMENT} (п. 9.2.4 Требований).")
                sw = [x for x in (pek.get("surface_water") or [])
                      if isinstance(x, dict)]
                if sw:
                    doc.add_paragraph("Программа наблюдений за водным объектом "
                                      "(фоновый и контрольный створы):")
                    _table(doc, ["Водный объект", "Створ / место",
                                 "Показатель", "Периодичность"],
                           [[x.get("water_body") or "—", x.get("location") or "—",
                             x.get("substance") or "—", x.get("period") or "—"]
                            for x in sw])
            else:
                doc.add_paragraph("Сбросы отсутствуют — контроль в области "
                                  "охраны водных объектов не ведётся.")

            # 9.3 отходы
            p93 = doc.add_paragraph()
            p93.add_run(f"{i}.3. Производственный контроль в области "
                        "обращения с отходами").bold = True
            if ctx.wastes or ctx.waste_acts:
                doc.add_paragraph(
                    "Учёт отходов ведётся по приказу Минприроды России от "
                    "08.12.2020 № 1028; данные учёта обобщаются по итогам "
                    "квартала и календарного года (п. 9.3 Требований). "
                    "Контроль мест накопления — визуальный, при каждом "
                    "обращении с отходами; предельный срок накопления — "
                    "11 месяцев.")
                if pek.get("oro"):
                    doc.add_paragraph(
                        "Мониторинг состояния и загрязнения окружающей среды "
                        "на объектах размещения отходов осуществляется по "
                        "программе, утверждённой в соответствии с приказом "
                        "Минприроды России от 08.12.2020 № 1030.")
                else:
                    doc.add_paragraph(
                        "Собственные объекты размещения отходов не "
                        "эксплуатируются — программа мониторинга ОРО "
                        "(приказ № 1030) не разрабатывается.")
                pts = _legacy_points(ctx, "отход")
                if pts:
                    _table(doc, ["Место контроля", "Показатели", "Периодичность"],
                           [[str(p.get("point", "")), str(p.get("indicators", "")),
                             str(p.get("frequency", ""))] for p in pts])
            else:
                doc.add_paragraph("Отходы не заведены — контроль в области "
                                  "обращения с отходами не описывается.")

            # 9.4–9.5 ППП и искусственные грунты
            p94 = doc.add_paragraph()
            p94.add_run(f"{i}.4. Производственный контроль в области обращения "
                        "с побочными продуктами производства").bold = True
            doc.add_paragraph(
                "Порядок учёта побочных продуктов производства"
                + (": ведётся по перечню раздела 5."
                   if (ctx.extra or {}).get("ppp")
                   else " не устанавливается — побочные продукты не образуются."))
            p95 = doc.add_paragraph()
            p95.add_run(f"{i}.5. Производственный контроль в области обращения "
                        "с искусственными грунтами").bold = True
            doc.add_paragraph(
                "Контроль производства искусственных грунтов из органической "
                "части ТКО"
                + (": ведётся по перечню раздела 6."
                   if (ctx.extra or {}).get("artificial_soil")
                   else " не осуществляется — грунты не производятся."))

        _marks(doc, gap.get(key, []))
        doc.add_paragraph()

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
