"""ТУ — письмо о технических условиях на приём/использование отходов и грунтов.

Типовой документ переписки: организация просит у оператора (полигона, объекта
рекультивации, завода) технические условия на приём конкретного отхода.
Содержательная часть — реквизиты сторон, отход (ФККО, класс, объём) и
подтверждающие документы; остальное — стандартные формулировки письма.
"""
from __future__ import annotations

from datetime import date
from decimal import Decimal
from pathlib import Path

from ecodoc.core.models import ReportContext

TITLE = "Запрос технических условий на приём отходов"


def _num(value) -> str:
    if value in (None, ""):
        return "—"
    try:
        return f"{Decimal(str(value).replace(',', '.')).normalize():f}"
    except Exception:
        return str(value)


def generate(ctx: ReportContext, out_path: str | Path,
             receiver: str = "", purpose: str = "") -> Path:
    """Письмо-запрос ТУ (.docx) по отходам объекта."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Cm, Pt

    from ecodoc.development.waste_inventory import collect

    org = ctx.organization
    rows = collect(ctx)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(3), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(2), Cm(2)

    head = doc.add_paragraph()
    head.alignment = AL.RIGHT
    head.add_run(receiver or "Руководителю организации-оператора\n"
                             "(наименование, адрес)").bold = True

    from_org = doc.add_paragraph()
    from_org.alignment = AL.LEFT
    from_org.add_run(f"{org.name or org.short_name}\n"
                     f"ИНН {org.inn or '—'}, ОГРН {org.ogrn or '—'}\n"
                     f"{org.address or ''}\n"
                     f"тел. {org.phone or '—'}, {org.email or ''}")

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = AL.CENTER
    title.add_run(TITLE).bold = True

    obj = ", ".join(o.code for o in ctx.objects) or "—"
    addr = "; ".join(o.address for o in ctx.objects if o.address) or (org.address or "—")
    doc.add_paragraph(
        f"Просим выдать технические условия на приём отходов, образующихся на "
        f"объекте НВОС {obj} по адресу: {addr}"
        + (f", для {purpose}." if purpose else "."))

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, text in enumerate(["№", "Наименование отхода", "Код ФККО",
                              "Класс опасности", "Объём, т/год"]):
        table.rows[0].cells[i].text = text
    for i, r in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = r["name"] or "—"
        cells[2].text = r["fkko"] or "—"
        cells[3].text = str(r["hazard"] or "—")
        cells[4].text = _num(r["generated"])

    doc.add_paragraph()
    docs_line = "Приложения: паспорта отходов, протоколы КХА/биотестирования " \
                "(при наличии), свидетельство о постановке объекта на учёт."
    doc.add_paragraph(docs_line)
    doc.add_paragraph(
        "Просим сообщить условия приёма, требования к транспортированию и "
        "оформлению сопроводительных документов, а также реквизиты лицензии на "
        "деятельность по обращению с отходами.")

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.add_run(f"{org.director_position or 'Руководитель'}\t\t"
                 f"_____________\t{org.director_name or ''}")
    doc.add_paragraph(f"«___» __________ {ctx.period.year or date.today().year} г.")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out
