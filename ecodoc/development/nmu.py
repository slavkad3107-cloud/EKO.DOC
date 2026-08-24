"""План мероприятий по снижению выбросов ЗВ в атмосферный воздух в периоды НМУ.

Действующая нормативная база (проверена на август 2026):
  * пп. 2 п. 12 ст. 19 Федерального закона от 04.05.1999 № 96-ФЗ «Об охране
    атмосферного воздуха» — обязанность разработки Плана мероприятий;
  * приказ Минприроды России от 28.11.2025 № 662 «Об утверждении требований
    к содержанию, составу, форме, порядку разработки, согласования
    и утверждения плана мероприятий по снижению выбросов загрязняющих веществ
    в атмосферный воздух в периоды неблагоприятных метеорологических условий»
    (Минюст 28.11.2025 № 84375; действует с 01.03.2026 до 01.03.2032).
    Форма — рекомендуемый образец-приложение к № 662: РОВНО 9 пунктов
    (пп. 1–6 — лицо, объект ОНВОС, вид прогноза; п. 7 — таблица на 6 граф
    со строкой номеров граф «1…6»; п. 8 — результаты расчётов рассеивания
    (приказ № 273, п. 35 Методики № 581); п. 9 — метод контроля (п. 18
    Порядка инвентаризации № 871)). Названия пунктов и заголовки граф
    взяты из образца дословно (константы ниже) — чтобы согласующий орган
    видел знакомую форму;
  * приказ Минприроды России от 26.11.2025 № 651 (Минюст 28.11.2025 № 84372;
    01.03.2026–01.03.2032) — требования к самим мероприятиям:
      - п. 3: мероприятия разрабатываются только для КОНТРОЛИРУЕМЫХ веществ —
        веществ из перечня Расп. Правительства РФ от 20.10.2023 № 2909-р,
        расчётные приземные концентрации которых при увеличении на проценты
        пп. 10/12/13 превысят ПДК в контрольных точках; определяются анализом
        расчётов рассеивания, там же считаются вклады источников (%);
      - п. 4: только на источниках, содержащих контролируемые вещества;
      - п. 5: если источник с непрерывным циклом снизить нельзя — снижение
        обеспечивается на других источниках объекта;
      - пп. 10/12/13: проценты (общий прогноз 20/15; степени НМУ 15/20/40,
        для регулируемых видов деятельности 5/10/20) — это снижение ВКЛАДОВ
        в приземные концентрации контролируемых веществ в контрольной точке,
        а НЕ доля выброса г/с. Поэтому графа 6 «выброс после мероприятия»
        НИКОГДА не вычисляется по этим процентам;
      - п. 8: ответственные — в каждом структурном подразделении;
      - п. 9: перечень запретов при НМУ (в документ — дословно);
  * приказ Минприроды России от 26.11.2025 № 652 — специализированные
    прогнозы НМУ (для объектов I категории, кроме регулируемых видов).

Приказ Минприроды России от 28.11.2019 № 811 УТРАТИЛ СИЛУ с 01.03.2026
(старые «режимы 1/2/3» и «продувка и чистка» — история); ссылаться на него
в новых планах нельзя.

Откуда данные (ничего не выдумываем):
  * пп. 1–5 формы — ctx.organization и ctx.objects (код, категория, адрес);
  * источники и их вещества (г/с) — ctx.extra['emission_sources'];
  * всё, что знает только эколог/технолог, — ctx.extra['nmu']:
      forecast_kind:   «общий» | «специализированный» (п. 6 формы);
      regulated:       True для регулируемых видов деятельности (ТЭК/ЖКХ);
      controlled_substances: [код, …] — контролируемые вещества по расчёту
                       рассеивания (п. 3 № 651). Без них таблица п. 7
                       не строится — печатается «[требуется: …]»;
      control_points:  [{name, code, sources: [{number, contribution_pct}]}]
                       — контрольные точки и вклады источников, % (п. 3);
      measures:        [{mode: 1|2|3, text, source, substance,
                         reduction_pct, after}] — мероприятия технолога;
                       графа 6 берётся ТОЛЬКО из after (г/с после) или
                       reduction_pct (снижение выброса именно этим
                       мероприятием, задано технологом); иначе — пометка;
      dispersion:      текст о выполненных расчётах рассеивания (п. 8 формы);
      control_method:  «инструментальный» | «расчетный» (п. 9 формы);
      responsible:     строка «должность ФИО» или список
                       [{unit, position, name}] по подразделениям (п. 8 № 651);
      no_measures_required: True — по п. 5 требований № 662 (превышений ПДК
                       в периоды НМУ по расчётам нет, мероприятия не нужны).
Если своих мероприятий нет — по каждому источнику с контролируемыми
веществами подставляются типовые формулировки (по виду источника), это явно
помечается, а графа 6 остаётся «[требуется: …]». Всё, что машина взять не
может, помечено «[требуется: …]», и тот же список возвращает gaps(ctx).
Пояснения (требование о снижении вкладов, запреты п. 9 № 651,
ответственные, журнал) печатаются ПОСЛЕ формы как пояснительная записка —
без продолжения нумерации образца.

Региональные варианты (ecodoc.development.nmu_regions, сверены с ПРИНЯТЫМИ
планами из Формы/Разработка/НМУ): регион объекта 77 — Москва (таблица п. 7
на 8 граф как в плане, согласованном ДПиООС 24.04.2026; приложения по
Порядку 231-ПП — пояснительная записка п. 17, журнал прил. 1, план-график
контроля прил. 3, состав заявления п. 16), 78 — Санкт-Петербург (таблица
на 9 граф с «экологическим эффектом, %» и блок «Должностное лицо,
ответственное…», как в принятом перечне; согласует Комитет по
природопользованию), иначе — федеральный образец № 662. Ядро (пп. 1–9
формы № 662) одинаково во всех вариантах. Дополнительно из extra['nmu']:
      below_0_1_pdk:   True — приземные концентрации при штатном режиме
                       < 0,1 ПДК на границе СЗЗ/жилой застройки (п. 14/17
                       Порядка 231-ПП — достаточно мероприятий общего
                       характера, пояснительная записка из 3 разделов);
      processes:       характеристика производственных процессов (текст
                       для пояснительной записки, п. 17 Порядка 231-ПП);
      control_schedule:[{source, unit, substance, periodicity, method,
                       norm, org}] — план-график контроля (прил. 3 231-ПП).
"""
from __future__ import annotations

from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

from ecodoc.core.models import Medium, ReportContext

TITLE = ("План мероприятий по снижению выбросов загрязняющих веществ "
         "в атмосферный воздух в периоды неблагоприятных метеорологических "
         "условий")

NPA = ("приказ Минприроды России от 26.11.2025 № 651, приказ Минприроды "
       "России от 28.11.2025 № 662; пп. 2 п. 12 ст. 19 Федерального закона "
       "от 04.05.1999 № 96-ФЗ «Об охране атмосферного воздуха»")

# Названия пунктов формы — дословно из рекомендуемого образца № 662
# (почему константы: согласующий орган сверяет с образцом построчно).
P1 = ("1. Наименование юридического лица или фамилия, имя, отчество "
      "(при наличии) индивидуального предпринимателя, осуществляющего "
      "хозяйственную и (или) иную деятельность")
P2 = ("2. Наименование объекта, оказывающего негативное воздействие "
      "на окружающую среду")
P3 = ("3. Сведения о фактическом месте нахождения объекта, оказывающего "
      "негативное воздействие на окружающую среду")
P4 = ("4. Категория объекта, оказывающего негативное воздействие "
      "на окружающую среду")
P5 = "5. Код объекта, оказывающего негативное воздействие на окружающую среду"
P6 = ("6. Вид прогноза неблагоприятных метеорологических условий (далее - "
      "НМУ), по которому работает объект, оказывающий негативное воздействие "
      "на окружающую среду (общий или специализированный)")
P7 = "7. Перечень мероприятий по снижению выбросов в периоды НМУ:"
P8 = ("8. Результаты расчетов рассеивания выбросов, выполненных в соответствии "
      "с методами расчетов рассеивания выбросов вредных (загрязняющих) "
      "веществ в атмосферном воздухе, утвержденных приказом Минприроды "
      "России 6 июня 2017 г. N 273 (далее - расчеты рассеивания), при "
      "проведении мероприятий по снижению выбросов в периоды НМУ, "
      "обосновывающие эффективность мероприятий при НМУ, включенных в план "
      "мероприятий по снижению выбросов загрязняющих веществ в атмосферный "
      "воздух в периоды неблагоприятных метеорологических условий "
      "с приложением документарного подтверждения проведенных расчетов "
      "рассеивания, в соответствии с требованиями пункта 35 Методики "
      "разработки (расчета) и установления нормативов допустимых выбросов "
      "загрязняющих веществ в атмосферный воздух, утвержденной приказом "
      "Минприроды России от 11 августа 2020 г. N 581")
P9 = ("9. Информация о методе контроля (инструментальный или расчетный), "
      "определенном при проведении инвентаризации стационарных источников "
      "и выбросов в соответствии с пунктом 18 Порядка проведения "
      "инвентаризации стационарных источников и выбросов загрязняющих "
      "веществ в атмосферный воздух, корректировки ее данных, "
      "документирования и хранения данных, полученных в результате "
      "проведения таких инвентаризации и корректировки, утвержденного "
      "приказом Минприроды России от 19 ноября 2021 г. N 871")

# Заголовки граф таблицы п. 7 — дословно из образца № 662
HEADER = ["N п/п",
          "Номер источника (источников) выбросов загрязняющих веществ "
          "в атмосферный воздух",
          "Наименование мероприятия по уменьшению выбросов загрязняющих "
          "веществ в атмосферный воздух в периоды НМУ (далее - мероприятие)",
          "Наименование загрязняющего вещества",
          "Величины выбросов до мероприятия г/с",
          "Величины выбросов после мероприятия г/с"]

# Пометка для графы 6, когда технолог не задал величину «после»
AFTER_REQUIRED = ("[требуется: величина выброса после мероприятия "
                  "по данным технолога]")
# Пометка вместо таблицы п. 7, когда нет перечня контролируемых веществ
CONTROLLED_REQUIRED = ("требуется: перечень контролируемых веществ и вклады "
                       "источников по расчёту рассеивания (пп. 3–5 приказа "
                       "№ 651)")

# Подсказка, пока в Формы/Разработка/НМУ нет образца пользователя
SAMPLE_HINT = ("подсказка: положите актуальный бланк или принятый план по "
               "приказу № 662 в папку Формы/Разработка/НМУ (сейчас пуста) — "
               "сверка форм будет идти по нему")


def _user_sample_exists() -> bool:
    """Есть ли у пользователя свой образец плана НМУ в папке бланков."""
    try:
        from ecodoc.core.forms_registry import sample_for
        return sample_for("nmu") is not None
    except Exception:   # реестр бланков недоступен — считаем, что образца нет
        return False


# Запреты при НМУ — ДОСЛОВНО п. 9 требований № 651
PROHIBITIONS_651 = (
    "При проведении мероприятий в периоды НМУ на объектах ОНВОС соблюдаются "
    "технологические регламенты работ всех производств, оборудования "
    "и установок, а также запрещаются остановки газопылеулавливающих "
    "сооружений для выполнения профилактических работ, залповые выбросы "
    "вредных веществ в атмосферный воздух (кроме случаев, когда уже "
    "проводятся технологические операции по подготовке к проведению "
    "залповых выбросов), запрещается проведение пусконаладочных работ "
    "и испытаний оборудования (п. 9 требований, утв. приказом Минприроды "
    "России от 26.11.2025 № 651).")

# Степени НМУ (специализированный прогноз, приказ № 652);
# ключ 0 — общий прогноз НМУ (без градации по степеням).
MODES = {
    0: "Мероприятия при поступлении общего прогноза НМУ",
    1: "Мероприятия при НМУ 1 степени",
    2: "Мероприятия при НМУ 2 степени",
    3: "Мероприятия при НМУ 3 степени",
}

# Снижение ВКЛАДОВ в приземные концентрации в контрольной точке, %
# (пп. 10/12/13 приказа № 651): {mode: (прочие, регулируемые виды деятельности)}.
# К величинам г/с в графе 6 НЕ применяется — см. docstring модуля.
TARGETS: dict[int, tuple[int, int]] = {
    0: (20, 15),
    1: (15, 5),
    2: (20, 10),
    3: (40, 20),
}

# типовые формулировки мероприятий по виду источника: (ключевые слова, тексты
# по степеням 1/2/3). Первая степень — организационно-технические без снижения
# производительности; вторая — с частичным снижением нагрузки; третья —
# частичная остановка производств, не имеющих непрерывного цикла.
# «Продувка и чистка» из № 811 убраны — в № 651 такого запрета нет.
_TYPICAL: list[tuple[tuple[str, ...], dict[int, str]]] = [
    (("котел", "котель", "печь", "тэц", "теплогенератор", "сушил"), {
        1: "усиление контроля топочного режима и герметичности газоходов",
        2: "перевод котлоагрегатов на резервное газообразное/малосернистое "
           "топливо, снижение нагрузки котлоагрегатов",
        3: "снижение нагрузки котлоагрегатов до технологического минимума, "
           "остановка резервных котлоагрегатов"}),
    (("сварк", "свароч", "резк", "металлообраб"), {
        1: "усиление контроля режимов сварки, работа только на исправном "
           "оборудовании с местной вытяжкой",
        2: "ограничение объёма сварочных и газорезательных работ",
        3: "запрет сварочных и газорезательных работ, кроме аварийных"}),
    (("окрас", "покрас", "лакокрас", "грунтов"), {
        1: "запрет окрасочных работ на открытых площадках",
        2: "ограничение окрасочных работ, перенос на период после отмены НМУ",
        3: "полное прекращение окрасочных работ"}),
    (("пыл", "дроб", "пересып", "сыпуч", "щебен", "цемент", "бетон",
      "грунт", "склад"), {
        1: "интенсификация орошения (гидрообеспыливания) мест пересыпки "
           "и открытых складов пылящих материалов",
        2: "ограничение погрузочно-разгрузочных работ с пылящими материалами",
        3: "прекращение погрузочно-разгрузочных работ с пылящими материалами"}),
    (("транспорт", "автотранспорт", "стоянк", "двигател", "техник",
      "дорог", "дэс", "дизель"), {
        1: "запрет работы двигателей автотранспорта и спецтехники "
           "на холостом ходу",
        2: "ограничение движения и работы автотранспорта и спецтехники "
           "на территории объекта",
        3: "остановка работы техники, не занятой в непрерывном цикле"}),
]

_TYPICAL_DEFAULT = {
    1: "усиление контроля за соблюдением технологического регламента "
       "и эффективностью газоочистного оборудования, запрет пусконаладочных "
       "работ и работ с залповыми выбросами",
    2: "снижение производительности источника с соответствующим уменьшением "
       "выброса загрязняющих веществ",
    3: "частичная остановка источника (производства, не имеющего "
       "непрерывного цикла), снижение нагрузки до технологического минимума",
}


# ── числа ────────────────────────────────────────────────────────────────────

def _dec(value) -> Decimal | None:
    if value in (None, ""):
        return None
    try:
        return Decimal(str(value).replace(",", "."))
    except Exception:
        return None


def _fmt(value: Decimal | None) -> str:
    """Число как в принятых планах: запятая и 7 знаков («0,0640000»), ноль — «0»."""
    if value is None:
        return "—"
    if not value:
        return "0"
    text = f"{value.quantize(Decimal('0.0000001'), rounding=ROUND_HALF_UP):f}"
    return text.replace(".", ",")


def _after(before: Decimal | None, pct: Decimal | None) -> Decimal | None:
    """Выброс после мероприятия по заданному ТЕХНОЛОГОМ проценту снижения
    выброса этим мероприятием (не по процентам приказа № 651)."""
    if before is None or pct is None:
        return None
    return before * (Decimal(100) - pct) / Decimal(100)


def _dot(text: str) -> str:
    """Точка в конце, если её ещё нет (чтобы не было «Иванов И.И..»)."""
    text = text.rstrip()
    return text if text.endswith(".") else text + "."


# ── сбор данных ──────────────────────────────────────────────────────────────

def _cfg(ctx: ReportContext) -> dict:
    extra = ctx.extra if isinstance(ctx.extra, dict) else {}
    cfg = extra.get("nmu")
    return cfg if isinstance(cfg, dict) else {}


def sources(ctx: ReportContext) -> list[dict]:
    """Источники выбросов — та же выборка, что в инвентаризации (№ 871)."""
    from ecodoc.development.air_inventory import sources as air_sources
    return air_sources(ctx)


def substances(ctx: ReportContext) -> list[dict]:
    """Сводный перечень веществ: {code, name, g_s, t_year}.

    Суммы г/с и т/год — по источникам из extra['emission_sources'];
    валовые т/год дополняются из ctx.pollutants (medium=air), если по
    источникам масса не набралась. Вещества, известные только из
    ctx.pollutants, тоже попадают в перечень (без г/с).
    """
    agg: dict[str, dict] = {}

    def slot(code: str, name: str) -> dict:
        key = code or name.lower()
        return agg.setdefault(key, {"code": code, "name": name,
                                    "g_s": None, "t_year": None})

    for s in sources(ctx):
        for p in s["pollutants"]:
            code = str(p.get("code") or "").strip()
            name = str(p.get("name") or "").strip()
            if not (code or name):
                continue
            r = slot(code, name)
            r["name"] = r["name"] or name
            for src_key, dst_key in (("g_s", "g_s"), ("t_year", "t_year")):
                d = _dec(p.get(src_key))
                if d is not None:
                    r[dst_key] = (r[dst_key] or Decimal(0)) + d

    for p in ctx.pollutants:
        if p.medium != Medium.AIR:
            continue
        total = ((_dec(p.mass_norm) or Decimal(0))
                 + (_dec(p.mass_limit) or Decimal(0))
                 + (_dec(p.mass_over) or Decimal(0)))
        r = slot(str(p.code or "").strip(), str(p.name or "").strip())
        r["name"] = r["name"] or p.name
        if r["t_year"] is None and total:
            r["t_year"] = total
    return list(agg.values())


def controlled_substances(ctx: ReportContext) -> set[str]:
    """Коды контролируемых веществ (п. 3 № 651) из extra['nmu'].

    Принимаем список кодов/названий или список словарей {code, name};
    сравнение — по коду либо по названию в нижнем регистре.
    """
    out: set[str] = set()
    for item in _cfg(ctx).get("controlled_substances", []) or []:
        if isinstance(item, dict):
            for key in ("code", "name"):
                v = str(item.get(key) or "").strip()
                if v:
                    out.add(v.lower())
        else:
            v = str(item or "").strip()
            if v:
                out.add(v.lower())
    return out


def _is_controlled(p: dict, controlled: set[str]) -> bool:
    code = str(p.get("code") or "").strip().lower()
    name = str(p.get("name") or "").strip().lower()
    return bool((code and code in controlled) or (name and name in controlled))


def controlled_sources(ctx: ReportContext) -> list[dict]:
    """Источники, в выбросах которых есть контролируемые вещества (п. 4 № 651);
    у каждого оставлены только контролируемые вещества."""
    controlled = controlled_substances(ctx)
    if not controlled:
        return []
    out = []
    for s in sources(ctx):
        polls = [p for p in s["pollutants"] if _is_controlled(p, controlled)]
        if polls:
            out.append({**s, "pollutants": polls})
    return out


def control_points(ctx: ReportContext) -> list[dict]:
    """Контрольные точки и вклады источников (п. 3 № 651):
    [{name, code, sources: [{number, contribution_pct}]}]."""
    out = []
    for cp in _cfg(ctx).get("control_points", []) or []:
        if isinstance(cp, dict):
            out.append(cp)
    return out


def forecast_kind(ctx: ReportContext) -> str:
    """Вид прогноза НМУ (п. 6 формы): из extra['nmu'] или по категории.

    Объекты I категории (кроме регулируемых видов деятельности) обязаны
    получать специализированный прогноз (приказ № 652); II–III категории
    работают по общему прогнозу. Пустая строка — вид не определить.
    """
    kind = str(_cfg(ctx).get("forecast_kind") or "").strip().lower()
    if kind in ("общий", "специализированный"):
        return kind
    cats = {str(o.category or "").strip().upper() for o in ctx.objects}
    if "I" in cats and not _cfg(ctx).get("regulated"):
        return "специализированный"
    if cats & {"II", "III"}:
        return "общий"
    return ""


def plan_modes(ctx: ReportContext) -> list[int]:
    """Какие таблицы мероприятий входят в План: [0] — общий прогноз,
    [1, 2, 3] — специализированный (и по умолчанию, пока вид не указан:
    полный вариант, чтобы ничего не потерять)."""
    return [0] if forecast_kind(ctx) == "общий" else [1, 2, 3]


def target_pct(ctx: ReportContext, mode: int) -> int:
    """Требуемое снижение вкладов в приземные концентрации в контрольной
    точке по приказу № 651 для степени/вида прогноза, %."""
    other, regulated = TARGETS[mode]
    return regulated if _cfg(ctx).get("regulated") else other


def _typical_text(source_name: str, mode: int) -> str:
    low = (source_name or "").lower()
    step = 1 if mode == 0 else mode   # общий прогноз — орг.-тех. комплекс
    for keys, texts in _TYPICAL:
        if any(k in low for k in keys):
            return texts[step]
    return _TYPICAL_DEFAULT[step]


def user_measures(ctx: ReportContext, mode: int) -> list[dict]:
    """Мероприятия, заданные пользователем для степени (mode 0 → берём 1)."""
    step = 1 if mode == 0 else mode
    out = []
    for m in _cfg(ctx).get("measures", []) or []:
        if isinstance(m, dict) and int(m.get("mode", 0) or 0) == step:
            out.append(m)
    return out


def _after_cell(before: Decimal | None, m: dict | None) -> str:
    """Графа 6: только из данных технолога по мероприятию — явная величина
    after (г/с) либо reduction_pct; нормативный процент № 651 не применяется."""
    if m:
        after = _dec(m.get("after"))
        if after is not None:
            return _fmt(after)
        pct = _dec(m.get("reduction_pct"))
        if pct is not None and before is not None:
            return _fmt(_after(before, pct))
    return AFTER_REQUIRED


def measure_rows(ctx: ReportContext, mode: int) -> tuple[list[list[str]], bool]:
    """Строки таблицы п. 7 (графы образца № 662) для одной степени.

    Только контролируемые вещества и источники с ними (пп. 3–4 № 651).
    Возвращает (строки, типовые_ли): графы — N п/п, номер источника,
    мероприятие, вещество, выброс до г/с, выброс после г/с.
    Пустой список при отсутствии перечня контролируемых веществ.
    """
    srcs = controlled_sources(ctx)
    own = user_measures(ctx, mode)
    rows: list[list[str]] = []
    if not srcs:
        # перечня контролируемых веществ нет — типовые строки по всем
        # источникам не строим (пп. 3–4 № 651); но собственные мероприятия
        # пользователя не теряем: печатаем их без величин г/с (графы 5–6 —
        # только то, что задано явно), рядом остаётся пометка «[требуется…]»
        for n, m in enumerate(own, 1):
            rows.append([str(n),
                         str(m.get("source") or "") or "по объекту в целом",
                         str(m.get("text") or ""),
                         str(m.get("substance") or "") or
                         "[требуется: контролируемое вещество]",
                         "—", _after_cell(None, m)])
        return rows, False
    by_number = {s["number"]: s for s in srcs if s["number"]}

    def label(p: dict) -> str:
        return f"{p.get('code', '')} {p.get('name', '')}".strip() or "—"

    if own:
        n = 0
        for m in own:
            text = str(m.get("text") or "")
            want = str(m.get("substance") or "").strip().lower()
            src = by_number.get(str(m.get("source") or "").strip())
            if src:
                polls = [p for p in src["pollutants"]
                         if not want or _is_controlled(p, {want})]
                for p in polls:
                    n += 1
                    before = _dec(p.get("g_s"))
                    rows.append([str(n), src["number"], text, label(p),
                                 _fmt(before), _after_cell(before, m)])
            else:
                # мероприятие без привязки к источнику — по объекту в целом:
                # «до» = сумма г/с контролируемых веществ (по веществу, если
                # оно указано), «после» — только из данных технолога
                n += 1
                total = None
                for s in srcs:
                    for p in s["pollutants"]:
                        if want and not _is_controlled(p, {want}):
                            continue
                        d = _dec(p.get("g_s"))
                        if d is not None:
                            total = (total or Decimal(0)) + d
                rows.append([str(n),
                             str(m.get("source") or "") or "по объекту в целом",
                             text,
                             str(m.get("substance") or "") or
                             "контролируемые вещества",
                             _fmt(total), _after_cell(total, m)])
        return rows, False

    n = 0
    for s in srcs:
        for p in s["pollutants"]:
            n += 1
            before = _dec(p.get("g_s"))
            rows.append([str(n), s["number"] or "—",
                         _typical_text(s["name"], mode), label(p),
                         _fmt(before), AFTER_REQUIRED])
    return rows, True


def units_by_source(ctx: ReportContext) -> dict[str, str]:
    """Номер источника → «цех. участок» (графа «Структурное подразделение
    (цех)» региональных таблиц). Берём workshop/site инвентаризации;
    если их нет — пусто (дальше печатается пометка, название источника
    цехом не подменяем: в принятых планах это именно подразделение)."""
    out: dict[str, str] = {}
    for s in sources(ctx):
        parts = [str(s.get("workshop") or "").strip(),
                 str(s.get("site") or "").strip()]
        out[s["number"]] = ". ".join(p for p in parts if p)
    return out


def regional_rows(ctx: ReportContext, prof: dict) -> list[list[str]]:
    """Строки таблицы п. 7 в региональной форме (Москва — 8 граф, СПб — 9):
    одна таблица на все степени, степень — отдельной графой (так в обоих
    принятых планах). Основа — measure_rows() по образцу № 662, графы
    «степень», «цех» и «эффект, %» добавляются, ничего не пересчитывая.
    """
    from ecodoc.development.nmu_regions import (UNIT_REQUIRED, effect_pct)
    units = units_by_source(ctx)
    rows: list[list[str]] = []
    n = 0
    for mode in plan_modes(ctx):
        base, _ = measure_rows(ctx, mode)
        degree = (prof["degree_general"] if mode == 0
                  else prof["degree_fmt"].format(n=mode))
        for r in base:
            n += 1
            unit = units.get(r[1], "") or UNIT_REQUIRED
            row = [str(n), degree, unit, r[1], r[2], r[3], r[4], r[5]]
            if prof["effect_col"]:
                row.append(effect_pct(r[4], r[5]) if r[5] != AFTER_REQUIRED
                           else effect_pct("x", "x"))
            rows.append(row)
    return rows


def control_schedule(ctx: ReportContext) -> list[dict]:
    """План-график контроля в периоды НМУ (прил. 3 к 231-ПП) — только
    из extra['nmu']['control_schedule']; периодичность и методику
    машина не знает."""
    return [r for r in (_cfg(ctx).get("control_schedule") or [])
            if isinstance(r, dict)]


def responsible_list(ctx: ReportContext) -> list[dict]:
    """Ответственные по подразделениям (п. 8 № 651): [{unit, position, name}].

    Принимаем строку («должность ФИО» — одно лицо по объекту) или список
    словарей; если ничего нет — руководитель из карточки организации.
    """
    raw = _cfg(ctx).get("responsible")
    org = ctx.organization
    out: list[dict] = []
    if isinstance(raw, (list, tuple)):
        for r in raw:
            if isinstance(r, dict):
                out.append({"unit": str(r.get("unit") or "").strip(),
                            "position": str(r.get("position") or "").strip(),
                            "name": str(r.get("name") or "").strip()})
            elif str(r or "").strip():
                out.append({"unit": "", "position": "", "name": str(r).strip()})
    elif str(raw or "").strip():
        out.append({"unit": "", "position": "", "name": str(raw).strip()})
    if not out and org.director_name:
        out.append({"unit": "", "position": org.official_title or "Руководитель",
                    "name": org.director_name})
    return [r for r in out if r["name"] or r["position"]]


# ── чего не хватает ─────────────────────────────────────────────────────────

def gaps(ctx: ReportContext) -> list[str]:
    """Список того, что машина взять не может, — он же печатается в Плане."""
    out: list[str] = []
    org = ctx.organization
    cfg = _cfg(ctx)
    if not (org.name or org.short_name):
        out.append("требуется: наименование юридического лица / ИП (п. 1 Плана)")
    if not ctx.objects:
        out.append("требуется: объект ОНВОС — наименование, адрес, категория, "
                   "код (пп. 2–5 Плана); заведите объект на вкладке ОБЪЕКТ")
    else:
        for o in ctx.objects:
            label = o.code or o.name or "объект"
            if not o.address:
                out.append(f"требуется: адрес (место нахождения) объекта {label} "
                           "(п. 3 Плана)")
            if not o.category:
                out.append(f"требуется: категория объекта {label} (п. 4 Плана)")
            if not o.code:
                out.append(f"требуется: код объекта ОНВОС «{o.name or '—'}» "
                           "(п. 5 Плана)")
    if not forecast_kind(ctx):
        out.append("требуется: вид прогноза НМУ — общий или специализированный "
                   "(п. 6 Плана; extra['nmu']['forecast_kind'])")
    # образца/принятого плана по № 662 у пользователя пока нет — подсказываем,
    # куда положить, чтобы сверка форм его подхватила
    if not _user_sample_exists():
        out.append(SAMPLE_HINT)
    if cfg.get("no_measures_required"):
        out.append("требуется: подтверждающие материалы об отсутствии "
                   "превышений ПДК в периоды НМУ (п. 5 требований № 662)")
        return out
    srcs = sources(ctx)
    if not srcs:
        out.append("требуется: источники выбросов — загрузите отчёт об "
                   "инвентаризации (приказ № 871), раздел ООС или проект НДВ "
                   "(вкладка ЗАГРУЗКА, категория «воздух»)")
    if not controlled_substances(ctx):
        out.append(CONTROLLED_REQUIRED)
    elif not controlled_sources(ctx):
        out.append("требуется: ни у одного источника нет контролируемых "
                   "веществ из extra['nmu']['controlled_substances'] — "
                   "сверьте коды с инвентаризацией (п. 4 приказа № 651)")
    if controlled_substances(ctx) and not control_points(ctx):
        out.append("требуется: контрольные точки и вклады источников в "
                   "приземные концентрации, % (п. 3 приказа № 651; "
                   "extra['nmu']['control_points'])")
    for s in controlled_sources(ctx):
        if not s["number"]:
            out.append(f"требуется: номер источника «{s['name'] or '—'}» "
                       "(графа 2 перечня мероприятий)")
        for p in s["pollutants"]:
            if _dec(p.get("g_s")) is None:
                out.append("требуется: выброс г/с вещества "
                           f"«{p.get('name') or p.get('code') or '—'}» "
                           f"источника №{s['number'] or '—'} (графа 5)")
    for mode in plan_modes(ctx):
        own = user_measures(ctx, mode)
        if not own:
            out.append(f"требуется: проверка технологом типовых мероприятий — "
                       f"«{MODES[mode]}» (подставлены по виду источников)")
        rows, _ = measure_rows(ctx, mode)
        if any(r[5] == AFTER_REQUIRED for r in rows):
            out.append("требуется: величина выброса после мероприятия по данным "
                       f"технолога — «{MODES[mode]}» (графа 6; measures[].after "
                       "или reduction_pct)")
    if not cfg.get("dispersion"):
        out.append("требуется: результаты расчётов рассеивания (приказ МПР "
                   "от 06.06.2017 № 273, п. 35 Методики № 581), "
                   "обосновывающие эффективность мероприятий (п. 8 Плана)")
    if str(cfg.get("control_method") or "").strip().lower() not in (
            "инструментальный", "расчетный", "расчётный"):
        out.append("требуется: метод контроля — инструментальный или расчётный, "
                   "по п. 18 Порядка инвентаризации № 871 (п. 9 Плана)")
    if not responsible_list(ctx):
        out.append("требуется: ответственные за проведение мероприятий "
                   "по структурным подразделениям (должность, ФИО; "
                   "п. 8 приказа № 651)")
    out += regional_gaps(ctx)
    return out


def regional_gaps(ctx: ReportContext) -> list[str]:
    """Чего не хватает именно для региональной формы (Москва/СПб)."""
    from ecodoc.development import nmu_regions as R
    prof = R.profile(ctx)
    out: list[str] = []
    if not prof or _cfg(ctx).get("no_measures_required"):
        return out
    units = units_by_source(ctx)
    missing = sorted({s["number"] or "—" for s in controlled_sources(ctx)
                      if not units.get(s["number"])})
    if missing:
        out.append("требуется: структурное подразделение (цех) для "
                   f"источников {', '.join(missing)} (графа 3 формы "
                   f"{prof['name']}; workshop/site в инвентаризации)")
    if prof["effect_col"] and any(
            r[-1] == R.EFFECT_REQUIRED for r in regional_rows(ctx, prof)):
        out.append("требуется: достигаемый экологический эффект, % "
                   "(графа 9 перечня СПб) — считается из величины после "
                   "мероприятия, задайте measures[].after/reduction_pct")
    if prof["key"] == "msk":
        cfg = _cfg(ctx)
        if "below_0_1_pdk" not in cfg:
            out.append("требуется: признак «приземные концентрации при "
                       "штатном режиме < 0,1 ПДК на границе СЗЗ/жилой "
                       "застройки» (п. 14/17 Порядка 231-ПП; "
                       "extra['nmu']['below_0_1_pdk'])")
        if not cfg.get("processes"):
            out.append("требуется: характеристика производственных "
                       "процессов для пояснительной записки (п. 17 "
                       "Порядка 231-ПП; extra['nmu']['processes'])")
        if not control_schedule(ctx):
            out.append("требуется: план-график контроля выбросов в периоды "
                       "НМУ — периодичность, методика, организация "
                       "(приложение 3 к Порядку 231-ПП; "
                       "extra['nmu']['control_schedule'])")
    return out


# ── документ ────────────────────────────────────────────────────────────────

def generate(ctx: ReportContext, out_path: str | Path) -> Path:
    """Собрать План мероприятий НМУ (.docx) по форме приказа № 662."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Cm, Pt

    from docx.enum.section import WD_ORIENT
    from ecodoc.development import nmu_regions as R

    org = ctx.organization
    cfg = _cfg(ctx)
    kind = forecast_kind(ctx)
    modes = plan_modes(ctx)
    problems = gaps(ctx)
    prof = R.profile(ctx)          # None — федеральный образец № 662

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(3), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(2), Cm(2)
        if prof:
            # почему альбомная: обе региональные таблицы (8–9 граф) в
            # принятых планах напечатаны на альбомном листе
            s.orientation = WD_ORIENT.LANDSCAPE
            s.page_width, s.page_height = s.page_height, s.page_width
            s.left_margin, s.right_margin = Cm(2), Cm(1.5)

    # грифы по рекомендуемому образцу № 662: слева «Утверждено», справа
    # «Согласовано» — штамп уполномоченного органа в правом верхнем углу
    # (п. 8 требований № 662); для Москвы/СПб орган известен по имени
    authority = (prof["authority"] if prof else
                 "[требуется: уполномоченный орган субъекта РФ — "
                 "региональный экологический надзор]")
    grif = doc.add_table(rows=1, cols=2)
    grif.rows[0].cells[0].text = ("УТВЕРЖДЕНО\n"
                                  f"{org.official_title or 'Руководитель'}\n"
                                  f"____________ {org.director_name or '[ФИО]'}\n"
                                  "«___» __________ 20__ г.\nМесто для печати")
    grif.rows[0].cells[1].text = (f"СОГЛАСОВАНО\n{authority}\n"
                                  "____________ /______________/\n"
                                  "«___» ______ 20__ г.\n"
                                  "Место для печати")

    doc.add_paragraph()
    h = doc.add_paragraph()
    h.alignment = AL.CENTER
    run = h.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(14)
    base = doc.add_paragraph()
    base.alignment = AL.CENTER
    base.add_run(f"Разработан в соответствии с: {NPA}.").italic = True

    # ── пп. 1–6 формы (названия — дословно из образца) ──────────────────
    obj = ctx.objects[0] if ctx.objects else None
    doc.add_paragraph(_dot(
        f"{P1}: {org.name or org.short_name or '[требуется: наименование]'}"
        + (f", ИНН {org.inn}" if org.inn else "")
        + (f", ОГРН {org.ogrn}" if org.ogrn else "")))
    doc.add_paragraph(_dot(
        f"{P2}: {obj.name if obj and obj.name else '[требуется]'}"))
    doc.add_paragraph(_dot(
        f"{P3}: {obj.address if obj and obj.address else '[требуется: адрес]'}"))
    doc.add_paragraph(_dot(
        f"{P4}: {obj.category if obj and obj.category else '[требуется]'}"))
    doc.add_paragraph(_dot(
        f"{P5}: {obj.code if obj and obj.code else '[требуется: код ОНВОС]'}"))
    # п. 3 е) требований № 662: при специализированном прогнозе указывается
    # степень НМУ
    kind_text = kind or "[требуется: общий или специализированный]"
    if kind == "специализированный":
        kind_text += " (степени НМУ: 1, 2, 3)"
    doc.add_paragraph(_dot(f"{P6}: {kind_text}"))

    # ── п. 7: перечень мероприятий ──────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph().add_run(P7).bold = True
    if cfg.get("no_measures_required"):
        doc.add_paragraph(
            "По результатам инвентаризации стационарных источников и выбросов "
            "и расчётов рассеивания в выбросах объекта ОНВОС отсутствуют "
            "загрязняющие вещества, концентрация которых в периоды НМУ будет "
            "превышать предельно допустимые концентрации. Разработка "
            "мероприятий по снижению выбросов в периоды НМУ не требуется "
            "(п. 5 требований, утв. приказом № 662). "
            "[требуется: подтверждающие материалы об отсутствии превышений "
            "ПДК в периоды НМУ (п. 5 требований № 662)]")
    elif prof:
        # региональная форма: одна таблица на все степени, графы «степень»
        # и «цех» (+ «эффект, %» для СПб) — дословно как в принятых планах.
        # Регион проверяется РАНЬШЕ перечня контролируемых веществ: у
        # реального объекта перечня часто ещё нет, но форма Комитета/ДПиООС
        # от этого не меняется (поймано на базе Техностроя — петербургский
        # объект получал федеральные 6 граф вместо 9 граф СПб)
        if not controlled_sources(ctx):
            doc.add_paragraph(f"[{CONTROLLED_REQUIRED}]")
        _table(doc, prof["header"], regional_rows(ctx, prof), numbered=True,
               small=True)
        for mode in [m for m in modes if not user_measures(ctx, m)]:
            doc.add_paragraph(
                "Мероприятия подставлены по виду источников — "
                f"[требуется: проверка технологом типовых мероприятий — "
                f"«{MODES[mode]}» (подставлены по виду источников)]")
    elif not controlled_sources(ctx):
        # федеральная форма без перечня контролируемых веществ: таблицу по
        # всем источникам не строим — это противоречит пп. 3–4 № 651;
        # собственные мероприятия пользователя печатаются без величин г/с
        doc.add_paragraph(f"[{CONTROLLED_REQUIRED}]")
        for mode in modes:
            rows, _ = measure_rows(ctx, mode)
            if not rows and mode != modes[0]:
                continue
            if len(modes) > 1 and rows:
                doc.add_paragraph().add_run(f"{MODES[mode]}:").bold = True
            _table(doc, HEADER, rows, numbered=True)
    else:
        for mode in modes:
            if len(modes) > 1:
                # степень НМУ — отдельной таблицей с простой подписью,
                # без «целевого %» (он относится к вкладам, а не к г/с)
                doc.add_paragraph().add_run(f"{MODES[mode]}:").bold = True
            rows, typical = measure_rows(ctx, mode)
            _table(doc, HEADER, rows, numbered=True)
            if typical and rows:
                doc.add_paragraph(
                    "Мероприятия подставлены по виду источников — "
                    f"[требуется: проверка технологом типовых мероприятий — "
                    f"«{MODES[mode]}» (подставлены по виду источников)]")

    # ── п. 8: расчёты рассеивания ───────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph().add_run(P8 + ":").bold = True
    if cfg.get("dispersion"):
        doc.add_paragraph(str(cfg["dispersion"]))
    else:
        doc.add_paragraph(
            "[требуется: результаты расчётов рассеивания (приказ МПР "
            "от 06.06.2017 № 273, п. 35 Методики № 581), обосновывающие "
            "эффективность мероприятий (п. 8 Плана)]")
    cps = control_points(ctx)
    if cps:
        doc.add_paragraph("Контрольные точки и вклады источников в приземные "
                          "концентрации контролируемых веществ (п. 3 "
                          "требований № 651):")
        cp_rows = []
        for cp in cps:
            srcs_txt = "; ".join(
                f"№{s.get('number', '—')} — {s.get('contribution_pct', '—')} %"
                for s in (cp.get("sources") or []) if isinstance(s, dict))
            cp_rows.append([str(cp.get("name") or "—"),
                            str(cp.get("code") or "—"),
                            srcs_txt or "—"])
        _table(doc, ["Контрольная точка", "Контролируемое вещество",
                     "Вклады источников, % от суммарной концентрации"], cp_rows)

    # ── п. 9: метод контроля ────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph().add_run(P9 + ":").bold = True
    method = str(cfg.get("control_method") or "").strip()
    doc.add_paragraph(_dot(
        method if method else
        "[требуется: метод контроля — инструментальный или расчётный, "
        "по п. 18 Порядка инвентаризации № 871 (п. 9 Плана)]"))

    # ── СПб: блок ответственного лица — сразу под формой, как в принятом
    # перечне («Должностное лицо, ответственное за проведение мероприятий:
    # ФИО, подпись»)
    if prof and prof["key"] == "spb":
        doc.add_paragraph()
        resp_spb = responsible_list(ctx)
        t = doc.add_table(rows=max(1, len(resp_spb)), cols=2)
        t.rows[0].cells[0].text = R.SPB_RESPONSIBLE_LABEL
        for i in range(max(1, len(resp_spb))):
            who = (f"{resp_spb[i]['name']} ______________\nФИО, подпись"
                   if resp_spb else
                   "[требуется: ФИО ответственного] ______________\n"
                   "ФИО, подпись")
            t.rows[i].cells[1].text = who

    # ── пояснительная записка — ВНЕ формы образца № 662 (без нумерации) ──
    doc.add_page_break()
    if prof:
        _regional_note(doc, ctx, prof)
    doc.add_paragraph().add_run(
        "Пояснительная записка к Плану (приложение, в форму образца "
        "приказа № 662 не входит)").bold = True
    if not cfg.get("no_measures_required"):
        doc.add_paragraph().add_run("Требования к мероприятиям").bold = True
        if len(modes) == 1:
            doc.add_paragraph(
                "Мероприятия при поступлении общего прогноза НМУ должны "
                "обеспечивать снижение вкладов в приземные концентрации "
                "контролируемых веществ, создаваемых выбросами источников "
                "объекта ОНВОС, для рассматриваемой контрольной точки не менее "
                f"чем на {target_pct(ctx, 0)} % (п. 10 требований, утв. "
                "приказом Минприроды России от 26.11.2025 № 651).")
        else:
            p_ref = "13" if cfg.get("regulated") else "12"
            doc.add_paragraph(
                "Мероприятия при поступлении специализированного прогноза НМУ "
                "должны обеспечивать снижение вкладов в приземные концентрации "
                "контролируемых веществ, создаваемых выбросами источников "
                "объекта ОНВОС, для рассматриваемой контрольной точки не менее "
                f"чем: на {target_pct(ctx, 1)} % при НМУ 1 степени, "
                f"на {target_pct(ctx, 2)} % при НМУ 2 степени, "
                f"на {target_pct(ctx, 3)} % при НМУ 3 степени "
                f"(п. {p_ref} требований, утв. приказом Минприроды России "
                "от 26.11.2025 № 651).")
        doc.add_paragraph(
            "Достаточность снижения вкладов подтверждается расчётами "
            "рассеивания (п. 8 Плана); величины выбросов после мероприятий "
            "(графа 6 п. 7) приняты по данным технолога для каждого "
            "мероприятия. Если для источника с непрерывным циклом снижение "
            "невозможно, оно обеспечивается на других источниках объекта "
            "(п. 5 требований № 651).")

    doc.add_paragraph().add_run("Организация работ в периоды НМУ").bold = True
    resp = responsible_list(ctx)
    if resp:
        doc.add_paragraph("Ответственные за реализацию мероприятий в периоды "
                          "НМУ (п. 8 требований № 651):")
        _table(doc, ["Структурное подразделение", "Должность", "ФИО"],
               [[r["unit"] or "по объекту в целом", r["position"] or "—",
                 r["name"] or "—"] for r in resp])
    else:
        doc.add_paragraph(
            "[требуется: ответственные за проведение мероприятий по "
            "структурным подразделениям (должность, ФИО; п. 8 приказа № 651)]")
    doc.add_paragraph(
        "Мероприятия вводятся с момента поступления прогноза НМУ и "
        "проводятся в течение всего периода его действия.")
    doc.add_paragraph(PROHIBITIONS_651)
    if prof and prof["key"] == "msk":
        # Москва: журнал строго по приложению 1 к Порядку 231-ПП
        doc.add_paragraph(
            "Поступление информации о НМУ регистрируется ответственным "
            "лицом в журнале по приложению 1 к Порядку 231-ПП (п. 9 "
            "Порядка):")
        doc.add_paragraph().add_run(R.MSK_JOURNAL_TITLE).bold = True
        _table(doc, R.MSK_JOURNAL_HEADER, [[""] * 6], numbered=True,
               small=True)
        doc.add_paragraph("Примечания:")
        for note in R.MSK_JOURNAL_NOTES:
            doc.add_paragraph(note)
    else:
        doc.add_paragraph(
            "Поступление прогнозов и проведение мероприятий фиксируются "
            "в журнале учёта:")
        _table(doc, ["Дата, время получения прогноза",
                     "Вид прогноза / степень НМУ",
                     "Период действия НМУ",
                     "Проведённые мероприятия (N по Плану)",
                     "Начало / окончание работ", "Подпись ответственного"],
               [["", "", "", "", "", ""]])

    # ── порядок согласования: федеральный (№ 662) + региональный орган ──
    doc.add_paragraph().add_run("Порядок согласования и утверждения").bold = True
    doc.add_paragraph(
        "План направляется на согласование в уполномоченный орган субъекта "
        "РФ сопроводительным письмом со способами связи; срок рассмотрения — "
        "15 рабочих дней; при согласовании штамп проставляется в верхнем "
        "правом углу Плана; План утверждается не позднее 3 рабочих дней со "
        "дня получения согласования (пп. 6–12 требований, утв. приказом "
        "Минприроды России от 28.11.2025 № 662).")
    if prof:
        doc.add_paragraph(
            f"Уполномоченный орган для объекта в регионе «{prof['name']}» — "
            f"{prof['authority']}. Региональный акт: {prof['npa']}. "
            f"Подача: {prof['submit']}. Прогноз НМУ: {prof['forecast']}.")
        if prof["key"] == "msk":
            doc.add_paragraph(
                "К заявлению о согласовании прилагаются (п. 16 Порядка "
                "231-ПП):")
            for item in R.MSK_APPLICATION:
                doc.add_paragraph(f"— {item};")

    # ── чего не хватает (тот же список, что gaps) ───────────────────────
    if problems:
        doc.add_paragraph()
        doc.add_paragraph().add_run(
            "Чего не хватает (проверьте перед согласованием — "
            "п. 6 требований № 662: план согласовывает уполномоченный "
            "орган субъекта РФ в течение 15 рабочих дней)").bold = True
        for text in problems:
            doc.add_paragraph(f"— [{text}]")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _regional_note(doc, ctx: ReportContext, prof: dict) -> None:
    """Региональные приложения к Плану: Москва — пояснительная записка
    по п. 17 Порядка 231-ПП и план-график контроля (прил. 3);
    СПб — пояснение о форме перечня. Печатается перед общей пояснительной
    запиской, без нумерации образца № 662."""
    from ecodoc.development import nmu_regions as R
    cfg = _cfg(ctx)
    if prof["key"] != "msk":
        doc.add_paragraph().add_run(
            "Примечание о форме (Санкт-Петербург)").bold = True
        doc.add_paragraph(
            "Таблица п. 7 приведена в форме перечня, принимаемого Комитетом "
            "по природопользованию (графы «Степень опасности НМУ», "
            "«Структурное подразделение (цех)», «Достигаемый экологический "
            "эффект, %» сохранены из принятых перечней); пп. 1–9 и грифы — "
            "по образцу приказа Минприроды России от 28.11.2025 № 662. "
            "Эффект, % рассчитан как снижение выброса г/с по данным "
            "технолога.")
        doc.add_paragraph()
        return

    doc.add_paragraph().add_run(
        "Пояснительная записка к плану мероприятий (п. 17 Порядка, "
        "утв. постановлением Правительства Москвы от 25.04.2017 "
        "№ 231-ПП)").bold = True
    below = cfg.get("below_0_1_pdk")
    if below is None:
        doc.add_paragraph(
            "[требуется: признак «приземные концентрации при штатном "
            "режиме < 0,1 ПДК на границе СЗЗ/жилой застройки» (п. 14/17 "
            "Порядка 231-ПП; extra['nmu']['below_0_1_pdk'])] — ниже "
            "приведён полный состав разделов.")
    sections = R.MSK_NOTE_BELOW_01 if below else R.MSK_NOTE_FULL
    if below:
        doc.add_paragraph(
            "Максимальные приземные концентрации загрязняющих веществ от "
            "источников производственной территории при нормальных "
            "условиях составляют менее 0,1 ПДК на границе нормативной "
            "СЗЗ и ближайшей жилой застройки — достаточно мероприятий "
            "общего характера (п. 14 Порядка 231-ПП).")
    units = units_by_source(ctx)
    for i, title in enumerate(sections, 1):
        doc.add_paragraph().add_run(
            f"{i}. {title[:1].upper()}{title[1:]}").bold = True
        if i == 1:
            doc.add_paragraph(str(cfg.get("processes") or
                                  "[требуется: характеристика "
                                  "производственных процессов для "
                                  "пояснительной записки (п. 17 Порядка "
                                  "231-ПП; extra['nmu']['processes'])]"))
        elif "инвентаризация" in title:
            _table(doc, ["Номер источника", "Наименование",
                         "Цех, участок", "Вещества (код — г/с)"],
                   [[s["number"] or "—", s["name"] or "—",
                     units.get(s["number"]) or "—",
                     "; ".join(f"{p.get('code', '')} — "
                               f"{_fmt(_dec(p.get('g_s')))}"
                               for p in s["pollutants"])]
                    for s in sources(ctx)], small=True)
        elif "план-график" in title:
            doc.add_paragraph("См. план-график контроля ниже.")
        elif "расчеты рассеивания" in title:
            doc.add_paragraph(str(cfg.get("dispersion") or
                                  "[требуется: результаты расчётов "
                                  "рассеивания (п. 8 Плана)]"))
        else:
            doc.add_paragraph("[требуется: обоснование и расчёты "
                              "эффективности по данным технолога "
                              f"(раздел «{title}», п. 17 Порядка 231-ПП)]")

    # план-график контроля — приложение 3 к Порядку (графы дословно)
    doc.add_paragraph()
    doc.add_paragraph().add_run(
        R.MSK_SCHEDULE_TITLE + " (приложение 3 к Порядку 231-ПП)").bold = True
    obj = ctx.objects[0] if ctx.objects else None
    addr = obj.address if obj and obj.address else "[требуется]"
    doc.add_paragraph(f"для {ctx.organization.name or '[требуется]'} "
                      f"по адресу: {addr}")
    rows = [[str(r.get("source") or "—"), str(r.get("unit") or "—"),
             str(r.get("substance") or "—"),
             str(r.get("periodicity") or "—"), str(r.get("method") or "—"),
             str(r.get("norm") or "—"), "", str(r.get("org") or "—")]
            for r in control_schedule(ctx)]
    if not rows:
        rows = [[s["number"] or "—", units.get(s["number"]) or "—",
                 f"{p.get('code', '')} {p.get('name', '')}".strip(),
                 "[требуется]", "[требуется]",
                 f"{_fmt(_dec(p.get('g_s')))} г/с", "", "[требуется]"]
                for s in controlled_sources(ctx) for p in s["pollutants"]]
    _table(doc, R.MSK_SCHEDULE_HEADER, rows, numbered=True, small=True)
    doc.add_paragraph()


def _table(doc, header: list[str], rows: list[list[str]],
           numbered: bool = False, small: bool = False) -> None:
    """Таблица с шапкой; numbered=True — строка номеров граф «1 2 … n»,
    как в образце № 662; small=True — 9 pt для широких региональных
    таблиц (8–9 граф на альбомном листе)."""
    from docx.shared import Pt
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, text in enumerate(header):
        table.rows[0].cells[i].text = text
    if numbered:
        cells = table.add_row().cells
        for i in range(len(header)):
            cells[i].text = str(i + 1)
    for r in rows:
        cells = table.add_row().cells
        for i, text in enumerate(r):
            cells[i].text = text
    if not rows:
        cells = table.add_row().cells
        cells[0].text = ("[требуется: данные не заведены — источники и "
                         "вещества]")
    if small:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.font.size = Pt(9)
