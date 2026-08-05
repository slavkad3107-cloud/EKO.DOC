"""Расчёт класса опасности отхода по критериям Минприроды России.

Действующие критерии — приказ Минприроды от 31.03.2025 № 158 (в силе
с 01.09.2025, заменил приказ № 536 от 04.12.2014; шкала K не менялась).

Компонентный метод: по каждому компоненту отхода известен коэффициент
степени опасности Wi (мг/кг), рассчитывается K = Σ(Ci / Wi), где Ci —
концентрация компонента (мг/кг). Класс — по приложению № 1 к Критериям:
  10^6 ≥ K > 10^4 → I класс (чрезвычайно опасные)
  10^4 ≥ K > 10^3 → II класс (высокоопасные)
  10^3 ≥ K > 10^2 → III класс (умеренно опасные)
  10^2 ≥ K > 10   → IV класс (малоопасные)
  K ≤ 10          → V класс (практически неопасные)
Верхняя граница каждого диапазона включается, нижняя — нет:
K = 10^4 — это II класс, K = 10 — это V.

Wi компонентов берётся из аттестованных источников/протоколов; здесь —
машина расчёта K и класса по заданным (Ci, Wi). Биотестирование (для
подтверждения V класса) в расчёт не входит — это отдельная процедура.
"""
from __future__ import annotations

import math
from dataclasses import dataclass, field


@dataclass
class Component:
    name: str
    ci: float          # концентрация компонента в отходе, мг/кг
    wi: float          # коэффициент степени опасности Wi, мг/кг


@dataclass
class HazardResult:
    k_total: float
    hazard_class: int
    components: list[dict] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


# официальное имя действующих критериев — для документов и сообщений
NPA = "приказ Минприроды России от 31.03.2025 № 158"


def _class_by_k(k: float) -> int:
    # приложение № 1 к Критериям: верхняя граница включается, нижняя — нет
    if k > 1e4:
        return 1                       # 10^6 >= K > 10^4
    if k > 1e3:
        return 2                       # 10^4 >= K > 10^3
    if k > 1e2:
        return 3                       # 10^3 >= K > 10^2
    if k > 10:
        return 4                       # 10^2 >= K > 10
    return 5                           # K <= 10


def calculate(components: list[Component]) -> HazardResult:
    """K = Σ(Ci/Wi) и класс опасности по действующим Критериям (пр. № 158)."""
    res = HazardResult(k_total=0.0, hazard_class=5)
    total_ci = sum(c.ci for c in components)
    if total_ci and abs(total_ci - 1_000_000) / 1_000_000 > 0.05:
        res.warnings.append(
            f"Сумма концентраций компонентов {total_ci:.0f} мг/кг ≠ 1 000 000 "
            f"(100%) — проверьте состав отхода")
    k = 0.0
    for c in components:
        if c.wi <= 0:
            res.warnings.append(f"{c.name}: Wi ≤ 0 — компонент пропущен")
            ki = 0.0
        else:
            ki = c.ci / c.wi
            k += ki
        res.components.append({"name": c.name, "ci": c.ci, "wi": c.wi,
                               "ki": round(ki, 4)})
    res.k_total = k
    res.hazard_class = _class_by_k(k)
    return res


def wi_from_logk(log_k: float) -> float:
    """Wi из унифицированного показателя lg(Wi) (приложения к Критериям):
    Wi = 10^(lg Wi). Хелпер, если известен lg Wi компонента."""
    return math.pow(10, log_k)


def generate(components: list[Component], out_path,
             waste_name: str = "", fkko: str = "", org_name: str = "",
             basis: str = "") -> "Path":
    """Оформить расчёт документом (.docx): такой расчёт прикладывается к
    паспорту отхода и к материалам отнесения отхода к классу опасности.

    Машина считает K и класс; откуда взяты Ci (протоколы КХА) и Wi
    (приложения к Критериям / БДО) — пишет пользователь в поле basis."""
    from pathlib import Path

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Cm, Pt

    r = calculate(components)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(3), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(2), Cm(2)

    if org_name:
        head = doc.add_paragraph()
        head.alignment = AL.CENTER
        head.add_run(org_name).bold = True
    title = doc.add_paragraph()
    title.alignment = AL.CENTER
    run = title.add_run(
        "РАСЧЁТ класса опасности отхода\n"
        f"(критерии — {NPA})")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Отход: {waste_name or '[наименование отхода]'}"
                      + (f", код ФККО {fkko}" if fkko else ""))
    doc.add_paragraph(
        "Метод: компонентный. Показатель степени опасности отхода "
        "K = Σ(Ci / Wi), где Ci — концентрация i-го компонента (мг/кг), "
        "Wi — коэффициент степени опасности компонента (мг/кг).")
    if basis:
        doc.add_paragraph(f"Исходные данные: {basis}")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for i, text in enumerate(["№", "Компонент", "Ci, мг/кг", "Wi, мг/кг",
                              "Ki = Ci/Wi"]):
        table.rows[0].cells[i].text = text
    for i, c in enumerate(r.components, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = c["name"]
        cells[2].text = f"{c['ci']:g}"
        cells[3].text = f"{c['wi']:g}"
        cells[4].text = f"{c['ki']:g}"

    doc.add_paragraph()
    doc.add_paragraph(f"K = Σ(Ci/Wi) = {r.k_total:.4g}")
    bands = ("10⁶ ≥ K > 10⁴ — I класс; 10⁴ ≥ K > 10³ — II; "
             "10³ ≥ K > 10² — III; 10² ≥ K > 10 — IV; K ≤ 10 — V")
    doc.add_paragraph(f"Границы классов (приложение № 1 к Критериям): {bands}.")
    concl = doc.add_paragraph()
    concl.add_run(f"ВЫВОД: отход относится к {r.hazard_class} классу "
                  f"опасности.").bold = True
    if r.hazard_class == 5:
        doc.add_paragraph(
            "Примечание: отнесение к V классу подлежит подтверждению "
            "биотестированием водной вытяжки на двух тест-объектах из разных "
            "систематических групп (раздел III Критериев, пр. № 158).")
    for w in r.warnings:
        doc.add_paragraph(f"⚠ {w}")

    doc.add_paragraph()
    doc.add_paragraph("Расчёт выполнил: ____________________ "
                      "/______________________/")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def report(components: list[Component]) -> str:
    r = calculate(components)
    lines = [f"── Расчёт класса опасности отхода ({NPA}) ──"]
    for c in r.components:
        lines.append(f"  {c['name']}: Ci={c['ci']} мг/кг, Wi={c['wi']} → "
                     f"Ki={c['ki']}")
    lines.append(f"K = Σ(Ci/Wi) = {r.k_total:.4g}")
    lines.append(f"КЛАСС ОПАСНОСТИ: {r.hazard_class}")
    if r.hazard_class == 5:
        lines.append("⚠ V класс требует подтверждения биотестированием.")
    for w in r.warnings:
        lines.append(f"  ⚠ {w}")
    return "\n".join(lines)
