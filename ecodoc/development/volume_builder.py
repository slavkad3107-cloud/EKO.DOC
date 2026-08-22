"""Сборка томов НДВ / НДС / СЗЗ.

Идея автоматизации: сам расчёт рассеивания делает специализированное ПО
(«Эколог»/УПРЗА «Эколог», «Эколог-Шум» для физфакторов) — его ЭкоДок не
заменяет. Но ~70% объёма тома — это типовой текст и таблицы, которые собираются
автоматически:

  • титул, аннотация, сведения о предприятии, разделы-обоснования, выводы —
    генерируются из ReportContext по шаблону;
  • таблицы параметров источников и результаты рассеивания — ИМПОРТИРУЮТСЯ из
    выгрузки «Эколога» (Excel/RTF) одной функцией ingest_*;
  • протоколы КХА/биотестирования и карты изолиний — подгружаются как исходники
    и попадают в приложения (parsers.text_extract уже умеет их читать).

Так инвентаризация источников и реквизиты вводятся ОДИН раз (в ЭкоДок),
выгружаются для «Эколога», а его результаты возвращаются и вшиваются в готовый
том без ручного переписывания таблиц.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from ecodoc.core.models import ReportContext


@dataclass
class VolumeSources:
    """Исходники, приходящие извне ЭкоДок."""
    # таблица параметров источников выбросов (из инвентаризации/«Эколога»)
    sources_table: list[list] = field(default_factory=list)      # [[№, наименование, H, D, ...], ...]
    sources_header: list[str] = field(default_factory=list)
    # результаты расчёта рассеивания (из «Эколога»): макс. концентрации в долях ПДК
    dispersion_table: list[list] = field(default_factory=list)
    dispersion_header: list[str] = field(default_factory=list)
    # расчётные точки (код, X, Y, высота, тип, комментарий) — из «Эколога»
    points_table: list[list] = field(default_factory=list)
    points_header: list[str] = field(default_factory=list)
    # вклады источников в максимальные концентрации (таблица № 5 Методики № 581)
    contrib_table: list[list] = field(default_factory=list)
    contrib_header: list[str] = field(default_factory=list)
    # приложения-исходники (пути к протоколам КХА, картам изолиний)
    appendices: list[str] = field(default_factory=list)


# по каким словам в названии листа/шапке узнаём тип выгрузки «Эколога»:
# книга результатов обычно многолистовая, и пользователь грузит её целиком
_SHEET_KINDS = (
    ("points", ("расчетн", "точк")),       # «Расчетные точки»
    ("points", ("расчётн", "точк")),
    ("contrib", ("вклад",)),               # «Вклады источников»
    ("sources", ("источник",)),            # «Источники», «Параметры источников»
    ("dispersion", ("концентрац",)),       # «Максимальные концентрации»
    ("dispersion", ("результат",)),
)


def _sheet_kind(title: str, header: list[str]) -> str | None:
    text = (title + " " + " ".join(header)).lower()
    for kind, words in _SHEET_KINDS:
        if all(w in text for w in words):
            return kind
    return None


def ingest_ecolog_workbook(path: str | Path, src: "VolumeSources",
                           default: str = "dispersion") -> list[str]:
    """Разобрать многолистовую книгу «Эколога» по типам листов в VolumeSources.

    Почему: результаты рассеивания «Эколог» выгружает одной книгой — точки,
    вклады, концентрации на разных листах. Лист, тип которого не распознан,
    идёт в `default`, если этот слот ещё пуст. Возвращает заметки по листам.
    """
    from openpyxl import load_workbook

    notes = []
    wb = load_workbook(path, data_only=True)
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        header = [str(c) if c is not None else "" for c in rows[0]]
        body = [[("" if c is None else c) for c in r] for r in rows[1:]
                if any(c not in (None, "") for c in r)]
        kind = _sheet_kind(ws.title, header)
        if kind is None:
            if getattr(src, default + "_table"):
                notes.append(f"лист «{ws.title}»: тип не распознан, пропущен")
                continue
            kind = default
        setattr(src, kind + "_header", header)
        setattr(src, kind + "_table", body)
        notes.append(f"лист «{ws.title}» → {kind}: строк {len(body)}")
    return notes


def ingest_excel(path: str | Path, sheet: str | None = None,
                 header_row: int = 1) -> tuple[list[str], list[list]]:
    """Прочитать выгрузку «Эколога» (Excel) в (заголовок, строки).

    «Эколог» умеет экспортировать таблицы источников и результаты в Excel —
    эта функция превращает их в данные для тома.
    """
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    ws = wb[sheet] if sheet else wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return [], []
    header = [str(c) if c is not None else "" for c in rows[header_row - 1]]
    body = [[("" if c is None else c) for c in r] for r in rows[header_row:]]
    return header, body


def collect_appendices(paths: list[str | Path]) -> list[str]:
    """Зафиксировать приложения-исходники (КХА, биотестирование, карты).

    Извлечение текста — через parsers.text_extract (для проверки/реестра);
    в том они идут как приложения по списку.
    """
    return [str(Path(p).name) for p in paths]


_TYPES = {
    "ndv": ("ТОМ НОРМАТИВОВ ДОПУСТИМЫХ ВЫБРОСОВ (НДВ)",
            "выбросов загрязняющих веществ в атмосферный воздух"),
    "nds": ("ТОМ НОРМАТИВОВ ДОПУСТИМЫХ СБРОСОВ (НДС)",
            "сбросов загрязняющих веществ в водные объекты"),
    "szz": ("ПРОЕКТ САНИТАРНО-ЗАЩИТНОЙ ЗОНЫ (СЗЗ)",
            "обоснования размера санитарно-защитной зоны"),
}


def gaps(vtype: str, ctx: ReportContext, src: VolumeSources | None = None) -> list[str]:
    """Чего не хватает тому. Для НДВ — полный перечень по Методике № 581."""
    if vtype == "ndv":
        from ecodoc.development import ndv_project
        return ndv_project.gaps(ctx, _ndv_inputs(src or VolumeSources()))
    out = []
    if not (src and src.sources_table):
        out.append("нет таблицы источников")
    if not (src and src.dispersion_table):
        out.append("нет результатов расчёта/обоснования")
    return out


def _ndv_inputs(src: VolumeSources):
    from ecodoc.development.ndv_project import NdvInputs
    return NdvInputs(sources_header=src.sources_header, sources_table=src.sources_table,
                     dispersion_header=src.dispersion_header,
                     dispersion_table=src.dispersion_table,
                     points_header=src.points_header, points_table=src.points_table,
                     contrib_header=src.contrib_header, contrib_table=src.contrib_table,
                     appendices=list(src.appendices))


def build(vtype: str, ctx: ReportContext, src: VolumeSources,
          out_path: str | Path) -> Path:
    """Собрать том заданного типа (ndv|nds|szz) в .docx.

    НДВ собирается полноценным проектом (ndv_project — структура реального
    проекта по Методике № 581); НДС и СЗЗ — пока каркасом, как раньше.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    if vtype not in _TYPES:
        raise ValueError(f"Тип тома: {list(_TYPES)}")
    if vtype == "ndv":
        from ecodoc.development import ndv_project
        return ndv_project.build(ctx, _ndv_inputs(src), out_path)
    title, subject = _TYPES[vtype]
    org = ctx.organization

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    # — титул —
    h = doc.add_paragraph(title)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].bold = True
    p = doc.add_paragraph(f"для {org.name}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ob in ctx.objects:
        po = doc.add_paragraph(f"объект НВОС {ob.code} — {ob.name}")
        po.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    _h(doc, "1. Аннотация")
    doc.add_paragraph(
        f"Настоящий том разработан в целях установления нормативов {subject} "
        f"для {org.name} (ИНН {org.inn}). Расчёт рассеивания/обоснование выполнены "
        f"в специализированном ПО; результаты приведены в разделе 4.")

    _h(doc, "2. Общие сведения о предприятии")
    for k, v in [("Наименование", org.name), ("ИНН/КПП", f"{org.inn}/{org.kpp}"),
                 ("ОГРН", org.ogrn), ("Адрес", org.address),
                 ("ОКТМО", org.oktmo), ("ОКВЭД", org.okved)]:
        doc.add_paragraph(f"{k}: {v or '—'}")

    _h(doc, "3. Характеристика источников")
    _table(doc, src.sources_header, src.sources_table,
           empty="‹импортировать таблицу источников из «Эколога» (ingest_excel)›")

    _h(doc, "4. Результаты расчёта рассеивания / обоснования")
    _table(doc, src.dispersion_header, src.dispersion_table,
           empty="‹импортировать результаты расчёта из «Эколога»›")

    _h(doc, "5. Предлагаемые нормативы")
    doc.add_paragraph("Нормативы устанавливаются на уровне фактических/расчётных "
                      "значений, не превышающих гигиенических нормативов на границе "
                      "СЗЗ/жилой зоны (см. раздел 4).")

    _h(doc, "6. Мероприятия по снижению воздействия")
    doc.add_paragraph("‹перечень мероприятий — при превышении нормативов›")

    _h(doc, "7. Выводы")
    doc.add_paragraph("На основании выполненных расчётов нормативы " + subject +
                      " считать обоснованными.")

    _h(doc, "Приложения (исходники)")
    if src.appendices:
        for i, a in enumerate(src.appendices, 1):
            doc.add_paragraph(f"Приложение {i}. {a}")
    else:
        doc.add_paragraph("‹протоколы КХА/биотестирования, карты изолиний›")

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path


def _h(doc, text: str):
    p = doc.add_paragraph(text)
    p.runs[0].bold = True


def _table(doc, header, rows, empty: str):
    if not rows:
        doc.add_paragraph(empty)
        return
    ncol = max(len(header) if header else 0, max(len(r) for r in rows))
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    if header:
        cells = t.add_row().cells
        for i, c in enumerate(header[:ncol]):
            cells[i].text = str(c)
            cells[i].paragraphs[0].runs and setattr(cells[i].paragraphs[0].runs[0], "bold", True)
    for r in rows:
        cells = t.add_row().cells
        for i, c in enumerate(r[:ncol]):
            cells[i].text = "" if c is None else str(c)
