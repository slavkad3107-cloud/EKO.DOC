"""Справка о компонентном составе отхода по данным ООС/ПНООЛР — проект
для протокола КХА.

Замечание эколога: «ООС — основополагающий документ по отходам; протоколы
для паспортов должны быть сделаны на основе данных из него». Лаборатории
для протокола нужен заявленный состав — вот он, по каждому отходу, с
указанием файла и листа, откуда взят.

Данные: ctx.extra['waste_passports'] (состав, происхождение, агрегатное
состояние из ИИ-разбора ООС/ПНООЛР/паспортов), ctx.wastes (класс,
наименование), кандидаты площадки (файл и лист-источник строки отхода).
Оформление — как в waste_passport (Times New Roman 12, «Table Grid»).
"""
from __future__ import annotations

from pathlib import Path

from ecodoc.core import fkko as _fkko
from ecodoc.core.models import ReportContext
from ecodoc.core.waste_agg import norm_fkko
from ecodoc.development.waste_passport import (_AGG, _fix_widths, _set,
                                               _sorted_components,
                                               aggregate_state)

TITLE = ("СПРАВКА О КОМПОНЕНТНОМ СОСТАВЕ ОТХОДА (по данным ООС/ПНООЛР) — "
         "проект для протокола")


def _passport_for(ctx: ReportContext, code: str) -> dict:
    extra = ctx.extra if isinstance(ctx.extra, dict) else {}
    for p in extra.get("waste_passports") or []:
        if isinstance(p, dict) and norm_fkko(p.get("fkko")) == code:
            return p
    return {}


def _components(p: dict) -> list[dict]:
    return [c for c in (p.get("components") or [])
            if isinstance(c, dict) and str(c.get("name") or "").strip()]


def _pct(v):
    try:
        return float(str(v).replace(",", ".").replace(" ", "").strip("%"))
    except (TypeError, ValueError):
        return None


def _source_lines(ctx: ReportContext, code: str, p: dict,
                  site_dir: str | Path | None) -> list[str]:
    """«файл (лист N)» — из кандидатов wastes[fkko=…] или из _src паспорта."""
    out: list[str] = []
    if site_dir:
        from ecodoc.intake import candidates
        for c in candidates.Store(site_dir).items:
            coll, sel, _attr = candidates.parse_key(c.key)
            if coll == "wastes" and norm_fkko(sel.get("fkko")) == code and c.file:
                line = c.file + (f" (лист {c.page})" if c.page else "")
                if line not in out:
                    out.append(line)
    src = str(p.get("_src") or "").strip()
    if src and src not in out:
        out.append(src)
    return out


def _targets(ctx: ReportContext) -> list[tuple[str, str, int, dict]]:
    """(код, наименование, класс, паспорт-словарь) — отходы I–IV класса и
    V класса с известным составом."""
    seen: set[str] = set()
    out = []
    for w in ctx.wastes:
        code = norm_fkko(w.fkko_code)
        if not code or code in seen:
            continue
        seen.add(code)
        try:
            hazard = int(w.hazard_class)
        except (TypeError, ValueError):
            hazard = 0
        p = _passport_for(ctx, code)
        out.append((code, w.name or str(p.get("name") or ""), hazard, p))
    # отходы, которые есть только в справочнике паспортов (ООС загружен, актов нет)
    extra = ctx.extra if isinstance(ctx.extra, dict) else {}
    for p in extra.get("waste_passports") or []:
        if not isinstance(p, dict):
            continue
        code = norm_fkko(p.get("fkko"))
        if not code or code in seen:
            continue
        seen.add(code)
        try:
            hazard = int(p.get("hazard_class") or 0)
        except (TypeError, ValueError):
            hazard = 0
        out.append((code, str(p.get("name") or ""), hazard, p))
    return out


def generate(ctx: ReportContext, out_dir: str | Path,
             site_dir: str | Path | None = None) -> list[Path]:
    """По одному .docx на отход I–IV класса (и V — если есть состав) с
    компонентным составом из extra.waste_passports."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Cm, Pt

    out_dir = Path(out_dir)
    paths: list[Path] = []
    org = ctx.organization
    obj = ctx.objects[0] if ctx.objects else None
    for code, name, hazard, p in _targets(ctx):
        comps = _sorted_components(_components(p))
        if not comps:
            continue
        if not (1 <= hazard <= 4) and hazard != 5:
            continue
        out_dir.mkdir(parents=True, exist_ok=True)
        doc = Document()
        st = doc.styles["Normal"]
        st.font.name, st.font.size = "Times New Roman", Pt(12)
        st.paragraph_format.space_after = Pt(0)
        for s in doc.sections:
            s.left_margin, s.right_margin = Cm(2), Cm(1.5)

        # шапка организации
        for text, bold in ((org.name or "‹организация›", True),
                           (f"ИНН {org.inn}" if org.inn else "", False),
                           (org.address or "", False)):
            if text:
                par = doc.add_paragraph()
                par.alignment = AL.CENTER
                par.add_run(text).bold = bold
        if obj is not None:
            par = doc.add_paragraph()
            par.alignment = AL.CENTER
            par.add_run(" ".join(x for x in (
                "Объект НВОС:", obj.code, obj.name, obj.address) if x))
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.alignment = AL.CENTER
        h.add_run(TITLE).bold = True
        doc.add_paragraph()

        agg = (str(p.get("aggregate_state") or "").strip()
               or str(p.get(_AGG) or "").strip() or aggregate_state(code)
               or "‹указать›")
        rows = [("Код вида отходов по ФККО", _fkko.fmt(code)),
                ("Наименование вида отходов по ФККО", name or "—"),
                ("Класс опасности", _fkko.roman(hazard) or "—"),
                ("Агрегатное состояние и физическая форма", agg),
                ("Происхождение (технологический процесс)",
                 str(p.get("origin") or "").strip() or "‹указать по ООС›")]
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        _fix_widths(t, (Cm(7.5), Cm(10)))
        for i, (k, v) in enumerate(rows):
            _set(t.cell(i, 0), k, AL.LEFT)
            _set(t.cell(i, 1), v, AL.LEFT)
        doc.add_paragraph()

        # компоненты
        total = 0.0
        known = True
        t2 = doc.add_table(rows=len(comps) + 2, cols=3)
        t2.style = "Table Grid"
        _fix_widths(t2, (Cm(1.5), Cm(11), Cm(5)))
        for j, head in enumerate(("№", "Компонент", "Содержание, % масс.")):
            _set(t2.cell(0, j), head, AL.CENTER)
        for i, c in enumerate(comps, 1):
            pct = _pct(c.get("percent"))
            if pct is None:
                known = False
            else:
                total += pct
            _set(t2.cell(i, 0), str(i), AL.CENTER)
            _set(t2.cell(i, 1), str(c.get("name", "")), AL.LEFT)
            _set(t2.cell(i, 2), str(c.get("percent", "")).replace(".", ","), AL.CENTER)
        last = len(comps) + 1
        _set(t2.cell(last, 1), "Итого", AL.RIGHT)
        _set(t2.cell(last, 2), f"{total:.2f}".rstrip("0").rstrip(".").replace(".", ","),
             AL.CENTER)
        ctl = doc.add_paragraph()
        if not known:
            ctl.add_run("Контроль: у части компонентов нет процента — уточнить").bold = True
        elif not 99.0 <= total <= 101.0:
            ctl.add_run("Контроль: сумма состава ≠ 100 % — уточнить").bold = True
        else:
            ctl.add_run("Контроль: сумма состава 100 % — сходится")
        doc.add_paragraph()

        src = _source_lines(ctx, code, p, site_dir)
        doc.add_paragraph("Источник данных: " + ("; ".join(src) if src
                                                  else "‹не указан›"))
        doc.add_paragraph("Назначение: заявленный состав для протокола "
                          "количественного химического анализа (КХА) и "
                          "паспорта отхода.")
        doc.add_paragraph()
        doc.add_paragraph("Ответственный за обращение с отходами "
                          "_______________ /_________________/")
        path = out_dir / f"Состав_{code}.docx"
        doc.save(path)
        paths.append(path)
    return paths


def gaps(ctx: ReportContext) -> list[str]:
    """Отходы I–IV класса без компонентного состава."""
    out: list[str] = []
    for code, name, hazard, p in _targets(ctx):
        if 1 <= hazard <= 4 and not _components(p):
            out.append(f"{_fkko.fmt(code)} {name or ''}: нет состава — нужен "
                       f"ООС/ПНООЛР или протокол КХА".replace("  ", " "))
    return out
