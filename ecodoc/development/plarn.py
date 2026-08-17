"""ПЛАРН — план предупреждения и ликвидации разливов нефти и нефтепродуктов.

Нормативная основа (сверено по открытым источникам, август 2026):

  * ст. 46 Федерального закона от 10.01.2002 № 7-ФЗ «Об охране окружающей
    среды» (в ред. ФЗ от 13.07.2020 № 207-ФЗ): организации, осуществляющие
    геологическое изучение, разведку и добычу углеводородного сырья, а также
    переработку (производство), транспортировку, хранение и реализацию
    углеводородного сырья и произведённой из него продукции, обязаны иметь
    план предупреждения и ликвидации разливов нефти и нефтепродуктов,
    финансовое обеспечение его мероприятий и собственные аттестованные
    аварийно-спасательные формирования (АСФ) либо договоры с профессиональными
    АСФ. Обращение с нефтепродуктами не обязано быть основной деятельностью:
    котельные, ДГУ, склады ГСМ автопарков — тоже подпадают.

  * Постановление Правительства РФ от 31.12.2020 № 2451 «Об утверждении
    Правил организации мероприятий по предупреждению и ликвидации разливов
    нефти и нефтепродуктов на территории Российской Федерации, за исключением
    внутренних морских вод Российской Федерации и территориального моря
    Российской Федерации…» — ДЕЙСТВУЕТ (в силе с 01.01.2021). Изменения:
      - ПП РФ от 11.12.2023 № 2122 (с 01.09.2024): план направляется на
        согласование в Росприроднадзор (молчаливое согласование — 16 рабочих
        дней, при повторном поступлении 8/5); уточнены критерии; из порога на
        суше исключены АЗС (надземные резервуары до 100 т или подземные)
        и спецтранспорт для перевозки нефтепродуктов;
      - ПП РФ от 21.02.2026 № 177 (с 01.03.2026): пп. «г» п. 5, пп. 3.1, 10,
        добавлен п. 10.1;
      - ПП РФ от 02.06.2026 № 684 (с 01.09.2026): в план включаются сведения
        об отходах, образующихся при ликвидации разлива, об используемых
        сорбентах и о дальнейшем обращении с собранными нефтепродуктами и
        отходами; оповещение о разливе — в течение 30 минут; срок действия
        Правил продлён до 01.09.2032.

Кому нужен план (п. 4 Правил № 2451) — по максимальному расчётному объёму
разлива: 0,5 т и более для объектов на поверхностных водных объектах;
3 т и более на сухопутной части (кроме АЗС и спецтранспорта — см. выше).

Максимальный расчётный объём разлива (п. 7 Правил), в частности:
  * стационарные объекты хранения нефти и нефтепродуктов (склады, ёмкости) —
    100 % объёма ОДНОЙ НАИБОЛЬШЕЙ ёмкости;
  * железнодорожный состав — 50 % общего объёма цистерн;
  * трубопровод при порыве — 25 % максимального объёма прокачки за 6 часов
    плюс объём нефти между запорными задвижками; при проколе — 2 %.

Время локализации разлива с момента обнаружения (Правила № 2451): не более
4 часов на водном объекте, не более 6 часов на сухопутной части.

Что делает машина: титул, состав по п. 5 Правил, таблицы источников
(резервуаров) из ``ctx.extra['plarn']['tanks']``, расчёт максимального
расчётного объёма по правилу «100 % наибольшей ёмкости», вывод о пороге
3 т / 0,5 т, АСФ и финансовое обеспечение из данных, отходы ликвидации
разлива с подбором кодов по действующему каталогу ФККО. Чего машина НЕ
делает и не выдумывает: зоны распространения и площади разлива
(моделирование), время ликвидации и расчёт достаточности сил (подтверждает
АСФ), телефоны оповещения. Там стоит пометка «[требуется: …]», и та же
строка возвращается gaps().

Данные (`ctx.extra['plarn']`):
    tanks:   [{name, volume_m3, product, density?, mass_t?, location?}]
    asf:     {name, contract, date, certificate?}  или own_asf: "описание"
    finance: {kind, details?, amount?}
    notify:  [{to, phone}]
    operations: ["приём", "хранение", …];  on_water: bool;
    resources:  [{name, qty}] — собственные средства ЛРН (боны, сорбент…)
"""
from __future__ import annotations

import re
from datetime import date
from decimal import Decimal
from pathlib import Path

from ecodoc.core.models import Issue, ReportContext

TITLE = "План предупреждения и ликвидации разливов нефти и нефтепродуктов"

NPA = ("постановление Правительства РФ от 31.12.2020 № 2451 (в ред. "
       "постановлений от 11.12.2023 № 2122, от 21.02.2026 № 177, "
       "от 02.06.2026 № 684), ст. 46 Федерального закона от 10.01.2002 "
       "№ 7-ФЗ «Об охране окружающей среды»")

# Пороги обязательности плана — п. 4 Правил № 2451 (тонны)
THRESHOLD_LAND_T = Decimal("3")
THRESHOLD_WATER_T = Decimal("0.5")

# Время локализации с момента обнаружения разлива — Правила № 2451 (часы)
LOCALIZE_HOURS_WATER = 4
LOCALIZE_HOURS_LAND = 6

RULE_100 = ("для стационарных объектов хранения нефти и нефтепродуктов — "
            "100 % объёма одной наибольшей ёмкости (п. 7 Правил № 2451)")

# Состав плана — п. 5 Правил № 2451, подпункты «а»–«о» (формулировки
# приведены близко к тексту действующей редакции).
SECTIONS: list[tuple[str, str, str]] = [
    ("org", "а", "Общие сведения об эксплуатирующей организации и об основных "
                 "операциях с нефтью и нефтепродуктами"),
    ("sources", "б", "Сведения о потенциальных источниках разливов нефти "
                     "и нефтепродуктов"),
    ("volumes", "в", "Максимальные расчётные объёмы разливов нефти "
                     "и нефтепродуктов"),
    ("zones", "г", "Прогнозируемые зоны распространения разливов и описание "
                   "возможных негативных последствий"),
    ("actions", "д", "Первоочередные действия персонала при разливе"),
    ("time", "е", "Расчётное время локализации и ликвидации разлива"),
    ("forces", "ж", "Состав собственных сил и средств и расчёт их "
                    "достаточности для ликвидации максимального расчётного "
                    "разлива"),
    ("asf", "з", "Состав и порядок действий собственных и (или) привлекаемых "
                 "по договору аварийно-спасательных служб (формирований)"),
    ("extra_forces", "и", "Порядок привлечения дополнительных сил и средств"),
    ("notify", "к", "Схема оповещения и организация управления при разливе"),
    ("waste", "л", "Мероприятия по сбору, временному хранению, транспортировке "
                   "и утилизации собранных нефтепродуктов и отходов "
                   "ликвидации разлива"),
    ("finance", "м, н", "Финансовое обеспечение, объём и стоимость работ "
                        "по мероприятиям плана"),
    ("calendar", "о", "Календарный план оперативных мероприятий при разливе"),
    ("final", "", "Согласование, утверждение, учения и актуализация плана"),
]

# ── пометки «[требуется: …]» — единый текст для документа и gaps() ────────
REQ_TANKS = ("перечень ёмкостей/резервуаров с нефтепродуктами — вкладка "
             "данных, extra['plarn']['tanks'] (наименование, объём м³, "
             "нефтепродукт)")
REQ_DENSITY = ("плотность (density, т/м³) или масса (mass_t) нефтепродукта "
               "наибольшей ёмкости — без неё максимальный расчётный объём "
               "разлива не переводится в тонны и не сравнивается с порогом "
               "3 т (0,5 т на водном объекте)")
REQ_OPS = "перечень основных операций с нефтью и нефтепродуктами на объекте"
REQ_ZONES = ("прогнозируемые зоны распространения разлива и площади по "
             "сценариям — расчёт (моделирование) с учётом рельефа, покрытия "
             "и обвалования")
REQ_FORCES = ("расчётное время ликвидации разлива и расчёт достаточности "
              "сил и средств — подтверждается собственным или привлекаемым "
              "АСФ")
REQ_ASF = ("собственное аттестованное АСФ или договор с профессиональным "
           "АСФ (ст. 46 ФЗ-7) — extra['plarn']['asf'] (наименование, "
           "договор, дата)")
REQ_FINANCE = ("финансовое обеспечение мероприятий плана: банковская "
               "гарантия, договор страхования или резервный фонд "
               "(ст. 46 ФЗ-7) — extra['plarn']['finance']")
REQ_NOTIFY = ("телефоны оповещения (ЕДДС, ЦУКС ГУ МЧС, Росприроднадзор, "
              "орган субъекта РФ, ОМСУ) — extra['plarn']['notify']")

APPROVAL_NOTE = ("план подлежит направлению на согласование в Росприроднадзор "
                 "(считается согласованным через 16 рабочих дней без "
                 "замечаний) и утверждению руководителем организации; "
                 "готовность подтверждается комплексными учениями с "
                 "заключением комиссии — не реже 1 раза в 3 года")

# первоочередные действия персонала — типовой перечень (пп. «д» п. 5)
FIRST_ACTIONS = [
    "прекратить технологические операции с нефтепродуктами, перекрыть "
    "запорную арматуру, при необходимости обесточить оборудование",
    "оповестить руководителя работ и дежурные службы по схеме оповещения "
    "(сообщение о разливе — в течение 30 минут, ПП № 684)",
    "исключить источники открытого огня и искрообразования, выставить "
    "ограждение и знаки, удалить посторонних из зоны разлива",
    "локализовать разлив подручными и штатными средствами: обвалование "
    "грунтом, боновые заграждения, засыпка сорбентом",
    "применять средства индивидуальной защиты; при пожароопасной "
    "концентрации паров — действовать по плану пожаротушения",
    "приступить к сбору разлитого нефтепродукта и передать информацию "
    "прибывшему АСФ",
]

_OIL_RE = re.compile(r"нефт|масл|мазут|бензин|дизел|керосин", re.IGNORECASE)

# типовые отходы ликвидации разлива: коды подбираются по ДЕЙСТВУЮЩЕМУ
# каталогу ФККО, а не зашиты в код — каталог обновляется, коды меняются
TYPICAL_LRN_WASTES = [
    "грунт, загрязненный нефтью или нефтепродуктами (содержание нефти или "
    "нефтепродуктов 15 % и более)",
    "песок, загрязненный нефтью или нефтепродуктами (содержание нефти или "
    "нефтепродуктов 15 % и более)",
    "сорбенты из природных органических материалов отработанные, "
    "загрязненные нефтепродуктами",
    "обтирочный материал, загрязненный нефтью или нефтепродуктами "
    "(содержание нефти или нефтепродуктов менее 15 %)",
]


def _dec(value) -> Decimal | None:
    if value in (None, ""):
        return None
    try:
        return Decimal(str(value).replace(",", "."))
    except Exception:
        return None


def _num(value) -> str:
    if value in (None, ""):
        return "—"
    d = _dec(value)
    if d is None:
        return str(value)
    return f"{d.normalize():f}"


def _data(ctx: ReportContext) -> dict:
    d = (ctx.extra or {}).get("plarn")
    return d if isinstance(d, dict) else {}


def tanks(ctx: ReportContext) -> list[dict]:
    """Ёмкости/резервуары из данных; масса = mass_t или объём × плотность."""
    out: list[dict] = []
    for t in _data(ctx).get("tanks", []):
        if not isinstance(t, dict):
            continue
        vol, dens, mass = (_dec(t.get("volume_m3")), _dec(t.get("density")),
                           _dec(t.get("mass_t")))
        if mass is None and vol is not None and dens is not None:
            mass = vol * dens
        out.append({"name": str(t.get("name") or "ёмкость"),
                    "product": str(t.get("product") or ""),
                    "volume_m3": vol, "density": dens, "mass_t": mass,
                    "location": str(t.get("location") or "")})
    return out


def max_spill(ctx: ReportContext) -> dict:
    """Максимальный расчётный объём разлива: 100 % наибольшей ёмкости.

    «Наибольшая ёмкость» в Правилах — по объёму; массу считаем от неё же.
    Ничего не выдумываем: нет плотности — mass_t остаётся None.
    """
    ts = tanks(ctx)
    if not ts:
        return {}
    with_vol = [t for t in ts if t["volume_m3"] is not None]
    if with_vol:
        top = max(with_vol, key=lambda t: t["volume_m3"])
    else:
        with_mass = [t for t in ts if t["mass_t"] is not None]
        if not with_mass:
            return {"basis": RULE_100}
        top = max(with_mass, key=lambda t: t["mass_t"])
    return {"tank": top["name"], "product": top["product"],
            "volume_m3": top["volume_m3"], "mass_t": top["mass_t"],
            "basis": RULE_100}


def plan_required(ctx: ReportContext) -> tuple[bool | None, str]:
    """Обязателен ли план по порогам п. 4 Правил № 2451.

    True/False — когда масса посчитана; None — сравнить не с чем (нет
    ёмкостей или массы), причина в тексте.
    """
    on_water = bool(_data(ctx).get("on_water"))
    thr = THRESHOLD_WATER_T if on_water else THRESHOLD_LAND_T
    where = ("на поверхностном водном объекте" if on_water
             else "на сухопутной части")
    spill = max_spill(ctx)
    if not spill:
        return None, ("ёмкости с нефтепродуктами не заведены — максимальный "
                      "расчётный объём разлива не определён")
    mass = spill.get("mass_t")
    if mass is None:
        return None, ("масса нефтепродукта наибольшей ёмкости не определена "
                      "(нет плотности) — с порогом "
                      f"{_num(thr)} т не сравнить")
    if mass >= thr:
        return True, (f"максимальный расчётный объём разлива {_num(mass)} т — "
                      f"не менее {_num(thr)} т {where}: план ОБЯЗАТЕЛЕН "
                      "(п. 4 Правил № 2451)")
    return False, (f"максимальный расчётный объём разлива {_num(mass)} т — "
                   f"менее {_num(thr)} т {where}: обязанность разработки "
                   "плана по Правилам № 2451 не наступает (п. 4; проверьте "
                   "исключения — АЗС, спецтранспорт)")


def liquidation_wastes(ctx: ReportContext) -> list[dict]:
    """Отходы ликвидации разлива: из базы объекта + типовые по каталогу ФККО.

    С 01.09.2026 (ПП № 684) сведения об отходах ликвидации и сорбентах —
    обязательная часть плана. Берём: (1) нефтесодержащие отходы, уже
    заведённые по объекту (справки-акты и движение); (2) типовые отходы ЛРН,
    коды которых ПОДОБРАНЫ по действующему каталогу ФККО — не зашиты в код.
    """
    from ecodoc.core import fkko

    rows, seen = [], set()
    for w in list(ctx.waste_acts) + list(ctx.wastes):
        name = getattr(w, "name", "") or ""
        if not _OIL_RE.search(name):
            continue
        code = fkko.norm(getattr(w, "fkko_code", ""))
        key = code or name.lower()
        if key in seen:
            continue
        seen.add(key)
        rows.append({"name": name, "fkko": code,
                     "hazard": int(getattr(w, "hazard_class", 0) or 0),
                     "source": "отходы объекта (из базы)"})
    if fkko.codes():
        for typical in TYPICAL_LRN_WASTES:
            hits = fkko.suggest_by_name(typical, limit=1)
            if not hits or hits[0]["score"] < 0.5:
                continue
            hit = hits[0]
            if hit["code"] in seen:
                continue
            seen.add(hit["code"])
            rows.append({"name": hit["name"], "fkko": hit["code"],
                         "hazard": int(hit.get("class") or 0),
                         "source": "каталог ФККО (типовой отход ЛРН)"})
    return rows


def gaps(ctx: ReportContext) -> list[str]:
    """Чего не хватает для плана — тот же список, что и пометки в тексте."""
    from ecodoc.core import fkko

    out: list[str] = []
    org = ctx.organization
    data = _data(ctx)
    if not (org.name or org.short_name):
        out.append("не заполнено наименование эксплуатирующей организации")
    if not ctx.objects:
        out.append("не заведён объект: нет адреса и кода НВОС")
    ts = tanks(ctx)
    if not ts:
        out.append(f"требуется: {REQ_TANKS}")
    else:
        for t in ts:
            if t["volume_m3"] is None and t["mass_t"] is None:
                out.append(f"ёмкость «{t['name']}»: не указан объём (м³)")
        if max_spill(ctx).get("mass_t") is None:
            out.append(f"требуется: {REQ_DENSITY}")
    if not data.get("operations"):
        out.append(f"требуется: {REQ_OPS}")
    asf = data.get("asf") or {}
    if not (asf.get("name") or data.get("own_asf")):
        out.append(f"требуется: {REQ_ASF}")
    else:
        if asf.get("name") and not asf.get("contract"):
            out.append(f"АСФ «{asf['name']}»: не указаны реквизиты договора")
        if asf.get("name") and not asf.get("certificate"):
            out.append(f"АСФ «{asf['name']}»: не указано свидетельство об "
                       "аттестации на ведение аварийно-спасательных работ")
    if not data.get("finance"):
        out.append(f"требуется: {REQ_FINANCE}")
    if not data.get("notify"):
        out.append(f"требуется: {REQ_NOTIFY}")
    out.append(f"требуется: {REQ_ZONES}")
    out.append(f"требуется: {REQ_FORCES}")
    if not fkko.codes():
        out.append("каталог ФККО не загружен — коды отходов ликвидации "
                   "разлива не подобрать (кнопка «Обновить каталог ФККО»)")
    out.append(APPROVAL_NOTE)
    return out


def validate(ctx: ReportContext) -> list[Issue]:
    """Замечания перед генерацией; про согласование/учения — всегда честно."""
    issues: list[Issue] = []
    for g in gaps(ctx):
        level = "error" if (REQ_TANKS in g or "наименование" in g) else "warning"
        issues.append(Issue(level, "ПЛАРН", g))
    return issues


def generate(ctx: ReportContext, out_path: str | Path) -> Path:
    """Собрать ПЛАРН (.docx): состав по п. 5 Правил № 2451, данные из базы."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Cm, Pt

    org = ctx.organization
    obj = ctx.objects[0] if ctx.objects else None
    data = _data(ctx)
    ts = tanks(ctx)
    spill = max_spill(ctx)
    required, reason = plan_required(ctx)
    on_water = bool(data.get("on_water"))
    year = ctx.period.year or date.today().year

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(3), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(2), Cm(2)

    # ── титул ────────────────────────────────────────────────────────────
    grif = doc.add_paragraph()
    grif.alignment = AL.RIGHT
    grif.add_run(
        "СОГЛАСОВАНО\nТерриториальный орган Росприроднадзора\n"
        "[требуется: реквизиты согласования]\n\n"
        "УТВЕРЖДАЮ\n"
        f"{org.director_position or 'Руководитель'} "
        f"{org.name or org.short_name or ''}\n"
        f"____________ {org.director_name or '[ФИО]'}\n"
        f"«___» __________ {year} г.")
    doc.add_paragraph()
    h = doc.add_paragraph()
    h.alignment = AL.CENTER
    run = h.add_run(TITLE.upper())
    run.bold = True
    run.font.size = Pt(14)
    sub = doc.add_paragraph()
    sub.alignment = AL.CENTER
    sub.add_run(
        f"{org.name or org.short_name or '[требуется: организация]'}\n"
        f"Объект: {(obj.name or obj.code) if obj else '[требуется: объект]'}\n"
        f"Адрес: {obj.address if obj and obj.address else '[требуется: адрес]'}\n"
        f"{year} г.")
    doc.add_paragraph()
    base = doc.add_paragraph()
    base.alignment = AL.CENTER
    base.add_run(f"Разработан в соответствии с: {NPA}.").italic = True
    doc.add_page_break()

    # ── содержание ───────────────────────────────────────────────────────
    doc.add_paragraph().add_run("СОДЕРЖАНИЕ").bold = True
    for i, (_key, letter, title) in enumerate(SECTIONS, start=1):
        mark = f" (пп. «{letter}» п. 5 Правил № 2451)" if letter else ""
        doc.add_paragraph(f"{i}. {title}{mark}")
    doc.add_page_break()

    n = 0
    for key, letter, title in SECTIONS:
        n += 1
        doc.add_paragraph()
        head = doc.add_paragraph()
        head.add_run(f"{n}. {title}").bold = True

        if key == "org":
            doc.add_paragraph(
                f"Эксплуатирующая организация: "
                f"{org.name or org.short_name or '[требуется]'}, "
                f"ИНН {org.inn or '[требуется]'}, ОГРН {org.ogrn or '—'}, "
                f"адрес: {org.address or '[требуется]'}, "
                f"тел.: {org.phone or '—'}.")
            if obj:
                doc.add_paragraph(
                    f"Объект: {obj.name or '—'}, код объекта НВОС "
                    f"{obj.code or '—'}, категория {obj.category or '—'}, "
                    f"адрес: {obj.address or '[требуется]'}.")
            ops = data.get("operations") or []
            if ops:
                doc.add_paragraph("Основные операции с нефтью и "
                                  "нефтепродуктами: " + "; ".join(ops) + ".")
            else:
                doc.add_paragraph(f"[требуется: {REQ_OPS}]")
            doc.add_paragraph(
                "Размещение объекта: " +
                ("на поверхностном водном объекте (порог обязательности "
                 f"плана — {_num(THRESHOLD_WATER_T)} т)" if on_water else
                 "на сухопутной части территории РФ (порог обязательности "
                 f"плана — {_num(THRESHOLD_LAND_T)} т)") + ".")

        elif key == "sources":
            doc.add_paragraph(
                "Потенциальные источники разливов — ёмкости (резервуары) "
                f"хранения нефтепродуктов, всего {len(ts)}:")
            _table(doc, ["№", "Ёмкость (источник)", "Нефтепродукт",
                         "Объём, м³", "Масса, т", "Расположение"],
                   [[str(i), t["name"], t["product"] or "—",
                     _num(t["volume_m3"]), _num(t["mass_t"]),
                     t["location"] or "—"]
                    for i, t in enumerate(ts, start=1)])
            if not ts:
                doc.add_paragraph(f"[требуется: {REQ_TANKS}]")
            doc.add_paragraph(
                "Разливы возможны также при приёме/выдаче нефтепродуктов "
                "(разгерметизация сливо-наливных устройств, переполнение).")

        elif key == "volumes":
            doc.add_paragraph(
                "Максимальный расчётный объём разлива определён по п. 7 "
                f"Правил № 2451: {RULE_100}.")
            if spill.get("tank"):
                doc.add_paragraph(
                    f"Наибольшая ёмкость — «{spill['tank']}» "
                    f"({spill.get('product') or 'нефтепродукт'}): "
                    f"{_num(spill.get('volume_m3'))} м³. Максимальный "
                    "расчётный объём разлива — 100 % её объёма: "
                    f"{_num(spill.get('volume_m3'))} м³ "
                    f"({_num(spill.get('mass_t'))} т).")
            if ts and spill.get("mass_t") is None:
                doc.add_paragraph(f"[требуется: {REQ_DENSITY}]")
            doc.add_paragraph("Вывод: " + reason + ".")

        elif key == "zones":
            doc.add_paragraph(
                "Сценарии разливов приняты от полной разгерметизации "
                "наибольшей ёмкости (максимальный расчётный объём) до "
                "разгерметизации при сливо-наливных операциях:")
            _table(doc, ["№", "Сценарий", "Объём, м³", "Площадь разлива, м²"],
                   [[str(i),
                     f"Полная разгерметизация ёмкости «{t['name']}» "
                     f"({t['product'] or 'нефтепродукт'})",
                     _num(t["volume_m3"]), "[требуется: расчёт]"]
                    for i, t in enumerate(ts, start=1)])
            doc.add_paragraph(f"[требуется: {REQ_ZONES}]")

        elif key == "actions":
            doc.add_paragraph("Персонал объекта при обнаружении разлива:")
            for m in FIRST_ACTIONS:
                doc.add_paragraph(f"— {m};")

        elif key == "time":
            doc.add_paragraph(
                "Время локализации разлива с момента его обнаружения — "
                f"не более {LOCALIZE_HOURS_WATER} часов при разливе на "
                "поверхностном водном объекте и не более "
                f"{LOCALIZE_HOURS_LAND} часов при разливе на сухопутной "
                "части территории РФ (Правила № 2451).")
            doc.add_paragraph(
                "Расчётное время ликвидации определяется расчётом "
                "достаточности сил и средств (раздел "
                f"{next(i for i, s in enumerate(SECTIONS, 1) if s[0] == 'forces')}).")

        elif key == "forces":
            res = [r for r in data.get("resources", []) if isinstance(r, dict)]
            if res:
                doc.add_paragraph("Собственные средства ЛРН объекта:")
                _table(doc, ["№", "Средство", "Количество"],
                       [[str(i), str(r.get("name") or "—"),
                         str(r.get("qty") or "—")]
                        for i, r in enumerate(res, start=1)])
            doc.add_paragraph(f"[требуется: {REQ_FORCES}]")

        elif key == "asf":
            asf = data.get("asf") or {}
            if asf.get("name"):
                doc.add_paragraph(
                    "Ликвидация разлива обеспечивается привлекаемым на "
                    f"договорной основе профессиональным АСФ: {asf['name']}, "
                    f"договор {asf.get('contract') or '[требуется: договор]'}"
                    + (f" от {asf['date']}" if asf.get("date") else "")
                    + ", свидетельство об аттестации "
                    + (asf.get("certificate")
                       or "[требуется: свидетельство об аттестации]") + ".")
            elif data.get("own_asf"):
                doc.add_paragraph("Собственное аварийно-спасательное "
                                  f"формирование: {data['own_asf']}.")
            else:
                doc.add_paragraph(f"[требуется: {REQ_ASF}]")

        elif key == "extra_forces":
            doc.add_paragraph(
                "При угрозе превышения максимального расчётного объёма "
                "разлива либо недостаточности сил и средств дополнительные "
                "силы привлекаются через КЧС и ОПБ муниципального "
                "образования и субъекта РФ в порядке, установленном "
                "законодательством о единой государственной системе "
                "предупреждения и ликвидации ЧС (РСЧС).")

        elif key == "notify":
            doc.add_paragraph(
                "Оповещение о разливе — незамедлительно, сообщение о месте "
                "обнаружения разлива передаётся в течение 30 минут "
                "(ПП № 684). Адресаты оповещения:")
            notify = [x for x in data.get("notify", []) if isinstance(x, dict)]
            if notify:
                _table(doc, ["№", "Адресат", "Телефон"],
                       [[str(i), str(x.get("to") or "—"),
                         str(x.get("phone") or "—")]
                        for i, x in enumerate(notify, start=1)])
            else:
                _table(doc, ["№", "Адресат", "Телефон"],
                       [[str(i), to, "[требуется]"]
                        for i, to in enumerate(
                            ["ЕДДС муниципального образования",
                             "ЦУКС ГУ МЧС России по субъекту РФ",
                             "Территориальный орган Росприроднадзора",
                             "Орган исполнительной власти субъекта РФ",
                             "Орган местного самоуправления",
                             "Ростехнадзор (для опасных производственных "
                             "объектов)"], start=1)])
                doc.add_paragraph(f"[требуется: {REQ_NOTIFY}]")
            doc.add_paragraph(
                "В сообщении указываются: дата и время обнаружения, место и "
                "источник разлива, вид и ориентировочный объём "
                "нефтепродукта, принятые меры, данные передавшего.")

        elif key == "waste":
            rows = liquidation_wastes(ctx)
            doc.add_paragraph(
                "Собранные нефтепродукты, загрязнённый грунт и отработанные "
                "сорбенты накапливаются в герметичной таре на площадке с "
                "твёрдым покрытием и передаются организациям, имеющим "
                "лицензию на деятельность по обращению с отходами "
                "I–IV классов опасности. Сведения об отходах ликвидации "
                "разлива (обязательны с 01.09.2026, ПП № 684):")
            _table(doc, ["№", "Наименование отхода", "Код ФККО", "Класс",
                         "Откуда запись"],
                   [[str(i), r["name"], r["fkko"] or "—",
                     str(r["hazard"] or "—"), r["source"]]
                    for i, r in enumerate(rows, start=1)])
            doc.add_paragraph(
                "Коды ФККО подобраны по действующему каталогу "
                "(приказ Росприроднадзора от 22.05.2017 № 242 с "
                "изменениями) — уточните по фактическому составу отходов.")

        elif key == "finance":
            fin = data.get("finance") or {}
            doc.add_paragraph(
                "Финансовое обеспечение осуществления мероприятий плана "
                "предусмотрено ст. 46 ФЗ-7: банковская гарантия, договор "
                "страхования либо документ о резервировании собственных "
                "средств.")
            if fin:
                doc.add_paragraph(
                    f"Принятый способ: {fin.get('kind') or '—'}"
                    + (f", {fin['details']}" if fin.get("details") else "")
                    + (f", сумма {_num(fin['amount'])} руб."
                       if fin.get("amount") else "") + ".")
            else:
                doc.add_paragraph(f"[требуется: {REQ_FINANCE}]")
            doc.add_paragraph(
                "[требуется: объём работ и стоимость единицы объёма работ "
                "по каждому мероприятию плана — по смете/договору с АСФ]")

        elif key == "calendar":
            _table(doc, ["№", "Мероприятие", "Срок / норматив"],
                   [["1", "Обнаружение разлива, первоочередные действия "
                          "персонала", "незамедлительно"],
                    ["2", "Оповещение по схеме (сообщение о месте разлива)",
                     "в течение 30 минут"],
                    ["3", "Локализация разлива",
                     f"не более {LOCALIZE_HOURS_WATER} ч (водный объект) / "
                     f"{LOCALIZE_HOURS_LAND} ч (суша)"],
                    ["4", "Сбор нефтепродукта и загрязнённого грунта",
                     "[требуется: по расчёту достаточности сил]"],
                    ["5", "Вывоз отходов ликвидации, рекультивация",
                     "[требуется: по договору с АСФ/подрядчиком]"],
                    ["6", "Оценка последствий, отчёт о ликвидации",
                     "по завершении работ"]])

        elif key == "final":
            doc.add_paragraph("Порядок введения плана в действие: " +
                              APPROVAL_NOTE + ".")
            doc.add_paragraph(
                "План корректируется по результатам комплексных учений и "
                "фактических операций по ликвидации разливов. Правила "
                "№ 2451 действуют до 01.09.2032 (в ред. ПП от 02.06.2026 "
                "№ 684).")

    doc.add_paragraph()
    sign = doc.add_paragraph()
    sign.add_run(f"{org.director_position or 'Руководитель'}\t\t"
                 f"_____________\t{org.director_name or ''}")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def _table(doc, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, text in enumerate(header):
        table.rows[0].cells[i].text = text
    for r in rows:
        cells = table.add_row().cells
        for i, text in enumerate(r):
            cells[i].text = text
    if not rows:
        cells = table.add_row().cells
        cells[0].text = "[требуется: данные не заведены]"
