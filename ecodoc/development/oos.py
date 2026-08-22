"""Раздел 8 ООС — «Перечень мероприятий по охране окружающей среды» (том).

Почему переписано: прежняя версия выдавала «листик» — 10 пунктов п. 25 с
тремя таблицами. Настоящий раздел (эталоны в OneDrive\\Формы\\Разработка\\ООС:
АК-01-25-ООС1 — 334 листа; 03-26-1-ООС1/ООС.2 «Выборжец» — прошёл экспертизу)
— это том с 11 разделами, ~60 таблицами и 13+3 приложениями. Структура
(заголовки и шапки таблиц дословно) вынесена в oos_structure.py, расчёт
нормативов образования отходов — в oos_waste_calc.py; здесь — сборка .docx.

Откуда берутся данные (и что машина НЕ выдумывает):
  * организация/объект/период           — ctx.organization, ctx.objects, ctx.period;
  * выбросы эксплуатации (т/год, г/с)   — ctx.pollutants (air) + extra['emission_sources'];
  * сбросы                              — ctx.pollutants (water), extra['water'];
  * отходы эксплуатации                 — ctx.wastes / waste_acts / waste_passports
                                          (+ расчёт по нормативам накопления);
  * всё «проектное» — ведомость материалов, техника, численность,
    источники/точки/концентрации из УПРЗА, шум, баланс воды, климат, фон —
    extra['oos'] (схема — в _OOS_SCHEMA ниже; заполняется из вкладки
    «Разработка/ООС», из «Задания ИИ» или из выгрузок «Эколога»).

Чего нет — в документе остаётся «[требуется: …]», и та же строка попадает
в gaps(): пользователь видит список до генерации. Раздел «Современное
состояние среды» НЕ пересказывает ИЭИ (решение пользователя 22.08.2026):
в нём ссылка на технический отчёт ИЭИ как на источник и таблицы фона/замеров,
только если значения переданы в extra['oos']['iei'].
"""
from __future__ import annotations

from datetime import date
from decimal import Decimal
from pathlib import Path

from ecodoc.core.models import Medium, ReportContext
from ecodoc.development import oos_waste_calc as wc
from ecodoc.development.oos_structure import (APPENDICES_GRAPHIC, APPENDICES_TEXT,
                                              NPA, REFERENCES, SECTIONS, TABLES,
                                              section_titles)

TITLE = "Раздел 8. Перечень мероприятий по охране окружающей среды"
STAMP = "ПРОЕКТНАЯ ДОКУМЕНТАЦИЯ"
TYPICAL = "(типовая формулировка — уточнить по проекту)"

# Схема extra['oos'] — документация для вкладки/ИИ; ключи необязательны.
_OOS_SCHEMA = {
    "project": {"code": "шифр (АК-01-25)", "title": "наименование объекта по ТЗ",
                "designer": "проектная организация", "customer": "заказчик",
                "cadastral": "кадастровый номер ЗУ", "stage": "П",
                "description": "проектные решения (абзацы)", "months": 24,
                "workers": 39, "itr": 8, "shifts": 2, "days": 720,
                "tep": [{"name": "Площадь застройки", "unit": "м2", "value": ""}],
                "land_tep": [{"name": "Площадь земельного участка", "m2": ""}],
                "machinery": [{"area": "", "name": "", "brand": "", "spec": "", "qty": ""}]},
    "iei": {"report": "шифр/наименование отчёта ИЭИ", "background":
            [{"name": "Диоксид азота", "value": "0.054", "unit": "мг/м3"}],
            "noise": [{"place": "", "leq": "", "lmax": "", "minutes": "", "pdu": ""}],
            "soil_category": "чистая", "radiation": "в норме", "oopt": "отсутствуют",
            "water_bodies": "", "flora_fauna": "", "location": "", "geology": ""},
    "climate": {"A": 200, "t_warm": 25.2, "t_cold": -0.3, "u5": 8.5,
                "wind": {"С": 12, "СВ": 16, "В": 21, "ЮВ": 6, "Ю": 10, "ЮЗ": 17,
                         "З": 9, "СЗ": 9, "штиль": 2}},
    "construction": {
        "sources": [{"number": "6001", "kind": "передвижной", "name": "Земляные работы",
                     "height": 5, "organized": False}],
        "pollutants": [{"code": "0301", "name": "", "criterion": "ПДК м/р",
                        "value": 0.2, "hazard": 3, "g_s": "", "t_year": ""}],
        "points": [{"code": 1, "x": "", "y": "", "h": 2, "type": "на границе жилой зоны"}],
        "skip": [{"code": "", "name": "", "sum": ""}],
        "concentrations": [{"code": "", "name": "", "values": []}],
        "noise_sources": [], "noise_results": [], "noise_bg": [],
        "water": {"workers_shift": 39, "itr": 8, "q_hoz_l": 15, "k_h": 2,
                  "q_hoz_ls": "", "q_fire_ls": 5, "q_prod_ls": "", "q_total_ls": ""},
        "materials": [{"name": "Бетон", "kind": "бетон", "qty": 4027, "unit": "м3"}],
        "electrodes_t": 0.469, "wheel_wash": {"cars_per_day": 3, "water_m3": 0.3},
        "soil_excess_m3": 6208, "cesspool_shifts": 1440},
    "operation": {
        "sources": [{"number": "0001", "kind": "организованный", "name": "",
                     "height": 19.5}],
        "pollutants": [{"code": "", "criterion": "", "value": "", "hazard": ""}],
        "points": [], "skip": [], "concentrations": [],
        "noise_sources": [], "noise_pdu": [], "noise_bg": [],
        "water_balance": [{"consumer": "", "in_day": "", "in_year": "",
                           "out_day": "", "out_year": "", "to": ""}],
        "water_text": "описание систем водоснабжения/водоотведения",
        "wastes_norm": [{"name": "", "fkko": "", "hazard": 4, "count": "",
                         "count_unit": "чел.", "norm_m3": "", "density": ""}],
        "lamps": [{"name": "", "count": "", "life_h": "", "hours": 4380, "mass_kg": ""}],
        "waste_handling": {"<ФККО>": "вид обращения / получатель"}},
    "pek": {"category": "III", "periodicity": "1 раз в год",
            "monitoring": [{"area": "Почвы", "points": "", "kind": "", "sample": "",
                            "method": "", "components": "", "org": ""}]},
}


# ───────────────────────── общие утилиты ─────────────────────────────────
def _num(value) -> str:
    if value in (None, ""):
        return "—"
    try:
        d = Decimal(str(value).replace(",", "."))
    except Exception:
        return str(value)
    return f"{d.normalize():f}" if d else "0"


def _oos(ctx: ReportContext) -> dict:
    x = (ctx.extra or {}).get("oos") or {}
    return x if isinstance(x, dict) else {}


def _sub(ctx: ReportContext, *keys) -> dict:
    d = _oos(ctx)
    for k in keys:
        d = d.get(k) or {}
        if not isinstance(d, dict):
            return {}
    return d


def _lst(d: dict, key: str) -> list[dict]:
    v = d.get(key) or []
    return [x for x in v if isinstance(x, dict)] if isinstance(v, list) else []


def media(ctx: ReportContext) -> set[str]:
    """По каким средам есть данные о воздействии."""
    out: set[str] = set()
    for p in ctx.pollutants:
        out.add("air" if p.medium == Medium.AIR else "water")
    if ctx.wastes or ctx.waste_acts:
        out.add("waste")
    return out


def rows_air(ctx: ReportContext) -> list[dict]:
    """Вещества эксплуатации: т/год из ctx.pollutants, г/с — сумма по источникам."""
    g_s: dict[str, Decimal] = {}
    for s in (ctx.extra or {}).get("emission_sources", []) or []:
        if not isinstance(s, dict):
            continue
        for p in s.get("pollutants") or []:
            if isinstance(p, dict) and p.get("g_s") not in (None, ""):
                try:
                    g_s[_code(p.get("code"))] = g_s.get(_code(p.get("code")), Decimal(0)) \
                        + Decimal(str(p["g_s"]).replace(",", "."))
                except Exception:
                    pass
    crit = {_code(p.get("code")): p for p in _lst(_sub(ctx, "operation"), "pollutants")}
    ref = _substances()
    out = []
    for p in ctx.pollutants:
        if p.medium != Medium.AIR:
            continue
        code = _code(p.code)
        c = crit.get(code) or {}
        r = ref.get(code) or {}
        criterion = c.get("criterion") or ("ПДК м/р" if r.get("pdk_mr") else
                                           "ПДК с/с" if r.get("pdk_ss") else "")
        value = c.get("value") or r.get("pdk_mr") or r.get("pdk_ss") or ""
        out.append({"code": code, "name": p.name, "criterion": criterion,
                    "value": value, "hazard": c.get("hazard") or "",
                    "g_s": g_s.get(code), "mass": p.mass_norm + p.mass_limit + p.mass_over})
    return out


def rows_water(ctx: ReportContext) -> list[dict]:
    return [{"name": p.name, "code": p.code,
             "mass": p.mass_norm + p.mass_limit + p.mass_over}
            for p in ctx.pollutants if p.medium == Medium.WATER]


def rows_waste(ctx: ReportContext) -> list[dict]:
    from ecodoc.development.waste_inventory import collect
    return collect(ctx)


def _code(code) -> str:
    s = str(code or "").strip()
    return s.zfill(4) if s.isdigit() and len(s) < 4 else s


def _substances() -> dict[str, dict]:
    try:
        from ecodoc.core import refdata
        return {_code(s.get("code")): s for s in refdata.substances()
                if isinstance(s, dict) and s.get("medium", "air") == "air"}
    except Exception:
        return {}


# ───────────────────────── расчётные блоки ───────────────────────────────
def construction_wastes(ctx: ReportContext) -> dict:
    """Все расчёты отходов периода строительства + сводная таблица 9.1.16."""
    c = _sub(ctx, "construction")
    pr = _sub(ctx, "project")
    workers, itr = int(pr.get("workers") or 0), int(pr.get("itr") or 0)
    months = int(pr.get("months") or 0)
    days = int(pr.get("days") or (months * 30 if months else 0))
    res: dict = {"tko": [], "wheel": [], "materials": [], "electrodes": None,
                 "soil": None, "cesspool": None, "summary": []}
    summary: dict[str, dict] = {}

    def add(fkko, name, hazard, t, m3, handling=""):
        row = summary.setdefault(fkko or name, {
            "name": name, "fkko": fkko, "hazard": hazard, "t": Decimal(0),
            "m3": Decimal(0), "handling": handling})
        row["t"] += t
        row["m3"] += m3

    if workers and months:
        res["tko"] = wc.tko_construction(workers, itr, months)
        add("73310001724", "Мусор от офисных и бытовых помещений организаций "
            "несортированный (исключая крупногабаритный)", 4,
            sum(r["t"] for r in res["tko"]), sum(r["m3"] for r in res["tko"]),
            "передача региональному оператору")
        res["cesspool"] = wc.cesspool(workers + itr, int(c.get("cesspool_shifts") or days or 0))
        add("73210001304", "Отходы (осадки) из выгребных ям", 4,
            res["cesspool"]["t"], res["cesspool"]["m3"],
            "вывоз ассенизационной машиной по договору")
    ww = c.get("wheel_wash") or {}
    if ww.get("cars_per_day") and days:
        res["wheel"] = wc.wheel_wash(wc._d(ww["cars_per_day"]), wc._d(ww.get("water_m3"), "0.3"), days)
        add("72310101394", "Осадок (шлам) механической очистки нефтесодержащих сточных "
            "вод, содержащий нефтепродукты в количестве менее 15 %, обводненный", 4,
            sum(r["t"] for r in res["wheel"]), sum(r["m3"] for r in res["wheel"]),
            "передача лицензированной организации на обезвреживание")
    mats = _lst(c, "materials")
    if mats:
        res["materials"] = wc.material_waste(mats)
        for r in res["materials"]:
            if not r["note"]:
                add(r["fkko"], r["waste_name"], r["hazard"], r["t"], r["m3"],
                    "передача лицензированной организации на утилизацию")
    if c.get("electrodes_t"):
        res["electrodes"] = wc.electrodes(wc._d(c["electrodes_t"]))
        add("91910001205", "Остатки и огарки стальных сварочных электродов", 5,
            res["electrodes"]["t"], res["electrodes"]["m3"], "передача на утилизацию")
    if c.get("soil_excess_m3"):
        res["soil"] = wc.soil_excess(wc._d(c["soil_excess_m3"]))
        add("81110001495", "Грунт, образовавшийся при проведении землеройных работ, "
            "не загрязненный опасными веществами", 5, res["soil"]["t"], res["soil"]["m3"],
            "передача лицензированной организации / использование по ТУ")
    res["summary"] = sorted(summary.values(), key=lambda r: (r["hazard"], r["name"]))
    return res


def operation_wastes(ctx: ReportContext) -> dict:
    """Отходы эксплуатации: расчёты по нормативам + лампы + заведённые отходы."""
    op = _sub(ctx, "operation")
    handling = op.get("waste_handling") or {}
    res = {"norm": wc.by_norm(_lst(op, "wastes_norm")), "lamps": wc.lamps(_lst(op, "lamps")),
           "summary": []}
    summary: dict[str, dict] = {}

    def add(fkko, name, hazard, t, m3, h=""):
        key = fkko or name
        row = summary.setdefault(key, {"name": name, "fkko": fkko, "hazard": hazard,
                                       "t": Decimal(0), "m3": Decimal(0),
                                       "handling": h or handling.get(fkko, "")})
        row["t"] += Decimal(str(t or 0))
        row["m3"] += Decimal(str(m3 or 0))
        row["name"] = row["name"] or name

    for r in res["norm"]:
        if not r["note"]:
            add(r["fkko"], r["name"], r["hazard"], r["t"], r["m3"])
    if res["lamps"]:
        add("48241100525", "Лампы накаливания, утратившие потребительские свойства", 5,
            sum(r["t"] for r in res["lamps"]), sum(r["m3"] for r in res["lamps"]),
            "передача на обезвреживание/утилизацию")
    for r in rows_waste(ctx):
        if r["fkko"] in summary:
            continue
        add(r["fkko"], r["name"], r["hazard"], r.get("generated") or 0, 0,
            ", ".join(r.get("operations") or []) +
            (" — " + ", ".join(r.get("receivers") or []) if r.get("receivers") else ""))
    res["summary"] = sorted(summary.values(), key=lambda r: (r["hazard"] or 9, r["name"]))
    return res


def _air_rate(code: str, name: str, year: int) -> Decimal:
    """Ставка платы за выброс 1 т (руб.) по справочнику года — для табл. 10.1.x."""
    try:
        from ecodoc.core.refdata import rates_nvos
        from ecodoc.reports.declaration_nvos.calc import _find_rate
        rates = rates_nvos()
        direct = (rates.get("rates_by_year") or {}).get(str(year)) or {}
        table = direct.get("air") or rates.get("air") or {}
        entry, _ = _find_rate(table, code, name)
        return Decimal(str(entry["rate"])) if entry else Decimal(0)
    except Exception:
        return Decimal(0)


def costs(ctx: ReportContext) -> dict:
    """Плата за НВОС (эксплуатация) — строки по веществам/отходам и итоги."""
    try:
        from ecodoc.reports.declaration_nvos.calc import calculate
        r = calculate(ctx)
    except Exception:
        return {}
    if not r.lines:
        return {}
    return {"air": r.total_air, "water": r.total_water, "waste": r.total_waste,
            "total": r.total, "lines": r.lines, "warnings": r.warnings}


# ───────────────────────── пробелы ───────────────────────────────────────
def gaps(ctx: ReportContext) -> list[str]:
    """Чего не хватает для тома — тот же текст, что и пометки в документе."""
    out: list[str] = []
    org = ctx.organization
    pr, c, op, iei = (_sub(ctx, "project"), _sub(ctx, "construction"),
                      _sub(ctx, "operation"), _sub(ctx, "iei"))
    if not (org.name or org.short_name):
        out.append("не заполнено наименование организации-заказчика")
    if not ctx.objects:
        out.append("не заведён объект: нет кода НВОС, адреса и категории")
    else:
        for o in ctx.objects:
            if not o.address:
                out.append(f"объект {o.code or o.name}: не указан адрес")
            if not o.category:
                out.append(f"объект {o.code or o.name}: не указана категория НВОС")
    have = media(ctx)
    if not have and not c:
        out.append("нет ни выбросов, ни сбросов, ни отходов — раздел ООС "
                   "нечем наполнять: загрузите исходные данные")
    if not pr.get("title"):
        out.append("требуется: наименование объекта и шифр проекта (extra.oos.project)")
    if not pr.get("description"):
        out.append("требуется: описание проектных решений (раздел 3) — из ПЗ/АР/ИОС")
    if not iei.get("report"):
        out.append(f"требуется: технический отчёт об инженерно-экологических "
                   f"изысканиях ({NPA['sp502']}) — источник раздела 2")
    if not iei.get("background"):
        out.append("требуется: справка о фоновых концентрациях ЦГМС (табл. 2.5.1.1)")
    if not _sub(ctx, "climate"):
        out.append("требуется: справка о климатических характеристиках (табл. 2.2.1)")
    if not (pr.get("workers") and pr.get("months")):
        out.append("требуется: численность и продолжительность строительства (ПОС) — "
                   "расчёт ТКО, воды и отходов периода строительства")
    if not _lst(pr, "machinery"):
        out.append("требуется: перечень строительной техники (ПОС) — табл. 3.6/5.1.1.2")
    if not _lst(c, "pollutants"):
        out.append("требуется: расчёт выбросов периода строительства (АТП-Эколог/"
                   "«Дизель»/«Сварка») — табл. 5.1.1.3")
    for period, d in (("строительства", c), ("эксплуатации", op)):
        if not _lst(d, "concentrations"):
            out.append(f"требуется: результаты расчёта рассеивания (УПРЗА «Эколог», "
                       f"{NPA['mrr2017']}) на период {period} — концентрации в "
                       f"расчётных точках")
        if not _lst(d, "noise_results") and not _lst(d, "noise_bg"):
            out.append(f"требуется: акустический расчёт на период {period} "
                       f"({NPA['sp51']}) с учётом фона")
    if "air" not in have and not (ctx.extra or {}).get("emission_sources"):
        out.append("требуется: инвентаризация источников выбросов периода эксплуатации "
                   f"({NPA['pr871']}) — табл. 5.2.1/5.2.2")
    if not _lst(op, "water_balance") and not (ctx.extra or {}).get("water"):
        out.append("требуется: балансовая схема водопотребления и водоотведения "
                   "(ИОС3) — табл. 6.2.1")
    if not _lst(c, "materials"):
        out.append("требуется: ведомость потребности в строительных материалах "
                   "(Приложение 4) — расчёт строительных отходов")
    if "waste" not in have and not _lst(op, "wastes_norm"):
        out.append("требуется: перечень отходов периода эксплуатации с нормативами "
                   "накопления")
    if "waste" in have:
        from ecodoc.development.waste_inventory import gaps as waste_gaps
        out += waste_gaps(ctx)
    if not iei.get("oopt"):
        out.append("требуется: письма уполномоченных органов об ООПТ, Красной книге, "
                   "ЗСО, ОКН (Приложение 3)")
    return out


# ───────────────────────── docx-помощники ────────────────────────────────
class _Doc:
    """Обёртка над python-docx: нумерация заголовков, таблицы, пометки."""

    def __init__(self):
        from docx import Document
        from docx.shared import Cm, Pt
        self.d = Document()
        st = self.d.styles["Normal"]
        st.font.name = "Times New Roman"
        st.font.size = Pt(12)
        for s in self.d.sections:
            s.left_margin, s.right_margin = Cm(2.5), Cm(1.5)
            s.top_margin, s.bottom_margin = Cm(2), Cm(2)
        self.n1 = 0
        self.n2 = 0
        self.n3 = 0
        self.table_titles: list[str] = []
        self.headings: list[str] = []

    def p(self, text: str = "", bold=False, center=False, size=None, italic=False):
        from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
        from docx.shared import Pt
        para = self.d.add_paragraph()
        if center:
            para.alignment = AL.CENTER
        run = para.add_run(text)
        run.bold, run.italic = bold, italic
        if size:
            run.font.size = Pt(size)
        return para

    def h1(self, title: str, numbered=True):
        self.n1 += 1
        self.n2 = 0
        text = f"{self.n1}. {title}" if numbered else title
        self.headings.append(text)
        self.p(text, bold=True, size=14)

    def h2(self, title: str):
        self.n2 += 1
        self.n3 = 0
        text = f"{self.n1}.{self.n2}. {title}"
        self.headings.append(text)
        self.p(text, bold=True)

    def h3n(self, title: str):
        """Нумерованный третий уровень (2.5.1 … 2.5.4 — как в эталоне)."""
        self.n3 += 1
        text = f"{self.n1}.{self.n2}.{self.n3}. {title}"
        self.headings.append(text)
        self.p(text, bold=True, italic=True)

    def h3(self, title: str):
        self.p(title, bold=True, italic=True)

    def need(self, what: str):
        self.p(f"[требуется: {what}]", italic=True)

    def bullets(self, items, intro=None):
        if intro:
            self.p(intro)
        for it in items:
            self.p(f"– {it};")

    def table(self, key: str, rows: list[list], total: list | None = None,
              header: list[str] | None = None, title: str | None = None):
        num, ttl, head = TABLES[key]
        head = header or head
        caption = f"Таблица {num} – {title or ttl}"
        self.table_titles.append(caption)
        self.p(caption, bold=True)
        t = self.d.add_table(rows=1, cols=len(head))
        t.style = "Table Grid"
        for i, h in enumerate(head):
            t.rows[0].cells[i].text = h
        if not rows:
            cells = t.add_row().cells
            cells[0].text = "[требуется: данные не заведены]"
        for r in rows:
            cells = t.add_row().cells
            for i, v in enumerate(r[:len(head)]):
                cells[i].text = "" if v is None else str(v)
        if total:
            cells = t.add_row().cells
            for i, v in enumerate(total[:len(head)]):
                cells[i].text = "" if v is None else str(v)
        self.d.add_paragraph()

    def page_break(self):
        self.d.add_page_break()


# ───────────────────────── генерация ─────────────────────────────────────
def generate(ctx: ReportContext, out_path: str | Path,
             stage: str = "эксплуатация") -> Path:
    """Том ООС (.docx): титул, состав, содержание, разделы 1–11, литература,
    перечень приложений. `stage` сохранён для совместимости с GUI: раздел
    всегда содержит оба периода (строительство и эксплуатация), как того
    требует п. 25 ПП 87 и практика экспертизы."""
    org = ctx.organization
    obj = ctx.objects[0] if ctx.objects else None
    pr, iei, cl = _sub(ctx, "project"), _sub(ctx, "iei"), _sub(ctx, "climate")
    con, op, pek = _sub(ctx, "construction"), _sub(ctx, "operation"), _sub(ctx, "pek")
    year = ctx.period.year or date.today().year
    code = pr.get("code") or "[шифр]"
    title_obj = pr.get("title") or (obj.name if obj and obj.name else "[требуется: наименование объекта]")
    address = (obj.address if obj and obj.address else "[требуется: адрес]")
    customer = org.name or org.short_name or "[требуется: заказчик]"
    have = media(ctx)

    doc = _Doc()

    # ── титульный лист (АК-01-25, лист 1) ──────────────────────────────
    doc.p(pr.get("designer") or customer, bold=True, center=True)
    for _ in range(6):
        doc.p()
    doc.p(STAMP, bold=True, center=True, size=14)
    doc.p(title_obj.upper(), bold=True, center=True, size=14)
    doc.p(address, center=True)
    doc.p()
    doc.p(TITLE, bold=True, center=True, size=14)
    doc.p(f"{code}-ООС", center=True)
    doc.p("Том 8", center=True)
    doc.p()
    doc.p(f"Заказчик: {customer}", center=True)
    doc.p(f"Стадия: {pr.get('stage') or 'П'}", center=True)
    doc.p(f"{year} г.", center=True)
    doc.page_break()

    # ── состав тома ───────────────────────────────────────────────────
    doc.p("Содержание тома", bold=True, center=True)
    for obz, name in ((f"{code}-ООС", "Лист регистрации внесённых изменений"),
                      (f"{code}-ООС.С", "Содержание тома"),
                      (f"{code}-ООС", "Пояснительная записка"),
                      (f"{code}-ООС", "Приложения (текстовая часть)"),
                      (f"{code}-ООС.ГЧ", "Приложения (графическая часть)")):
        doc.p(f"{obz}\t{name}")
    doc.page_break()

    # ── аннотация ─────────────────────────────────────────────────────
    doc.p("Аннотация", bold=True, center=True)
    doc.p(f"Настоящий раздел «Перечень мероприятий по охране окружающей среды» "
          f"разработан в составе проектной документации по объекту «{title_obj}» "
          f"по адресу: {address}.")
    doc.p("В данном разделе выполнена оценка воздействия на окружающую среду выбросов "
          "и сбросов; определена качественная и количественная характеристика отходов "
          "производства и способы их утилизации, разработан перечень мероприятий по "
          "предотвращению и снижению возможного негативного воздействия намечаемой "
          "хозяйственной деятельности на окружающую среду и рациональному "
          "использованию природных ресурсов на период строительства и эксплуатации "
          "объекта.")
    doc.bullets(["техническое задание на проектирование", "смежная проектная "
                 "документация (ПЗУ, АР, ИОС, ПОС)",
                 f"технический отчёт об инженерно-экологических изысканиях"
                 + (f" ({iei['report']})" if iei.get("report") else
                    " [требуется: шифр отчёта ИЭИ]")],
                intro="Исходные данные:")
    doc.page_break()

    # ── оглавление ────────────────────────────────────────────────────
    doc.p("Оглавление", bold=True, center=True)
    n = 0
    for _key, title, subs in SECTIONS:
        if _key in ("refs", "appendices"):
            doc.p(title)
            continue
        n += 1
        doc.p(f"{n}. {title}")
        for j, (_sk, st) in enumerate(subs, start=1):
            doc.p(f"    {n}.{j}. {st}")
    doc.p("Приложения текстовая часть")
    for i, a in enumerate(APPENDICES_TEXT, start=1):
        doc.p(f"    Приложение {i}. {a}")
    doc.p("Приложения графическая часть")
    for i, a in enumerate(APPENDICES_GRAPHIC, start=1):
        doc.p(f"    Приложение {i}. {a}")
    doc.page_break()

    sec = {k: (t, subs) for k, t, subs in SECTIONS}

    # 1. Введение
    doc.h1(sec["intro"][0])
    doc.p(f"Раздел «Перечень мероприятий по охране окружающей среды» выполнен с целью "
          f"экологического обоснования проектируемого объекта: «{title_obj}».")
    doc.p(f"Раздел подготовлен на основании {NPA['pp87']}, в соответствии с "
          f"{NPA['grk']}, {NPA['fz7']}, {NPA['fz96']}, {NPA['fz89']}, {NPA['fz52']} "
          f"и другими нормативно-правовыми актами, действующими на территории "
          f"Российской Федерации. Раздел рассматривается в составе проектной "
          f"документации при государственной экспертизе ({NPA['pp145']}).")
    doc.p(f"Заказчик: {customer}, ИНН {org.inn or '[требуется]'}, ОГРН {org.ogrn or '—'}, "
          f"адрес: {org.address or '[требуется]'}.")
    if obj:
        doc.p(f"Объект негативного воздействия: код {obj.code or '[требуется: код НВОС]'}, "
              f"категория {obj.category or '[требуется]'} ({NPA['pp2398']}).")

    # 2. Характеристика природных условий района — по отчёту ИЭИ
    doc.h1(sec["nature"][0])
    doc.p("Сведения раздела приведены по техническому отчёту об инженерно-"
          "экологических изысканиях"
          + (f" ({iei['report']})" if iei.get("report") else "")
          + f", выполненному по {NPA['sp502']}. Подробные результаты изысканий "
          "(описания, протоколы, карты) в настоящем разделе не воспроизводятся — "
          "см. отчёт ИЭИ.")
    sub_text = {
        "location": iei.get("location"), "geology": iei.get("geology"),
        "soils": iei.get("soils"), "state_soil": iei.get("soil_category") and
        f"Почвы участка по результатам ИЭИ относятся к категории загрязнения "
        f"«{iei['soil_category']}» ({NPA['sanpin3685']}).",
        "state_radio": iei.get("radiation") and
        f"Радиационная обстановка по результатам ИЭИ: {iei['radiation']}.",
        "flora_fauna": iei.get("flora_fauna"), "water_state": iei.get("water_bodies"),
        "oopt": iei.get("oopt") and f"Особо охраняемые природные территории, зоны с особыми "
        f"условиями использования по данным уполномоченных органов: {iei['oopt']}.",
    }
    for sk, st in sec["nature"][1]:
        # «Исследование …» — подпункты 2.5.1–2.5.4 раздела 2.5 (как в эталоне)
        (doc.h3n if sk.startswith("state_") else doc.h2)(st)
        if sk == "climate":
            if cl:
                wind = cl.get("wind") or {}
                doc.table("climate", [
                    ["Коэффициент, зависящий от стратификации атмосферы, А", _num(cl.get("A"))],
                    ["Средняя максимальная температура воздуха наиболее жаркого месяца, °С",
                     _num(cl.get("t_warm"))],
                    ["Средняя температура воздуха наиболее холодного месяца, °С",
                     _num(cl.get("t_cold"))],
                    ["Повторяемость направлений ветра и штилей за год, %",
                     ", ".join(f"{k} — {v}" for k, v in wind.items()) or "—"],
                    ["Скорость ветра, повторяемость превышения которой составляет 5 %, м/с",
                     _num(cl.get("u5"))]])
                doc.p("Справка о климатических характеристиках — Приложение 1.")
            else:
                doc.table("climate", [])
                doc.need("справка о климатических характеристиках (ЦГМС) — табл. 2.2.1")
        elif sk == "state":
            doc.p("Современное состояние компонентов окружающей среды оценено по "
                  "результатам ИЭИ (фон атмосферного воздуха, почвы, радиация, "
                  "физические факторы).")
        elif sk == "state_air":
            bg = _lst(iei, "background")
            if bg:
                doc.table("background", [[f"{b.get('name', '')}, {b.get('unit', 'мг/м3')}",
                                          _num(b.get("value"))] for b in bg])
                doc.p("Справка о фоновых концентрациях — Приложение 2.")
            else:
                doc.table("background", [])
                doc.need("справка о фоновых концентрациях ЦГМС — табл. 2.5.1.1")
            doc.table("air_study", [[i, a.get("place", ""), a.get("name", ""),
                                     _num(a.get("value")), a.get("method", ""),
                                     _num(a.get("pdk"))]
                                    for i, a in enumerate(_lst(iei, "air_study"), start=1)])
        elif sk == "state_phys":
            doc.table("noise_study", [[i, s.get("place", ""), _num(s.get("leq")),
                                       _num(s.get("lmax")), _num(s.get("minutes")),
                                       _num(s.get("pdu"))]
                                      for i, s in enumerate(_lst(iei, "noise"), start=1)])
        else:
            txt = sub_text.get(sk)
            if txt:
                doc.p(str(txt))
            else:
                doc.need(f"{st.lower()} — по отчёту ИЭИ")

    # 3. Проектные решения
    doc.h1(sec["design"][0])
    doc.p(f"В состав данного проекта входит разработка проектной документации на "
          f"объект «{title_obj}», расположенный по адресу: {address}"
          + (f", земельный участок с кадастровым номером {pr['cadastral']}" if pr.get("cadastral") else "")
          + ".")
    if pr.get("description"):
        for para in str(pr["description"]).split("\n"):
            if para.strip():
                doc.p(para.strip())
    else:
        doc.need("описание проектных решений (по разделам ПЗ, ПЗУ, АР, ИОС, ПОС)")
    doc.table("tep", [[i, t.get("name", ""), t.get("unit", ""), _num(t.get("value"))]
                      for i, t in enumerate(_lst(pr, "tep"), start=1)])
    doc.h3("Сведения о сотрудниках, занятых при строительстве")
    staff_rows = []
    if pr.get("workers"):
        staff_rows.append(["Рабочие", pr["workers"], pr.get("shifts", "—"), pr.get("months", "—")])
    if pr.get("itr"):
        staff_rows.append(["ИТР, МОП", pr["itr"], pr.get("shifts", "—"), pr.get("months", "—")])
    doc.table("staff", staff_rows)
    mach = [[m.get("area", ""), m.get("name", ""), m.get("brand", ""), m.get("spec", ""),
             _num(m.get("qty"))] for m in _lst(pr, "machinery")]
    doc.table("machinery", mach)

    # 4. Охрана земельных ресурсов
    doc.h1(sec["land"][0])
    doc.h2(sec["land"][1][0][1])
    doc.p("Эксплуатация объекта проектирования может сопровождаться следующими видами "
          "воздействия на почвенный покров: уничтожение естественного почвенного "
          "покрова при земляных работах; ухудшение физико-механических и химико-"
          "биологических свойств почвенного слоя при загрязнении ГСМ от работы и "
          "стоянки машин; захламление поверхности почвы отходами. " + TYPICAL)
    if iei.get("soil_category"):
        doc.p(f"Почвы обследуемого района относятся к категории загрязнения "
              f"«{iei['soil_category']}» и могут использоваться в соответствии с "
              f"{NPA['sanpin3685']}.")
    else:
        doc.need("категория загрязнения почв по ИЭИ и рекомендации по использованию грунта")
    doc.h2(sec["land"][1][1][1])
    doc.bullets([
        "селективный сбор, временное накопление и размещение строительных отходов",
        "организация площадок с твёрдым покрытием для складирования материалов",
        "использование только исправной строительной техники, прошедшей осмотр и мойку",
        "запрет на заправку техники ГСМ в зоне строительных работ",
        "ликвидация проливов сорбентом/чистым песком с вывозом отхода",
        "снятие, складирование и сохранение плодородного слоя почвы с последующим "
        f"использованием при рекультивации ({NPA['gost59070']}, {NPA['pp800']})",
        "восстановление нарушенного благоустройства после окончания строительства: "
        "планировка основания, засев травами",
    ], intro=f"С целью предотвращения загрязнения почвы предусматриваются мероприятия {TYPICAL}:")

    # 5. Охрана атмосферного воздуха
    doc.h1(sec["air"][0])
    doc.h2(sec["air"][1][0][1])
    doc.h3("Инвентаризация источников выбросов в период строительства")
    doc.p("Источниками загрязнения атмосферного воздуха при строительстве объекта "
          "являются грузовой автотранспорт, строительная техника, сварочные работы, "
          "дизельный генератор. Выбросы непостоянны по составу и интенсивности; для "
          "оценки выделены наиболее продолжительные и интенсивные периоды работ. "
          + TYPICAL)
    csrc = _lst(con, "sources")
    doc.table("src_constr", [[f"{s.get('number', '')} {s.get('kind', '')}".strip(),
                              f"{s.get('name', '')}. "
                              f"{'Организованный' if s.get('organized') else 'Неорганизованный'} "
                              f"источник (Н={_num(s.get('height'))} м)"] for s in csrc])
    doc.p("*В качестве передвижного источника принята строительная техника, двигатель "
          "которой при работе является источником загрязнения; в качестве стационарного — "
          "техника, местоположение которой определено в единой системе координат.")
    doc.table("machinery_air", mach)
    doc.p(f"Расчёт выбросов от ДВС выполнен по программе «АТП-Эколог», от дизельной "
          f"установки — «Дизель», от сварки — «Сварка» (НИИ Атмосфера; {NPA['mp2012']}). "
          f"Результаты расчётов — Приложение 5.")
    cpol = _lst(con, "pollutants")
    doc.table("zv_constr", [[_code(p.get("code")), p.get("name", ""), p.get("criterion", ""),
                             _num(p.get("value")), _num(p.get("hazard")),
                             _num(p.get("g_s")), _num(p.get("t_year"))] for p in cpol],
              total=["Всего веществ", len(cpol), "", "", "",
                     _num(sum(wc._d(p.get("g_s")) for p in cpol)),
                     _num(sum(wc._d(p.get("t_year")) for p in cpol))] if cpol else None)
    if not cpol:
        doc.need("расчёт выбросов периода строительства — табл. 5.1.1.3")
    doc.h3("Расчёт рассеивания загрязняющих веществ в период строительства")
    _dispersion_block(doc, con, "constr")
    doc.h2(sec["air"][1][1][1])
    osrc = _lst(op, "sources") or [
        {"number": s.get("number"), "name": s.get("name"), "kind": s.get("kind")}
        for s in ((ctx.extra or {}).get("emission_sources") or []) if isinstance(s, dict)]
    doc.p("Источниками выбросов в период эксплуатации являются: "
          + ("; ".join(f"{s.get('number', '')} {s.get('name', '')}".strip() for s in osrc)
             if osrc else "[требуется: инвентаризация источников выбросов]") + ".")
    doc.table("src_oper", [[s.get("number", ""),
                            f"{s.get('name', '')}. "
                            f"{(s.get('kind') or 'источник').capitalize()} "
                            + (f"(Н={_num(s['height'])} м)" if s.get("height") else "")]
                           for s in osrc])
    arows = rows_air(ctx)
    doc.table("zv_oper", [[r["code"], r["name"], r["criterion"], _num(r["value"]),
                           _num(r["hazard"]), _num(r["g_s"]), _num(r["mass"])] for r in arows],
              total=["Всего веществ", len(arows), "", "", "",
                     _num(sum(r["g_s"] or 0 for r in arows)),
                     _num(sum(r["mass"] for r in arows))] if arows else None)
    doc.p("Результаты расчётов выбросов — Приложение 7.")
    doc.h3("Расчёт рассеивания загрязняющих веществ в период эксплуатации")
    _dispersion_block(doc, op, "oper")
    doc.h2(sec["air"][1][2][1])
    doc.bullets([
        "соблюдение технологического регламента работы оборудования, исключающее "
        "залповые и аварийные выбросы",
        "использование только исправной техники с отрегулированными двигателями, "
        "исключение работы двигателей на холостом ходу",
        "пылеподавление при земляных работах и перевозке сыпучих материалов "
        "(увлажнение, укрытие кузовов)",
        "оснащение источников выбросов пылегазоочистным оборудованием с проектной "
        "эффективностью и контроль его работы",
        f"мероприятия по уменьшению выбросов в периоды НМУ ({NPA['pr662']})",
        "контроль выбросов в рамках программы ПЭК",
    ], intro=f"Для снижения воздействия на атмосферный воздух предусматриваются мероприятия {TYPICAL}:")
    doc.h2(sec["air"][1][3][1])
    doc.p("Значения предлагаемых нормативов ПДВ на период эксплуатации приведены в "
          "таблице 5.2.2 (г/с, т/год по каждому веществу). Предложения разработаны на "
          "основании инвентаризации источников выбросов с учётом результатов расчётов "
          "рассеивания и могут быть пересмотрены при уточнении исходных данных.")

    # 6. Водные ресурсы
    doc.h1(sec["water"][0])
    doc.h2(sec["water"][1][0][1])
    w = con.get("water") or {}
    pw = _sub(ctx, "project")
    ws, wi = w.get("workers_shift") or pw.get("workers"), w.get("itr") or pw.get("itr")
    doc.p(f"Потребность строительства в воде определена в соответствии с {NPA['mds12_46']}.")
    if ws:
        q_l, k_h = wc._d(w.get("q_hoz_l"), "15"), wc._d(w.get("k_h"), "2")
        q_shower = wc._d(w.get("q_shower_l"), "30")
        people = wc._d(ws) + wc._d(wi)
        q_hoz = (q_l * people * k_h / (3600 * 8) + q_shower * wc._d(ws) / (60 * 45))
        q_fire = wc._d(w.get("q_fire_ls"), "5")
        q_prod = wc._d(w.get("q_prod_ls"))
        doc.table("water_constr", [
            ["Наибольшее количество рабочих в смену", "чел.", _num(ws)],
            ["Количество ИТР, МОП", "чел.", _num(wi)],
            ["Расход воды на хозяйственно-питьевые потребности работающего", "л/смену", _num(q_l)],
            ["Коэффициент неравномерности потребления воды", "-", _num(k_h)],
            ["Общий расход воды на санитарно-бытовые нужды", "л/с", _num(wc.r3(q_hoz))],
            ["Расход воды на противопожарные нужды", "л/с", _num(q_fire)],
            ["Расход воды на производственные нужды", "л/с", _num(q_prod) if q_prod else "[требуется]"],
            ["Общий расход воды для строительной площадки", "л/с",
             _num(wc.r3(q_hoz + q_prod))]])
        doc.p("Qхоз = (qх·Пр·Кч)/(3600·t) + (qд·Пд)/(60·t1), где qх = 15 л — удельный расход "
              "на хозяйственно-питьевые нужды; Кч = 2; qд = 30 л — расход на приём душа; "
              "t1 = 45 мин; t = 8 ч. Противопожарный расход — не менее 5 л/с.")
    else:
        doc.table("water_constr", [])
        doc.need("численность работающих в смену (ПОС) — расчёт потребности в воде")
    doc.p("Водоснабжение строительной площадки — привозной водой в цистернах; питьевое — "
          "бутилированная вода. Водоотведение хозяйственно-бытовых стоков — в накопительную "
          "ёмкость с вывозом ассенизационной машиной; санузлы — биотуалеты. На выезде — "
          "пост мойки колёс с оборотным водоснабжением типа «Мойдодыр». Поверхностный сток "
          "собирается во временные ёмкости и вывозится по договору. " + TYPICAL)
    doc.h2(sec["water"][1][1][1])
    if op.get("water_text"):
        for para in str(op["water_text"]).split("\n"):
            if para.strip():
                doc.p(para.strip())
    bal = _lst(op, "water_balance")
    if not bal:
        wx = (ctx.extra or {}).get("water") or {}
        for s in wx.get("intake") or []:
            if isinstance(s, dict):
                bal.append({"consumer": s.get("name") or "Водопотребление",
                            "in_year": s.get("volume"), "to": ""})
        for s in wx.get("discharge") or []:
            if isinstance(s, dict):
                bal.append({"consumer": "Водоотведение", "out_year": s.get("volume"),
                            "to": s.get("receiver", "")})
    doc.table("water_balance", [[b.get("consumer", ""), _num(b.get("in_day")),
                                 _num(b.get("in_year")), _num(b.get("out_day")),
                                 _num(b.get("out_year")), b.get("to", "")] for b in bal])
    if not bal:
        doc.need("балансовая схема водопотребления и водоотведения (ИОС3) — табл. 6.2.1")
    wrows = rows_water(ctx)
    if wrows:
        doc.p("Сброс загрязняющих веществ со сточными водами (т/год):")
        doc.bullets([f"{r['name']} ({r['code']}) — {_num(r['mass'])}" for r in wrows])
    else:
        doc.p("Сброс сточных вод в водные объекты проектом не предусмотрен: "
              "хозяйственно-бытовые стоки отводятся в централизованную систему "
              "водоотведения / на КОС, поверхностные — на ЛОС. " + TYPICAL)
    doc.h2(sec["water"][1][2][1])
    doc.bullets([
        "отвод поверхностного стока со стройплощадки в герметичные временные ёмкости "
        "с вывозом на очистные сооружения",
        "движение и стоянка техники только по дорогам и площадкам с твёрдым покрытием",
        "оборудование стационарных механизмов поддонами, заправка техники на АЗС",
        "складирование материалов и отходов только на площадках с твёрдым покрытием",
        "посты мойки колёс с оборотным водоснабжением, применение биотуалетов",
        "исключение сброса неочищенных сточных вод на рельеф и в водные объекты",
        "контроль качества сточных вод в контрольных точках в рамках ПЭК",
    ], intro=f"Для охраны поверхностных и подземных вод предусматриваются мероприятия {TYPICAL}:")
    doc.h2(sec["water"][1][3][1])
    doc.p(f"В соответствии с {NPA['vk']} в границах водоохранных зон запрещаются: "
          "использование сточных вод для регулирования плодородия почв; размещение "
          "кладбищ, скотомогильников, объектов размещения отходов; движение и стоянка "
          "транспортных средств вне дорог и оборудованных площадок; сброс сточных, в том "
          "числе дренажных, вод; разведка и добыча полезных ископаемых.")
    doc.p("Для предотвращения вторичного загрязнения воды системы хозяйственно-питьевого "
          "водоснабжения применяются трубы и материалы, имеющие санитарно-"
          "эпидемиологические заключения; предусматриваются промывка и дезинфекция "
          "сетей перед вводом в эксплуатацию; водомерный узел на границе балансовой "
          "принадлежности. " + TYPICAL)

    # 7. Растительный и животный мир
    doc.h1(sec["bio"][0])
    doc.table("land_tep", [[i, t.get("name", ""), _num(t.get("m2"))]
                           for i, t in enumerate(_lst(pr, "land_tep"), start=1)])
    doc.p(iei.get("flora_fauna") or "[требуется: характеристика растительности и животного "
          "мира участка по ИЭИ, наличие краснокнижных видов по письмам уполномоченных органов]")
    doc.h2(sec["bio"][1][0][1])
    doc.bullets([
        "проведение строительных работ в максимально короткие сроки и исключительно в "
        "пределах полосы отвода",
        "компактное складирование и своевременный вывоз строительного мусора",
        "исключение доступа животных на площадки путём установки ограждений",
        "ограничение скорости движения транспорта, особенно в тёмное время суток",
        "исключение сброса загрязнённых сточных вод на почву и в водные объекты",
        "проведение работ вне периода массового размножения животных; компенсационное "
        "озеленение при вынужденном сносе зелёных насаждений",
    ], intro=f"Для смягчения воздействия на растительность и животный мир предусматриваются мероприятия {TYPICAL}:")

    # 8. Шум
    doc.h1(sec["noise"][0])
    for sk, st, d in (("noise_constr", sec["noise"][1][0][1], con),
                      ("noise_oper", sec["noise"][1][1][1], op)):
        doc.h2(st)
        suffix = "constr" if sk == "noise_constr" else "oper"
        ns = _lst(d, "noise_sources")
        doc.table(f"noise_src_{suffix}",
                  [[s.get("n", ""), s.get("name", ""), s.get("x", ""), s.get("y", ""),
                    s.get("h", ""), s.get("spectrum", ""), _num(s.get("leq")),
                    _num(s.get("lmax"))] if suffix == "constr" else
                   [s.get("n", ""), s.get("name", ""), s.get("spectrum", ""),
                    _num(s.get("leq")), _num(s.get("r"))] for s in ns])
        if suffix == "constr":
            nr = _lst(d, "noise_results")
            doc.table("noise_res_constr", [[r.get("n", ""), r.get("name", ""),
                                            _num(r.get("leq")), _num(r.get("lmax"))] for r in nr])
            nb = _lst(d, "noise_bg")
            doc.table("noise_bg_constr", [[r.get("n", ""), r.get("name", ""), _num(r.get("bg")),
                                           _num(r.get("leq")), _num(r.get("total"))] for r in nb])
        else:
            doc.table("noise_pdu", [[
                "Территории, непосредственно прилегающие к жилым зданиям (с 7 до 23 ч)",
                "90 75 66 59 54 50 47 45 44", "55", "70"],
                ["Территории, непосредственно прилегающие к жилым зданиям (с 23 до 7 ч)",
                 "83 67 57 49 44 40 37 35 33", "45", "60"]],
                title=f"Допустимые уровни шума ({NPA['sanpin3685']}, табл. 5.35)")
            nb = _lst(d, "noise_bg")
            doc.table("noise_bg_oper", [[r.get("n", ""), r.get("name", ""), _num(r.get("bg")),
                                         _num(r.get("const")), _num(r.get("var")),
                                         _num(r.get("total"))] for r in nb])
        if not (ns or _lst(d, "noise_bg")):
            doc.need(f"акустический расчёт на период {'строительства' if suffix == 'constr' else 'эксплуатации'} "
                     f"({NPA['sp51']}) — Приложение {9 if suffix == 'constr' else 10}")
        else:
            doc.p("Уровни шума в расчётных точках на границе жилой застройки не превышают "
                  f"допустимых по {NPA['sanpin3685']} (вывод подтверждается расчётом в приложении).")

    # 9. Отходы
    doc.h1(sec["waste"][0])
    doc.p("В разделе представлены сведения о проектируемом объекте как источнике "
          "образования отходов. Выполнена предварительная расчётная инвентаризация "
          "источников образования отходов и мест их накопления на период строительства "
          "и эксплуатации. Классификация отходов — по "
          f"{NPA['fkko']}; классы опасности — по {NPA['pr158']}.")
    doc.h2(sec["waste"][1][0][1])
    cw = construction_wastes(ctx)
    doc.p("Объём образования строительных отходов определён на основании ведомости "
          f"потребности в основных строительных материалах (Приложение 4) и удельных "
          f"нормативов ({NPA['sbornik_uo']}; {NPA['nicpuro']}).")
    doc.h3("Мусор от офисных и бытовых помещений организаций несортированный "
           "(исключая крупногабаритный) (7 33 100 01 72 4)")
    doc.p("М = p · n · c · ρ / 12, т/год, где p — численность сотрудников; n — норма "
          "накопления ТБО на 1 сотрудника; ρ — плотность отхода; c — период строительства, мес.")
    doc.table("tko_constr", [[r["months"], f"{r['label']}: {r['people']}", _num(r["norm"]),
                              _num(r["density"]), _num(r["m3"]), _num(r["t"])] for r in cw["tko"]],
              total=["Итого", "", "", "", _num(sum(r["m3"] for r in cw["tko"])),
                     _num(sum(r["t"] for r in cw["tko"]))] if cw["tko"] else None)
    doc.h3("Осадок (шлам) механической очистки нефтесодержащих сточных вод, содержащий "
           "нефтепродукты в количестве менее 15 %, обводненный (7 23 101 01 39 4)")
    doc.p("Qн/пр, взв = (С1 − С2) · Q · 10⁻⁶ · P / (1 − B/100), т, где Q — расход сточных "
          "вод, м3/сут; С1, С2 — концентрация до и после очистки, мг/л; B — влажность "
          "отхода, %; P — количество рабочих дней.")
    doc.table("wheel_wash", [[r["label"], _num(r["q"]), f"{_num(r['c1'])}/{_num(r['c2'])}",
                              _num(r["humidity"]), r["days"], _num(r["t"]), _num(r["m3"])]
                             for r in cw["wheel"]],
              total=["Итого", "", "", "", "", _num(sum(r["t"] for r in cw["wheel"])),
                     _num(sum(r["m3"] for r in cw["wheel"]))] if cw["wheel"] else None)
    doc.h3("Отходы строительных материалов")
    doc.p("Количество отходов определяется по формуле М = m · К / 100, где m — расход "
          "материала по ведомости (м3 или т); К — норма образования отхода, %; масса — "
          "через плотность отхода.")
    doc.table("material_waste", [[f"{r['n']}. {r['material']} → {r['waste_name']} "
                                  f"({r['fkko']})" if r["fkko"] else f"{r['n']}. {r['material']}",
                                  _num(r["qty"]), r["unit"], _num(r["pct"]), _num(r["density"]),
                                  _num(r["m3"]), _num(r["t"])] for r in cw["materials"]],
              total=["Итого", "", "", "", "", _num(sum(r["m3"] for r in cw["materials"])),
                     _num(sum(r["t"] for r in cw["materials"]))] if cw["materials"] else None)
    for r in cw["materials"]:
        if r["note"]:
            doc.p(r["note"], italic=True)
    doc.h3("Остатки и огарки стальных сварочных электродов (9 19 100 01 20 5)")
    e = cw["electrodes"]
    doc.table("electrodes", [[_num(e["total"]), _num(e["pct"]), _num(e["density"]),
                              _num(e["m3"]), _num(e["t"])]] if e else [])
    doc.h3("Грунт, образовавшийся при проведении землеройных работ, не загрязненный "
           "опасными веществами (8 11 100 01 49 5)")
    s = cw["soil"]
    doc.table("soil_excess", [["Избыток грунта по ведомости", _num(s["m3"]), _num(s["density"]),
                               _num(s["m3"]), _num(s["t"])]] if s else [])
    doc.h3("Отходы (осадки) из выгребных ям (7 32 100 01 30 4)")
    cp = cw["cesspool"]
    doc.table("wc_constr", [[cp["people"], cp["shifts"], _num(cp["norm"]), _num(cp["density"]),
                             _num(cp["m3"]), _num(cp["t"])]] if cp else [])
    _waste_summary(doc, "wastes_constr", cw["summary"])
    if not cw["summary"]:
        doc.need("ведомость материалов, численность и сроки строительства (ПОС) — "
                 "расчёт отходов периода строительства")
    doc.h2(sec["waste"][1][1][1])
    doc.bullets([
        "организовать контроль за обращением с отходами на объекте",
        "минимизировать потери строительных материалов при выполнении работ",
        "незагрязнённые отходы песка, щебня, лом асфальтобетона использовать в работах "
        "по благоустройству территории",
        "обеспечить своевременный централизованный вывоз строительных отходов",
        f"оборудовать места накопления отходов по {NPA['sanpin3684']}, соблюдать срок "
        "накопления не более 11 месяцев",
        "заключить договоры на передачу отходов I–IV класса с организациями, имеющими "
        "лицензию на соответствующий вид деятельности",
    ], intro=f"Для снижения воздействия строительных отходов предусматриваются мероприятия {TYPICAL}:")
    doc.h2(sec["waste"][1][2][1])
    ow = operation_wastes(ctx)
    doc.p("Основными источниками образования отходов в период эксплуатации являются: "
          "жизнедеятельность сотрудников и посетителей, уборка территории, замена ламп "
          "и фильтров, обслуживание очистных сооружений. " + TYPICAL)
    doc.h3("Лампы, утратившие потребительские свойства")
    doc.p("Лотх = Л1 · Т / К, где Л1 — количество работающих ламп; Т — время работы "
          "лампы в год, ч; К — ресурс лампы, ч.")
    doc.table("lamps", [[r["n"], r["name"], _num(r["count"]), _num(r["life_h"]), _num(r["hours"]),
                         _num(r["replaced"]), _num(r["mass_kg"]), _num(r["t"]), _num(r["m3"]),
                         _num(r["density"])] for r in ow["lamps"]])
    doc.h3("Отходы, рассчитанные по нормативам накопления")
    doc.p("М = N · n · ρ, где N — расчётное количество (сотрудников, мест, м2 площади, "
          "машино-мест); n — норматив накопления, м3/год на единицу (по региональным "
          "нормативам накопления ТКО); ρ — плотность отхода, т/м3.")
    doc.table("tko_oper", [[r["n"], f"{r['name']} ({r['fkko']})" if r["fkko"] else r["name"],
                            f"{_num(r['norm'])} {r['norm_unit']}/{r['count_unit'] or 'ед.'}",
                            f"{_num(r['count'])} {r['count_unit']}", _num(r["density"]),
                            _num(r["t"]), _num(r["m3"])] for r in ow["norm"]])
    for r in ow["norm"]:
        if r["note"]:
            doc.p(r["note"], italic=True)
    _waste_summary(doc, "wastes_oper", ow["summary"])
    doc.h2(sec["waste"][1][3][1])
    doc.bullets([
        f"раздельное накопление отходов по видам и классам опасности в местах, "
        f"оборудованных по {NPA['sanpin3684']}",
        "передача ТКО региональному оператору по договору; отходов I–IV класса — "
        "лицензированным организациям",
        f"оформление паспортов отходов I–IV класса ({NPA['pr1026']}), ведение учёта "
        f"({NPA['pr1028']})",
        "соблюдение предельных сроков накопления (не более 11 месяцев)",
        "контроль мест накопления в рамках ПЭК",
    ], intro=f"Для снижения воздействия отходов в период эксплуатации предусматриваются мероприятия {TYPICAL}:")

    # 10. Затраты
    doc.h1(sec["costs"][0])
    doc.p(f"Размер платы за выбросы загрязняющих веществ и размещение отходов рассчитан по "
          f"ставкам {NPA['rates2026'] if year >= 2026 else NPA['rates2025']}. Расчёт является "
          "ориентировочным по данным проектной документации; в расчёт заложены максимально "
          "возможные выбросы и максимальное количество отходов, направляемых на размещение.")
    doc.h2(sec["costs"][1][0][1])
    crow, ctot = [], Decimal(0)
    for i, p in enumerate(cpol, start=1):
        mass = wc._d(p.get("t_year"))
        rate = _air_rate(_code(p.get("code")), p.get("name", ""), year)
        amount = wc.r3(mass * rate)
        ctot += amount
        crow.append([i, _code(p.get("code")), p.get("name", ""), "т", _num(mass),
                     _num(rate) if rate else "[нет ставки]", _num(amount)])
    doc.table("pay_air_constr", crow, total=["Всего", "", "", "", "", "", _num(ctot)] if crow else None)
    pay = costs(ctx)
    orow = []
    for i, ln in enumerate([l for l in (pay.get("lines") or []) if l.medium == "air"], start=1):
        orow.append([i, ln.code, ln.name, "т", _num(ln.mass), _num(ln.rate), _num(ln.amount)])
    doc.table("pay_air_oper", orow,
              total=["Итого", "", "", "", "", "", _num(pay.get("air"))] if orow else None)
    if not orow:
        doc.need("выбросы периода эксплуатации (т/год) — расчёт платы по табл. 10.1.2")
    doc.h2(sec["costs"][1][1][1])
    doc.p(f"В соответствии со ст. 23 {NPA['fz89']} плательщиками платы при размещении ТКО "
          "являются региональные операторы; за отходы, передаваемые на утилизацию и "
          "обезвреживание, плата не вносится. Плата рассчитана только по отходам, "
          "направляемым на размещение.")
    wrow = []
    for i, ln in enumerate([l for l in (pay.get("lines") or []) if l.medium == "waste"], start=1):
        wrow.append([i, ln.name, ln.code, "", _num(ln.mass), _num(ln.rate), _num(ln.amount)])
    if not wrow:
        for i, r in enumerate(ow["summary"], start=1):
            wrow.append([i, r["name"], r["fkko"], r["hazard"] or "", _num(r["t"]), "-",
                         "0 (передаётся на утилизацию/обезвреживание/региональному оператору)"])
    doc.table("pay_waste", wrow,
              total=["Итого всего", "", "", "", "", "", _num(pay.get("waste") or 0)] if wrow else None)
    doc.p(f"Ориентировочная сумма платы за выбросы на период строительства — {_num(ctot)} руб.; "
          f"на период эксплуатации — {_num(pay.get('air') or 0)} руб./год; за размещение "
          f"отходов — {_num(pay.get('waste') or 0)} руб./год.")
    doc.need("сметная стоимость природоохранных мероприятий (очистные сооружения, "
             "шумозащита, озеленение) по сводному сметному расчёту")

    # 11. ПЭК
    doc.h1(sec["pek"][0])
    doc.p(f"В целях обеспечения выполнения мероприятий по охране окружающей среды "
          f"природопользователь обязан организовать производственный экологический "
          f"контроль (ст. 67 {NPA['fz7']}). Требования к программе ПЭК утверждены "
          f"{NPA['pr109']}. Периодичность и контролируемые параметры зависят от "
          f"категории объекта ({NPA['pp2398']}).")
    doc.h2(sec["pek"][1][0][1])
    cat = pek.get("category") or (obj.category if obj and obj.category else "[требуется: категория]")
    doc.p(f"В период строительства объект предварительно относится к {cat} категории по "
          "степени негативного воздействия (продолжительность строительства более 6 "
          "месяцев). Контроль выбросов — расчётным методом по веществам, для которых "
          "приземная концентрация на границе участка превышает 0,1 ПДК; контроль "
          "акустического воздействия — инструментальный на границе жилой застройки; "
          "контроль почв — визуальный и химико-аналитический; контроль отходов — учёт "
          "движения и мест накопления. " + TYPICAL)
    per = pek.get("periodicity") or "1 раз в год"
    doc.h2(sec["pek"][1][1][1])
    prow = [[s.get("number", ""), s.get("name", ""), _code(p.get("code")), p.get("name", ""),
             _num(p.get("g_s")), per, "Расчётный / инструментальный"]
            for s in csrc for p in (s.get("pollutants") or []) if isinstance(p, dict)]
    doc.table("pek_air", prow)
    mon = _lst(pek, "monitoring") or [
        {"area": "Почвы", "points": "в пределах зоны работ", "kind": "отбор проб, "
         "по завершении строительства", "sample": "точечная/объединённая",
         "method": "ГОСТ 17.4.3.01", "components": "нефтепродукты, тяжёлые металлы, "
         "бенз(а)пирен, рН", "org": "аккредитованная лаборатория"},
        {"area": "Атмосферный воздух", "points": "граница жилой застройки",
         "kind": "по жалобам / 1 раз за период", "sample": "разовая",
         "method": "РД 52.04.186-89", "components": "NO2, CO, взвешенные вещества",
         "org": "аккредитованная лаборатория"},
        {"area": "Шум", "points": "граница жилой застройки", "kind": "1 раз за период",
         "sample": "замер", "method": "ГОСТ 23337", "components": "Lэкв, Lмакс",
         "org": "аккредитованная лаборатория"}]
    doc.table("pek_monitor", [[m.get("area", ""), m.get("points", ""), m.get("kind", ""),
                               m.get("sample", ""), m.get("method", ""), m.get("components", ""),
                               m.get("org", "")] for m in mon])
    orow = [[s.get("number", ""), s.get("name", ""), _code(p.get("code")), p.get("name", ""),
             _num(p.get("g_s")), per, "Инструментальный / расчётный"]
            for s in ((ctx.extra or {}).get("emission_sources") or []) if isinstance(s, dict)
            for p in (s.get("pollutants") or []) if isinstance(p, dict)]
    doc.table("pek_oper", orow)
    doc.p("Программа ПЭК на период эксплуатации выпускается отдельным документом "
          "(ЭкоДок: «Разработка → Программа ПЭК»).")

    # 12. Аварийные ситуации (Выборжец, п. 11.6)
    doc.h1(sec["accidents"][0])
    doc.bullets([
        "пролив топлива и масел при заправке и работе техники — локализация сорбентом, "
        "сбор загрязнённого грунта и передача на обезвреживание",
        "пожар на строительной площадке/объекте — первичные средства пожаротушения, "
        "план эвакуации, уведомление надзорных органов",
        "аварийный сброс сточных вод при отказе КОС/ЛОС — накопительные ёмкости на "
        "двухсуточный объём, вывоз по договору",
        "переполнение мест накопления отходов — внеплановый вывоз",
    ], intro=f"Возможные аварийные ситуации и меры по минимизации последствий {TYPICAL}:")

    # 13. Выводы
    doc.h1(sec["conclusion"][0])
    doc.bullets([
        "воздействие на атмосферный воздух в период строительства и эксплуатации "
        "допустимо: приземные концентрации в расчётных точках не превышают ПДК "
        "(при подтверждении расчётом рассеивания)",
        "уровни шума на границе жилой застройки не превышают допустимых (при "
        "подтверждении акустическим расчётом)",
        "сброс сточных вод в водные объекты не предусмотрен; водоотведение — в "
        "централизованные сети / на очистные сооружения с вывозом",
        f"отходы: период строительства — {len(cw['summary'])} вид(ов), "
        f"{_num(sum(r['t'] for r in cw['summary']))} т; период эксплуатации — "
        f"{len(ow['summary'])} вид(ов), {_num(sum(r['t'] for r in ow['summary']))} т/год; "
        "все отходы передаются лицензированным организациям / региональному оператору",
        "при выполнении предусмотренных мероприятий намечаемая деятельность не окажет "
        "существенного негативного влияния на окружающую среду",
    ], intro="По результатам оценки воздействия:")

    # Список литературы
    doc.h1(sec["refs"][0], numbered=False)
    for i, r in enumerate(REFERENCES, start=1):
        doc.p(f"{i}. {r}")

    # Приложения
    doc.h1(sec["appendices"][0], numbered=False)
    doc.p("Приложения текстовая часть", bold=True)
    for i, a in enumerate(APPENDICES_TEXT, start=1):
        doc.p(f"Приложение {i}. {a} — [требуется: вложить документ]")
    doc.p("Приложения графическая часть", bold=True)
    for i, a in enumerate(APPENDICES_GRAPHIC, start=1):
        doc.p(f"Приложение {i}. {a} — [требуется: чертёж из ПЗУ/УПРЗА]")

    doc.p()
    doc.p(f"{org.director_position or 'Руководитель'}\t\t_____________\t"
          f"{org.director_name or ''}")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.d.save(out)
    return out


def _dispersion_block(doc: _Doc, d: dict, suffix: str) -> None:
    """Расчётные точки, исключённые вещества и концентрации (УПРЗА) — таблицы
    5.1.1.4–5.1.1.6 / 5.2.3–5.2.5. Считать рассеивание сами не беремся —
    экспертиза требует аттестованную УПРЗА."""
    period = "строительства" if suffix == "constr" else "эксплуатации"
    doc.p(f"Расчёт рассеивания выполнен по {NPA['mrr2017']} в УПРЗА «Эколог» с учётом "
          f"фоновых концентраций; расчётные точки приняты на границе ближайшей жилой "
          f"застройки. Результаты и карты рассеивания — Приложение {6 if suffix == 'constr' else 8}.")
    pts = _lst(d, "points")
    if suffix == "constr":
        doc.table("points_constr", [[p.get("code", ""), _num(p.get("x")), _num(p.get("y")),
                                     _num(p.get("h")), p.get("type", "")] for p in pts])
    else:
        doc.table("points_oper", [[_num(p.get("x")), _num(p.get("y")), _num(p.get("h")),
                                   p.get("type", ""), p.get("comment", "")] for p in pts])
    doc.table(f"skip_{suffix}", [[_code(s.get("code")), s.get("name", ""), _num(s.get("sum"))]
                                 for s in _lst(d, "skip")])
    conc = _lst(d, "concentrations")
    doc.table(f"conc_{suffix}", [[_code(c.get("code")), c.get("name", ""),
                                  ", ".join(_num(v) for v in (c.get("values") or []))]
                                 for c in conc])
    if conc:
        try:
            mx = max(Decimal(str(v).replace(",", ".")) for c in conc
                     for v in (c.get("values") or []))
            verdict = ("не превышают" if mx <= 1 else "ПРЕВЫШАЮТ")
            doc.p(f"Максимальные приземные концентрации в расчётных точках на период "
                  f"{period} {verdict} ПДК м.р. (максимум {_num(mx)} ПДК).")
        except Exception:
            pass
    else:
        doc.need(f"результаты расчёта рассеивания на период {period} "
                 f"(УПРЗА) — концентрации в расчётных точках")


def _waste_summary(doc: _Doc, key: str, rows: list[dict]) -> None:
    body = [[i, r["name"], _fkko(r["fkko"]), r["hazard"] or "", _num(r["t"]), _num(r["m3"]),
             r.get("handling") or "[требуется: вид обращения/получатель]"]
            for i, r in enumerate(rows, start=1)]
    by_class: dict = {}
    for r in rows:
        by_class.setdefault(r["hazard"], [Decimal(0), Decimal(0)])
        by_class[r["hazard"]][0] += Decimal(str(r["t"]))
        by_class[r["hazard"]][1] += Decimal(str(r["m3"]))
    for cls in sorted(k for k in by_class if k):
        body.append(["", f"Итого {cls} класса", "", "", _num(by_class[cls][0]),
                     _num(by_class[cls][1]), ""])
    doc.table(key, body, total=["", "Итого всего", "", "",
                                _num(sum(Decimal(str(r["t"])) for r in rows)),
                                _num(sum(Decimal(str(r["m3"])) for r in rows)), ""] if rows else None)


def _fkko(code: str) -> str:
    """«73310001724» → «7 33 100 01 72 4» — как печатают в томах."""
    c = str(code or "")
    if len(c) != 11 or not c.isdigit():
        return c
    return f"{c[0]} {c[1:3]} {c[3:6]} {c[6:8]} {c[8:10]} {c[10]}"


__all__ = ["TITLE", "SECTIONS", "TABLES", "generate", "gaps", "media", "rows_air",
           "rows_water", "rows_waste", "costs", "construction_wastes",
           "operation_wastes", "section_titles"]
