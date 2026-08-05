"""Оценка грунта двумя путями: как ОТХОДА и как ПОЧВЫ.

Требование ТЗ (ИЭИ): «класс грунта по МПР и по СанПиН». Это два разных
норматива, и они дают разные ответы на разные вопросы:

  1. Как отход — класс опасности I–V по критериям Минприроды (действует
     приказ № 158 от 31.03.2025, заменил № 536): K = Σ Ci/Wi.
     Нужен для передачи грунта на размещение/утилизацию (ФККО 8 11 100 01 49 5
     и близкие). Считает ecodoc.development.hazard_class.

  2. Как почва — категория загрязнения по СанПиН 1.2.3685-21 и
     МУ 2.1.7.730-99: «чистая» / «допустимая» / «умеренно опасная» /
     «опасная» / «чрезвычайно опасная». Определяет, куда грунт можно
     использовать (рекультивация, подсыпка, только полигон).
     Показатели: кратность превышения ПДК/ОДК и суммарный показатель
     загрязнения Zc = Σ(Ci/Cфi) − (n − 1) по веществам с Ci > Cф.

Справочные ПДК/ОДК ниже — из СанПиН 1.2.3685-21 (разд. III, почвы);
ОДК валовых форм металлов даны для суглинистых/глинистых почв с pH > 5,5 —
при другой почве пользователь правит норматив в форме. Фон обязателен для Zc;
если фона нет, Zc не считается (и об этом говорится прямо).
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

# вещество → (ПДК/ОДК валовая, мг/кг) — СанПиН 1.2.3685-21, почвы
SOIL_NORMS: dict[str, float] = {
    "свинец": 130.0,          # ОДК (сугл., pH > 5,5)
    "кадмий": 2.0,            # ОДК
    "цинк": 220.0,            # ОДК
    "медь": 132.0,            # ОДК
    "никель": 80.0,           # ОДК
    "мышьяк": 10.0,           # ОДК
    "ртуть": 2.1,             # ПДК
    "бенз(а)пирен": 0.02,     # ПДК
    "нефтепродукты": 1000.0,  # допустимый уровень (регион. нормативы)
}

CATEGORIES = ["чистая", "допустимая", "умеренно опасная",
              "опасная", "чрезвычайно опасная"]

# рекомендации по использованию — СанПиН 1.2.3685-21, табл. 4.6 (по смыслу)
USE_BY_CATEGORY = {
    "чистая": "использование без ограничений",
    "допустимая": "использование без ограничений, кроме объектов повышенного "
                  "риска; под отсыпку — с контролем",
    "умеренно опасная": "использование под отсыпки котлованов и выемок, с "
                        "перекрытием чистым грунтом не менее 0,2 м",
    "опасная": "ограниченное использование под отсыпки с перекрытием чистым "
               "грунтом не менее 0,5 м; вывоз — на специализированный полигон",
    "чрезвычайно опасная": "вывоз и утилизация на специализированных "
                           "полигонах — использование запрещено",
}


@dataclass
class SoilComponent:
    name: str
    ci: float                 # измеренная концентрация, мг/кг
    norm: float = 0.0         # ПДК/ОДК, мг/кг (0 — взять из справочника)
    background: float = 0.0   # фоновая концентрация Cф, мг/кг (для Zc)
    wi: float = 0.0           # Wi по Критериям (пр. № 158) — для класса отхода


@dataclass
class SoilResult:
    category: str = ""            # категория загрязнения почвы (СанПиН)
    zc: float | None = None       # суммарный показатель загрязнения
    exceedances: list = field(default_factory=list)   # [(имя, крат. ПДК)]
    hazard_class: int = 0         # класс отхода по Критериям (0 — не считался)
    k_total: float = 0.0
    use: str = ""                 # что можно делать с таким грунтом
    warnings: list = field(default_factory=list)


def _norm_for(c: SoilComponent) -> float:
    if c.norm > 0:
        return c.norm
    return SOIL_NORMS.get(c.name.strip().lower(), 0.0)


def assess(components: list[SoilComponent]) -> SoilResult:
    """Категория почвы по СанПиН + (если заданы Wi) класс отхода (пр. № 158)."""
    res = SoilResult()

    # ── кратности превышения ПДК/ОДК ─────────────────────────────────
    max_ratio = 0.0
    for c in components:
        norm = _norm_for(c)
        if norm <= 0:
            res.warnings.append(f"{c.name}: нет ПДК/ОДК — вещество не участвует "
                                "в оценке по СанПиН (задайте норматив)")
            continue
        ratio = c.ci / norm
        if ratio > 1:
            res.exceedances.append((c.name, round(ratio, 2)))
        max_ratio = max(max_ratio, ratio)

    # ── Zc по фону ────────────────────────────────────────────────────
    over_bg = [(c.ci / c.background) for c in components
               if c.background > 0 and c.ci > c.background]
    if any(c.background > 0 for c in components):
        if over_bg:
            res.zc = round(sum(over_bg) - (len(over_bg) - 1), 2)
        else:
            res.zc = 0.0
    else:
        res.warnings.append("фоновые концентрации не заданы — Zc не считается, "
                            "категория оценена только по кратности ПДК/ОДК")

    # ── категория (МУ 2.1.7.730-99, СанПиН 1.2.3685-21) ──────────────
    zc = res.zc or 0.0
    if zc >= 128:
        res.category = "чрезвычайно опасная"
    elif zc >= 32 or max_ratio > 5:
        res.category = "опасная"
    elif zc >= 16 or max_ratio > 1:
        res.category = "умеренно опасная"
    elif zc > 1:
        res.category = "допустимая"       # выше фона, но в пределах ПДК/ОДК
    elif res.zc is not None:
        res.category = "чистая"           # фон задан, превышений нет
    else:
        res.category = "допустимая"       # без фона «чистую» не присваиваем
    res.use = USE_BY_CATEGORY.get(res.category, "")

    # ── класс отхода по Критериям № 158 (если есть Wi) ───────────────
    with_wi = [c for c in components if c.wi > 0]
    if with_wi:
        from ecodoc.development.hazard_class import Component, calculate
        total = sum(c.ci for c in components)
        rest = 1_000_000 - total
        comps = [Component(c.name, c.ci, c.wi) for c in with_wi]
        if rest > 0:
            # природная минеральная основа грунта — практически неопасна
            comps.append(Component("минеральная основа (по разности)",
                                   rest, 1_000_000))
        r = calculate(comps)
        res.hazard_class = r.hazard_class
        res.k_total = r.k_total
        res.warnings += [w for w in r.warnings if "Сумма концентраций" not in w]
    return res


def generate(components: list[SoilComponent], out_path: str | Path,
             site_label: str = "", org_name: str = "", basis: str = "") -> Path:
    """Заключение по грунту (.docx): обе оценки в одном документе."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.shared import Cm, Pt

    r = assess(components)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for s in doc.sections:
        s.left_margin, s.right_margin = Cm(3), Cm(1.5)
        s.top_margin, s.bottom_margin = Cm(2), Cm(2)

    if org_name:
        p = doc.add_paragraph()
        p.alignment = AL.CENTER
        p.add_run(org_name).bold = True
    t = doc.add_paragraph()
    t.alignment = AL.CENTER
    run = t.add_run("ОЦЕНКА загрязнённости грунта\n"
                    "(СанПиН 1.2.3685-21, МУ 2.1.7.730-99; "
                    "приказ Минприроды России № 158 от 31.03.2025)")
    run.bold = True
    run.font.size = Pt(14)
    if site_label:
        doc.add_paragraph(f"Участок / проба: {site_label}")
    if basis:
        doc.add_paragraph(f"Исходные данные: {basis}")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, text in enumerate(["Вещество", "Ci, мг/кг", "ПДК/ОДК, мг/кг",
                              "Крат. ПДК", "Фон, мг/кг", "Wi"]):
        table.rows[0].cells[i].text = text
    for c in components:
        norm = _norm_for(c)
        cells = table.add_row().cells
        cells[0].text = c.name
        cells[1].text = f"{c.ci:g}"
        cells[2].text = f"{norm:g}" if norm else "—"
        cells[3].text = f"{c.ci / norm:.2f}" if norm else "—"
        cells[4].text = f"{c.background:g}" if c.background else "—"
        cells[5].text = f"{c.wi:g}" if c.wi else "—"

    doc.add_paragraph()
    doc.add_paragraph("1. Оценка как почвы (СанПиН 1.2.3685-21):")
    if r.zc is not None:
        doc.add_paragraph(f"   Zc = Σ(Ci/Cфi) − (n − 1) = {r.zc:g}")
    for name, ratio in r.exceedances:
        doc.add_paragraph(f"   {name}: превышение ПДК/ОДК в {ratio} раза")
    p = doc.add_paragraph()
    p.add_run(f"   КАТЕГОРИЯ ЗАГРЯЗНЕНИЯ: {r.category}").bold = True
    doc.add_paragraph(f"   Использование: {r.use}")

    doc.add_paragraph("2. Оценка как отхода (приказ № 158):")
    if r.hazard_class:
        doc.add_paragraph(f"   K = Σ(Ci/Wi) = {r.k_total:.4g}")
        p = doc.add_paragraph()
        p.add_run(f"   КЛАСС ОПАСНОСТИ ОТХОДА: {r.hazard_class}").bold = True
        if r.hazard_class == 5:
            doc.add_paragraph("   V класс подлежит подтверждению "
                              "биотестированием (2 тест-объекта).")
    else:
        doc.add_paragraph("   Wi компонентов не заданы — класс отхода не "
                          "рассчитан. Для расчёта заполните Wi по приложению "
                          "к приказу № 158 или банку данных БДО.")
    for w in r.warnings:
        doc.add_paragraph(f"⚠ {w}")

    doc.add_paragraph()
    doc.add_paragraph("Оценку выполнил: ____________________ "
                      "/______________________/")

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out
