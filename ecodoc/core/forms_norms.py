"""Сверка бланков-образцов с действующими НПА.

Требование ТЗ: найденную форму «проверить на соответствие нормам в интернете».
Машина не может юридически подтвердить соответствие, но может честно сделать
три проверки и назвать их своими именами:

  1. Структурная: в бланке есть обязательные признаки формы из НПА
     (заголовок формы, ключевые графы). Нет признаков — это не тот бланк.
  2. По датам: бланк старше действующей редакции НПА — предупреждаем,
     что форму могли поменять после того, как бланк был сохранён.
  3. По интернету (online=True): страницы-источники из watch/sources.json
     скачиваются и сравниваются со снимком — изменение страницы = сигнал
     «форму могли обновить, проверьте».

Итог каждой проверки — вердикт с формулировкой, а решение «пользоваться ли
бланком» остаётся за пользователем.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

# код формы → чем она утверждена и по каким признакам узнать её бланк.
# effective — год вступления действующей редакции (0 = не фиксируем: форму
# ведомство обновляет часто, сверять надо по источнику из watch).
NORMS: dict[str, dict] = {
    # Реквизиты выверены сверкой 21–22.08.2026 (каждая форма — против текста
    # НПА и принятого отчёта). switch_date — дата, с которой действует новая
    # форма: бланк, сохранённый раньше неё, сделан ещё по старой редакции.
    "waste-movement": {
        "npa": "приказ Минприроды России от 08.12.2020 № 1028 в ред. от "
               "13.12.2023 № 825; с 01.09.2026 — приказ от 16.04.2026 № 227 "
               "(учёт отходов)",
        "effective": 2024,
        "switch_date": "2026-09-01",
        "marks": ["учет", "отход"],
    },
    "declaration-nvos": {
        "npa": "приказ Минприроды России от 10.12.2020 № 1043 в ред. от "
               "29.04.2025 № 241 (за 2025 г.); с 01.09.2026 — приказ от "
               "01.04.2026 № 182 — декларация о плате за НВОС",
        "effective": 2025,
        "switch_date": "2026-09-01",
        "marks": ["деклараци", "плат"],
    },
    "pek": {
        "npa": "приказ Минприроды России от 15.03.2024 № 173 в ред. от "
               "12.05.2025 № 262 — отчёт об организации и о результатах ПЭК",
        "effective": 2025,
        "marks": ["производственн", "контрол"],
    },
    "2tp-waste": {
        "npa": "приказ Росстата от 06.11.2025 № 614 — форма № 2-ТП (отходы), "
               "с отчёта за 2025 г.",
        "effective": 2026,
        "marks": ["2-тп", "отход"],
    },
    "2tp-air": {
        "npa": "приказ Росстата от 08.11.2018 № 661 — форма № 2-ТП (воздух), "
               "ОКУД 0609012",
        "effective": 2019,
        "marks": ["2-тп", "воздух"],
    },
    "2tp-water": {
        "npa": "приказ Росстата от 02.10.2024 № 445 в ред. от 06.08.2026 "
               "№ 473 — форма № 2-ТП (водхоз), ОКУД 0609060",
        "effective": 2025,
        "marks": ["2-тп", "водхоз"],
    },
    "4-oos": {
        "npa": "приказ Росстата № 346 (ред. 22.07.2025) — форма № 4-ОС",
        "effective": 2025,
        "marks": ["4-ос", "охран"],
    },
    "cadastre-spb": {
        "npa": "распоряжение Комитета по природопользованию СПб от 26.04.2023 "
               "№ 87-р в ред. от 14.02.2025 № 28-р — региональный кадастр "
               "отходов (срок 31 марта)",
        "effective": 2025,
        "marks": ["кадастр"],
    },
    "waste-passport": {
        "npa": "приказ Минприроды России от 15.05.2026 № 286 (с 01.09.2026; "
               "до неё — приказ от 08.12.2020 № 1026) — паспорт отходов "
               "I–IV классов",
        "effective": 2026,
        "switch_date": "2026-09-01",
        "marks": ["паспорт", "отход"],
    },
    "nmu": {
        "npa": "приказы Минприроды России от 28.11.2025 № 662 (форма плана) "
               "и от 26.11.2025 № 651 (требования), с 01.03.2026; № 811 "
               "утратил силу",
        "effective": 2026,
        "marks": ["нму", "мероприят"],
    },
    "pek-program": {
        "npa": "приказ Минприроды России от 18.02.2022 № 109 в ред. от "
               "12.05.2025 — программа ПЭК",
        "effective": 2025,
        "marks": ["программа", "контрол"],
    },
    "pnoolr": {
        "npa": "методические указания к ПНООЛР (приказ Минприроды России "
               "№ 1029 от 07.12.2020)",
        "effective": 2021,
        "marks": ["норматив", "отход"],
    },
    "waste-inventory": {
        "npa": "инвентаризация отходов — по объектовому перечню "
               "(единой федеральной формы нет)",
        "effective": 0,
        "marks": ["отход"],
    },
    "air-inventory": {
        "npa": "приказ Минприроды России от 19.11.2021 № 871 "
               "(инвентаризация выбросов)",
        "effective": 2022,
        "marks": ["инвентаризац", "выброс"],
    },
}


@dataclass
class Verdict:
    code: str
    title: str
    file: str = ""
    npa: str = ""
    ok: bool = True
    level: str = "ok"          # ok | warn | error
    notes: list = field(default_factory=list)


def _blank_text(path: Path, limit: int = 20000) -> str:
    """Текст бланка для структурной проверки (по возможности, без ИИ)."""
    suffix = path.suffix.lower()
    try:
        if suffix in (".xlsx", ".xlsm"):
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    parts += [str(v) for v in row if v is not None]
                    if sum(len(p) for p in parts) > limit:
                        break
            wb.close()
            return " ".join(parts)[:limit]
        if suffix == ".xls":
            import xlrd
            book = xlrd.open_workbook(str(path))
            parts = []
            for sh in book.sheets():
                for r in range(sh.nrows):
                    parts += [str(sh.cell_value(r, c)) for c in range(sh.ncols)
                              if sh.cell_value(r, c)]
                    if sum(len(p) for p in parts) > limit:
                        break
            return " ".join(parts)[:limit]
        if suffix == ".docx":
            from docx import Document
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for t in doc.tables:
                for row in t.rows:
                    parts += [c.text for c in row.cells]
            return " ".join(parts)[:limit]
        if suffix in (".html", ".htm"):
            # бланки, сохранённые из ЛКПП, приходят страницей: снимаем теги
            # сами — text_extract html не разбирает и вернул бы пустоту
            raw = path.read_bytes()
            for enc in ("utf-8", "cp1251"):
                try:
                    text = raw.decode(enc)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                text = raw.decode("utf-8", "replace")
            text = re.sub(r"(?is)<(script|style).*?</\1>", " ", text)
            text = re.sub(r"(?s)<[^>]+>", " ", text)
            return re.sub(r"\s+", " ", text).strip()[:limit]
        from ecodoc.parsers.text_extract import extract
        return (extract(path, ocr=False).text or "")[:limit]
    except Exception:
        return ""


def check_blank(code: str, path: Path) -> Verdict:
    """Один бланк против своего НПА."""
    norm = NORMS.get(code) or {}
    v = Verdict(code=code, title=path.name, file=str(path),
                npa=norm.get("npa", ""))
    if not norm:
        v.notes.append("для этой формы НПА в реестре не описан — "
                       "проверка только вручную")
        v.level = "warn"
        return v

    text = _blank_text(path).lower().replace("ё", "е")
    # у сканов PyMuPDF отдаёт «текст» из одних переносов строк: проверять
    # признаки формы не по чему, и объявлять бланк неверным нельзя
    if not re.search(r"[a-zа-я]", text):
        v.notes.append("в бланке нет текстового слоя (скан) — признаки формы "
                       "не проверить; смотрите глазами или прогоните OCR")
        v.level = "warn"
    else:
        missing = [m for m in norm.get("marks", []) if m not in text]
        if missing:
            v.ok = False
            v.level = "error"
            v.notes.append("в бланке не найдены обязательные признаки формы: "
                           + ", ".join(f"«{m}»" for m in missing)
                           + " — возможно, это другой документ")

    # смена формы внутри года: бланк, сохранённый до даты смены, сделан ещё
    # по старой редакции — предупреждаем отдельно, годовой проверке это не видно
    switch = norm.get("switch_date")
    if switch:
        try:
            switch_day = datetime.strptime(switch, "%Y-%m-%d")
            mtime = datetime.fromtimestamp(path.stat().st_mtime)
            if mtime < switch_day:
                v.level = "warn" if v.level == "ok" else v.level
                v.notes.append(
                    f"бланк сохранён {mtime:%d.%m.%Y}, а с "
                    f"{switch_day:%d.%m.%Y} форма сменилась ({norm.get('npa')}): "
                    "сверьте бланк с новой редакцией")
        except (OSError, ValueError):
            pass

    eff = int(norm.get("effective") or 0)
    if eff:
        try:
            file_year = datetime.fromtimestamp(path.stat().st_mtime).year
            if file_year < eff:
                v.level = "warn" if v.level == "ok" else v.level
                v.notes.append(
                    f"файл бланка {file_year} года, а действующая редакция "
                    f"формы — {eff} года: сверьте бланк с НПА")
        except OSError:
            pass
    else:
        v.notes.append("форма обновляется ведомством — актуальность "
                       "подтверждается только по источнику (проверка онлайн)")
    return v


def _watch_sources_for(codes: set[str]) -> list[dict]:
    from ecodoc.watch import watcher
    try:
        return [s for s in watcher.load_sources()
                if set(s.get("forms") or []) & codes]
    except Exception:
        return []


def check_all(online: bool = False) -> dict:
    """Все бланки из папки «Формы» против НПА; online — ещё и по интернету."""
    from ecodoc.core import forms_registry
    base = forms_registry.root()
    verdicts: list[Verdict] = []
    codes: set[str] = set()
    if base is not None:
        for slot in forms_registry.scan():
            if not slot.code or not slot.files:
                continue
            codes.add(slot.code)
            for f in slot.files[:3]:              # не перечитываем всю папку
                verdicts.append(check_blank(slot.code, base / f))

    changed: list[str] = []
    online_errors: list[str] = []
    if online and codes:
        from ecodoc.watch import watcher
        for src in _watch_sources_for(codes):
            try:
                # save=False: снимок не переписываем — «изменилось» должен
                # увидеть и штатный сторож форм (ecodoc watch check)
                res = watcher.check_source(src, save=False)
            except Exception as e:
                online_errors.append(f"{src.get('id')}: {e}")
                continue
            if res.get("status") == "changed":
                changed.append(f"{res.get('name', res.get('id'))} — "
                               f"{res.get('detail') or 'страница изменилась'}")
            elif res.get("status") == "error":
                online_errors.append(f"{res.get('id')}: {res.get('detail')}")

    return {"verdicts": [vars(v) for v in verdicts],
            "changed": changed, "online_errors": online_errors,
            "checked_online": bool(online)}
