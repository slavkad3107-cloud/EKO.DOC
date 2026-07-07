"""Извлечение текста из приложенных документов: pdf / docx / doc / jpg / png.

Все «тяжёлые» библиотеки импортируются лениво, чтобы отсутствие, например,
Tesseract не ломало работу с PDF/DOCX.
"""
from __future__ import annotations

from pathlib import Path


class ExtractedDoc:
    def __init__(self, path: Path, text: str, pages: list[str], method: str):
        self.path = Path(path)
        self.text = text
        self.pages = pages          # текст постранично (для провенанса)
        self.method = method        # как извлекли: pdf-text / pdf-ocr / docx / ocr

    def __repr__(self) -> str:
        return f"<ExtractedDoc {self.path.name} method={self.method} chars={len(self.text)}>"


def extract(path: str | Path, ocr: bool = True) -> ExtractedDoc:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(p, ocr=ocr)
    if suffix in (".docx",):
        return _extract_docx(p)
    if suffix in (".doc",):
        return _extract_doc(p)
    if suffix in (".xlsx", ".xlsm"):
        return _extract_xlsx(p)
    if suffix in (".xls",):
        return _extract_xls(p)
    if suffix in (".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"):
        return _extract_image(p)
    if suffix in (".txt", ".csv"):
        text = p.read_text(encoding="utf-8", errors="replace")
        return ExtractedDoc(p, text, [text], "txt")
    if suffix in (".xml",):
        return _extract_xml(p)
    if suffix in (".rtf",):
        return _extract_rtf(p)
    raise ValueError(f"Неподдерживаемый формат: {suffix}")


def _extract_xlsx(p: Path) -> ExtractedDoc:
    import openpyxl

    wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"═ Лист: {ws.title} ═")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c not in (None, "")]
            if cells:
                parts.append("\t".join(cells))
    wb.close()
    text = "\n".join(parts)
    return ExtractedDoc(p, text, [text], "xlsx")


def _extract_xls(p: Path) -> ExtractedDoc:
    import xlrd

    book = xlrd.open_workbook(str(p))
    parts = []
    for sh in book.sheets():
        parts.append(f"═ Лист: {sh.name} ═")
        for r in range(sh.nrows):
            cells = [str(sh.cell_value(r, c)) for c in range(sh.ncols)]
            cells = [c for c in cells if c not in ("", "0.0")]
            if cells:
                parts.append("\t".join(cells))
    text = "\n".join(parts)
    return ExtractedDoc(p, text, [text], "xls")


def _extract_xml(p: Path) -> ExtractedDoc:
    import re

    raw = p.read_text(encoding="utf-8", errors="replace")
    # оставляем текстовое содержимое + значения атрибутов (в них часто ИНН/КПП)
    text = re.sub(r"<[^>]+>", " ", raw)
    text = re.sub(r"\s+", " ", text).strip()
    return ExtractedDoc(p, text or raw, [text], "xml")


def _extract_rtf(p: Path) -> ExtractedDoc:
    raw = p.read_text(encoding="utf-8", errors="replace")
    try:
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(raw)
    except ImportError:
        import re
        # грубая очистка RTF без зависимости: убрать управляющие слова и группы
        text = re.sub(r"\\[a-z]+-?\d* ?", " ", raw)
        text = re.sub(r"[{}]", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
    return ExtractedDoc(p, text, [text], "rtf")


def _extract_pdf(p: Path, ocr: bool) -> ExtractedDoc:
    import fitz  # PyMuPDF

    doc = fitz.open(p)
    pages = [page.get_text("text") for page in doc]
    text = "\n".join(pages)
    method = "pdf-text"
    # скан без текстового слоя -> OCR постранично
    if ocr and len(text.strip()) < 40:
        ocr_pages = []
        for page in doc:
            pix = page.get_pixmap(dpi=300)
            ocr_pages.append(_ocr_pixmap(pix))
        pages = ocr_pages
        text = "\n".join(pages)
        method = "pdf-ocr"
    doc.close()
    return ExtractedDoc(p, text, pages, method)


def _extract_docx(p: Path) -> ExtractedDoc:
    from docx import Document

    d = Document(str(p))
    parts = [par.text for par in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(parts)
    return ExtractedDoc(p, text, [text], "docx")


def _extract_doc(p: Path) -> ExtractedDoc:
    # старый .doc: Word COM (Windows). Файл копируем в КОРОТКИЙ путь — Word
    # не открывает файлы с длинным путём (>~218 симв.), а папки площадок у нас
    # называются полным адресом. Если Word не сработал — пробуем LibreOffice.
    import shutil
    import tempfile

    short = Path(tempfile.gettempdir()) / f"ed_{abs(hash(p.name)) % 100000}.doc"
    try:
        shutil.copy2(p, short)
    except OSError:
        short = p
    try:
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        try:
            word.DisplayAlerts = False
        except Exception:
            pass
        try:
            doc = word.Documents.Open(str(short), ReadOnly=True,
                                      AddToRecentFiles=False, ConfirmConversions=False)
            text = doc.Content.Text
            doc.Close(False)
        finally:
            word.Quit()
        return ExtractedDoc(p, text, [text], "doc-com")
    except Exception as com_exc:
        # запасной путь: LibreOffice → txt
        txt = _soffice_to_txt(short)
        if txt is not None:
            return ExtractedDoc(p, txt, [txt], "doc-soffice")
        raise RuntimeError(
            f"Не удалось прочитать .doc ({com_exc}). "
            f"Сконвертируйте {p.name} в .docx или .pdf."
        ) from com_exc
    finally:
        if short != p:
            try:
                short.unlink()
            except OSError:
                pass


def _soffice_to_txt(src: Path) -> str | None:
    """Конвертировать .doc/.rtf в текст через LibreOffice, если он установлен."""
    import shutil
    import subprocess
    import tempfile

    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if not soffice:
        for c in (r"C:\Program Files\LibreOffice\program\soffice.exe",
                  r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"):
            if Path(c).exists():
                soffice = c
                break
    if not soffice:
        return None
    outdir = Path(tempfile.mkdtemp(prefix="ed_txt_"))
    try:
        subprocess.run([soffice, "--headless", "--convert-to", "txt:Text",
                        "--outdir", str(outdir), str(src)],
                       capture_output=True, timeout=120)
        produced = outdir / (src.stem + ".txt")
        if produced.exists():
            return produced.read_text(encoding="utf-8", errors="replace")
    except (subprocess.SubprocessError, OSError):
        pass
    finally:
        import shutil as _sh
        _sh.rmtree(outdir, ignore_errors=True)
    return None


def _extract_image(p: Path) -> ExtractedDoc:
    from PIL import Image

    text = _ocr_image(Image.open(p))
    return ExtractedDoc(p, text, [text], "ocr")


_TESS_READY = None


def _setup_tesseract():
    """Найти tesseract.exe где угодно и настроить pytesseract.

    Возвращает 'rus+eng' | 'eng' | None (не найден). Кэшируется.
    """
    global _TESS_READY
    if _TESS_READY is not None:
        return _TESS_READY
    import os
    import shutil

    try:
        import pytesseract
    except ImportError:
        _TESS_READY = None
        return None

    cmd = shutil.which("tesseract")
    if not cmd:
        for c in (r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                  r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                  os.path.expandvars(r"%LOCALAPPDATA%\Programs\Tesseract-OCR\tesseract.exe"),
                  os.path.expandvars(r"%LOCALAPPDATA%\Tesseract-OCR\tesseract.exe")):
            if Path(c).exists():
                cmd = c
                break
    if not cmd:
        _TESS_READY = None
        return None
    pytesseract.pytesseract.tesseract_cmd = cmd
    # русский язык может лежать в пользовательской папке (установка без прав
    # админа не пишет в Program Files) — указываем её через TESSDATA_PREFIX
    user_td = Path(os.path.expandvars(r"%LOCALAPPDATA%\EcoDoc\tessdata"))
    if (user_td / "rus.traineddata").exists():
        os.environ["TESSDATA_PREFIX"] = str(user_td)
    try:
        langs = pytesseract.get_languages(config="")
    except Exception:
        langs = []
    _TESS_READY = "rus+eng" if "rus" in langs else "eng"
    return _TESS_READY


def _ocr_image(img) -> str:
    lang = _setup_tesseract()
    if lang is None:
        raise RuntimeError(
            "Tesseract-OCR не установлен — сканы и фото не распознаются. "
            "Установите его (в установщике ЭКО.DOC это делается автоматически) "
            "или скачайте: https://github.com/UB-Mannheim/tesseract/wiki")
    import pytesseract
    return pytesseract.image_to_string(img, lang=lang)


def _ocr_pixmap(pix) -> str:
    from io import BytesIO

    from PIL import Image

    img = Image.open(BytesIO(pix.tobytes("png")))
    return _ocr_image(img)
